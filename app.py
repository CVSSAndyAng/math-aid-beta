from __future__ import annotations
from datetime import date
from urllib.parse import urlencode
from urllib.request import Request, urlopen
from types import SimpleNamespace
import copy

import ast
import base64
import hashlib
import inspect
import math
import json
import os
import re
pmatrix = "pmatrix"  # MathIO matrix environment token

import html
import secrets
import zipfile
from io import BytesIO
from collections import Counter
from datetime import datetime, timezone
from typing import Any

import pandas as pd
import streamlit as st
import streamlit.components.v1 as st_components_v1
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt, Cm
from pypdf import PdfReader, PdfWriter

from gemini_service import (
    DEFAULT_MODEL,
    GeminiAnalysis,
    GuidedSolution,
    GeminiTutorError,
    MathVerificationResult,
    PaperQuestionSolution,
    PaperMarkPoint,
    ExamPaperDraft,
    PracticeEvaluation,
    QuestionDetectionResult,
    QuestionFeasibilityResult,
    TargetedPracticeQuestion,
    VisualExplanationResult,
    UploadedAsset,
    analyze_submission,
    assess_question_feasibility,
    detect_questions_in_assets,
    evaluate_practice_attempt,
    generate_followup_practice_question,
    generate_guided_solution,
    generate_paper_question_solution,
    generate_exam_paper_draft,
    generate_visual_explanation,
    get_api_key,
    required_parts_for_question,
    verify_question_math,
)
from offline_engine import (
    TRACKS,
    AttemptResult,
    Question,
    analyze_own_algebra_question,
    evaluate_attempt,
    generate_question,
    generate_similar,
    official_topic_code,
    topics_for_track,
)
from syllabus_topics import paper_setting_topics, topic_notes_for_selection



# ---------------------------------------------------------------------------
# GeoGebra external graph renderer
# ---------------------------------------------------------------------------
# GeoGebra is loaded in the browser from its official deployggb.js endpoint.
# The mathematical expression is passed as data (not injected into HTML/JS).
# The component exports the active Graphics View back to Python as PNG base64.
_GEOGEBRA_COMPONENT = None
_GEOGEBRA_COMPONENT_AVAILABLE = False

try:
    if hasattr(st, "components") and hasattr(st.components, "v2") and hasattr(st.components.v2, "component"):
        _GEOGEBRA_COMPONENT = st.components.v2.component(
            "math_advisor_geogebra_graph_capture",
            html="""
                <div class="math-advisor-ggb">
                    <div class="ggb-canvas"></div>
                    <div class="ggb-note">Loading GeoGebra graph…</div>
                </div>
            """,
            css="""
                .math-advisor-ggb {
                    width: 100%;
                    font-family: var(--st-font);
                    color: var(--st-text-color);
                }
                .ggb-canvas {
                    width: 100%;
                    min-height: 430px;
                    background: white;
                    border: 1px solid rgba(49, 51, 63, .16);
                    border-radius: 8px;
                    overflow: hidden;
                }
                .ggb-note {
                    margin-top: .35rem;
                    font-size: .82rem;
                    color: var(--st-secondary-text-color);
                }
            """,
            js=r"""
                function loadScriptOnce(src) {
                    return new Promise((resolve, reject) => {
                        if (window.GGBApplet) {
                            resolve();
                            return;
                        }
                        const existing = document.querySelector(`script[data-math-advisor-ggb="${src}"]`);
                        if (existing) {
                            existing.addEventListener("load", resolve, {once: true});
                            existing.addEventListener("error", reject, {once: true});
                            return;
                        }
                        const script = document.createElement("script");
                        script.src = src;
                        script.async = true;
                        script.dataset.mathAdvisorGgb = src;
                        script.onload = resolve;
                        script.onerror = reject;
                        document.head.appendChild(script);
                    });
                }

                export default async function(component) {
                    const { parentElement, data, setStateValue } = component;
                    const root = parentElement.querySelector(".ggb-canvas");
                    const note = parentElement.querySelector(".ggb-note");
                    if (!root || !data) return;

                    const signature = String(data.signature || "");
                    if (root.dataset.signature === signature && root.dataset.ready === "true") {
                        return;
                    }
                    root.dataset.signature = signature;
                    root.dataset.ready = "false";
                    root.innerHTML = "";
                    note.textContent = "Loading GeoGebra graph…";

                    try {
                        await loadScriptOnce("https://www.geogebra.org/apps/deployggb.js");

                        const width = Math.max(620, Math.floor(root.getBoundingClientRect().width || 760));
                        const height = Number(data.height || 450);

                        const params = {
                            appName: "graphing",
                            width,
                            height,
                            showToolBar: false,
                            showAlgebraInput: false,
                            showMenuBar: false,
                            showResetIcon: false,
                            enableRightClick: false,
                            enableLabelDrags: false,
                            enableShiftDragZoom: true,
                            language: "en",
                            appletOnLoad: function(api) {
                                try {
                                    api.setErrorDialogsActive(false);
                                    api.setCoordSystem(
                                        Number(data.xmin),
                                        Number(data.xmax),
                                        Number(data.ymin),
                                        Number(data.ymax)
                                    );
                                    api.setAxesVisible(true, true);
                                    api.setGridVisible(Boolean(data.grid));

                                    const commands = Array.isArray(data.commands) ? data.commands : [];
                                    let ok = commands.length > 0;
                                    commands.forEach((cmd) => {
                                        const success = api.evalCommand(String(cmd));
                                        ok = ok && Boolean(success);
                                    });

                                    if (!ok) {
                                        note.textContent = "GeoGebra could not evaluate one or more functions. Local graph fallback will be used.";
                                        setStateValue("capture", {
                                            ok: false,
                                            status: "eval_failed",
                                            signature,
                                            png_base64: ""
                                        });
                                        return;
                                    }

                                    // Allow the graphics view to paint before exporting.
                                    window.setTimeout(() => {
                                        try {
                                            const png = api.getPNGBase64(2, false, 300);
                                            if (!png) throw new Error("Empty PNG export");
                                            root.dataset.ready = "true";
                                            note.textContent = "GeoGebra graph captured for the generated paper.";
                                            setStateValue("capture", {
                                                ok: true,
                                                status: "captured",
                                                signature,
                                                png_base64: png
                                            });
                                        } catch (exportError) {
                                            note.textContent = "GeoGebra export failed. Local graph fallback will be used.";
                                            setStateValue("capture", {
                                                ok: false,
                                                status: "export_failed",
                                                signature,
                                                png_base64: ""
                                            });
                                        }
                                    }, 550);
                                } catch (apiError) {
                                    note.textContent = "GeoGebra graph construction failed. Local graph fallback will be used.";
                                    setStateValue("capture", {
                                        ok: false,
                                        status: "construction_failed",
                                        signature,
                                        png_base64: ""
                                    });
                                }
                            }
                        };

                        const applet = new window.GGBApplet(params, true);
                        applet.inject(root);
                    } catch (loadError) {
                        note.textContent = "GeoGebra is unavailable. Local graph fallback will be used.";
                        setStateValue("capture", {
                            ok: false,
                            status: "load_failed",
                            signature,
                            png_base64: ""
                        });
                    }
                }
            """,
        )
        _GEOGEBRA_COMPONENT_AVAILABLE = True
except Exception:
    _GEOGEBRA_COMPONENT = None
    _GEOGEBRA_COMPONENT_AVAILABLE = False




# 2027 Singapore-Cambridge Secondary Education Certificate (SEC) mathematics tracks.
# SEAB 2027 subject codes:
# G1 Mathematics K110; G2 Mathematics K210; G3 Mathematics K310;
# G2 Additional Mathematics K232; G3 Additional Mathematics K341.
SEC_2027_TRACKS = {
    "2027 SEC · G1 Mathematics (K110)": {
        "engine_code": "NT",
        "subject_code": "K110",
        "subject": "Mathematics",
        "level": "G1",
        "year": 2027,
        "reference_2026": "4046",
        "strands": ["Number and Algebra", "Geometry and Measurement", "Statistics and Probability"],
        "offline_supported": True,
    },
    "2027 SEC · G2 Mathematics (K210)": {
        "engine_code": "NA",
        "subject_code": "K210",
        "subject": "Mathematics",
        "level": "G2",
        "year": 2027,
        "reference_2026": "4045",
        "strands": ["Number and Algebra", "Geometry and Measurement", "Statistics and Probability"],
        "offline_supported": True,
    },
    "2027 SEC · G3 Mathematics (K310)": {
        "engine_code": "O",
        "subject_code": "K310",
        "subject": "Mathematics",
        "level": "G3",
        "year": 2027,
        "reference_2026": "4052",
        "strands": ["Number and Algebra", "Geometry and Measurement", "Statistics and Probability"],
        "offline_supported": True,
    },
    "2027 SEC · G2 Additional Mathematics (K232)": {
        "engine_code": "G2A",
        "subject_code": "K232",
        "subject": "Additional Mathematics",
        "level": "G2",
        "year": 2027,
        "reference_2026": "4051",
        "strands": ["Algebra", "Geometry and Trigonometry", "Calculus"],
        "offline_supported": True,
    },
    "2027 SEC · G3 Additional Mathematics (K341)": {
        "engine_code": "G3A",
        "subject_code": "K341",
        "subject": "Additional Mathematics",
        "level": "G3",
        "year": 2027,
        "reference_2026": "4049",
        "strands": ["Algebra", "Geometry and Trigonometry", "Calculus"],
        "offline_supported": True,
    },
}

# Keep the existing 2026 O/N-Level tracks available during the transition.
LEGACY_TRACK_INFO = {
    label: {
        "engine_code": code,
        "subject_code": {"O": "4052", "NA": "4045", "NT": "4046"}.get(code, str(code)),
        "subject": "Mathematics",
        "level": {"O": "O-Level", "NA": "N(A)-Level", "NT": "N(T)-Level"}.get(code, "Legacy"),
        "year": 2026,
        "reference_2026": {"O": "4052", "NA": "4045", "NT": "4046"}.get(code, str(code)),
        "strands": ["Number and Algebra", "Geometry and Measurement", "Statistics and Probability"],
        "offline_supported": True,
    }
    for label, code in TRACKS.items()
}

APP_TRACKS = {**SEC_2027_TRACKS, **LEGACY_TRACK_INFO}


def selected_track_info(label: str) -> dict:
    return APP_TRACKS[label]


def is_additional_math_track(label: str) -> bool:
    return selected_track_info(label)["subject"] == "Additional Mathematics"



st.set_page_config(
    page_title="Singapore SEC / O-N Level Math Tutor — Gemini + Offline",
    page_icon="🇸🇬",
    layout="wide",
    initial_sidebar_state="auto",
)

st.markdown(
    """
<style>
:root {
  --omt-ink: #172033;
  --omt-muted: #667085;
  --omt-border: #e4e7ec;
  --omt-surface: #ffffff;
  --omt-soft: #f7f9fc;
  --omt-brand: #3b5ccc;
  --omt-brand-soft: #eef2ff;
  --omt-success: #12805c;
  --omt-success-soft: #ecfdf3;
  --omt-warn: #b25e09;
  --omt-warn-soft: #fff7ed;
  --omt-danger: #c2414b;
  --omt-danger-soft: #fff1f2;
  --omt-radius: 18px;
  --omt-shadow: 0 10px 30px rgba(23, 32, 51, .07);
}

html, body, [class*="css"] { color: var(--omt-ink); }
[data-testid="stAppViewContainer"] { background: #f6f8fc; }
[data-testid="stHeader"] { background: rgba(246,248,252,.86); backdrop-filter: blur(10px); }
.block-container { padding-top: 1.25rem; padding-bottom: 4rem; max-width: 1240px; }

h1, h2, h3, h4 { letter-spacing: -.025em; color: var(--omt-ink); }
h1 { font-weight: 800 !important; }
h2, h3 { font-weight: 740 !important; }
p, li { line-height: 1.58; }

[data-testid="stSidebar"] { background: linear-gradient(180deg, #f9fbff 0%, #f4f6fb 100%); border-right: 1px solid var(--omt-border); }
[data-testid="stSidebar"] .block-container { padding-top: 1rem; }

.omt-hero {
  background: radial-gradient(circle at 92% 12%, rgba(112, 132, 255, .18), transparent 28%),
              linear-gradient(135deg, #ffffff 0%, #f7f8ff 100%);
  border: 1px solid #e1e6f4;
  border-radius: 24px;
  padding: 1.55rem 1.7rem;
  box-shadow: var(--omt-shadow);
  margin: .25rem 0 1.25rem;
}
.omt-eyebrow { color: var(--omt-brand); font-weight: 750; font-size: .82rem; letter-spacing: .08em; text-transform: uppercase; margin-bottom: .35rem; }
.omt-hero h1 { margin: 0 0 .45rem; font-size: clamp(1.8rem, 3.3vw, 2.75rem); line-height: 1.08; }
.omt-hero p { color: var(--omt-muted); font-size: 1.02rem; margin: 0; max-width: 780px; }
.omt-chip-row { display:flex; flex-wrap:wrap; gap:.45rem; margin-top:1rem; }
.omt-chip { display:inline-flex; align-items:center; gap:.35rem; background:#fff; border:1px solid var(--omt-border); border-radius:999px; padding:.38rem .68rem; font-size:.82rem; color:#475467; }

.omt-side-brand { padding:.45rem .15rem .85rem; }
.omt-side-brand .title { font-size:1.2rem; font-weight:800; letter-spacing:-.02em; }
.omt-side-brand .sub { color:var(--omt-muted); font-size:.86rem; line-height:1.45; margin-top:.25rem; }
.omt-status-pill { display:flex; align-items:center; gap:.45rem; border-radius:12px; padding:.62rem .72rem; font-size:.84rem; margin:.35rem 0; }
.omt-status-pill.good { background:var(--omt-success-soft); color:#116149; border:1px solid #c9f1df; }
.omt-status-pill.neutral { background:var(--omt-brand-soft); color:#4353a3; border:1px solid #dfe4ff; }

.omt-section-kicker { color:var(--omt-brand); font-weight:750; font-size:.76rem; text-transform:uppercase; letter-spacing:.08em; }
.omt-section-title { font-size:1.45rem; font-weight:780; letter-spacing:-.025em; margin:.08rem 0 .25rem; }
.omt-section-copy { color:var(--omt-muted); margin-bottom:.9rem; }

/* Offline practice: keep instruction text and MathIO expression compact. */
[data-testid="stVerticalBlockBorderWrapper"] h1,
[data-testid="stVerticalBlockBorderWrapper"] h2,
[data-testid="stVerticalBlockBorderWrapper"] h3,
[data-testid="stVerticalBlockBorderWrapper"] h4,
[data-testid="stVerticalBlockBorderWrapper"] p {
  margin-bottom: .25rem;
}
[data-testid="stVerticalBlockBorderWrapper"] [data-testid="stCustomComponentV1"] {
  margin-top: 0 !important;
  margin-bottom: .15rem !important;
}

[data-testid="stVerticalBlockBorderWrapper"] [data-testid="stCustomComponentV1"] {
  margin-top: .05rem !important;
  margin-bottom: .05rem !important;
}
[data-testid="stVerticalBlockBorderWrapper"] p {
  margin-top: .15rem;
  margin-bottom: .3rem;
}

.omt-focus-card {
  background: linear-gradient(145deg, #ffffff, #fbfcff);
  border: 1px solid #dfe4ee;
  border-radius: 20px;
  box-shadow: 0 7px 22px rgba(23,32,51,.055);
  padding: 1rem 1.1rem;
}
.omt-focus-title { font-size:.8rem; text-transform:uppercase; letter-spacing:.07em; color:var(--omt-brand); font-weight:760; margin-bottom:.4rem; }
.omt-key-grid { display:grid; grid-template-columns:repeat(auto-fit,minmax(150px,1fr)); gap:.55rem; margin:.7rem 0; }
.omt-key-item { background:#f8faff; border:1px solid #e6eaf5; border-radius:12px; padding:.62rem .72rem; }

.omt-stage-row { display:grid; grid-template-columns:repeat(3,1fr); gap:.65rem; margin:.7rem 0 1.1rem; }
.omt-stage { border:1px solid var(--omt-border); background:#fff; border-radius:14px; padding:.72rem .8rem; min-height:66px; }
.omt-stage .name { font-weight:720; font-size:.92rem; }
.omt-stage .detail { color:var(--omt-muted); font-size:.78rem; margin-top:.18rem; }
.omt-stage.current { border-color:#aab8ff; background:#f4f6ff; box-shadow:0 0 0 2px rgba(59,92,204,.06); }
.omt-stage.done { border-color:#bfe8d8; background:#f1fbf6; }
.omt-stage.locked { opacity:.72; }

.omt-logic-break { background:var(--omt-warn-soft); border:1px solid #fed7aa; border-left:5px solid #f59e0b; border-radius:14px; padding:.8rem 1rem; margin:.7rem 0 1rem; }
.omt-success-card { background:var(--omt-success-soft); border:1px solid #cceedd; border-left:5px solid #1ca878; border-radius:14px; padding:.8rem 1rem; margin:.7rem 0 1rem; }

/* Streamlit cards and widgets */
[data-testid="stVerticalBlockBorderWrapper"] { border-color: var(--omt-border) !important; border-radius: var(--omt-radius) !important; background: var(--omt-surface); box-shadow: 0 4px 18px rgba(23,32,51,.035); }
[data-testid="stMetric"] { background:#fff; border:1px solid var(--omt-border); border-radius:15px; padding:.75rem .85rem; }
[data-testid="stMetricValue"] { font-weight:780; letter-spacing:-.03em; }
[data-testid="stExpander"] { border:1px solid var(--omt-border) !important; border-radius:14px !important; background:#fff; overflow:hidden; }
[data-testid="stFileUploader"] { border-radius:16px; }
.stTextArea textarea, .stTextInput input { border-radius:12px !important; }

.stButton > button { border-radius:12px; font-weight:680; border-color:#d8deea; transition:transform .12s ease, box-shadow .12s ease; }
.stButton > button:hover { transform:translateY(-1px); box-shadow:0 5px 14px rgba(23,32,51,.08); }
.stButton > button[kind="primary"] { background:linear-gradient(135deg,#425fd6,#5b6fe8); border:none; color:white; box-shadow:0 7px 16px rgba(66,95,214,.20); }

/* Tabs */
button[data-baseweb="tab"] { border-radius:12px 12px 0 0; padding:.7rem .9rem !important; font-weight:650; }
button[data-baseweb="tab"][aria-selected="true"] { background:#eef2ff; color:#334bb3; }

/* Alerts */
[data-testid="stAlert"] { border-radius:14px; }

@media (max-width: 1100px) {
  .block-container { max-width:100%; padding-left:1rem; padding-right:1rem; }
}
@media (max-width: 720px) {
  .omt-hero { padding:1.15rem 1.05rem; border-radius:18px; }
  .omt-stage-row { grid-template-columns:1fr; }
}
@media (pointer: coarse) {
  button, [role="button"] { min-height:46px; }
  input, textarea, select { font-size:16px !important; }
}

/* Generic question text/MathIO segmentation */
[data-testid="stVerticalBlockBorderWrapper"] p {
  margin-top: .08rem !important;
  margin-bottom: .16rem !important;
}
[data-testid="stVerticalBlockBorderWrapper"] [data-testid="stCustomComponentV1"] {
  margin-top: -.12rem !important;
  margin-bottom: .08rem !important;
}

/* Compact mixed question layout: avoid one Streamlit row per short maths fragment. */
[data-testid="stVerticalBlockBorderWrapper"] p {
  margin-top: 0 !important;
  margin-bottom: .22rem !important;
  line-height: 1.55 !important;
}
[data-testid="stVerticalBlockBorderWrapper"] [data-testid="stCustomComponentV1"] {
  margin-top: .05rem !important;
  margin-bottom: .12rem !important;
}
</style>
""",
    unsafe_allow_html=True,
)

MAX_FILE_BYTES = 12 * 1024 * 1024
MAX_TOTAL_BYTES = 30 * 1024 * 1024


_MATHIO_RENDER_SEQ = 0
_mathio_display_component = None
_mathio_rich_component = None


def _strip_math_transport_delimiters(text: str) -> str:
    """Remove model transport delimiters before sending maths to the MathIO view."""
    if not text:
        return ""
    value = str(text).strip()
    pairs = ((r"\(", r"\)"), (r"\[", r"\]"), ("$$", "$$"), ("$", "$"))
    for left, right in pairs:
        if value.startswith(left) and value.endswith(right) and len(value) >= len(left) + len(right):
            value = value[len(left): len(value) - len(right)]
            break
    return value.strip()


def _next_mathio_key(text: str) -> str:
    global _MATHIO_RENDER_SEQ
    _MATHIO_RENDER_SEQ += 1
    digest = hashlib.sha1(str(text).encode("utf-8", errors="ignore")).hexdigest()[:10]
    return f"mathio_view_{_MATHIO_RENDER_SEQ}_{digest}"



_MATH_UPRIGHT_WORDS = {
    "sin", "cos", "tan", "sec", "cosec", "cot",
    "log", "ln", "exp", "lim", "max", "min",
    "cm", "mm", "km", "m", "kg", "g", "s", "h",
    "rad", "deg",
}


def _normalise_math_variable_italics(source: str) -> str:
    """Normalise mathematical typography for MathIO/MathJax safely."""
    text = str(source or "")
    if not text.strip():
        return text

    function_map = {
        "sin": r"\sin",
        "cos": r"\cos",
        "tan": r"\tan",
        "log": r"\log",
        "ln": r"\ln",
        "exp": r"\exp",
    }
    for name, latex in function_map.items():
        text = re.sub(
            rf"(?<!\\)\b{name}\b",
            lambda _m, repl=latex: repl,
            text,
            flags=re.IGNORECASE,
        )

    text = re.sub(
        r"\\+(sin|cos|tan|log|ln|exp)\b",
        lambda m: "\\" + m.group(1),
        text,
    )

    unit_map = {
        "mm": r"\mathrm{mm}",
        "cm": r"\mathrm{cm}",
        "km": r"\mathrm{km}",
        "kg": r"\mathrm{kg}",
        "rad": r"\mathrm{rad}",
    }
    for unit, latex in unit_map.items():
        text = re.sub(
            rf"(?<![A-Za-z\\]){re.escape(unit)}\b",
            lambda _m, repl=latex: repl,
            text,
        )

    return text




def _normalise_unit_braces(text: str) -> str:
    """Remove stray literal braces around ordinary units in question text.

    Examples:
      12 {cm} -> 12 cm
      5{kg}   -> 5 kg
      30 {degrees} -> 30 degrees

    LaTeX commands such as \\text{cm}, \\mathrm{cm}, \\frac{...}{...}
    are preserved because their braces are preceded by a backslash command.
    """
    value = str(text or "")

    units = (
        "mm", "cm", "m", "km",
        "mg", "g", "kg",
        "ml", "mL", "l", "L",
        "s", "min", "h",
        "degrees", "degree",
    )

    pattern = re.compile(
        r"(?<!\\\\)(?<![A-Za-z])\\{\\s*(" + "|".join(re.escape(u) for u in units) + r")\\s*\\}",
        flags=re.IGNORECASE,
    )
    value = pattern.sub(lambda m: " " + m.group(1), value)

    # Tidy duplicate spacing introduced by replacement.
    value = re.sub(r"[ \t]{2,}", " ", value)
    return value


def render_mathio(text: str) -> None:
    """Render mathematics with the read-only MathIO/MathLive view; never expose source notation."""
    value = _strip_math_transport_delimiters(text)
    if not value:
        return
    value = _normalise_math_variable_italics(value)
    if _mathio_display_component is None:
        st.info("Equation view is temporarily unavailable. Reload the page to restore the maths display.")
        return
    _mathio_display_component(
        data={"math": value},
        default={},
        key=_next_mathio_key(value),
        width="stretch",
        height="content",
    )


_MATHIO_MIXED_PATTERN = re.compile(
    r"(\\\[[\s\S]*?\\\]|\\\([\s\S]*?\\\)|\$\$[\s\S]*?\$\$|\$[^$\n]+?\$)"
)


def _contains_raw_math_source(text: str) -> bool:
    """Detect source-style maths commands that should never be shown directly to students."""
    return bool(re.search(r"\\(?:frac|sqrt|times|div|cdot|theta|alpha|beta|gamma|pi|sin|cos|tan|log|ln|leq|geq|neq|pm|text|overline|bar|angle|circ)\b|\^\{|_\{", text or ""))



def _normalize_generated_math_text(value: str) -> str:
    """Repair common generated maths artifacts before display."""
    text = str(value or "").strip()
    if not text:
        return ""

    text = re.sub(r"\\+(?:textbullet|bullet|dots|ldots|cdots)\b\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\.{3,}", "", text)

    # Repair common text-maths leakage.
    text = re.sub(r"(\d+(?:\.\d+)?)\s*degrees\b", r"\1^{\\circ}", text, flags=re.IGNORECASE)
    text = re.sub(r"\bangle\s+([A-Z]{2,4})\b", r"\\angle \1", text)
    text = re.sub(r"\btheta\b", r"\\theta", text, flags=re.IGNORECASE)
    text = re.sub(r"(?<!\\)\barctan\s*\(", r"\\arctan(", text, flags=re.IGNORECASE)
    text = re.sub(r"(?<!\\)\barcsin\s*\(", r"\\arcsin(", text, flags=re.IGNORECASE)
    text = re.sub(r"(?<!\\)\barccos\s*\(", r"\\arccos(", text, flags=re.IGNORECASE)
    text = re.sub(r"\bpi\b", r"\\pi", text, flags=re.IGNORECASE)
    text = re.sub(r"\{(cm|mm|km|kg|m|g|s|h|ml|l)\}", r"\1", text, flags=re.IGNORECASE)
    text = re.sub(r"(?<=\d)(cm|mm|km|kg|m|g|s|h|ml|l)\b", r" \1", text, flags=re.IGNORECASE)

    # Repair fused English around mathematical tokens.
    text = re.sub(r"(?i)(\^\{\\circ\})and(?=\\angle)", r"\1 and ", text)
    text = re.sub(r"(?i)\band(?=\\angle)", "and ", text)

    return re.sub(r"\s{2,}", " ", text).strip()


_AUTO_MATHIO_FRAGMENT_RE = re.compile(
    r"""
    (
        \\angle\s*[A-Z]{2,4}
        (?:\s*=\s*[^,.;:\n]+)+
        |
        [A-Za-z][A-Za-z0-9_]*\s*=\s*[^,.;:\n]+
        |
        \([^()\n]{1,120}\)\^\{?[-+]?\d+\}?
        |
        (?<!\w)[A-Za-z0-9]+\^\{?[-+]?\d+\}?
        |
        (?<![\w\\])(?:-?\d+(?:\.\d+)?[A-Za-z]|[A-Za-z])(?:\^\{?[-+]?\d+\}?)?(?:\s*[+\-]\s*(?:-?\d+(?:\.\d+)?[A-Za-z]|[A-Za-z])(?:\^\{?[-+]?\d+\}?)?)+(?!\w)
        |
        \\frac\{[^{}\n]+\}\{[^{}\n]+\}
        |
        \\sqrt\{[^{}\n]+\}
        |
        \\(?:sin|cos|tan|log|ln)\s*(?:\\left)?\([^)\n]+(?:\\right)?\)
        |
        \\(?:pi|theta|alpha|beta|gamma)\b
        |
        \(?\s*-?\d+(?:\.\d+)?\s*\)?\s*\\times\s*10\^\{?[-+]?\d+\}?
        |
        (?<!\w)\d+\s*:\s*\d+(?!\w)
        |
        (?<!\w)-?\d+(?:\.\d+)?\s*(?:cm|mm|m|km|g|kg|s|h|ml|l|%)(?:\^2|\^3)?\b
        |
        (?<!\w)\d+(?:\.\d+)?\^\{?\\circ\}?
    )
    """,
    re.VERBOSE,
)


def _auto_mathio_markup(value: str) -> str:
    """Wrap undelimited mathematical fragments in MathIO inline delimiters."""
    text = _normalize_generated_math_text(value)
    if not text:
        return ""

    if _MATHIO_MIXED_PATTERN.search(text):
        return text

    pieces: list[str] = []
    cursor = 0
    for match in _AUTO_MATHIO_FRAGMENT_RE.finditer(text):
        if match.start() > cursor:
            pieces.append(text[cursor:match.start()])
        fragment = match.group(0).strip()
        fragment = re.sub(r"\bsin\s*\(", r"\\sin(", fragment)
        fragment = re.sub(r"\bcos\s*\(", r"\\cos(", fragment)
        fragment = re.sub(r"\btan\s*\(", r"\\tan(", fragment)
        pieces.append(r"\(" + fragment + r"\)")
        cursor = match.end()

    if cursor < len(text):
        pieces.append(text[cursor:])

    return "".join(pieces) if pieces else text


def _latex_leak_detected(value: str) -> bool:
    """True when raw maths source would otherwise be exposed as prose."""
    return bool(
        re.search(
            r"\\(?:frac|sqrt|angle|theta|alpha|beta|gamma|pi|sin|cos|tan|arcsin|arccos|arctan|log|ln|circ|times|div|leq|geq|pm)\b"
            r"|\^\{?|_\{?",
            value or "",
        )
    )


def render_mathio_mixed(text: str) -> None:
    """Render prose normally and every mathematical fragment through MathIO."""
    if not text:
        return

    value = _auto_mathio_markup(str(text))
    if not value:
        return

    if _mathio_rich_component is not None and _MATHIO_MIXED_PATTERN.search(value):
        _mathio_rich_component(
            data={"text": value},
            default={},
            key=_next_mathio_key("mixed:" + value),
            width="stretch",
            height="content",
        )
        return

    if _latex_leak_detected(value):
        # Never expose raw LaTeX as ordinary markdown.
        stripped = re.sub(r"\\\(|\\\)|\\\[|\\\]|\$\$", "", value)
        render_mathio(stripped)
        return

    st.markdown(value)



def render_math_text(text: str) -> None:
    """Compatibility wrapper: all student-facing maths now goes through MathIO."""
    render_mathio_mixed(text)


MATHLIVE_VERSION = "0.110.0"  # Patched MathLive release used by the visual equation editor.

_EQUATION_EDITOR_HTML = """
<div class="omt-math-editor">
  <div class="omt-editor-label"></div>
  <div class="omt-editor-help">Type directly into each maths box. Use the quick equation tools or tap <b>Math keyboard</b>. The full keyboard opens directly below the maths box.</div>
  <div class="omt-math-toolbar" role="toolbar" aria-label="Math equation tools">
    <button type="button" data-insert="\\frac{#@}{#?}" title="Fraction">a⁄b</button>
    <button type="button" data-insert="\\sqrt{#0}" title="Square root">√</button>
    <button type="button" data-insert="#@^{#?}" title="Power">xʸ</button>
    <button type="button" data-insert="#@_{#?}" title="Subscript">xₙ</button>
    <button type="button" data-insert="\\sin\\left(#0\\right)">sin</button>
    <button type="button" data-insert="\\cos\\left(#0\\right)">cos</button>
    <button type="button" data-insert="\\tan\\left(#0\\right)">tan</button>
    <button type="button" data-insert="\\log_{#?}\\left(#0\\right)">log</button>
    <button type="button" data-insert="\\ln\\left(#0\\right)">ln</button>
    <button type="button" data-insert="\\pi">π</button>
    <button type="button" data-insert="\\theta">θ</button>
    <button type="button" data-insert="^{\\circ}">°</button>
    <button type="button" data-insert="\\le">≤</button>
    <button type="button" data-insert="\\ge">≥</button>
    <button type="button" data-insert="\\pm">±</button>
    <button type="button" class="omt-keyboard-toggle">⌨ Math keyboard</button>
  </div>
  <div class="omt-editor-rows"></div>

  <div class="omt-full-keyboard" hidden aria-label="Full mathematics keyboard">
    <div class="omt-keyboard-row">
      <button type="button" data-kb="7">7</button><button type="button" data-kb="8">8</button><button type="button" data-kb="9">9</button>
      <button type="button" data-kb="\\times">×</button><button type="button" data-kb="\\div">÷</button><button type="button" data-kb="\\frac{#@}{#?}">a⁄b</button>
    </div>
    <div class="omt-keyboard-row">
      <button type="button" data-kb="4">4</button><button type="button" data-kb="5">5</button><button type="button" data-kb="6">6</button>
      <button type="button" data-kb="+">+</button><button type="button" data-kb="-">−</button><button type="button" data-kb="\\sqrt{#0}">√</button>
    </div>
    <div class="omt-keyboard-row">
      <button type="button" data-kb="1">1</button><button type="button" data-kb="2">2</button><button type="button" data-kb="3">3</button>
      <button type="button" data-kb="=">=</button><button type="button" data-kb="#@^{#?}">xʸ</button><button type="button" data-kb="#@_{#?}">xₙ</button>
    </div>
    <div class="omt-keyboard-row">
      <button type="button" data-kb="0">0</button><button type="button" data-kb=".">.</button><button type="button" data-kb="(">(</button>
      <button type="button" data-kb=")">)</button><button type="button" data-kb="\\pi">π</button><button type="button" data-kb="\\theta">θ</button>
    </div>
    <div class="omt-keyboard-row">
      <button type="button" data-kb="\\sin\\left(#0\\right)">sin</button><button type="button" data-kb="\\cos\\left(#0\\right)">cos</button>
      <button type="button" data-kb="\\tan\\left(#0\\right)">tan</button><button type="button" data-kb="\\log_{#?}\\left(#0\\right)">log</button>
      <button type="button" data-kb="\\ln\\left(#0\\right)">ln</button><button type="button" data-kb="^{\\circ}">°</button>
    </div>
    <div class="omt-keyboard-row">
      <button type="button" data-kb="\\le">≤</button><button type="button" data-kb="\\ge">≥</button><button type="button" data-kb="\\ne">≠</button>
      <button type="button" data-kb="\\pm">±</button><button type="button" data-kb="\\infty">∞</button><button type="button" class="omt-kb-backspace">⌫</button>
    </div>
    <div class="omt-keyboard-row">
      <button type="button" data-kb="x">x</button><button type="button" data-kb="y">y</button><button type="button" data-kb="a">a</button>
      <button type="button" data-kb="b">b</button><button type="button" data-kb="c">c</button><button type="button" class="omt-kb-close">Close</button>
    </div>
  </div>

  <div class="omt-editor-actions">
    <button type="button" class="omt-add-step">＋ Add step</button>
    <button type="button" class="omt-save-working">💾 Save working</button>
  </div>
  <div class="omt-editor-status" aria-live="polite"></div>
</div>
"""

_EQUATION_EDITOR_CSS = """
.omt-math-editor { width: 100%; font-family: var(--st-font, sans-serif); }
.omt-editor-label { font-weight: 600; margin-bottom: .3rem; }

.omt-math-toolbar { display:flex; flex-wrap:wrap; gap:.35rem; margin:.35rem 0 .65rem; align-items:center; }
.omt-math-toolbar button { border:1px solid rgba(128,128,128,.38); border-radius:.45rem; background:transparent; color:var(--st-text-color,#222); min-height:2.25rem; padding:.3rem .55rem; cursor:pointer; font-size:.92rem; }
.omt-math-toolbar button:hover, .omt-math-toolbar button:focus-visible { border-color:var(--st-primary-color,#ff4b4b); outline:none; }
.omt-math-toolbar .omt-keyboard-toggle { font-weight:600; padding-inline:.75rem; }
math-field::part(virtual-keyboard-toggle) { display:none; }
body { --keyboard-zindex: 999999; }

.omt-editor-help { color: var(--st-text-color); opacity: .72; font-size: .86rem; margin-bottom: .65rem; }
.omt-editor-row { display: grid; grid-template-columns: 3.2rem minmax(0,1fr) 2.5rem; align-items: center; gap: .45rem; margin: .45rem 0; }
.omt-step-label { font-size: .83rem; opacity: .75; }
.omt-editor-row math-field { width: 100%; min-height: 3.1rem; box-sizing: border-box; border: 1px solid rgba(128,128,128,.45); border-radius: .55rem; padding: .45rem .6rem; font-size: 1.12rem; background: var(--st-background-color, white); color: var(--st-text-color, #222); --caret-color: var(--st-primary-color, #ff4b4b); --selection-background-color: color-mix(in srgb, var(--st-primary-color, #ff4b4b) 20%, transparent); }
.omt-editor-row math-field:focus-within { outline: 2px solid color-mix(in srgb, var(--st-primary-color, #ff4b4b) 45%, transparent); outline-offset: 1px; }
.omt-remove-step { border: 0; background: transparent; cursor: pointer; font-size: 1.05rem; opacity: .65; padding: .3rem; }
.omt-remove-step:hover { opacity: 1; }
.omt-editor-actions { margin-top: .5rem; }
.omt-add-step { border: 1px solid rgba(128,128,128,.38); border-radius: .45rem; background: transparent; color: var(--st-text-color, #222); padding: .4rem .7rem; cursor: pointer; }
.omt-add-step:hover { border-color: var(--st-primary-color, #ff4b4b); }
.omt-editor-status { font-size: .78rem; opacity: .7; margin-top: .4rem; min-height: 1rem; }
@media (max-width: 640px) {
  .omt-editor-row { grid-template-columns: 2.7rem minmax(0,1fr) 2.2rem; }
  .omt-editor-row math-field { font-size: 1.05rem; }
}
@media (pointer: coarse) {
  .omt-editor-row math-field { min-height: 4rem; font-size: 1.2rem; padding: .7rem .75rem; }
  .omt-add-step, .omt-save-working, .omt-remove-step, .omt-math-toolbar button { min-height: 44px; min-width: 44px; }
}

/* Mobile/tablet equation editor improvements */
.omt-math-toolbar {
  position: sticky;
  top: 0;
  z-index: 20;
  background: var(--st-background-color, #fff);
  padding: .35rem 0;
}
.omt-keyboard-toggle { margin-left: auto; }

@media (max-width: 640px), (pointer: coarse) {
  .omt-editor-help { font-size: .92rem; }
  .omt-math-toolbar {
    flex-wrap: nowrap;
    overflow-x: auto;
    -webkit-overflow-scrolling: touch;
    scrollbar-width: thin;
    padding: .45rem 0 .55rem;
  }
  .omt-math-toolbar button {
    flex: 0 0 auto;
    min-width: 50px;
    min-height: 50px;
    font-size: 1rem;
  }
  .omt-math-toolbar .omt-keyboard-toggle {
    position: sticky;
    right: 0;
    min-width: 145px;
    background: var(--st-primary-color, #ff4b4b);
    color: #fff;
    border-color: var(--st-primary-color, #ff4b4b);
  }
  .omt-editor-row {
    grid-template-columns: 42px minmax(0, 1fr) 48px;
    gap: .35rem;
  }
  .omt-editor-row math-field {
    min-height: 58px;
    font-size: 1.18rem;
  }
}

.omt-full-keyboard{margin:.65rem 0 .75rem;padding:.55rem;border:1px solid rgba(128,128,128,.32);border-radius:.75rem;background:var(--st-background-color,#fff);box-shadow:0 6px 22px rgba(0,0,0,.08)}
.omt-full-keyboard[hidden]{display:none!important}
.omt-keyboard-row{display:grid;grid-template-columns:repeat(6,minmax(0,1fr));gap:.38rem;margin-bottom:.38rem}
.omt-keyboard-row:last-child{margin-bottom:0}
.omt-keyboard-row button{min-height:46px;border:1px solid rgba(128,128,128,.38);border-radius:.5rem;background:var(--st-secondary-background-color,#f6f7f9);color:var(--st-text-color,#222);font-size:1rem;cursor:pointer;touch-action:manipulation}
@media (max-width:640px),(pointer:coarse){
  .omt-full-keyboard{position:sticky;bottom:0;z-index:40;max-height:52vh;overflow-y:auto;overscroll-behavior:contain}
  .omt-keyboard-row button{min-height:50px;font-size:1.05rem}
}
"""

_EQUATION_EDITOR_JS = f"""
const MATHLIVE_URL = 'https://cdn.jsdelivr.net/npm/mathlive@{MATHLIVE_VERSION}/+esm';

async function ensureMathLive() {{
  if (!globalThis.__omtMathLivePromise) globalThis.__omtMathLivePromise = import(MATHLIVE_URL);
  const module = await globalThis.__omtMathLivePromise;
  if (!customElements.get('math-field')) await customElements.whenDefined('math-field');
  return module;
}}

function normalizedPayload(raw) {{
  const latex = Array.isArray(raw?.latex) ? raw.latex.map(x => String(x ?? '')) : [''];
  const ascii = Array.isArray(raw?.ascii) ? raw.ascii.map(x => String(x ?? '')) : latex.map(() => '');
  if (latex.length === 0) latex.push('');
  while (ascii.length < latex.length) ascii.push('');
  return {{ latex: latex.slice(0, 20), ascii: ascii.slice(0, 20) }};
}}

export default async function(component) {{
  const {{ parentElement, data, setStateValue }} = component;
  const label = parentElement.querySelector('.omt-editor-label');
  const rows = parentElement.querySelector('.omt-editor-rows');
  const addButton = parentElement.querySelector('.omt-add-step');
  const saveButton = parentElement.querySelector('.omt-save-working');
  const status = parentElement.querySelector('.omt-editor-status');
  const toolbar = parentElement.querySelector('.omt-math-toolbar');
  const keyboardButton = parentElement.querySelector('.omt-keyboard-toggle');
  const keyboardPanel = parentElement.querySelector('.omt-full-keyboard');
  const keyboardClose = parentElement.querySelector('.omt-kb-close');
  const keyboardBackspace = parentElement.querySelector('.omt-kb-backspace');
  const singleQuestion = Boolean(data?.singleQuestion);
  label.textContent = data?.label || (singleQuestion ? 'Question mathematics' : 'Student working');
  if (singleQuestion) {{
    if (addButton) addButton.hidden = true;
    if (saveButton) saveButton.textContent = '💾 Save question';
  }}

  let module;
  try {{ module = await ensureMathLive(); }}
  catch (err) {{ status.textContent = 'Equation editor could not load. Reload the page.'; return; }}

  const vk = module?.mathVirtualKeyboard || globalThis.mathVirtualKeyboard || null;
  if (vk) {{
    vk.layouts = [
      {{
        id: 'sg-math',
        label: 'Math tools',
        tooltip: 'Common school mathematics equations and symbols',
        rows: [
          [
            {{ latex: '\\\\frac{{#@}}{{#?}}', label: '\\\\frac{{a}}{{b}}' }},
            {{ latex: '\\\\sqrt{{#0}}', label: '\\\\sqrt{{x}}' }},
            {{ latex: '#@^{{#?}}', label: 'x^n' }},
            {{ latex: '#@_{{#?}}', label: 'x_n' }},
            '=', '+', '-', '\\\\times', '\\\\div'
          ],
          [
            {{ latex: '\\\\sin\\\\left(#0\\\\right)', label: '\\\\sin' }},
            {{ latex: '\\\\cos\\\\left(#0\\\\right)', label: '\\\\cos' }},
            {{ latex: '\\\\tan\\\\left(#0\\\\right)', label: '\\\\tan' }},
            {{ latex: '\\\\log_{{#?}}\\\\left(#0\\\\right)', label: '\\\\log' }},
            {{ latex: '\\\\ln\\\\left(#0\\\\right)', label: '\\\\ln' }},
            '\\\\pi', '\\\\theta', {{ latex: '^{{\\\\circ}}', label: '90^\\\\circ' }}, '\\\\pm'
          ],
          ['[7]','[8]','[9]','[(]','[)]','\\\\lt','\\\\le','\\\\gt','\\\\ge'],
          ['[4]','[5]','[6]','[1]','[2]','[3]','[0]','[.]',{{ label:'[backspace]', width:1.5 }}]
        ]
      }},
      'functions', 'symbols', 'alphabetic', 'greek'
    ];
  }}

  const incoming = normalizedPayload(data?.payload);
  if (singleQuestion) {{
    incoming.latex = [incoming.latex[0] || ''];
    incoming.ascii = [incoming.ascii[0] || ''];
  }}
  const state = parentElement.__omtState || {{
    payload: incoming,
    timer: null,
    active: null,
    keyboardOpen: false,
  }};
  parentElement.__omtState = state;

  // Only accept Streamlit's incoming payload on the first render. On subsequent
  // rerenders, the browser-local payload is newer and must win, otherwise a delayed
  // save can restore stale values and appear to "reset" the student's typing.
  if (!state.initialized) {{
    state.payload = incoming;
    state.initialized = true;
  }}
  if (typeof state.keyboardOpen !== 'boolean') state.keyboardOpen = false;

  const currentField = () => {{
    const mf = state.active && rows.contains(state.active) ? state.active : rows.querySelector('math-field');
    if (mf) state.active = mf;
    return mf;
  }};

  const captureLocal = () => {{
    const editors = Array.from(rows.querySelectorAll('math-field'));
    state.payload = {{
      latex: editors.map(mf => mf.value || ''),
      ascii: editors.map(mf => {{ try {{ return mf.getValue('ascii-math') || ''; }} catch (_) {{ return ''; }} }})
    }};
  }};

  const emit = () => {{
    captureLocal();
    if (state.timer) {{
      clearTimeout(state.timer);
      state.timer = null;
    }}
    setStateValue('payload', state.payload);
    status.textContent = singleQuestion ? 'Question saved' : 'Working saved';
  }};

  const scheduleEmit = () => {{
    // Deliberately DO NOT call setStateValue() while typing. Any component update
    // reruns Streamlit and can rebuild MathLive. Keep edits local until explicit save.
    captureLocal();
    status.textContent = singleQuestion ? 'Editing locally · tap Save question when finished' : 'Editing locally · tap Save working when finished';
  }};

  const showKeyboard = () => {{
    const mf = currentField();
    if (!mf) {{ status.textContent='Tap a maths box first.'; return; }}
    state.active = mf;
    state.keyboardOpen = true;
    mf.focus();
    keyboardPanel.hidden = false;
    keyboardButton.textContent = '⌨ Hide keyboard';
    keyboardButton.setAttribute('aria-expanded','true');
    status.textContent = 'Math keyboard open';
  }};

  const hideKeyboard = () => {{
    state.keyboardOpen = false;
    keyboardPanel.hidden = true;
    keyboardButton.textContent = '⌨ Math keyboard';
    keyboardButton.setAttribute('aria-expanded','false');
    status.textContent = 'Math keyboard hidden';
  }};

  const toggleKeyboard = () => {{
    if (keyboardPanel.hidden) showKeyboard(); else hideKeyboard();
  }};

  const renderRows = () => {{
    let activeIndex = 0;
    if (state.active && rows.contains(state.active)) {{
      const editorsBefore = Array.from(rows.querySelectorAll('math-field'));
      const found = editorsBefore.indexOf(state.active);
      if (found >= 0) activeIndex = found;
    }}

    rows.replaceChildren();
    state.payload.latex.forEach((value, index) => {{
      const row = document.createElement('div');
      row.className = 'omt-editor-row';

      const step = document.createElement('span');
      step.className = 'omt-step-label';
      step.textContent = singleQuestion ? 'Question' : `Step ${{index + 1}}`;

      const mf = document.createElement('math-field');
      mf.value = value || '';
      mf.mathVirtualKeyboardPolicy = 'manual';
      try {{ mf.virtualKeyboardMode = 'manual'; }} catch (_) {{}}
      mf.setAttribute('math-virtual-keyboard-policy', 'manual');
      mf.setAttribute('virtual-keyboard-mode', 'manual');
      mf.setAttribute('smart-fence', '');
      mf.setAttribute('aria-label', singleQuestion ? 'Question mathematics' : `Mathematics working step ${{index + 1}}`);
      mf.addEventListener('focusin', () => {{
        state.active = mf;
        const coarse = globalThis.matchMedia && globalThis.matchMedia('(pointer: coarse)').matches;
        const narrow = globalThis.innerWidth <= 700;
        if ((coarse || narrow) && !state.keyboardOpen) {{
          state.keyboardOpen = true;
          keyboardPanel.hidden = false;
          keyboardButton.textContent = '⌨ Hide keyboard';
          keyboardButton.setAttribute('aria-expanded','true');
          status.textContent = 'Math keyboard open';
        }}
      }});
      mf.addEventListener('input', scheduleEmit);
      mf.addEventListener('change', captureLocal);
      mf.addEventListener('blur', captureLocal);

      const remove = document.createElement('button');
      remove.type = 'button';
      remove.className = 'omt-remove-step';
      remove.textContent = '✕';
      remove.title = 'Remove this step';
      remove.disabled = state.payload.latex.length <= 1;
      if (singleQuestion) remove.style.display = 'none';
      remove.onclick = () => {{
        captureLocal();
        if (state.payload.latex.length <= 1) return;
        state.payload.latex.splice(index, 1);
        state.payload.ascii.splice(index, 1);
        renderRows();
        emit();
      }};

      row.append(step, mf, remove);
      rows.appendChild(row);
    }});

    const editorsAfter = rows.querySelectorAll('math-field');
    if (editorsAfter.length) {{
      state.active = editorsAfter[Math.min(activeIndex, editorsAfter.length - 1)];
    }}
  }};

  toolbar.querySelectorAll('button[data-insert]').forEach(button => {{
    button.onclick = () => {{
      const mf = currentField();
      if (!mf) return;
      state.active = mf;
      mf.focus();
      const latex = button.dataset.insert || '';
      try {{ mf.insert(latex, {{ insertionMode:'replaceSelection', selectionMode:'placeholder' }}); }}
      catch (_) {{ try {{ mf.executeCommand(['insert', latex]); }} catch (_) {{}} }}
      scheduleEmit();
    }};
  }});

  keyboardButton.onclick = toggleKeyboard;

  keyboardPanel.addEventListener('pointerdown', (event) => {{
    if (event.target && event.target.closest('button')) event.preventDefault();
  }});

  parentElement.querySelectorAll('button[data-kb]').forEach(button => {{
    button.onclick = () => {{
      const mf = currentField(); if (!mf) return;
      state.active = mf; mf.focus();
      const latex = button.dataset.kb || '';
      try {{ mf.insert(latex, {{ insertionMode:'replaceSelection', selectionMode:'placeholder' }}); }}
      catch (_) {{ try {{ mf.executeCommand(['insert',latex]); }} catch (_) {{}} }}
      captureLocal();
      status.textContent = singleQuestion ? 'Editing locally · tap Save question when finished' : 'Editing locally · tap Save working when finished';
    }};
  }});

  keyboardBackspace.onclick = () => {{
    const mf=currentField(); if(!mf) return; mf.focus();
    try {{ mf.executeCommand('deleteBackward'); }} catch(_) {{}}
    captureLocal();
    status.textContent = singleQuestion ? 'Editing locally · tap Save question when finished' : 'Editing locally · tap Save working when finished';
  }};

  keyboardClose.onclick = hideKeyboard;

  if (saveButton) {{
    saveButton.onclick = () => {{
      emit();
      const mf = currentField();
      if (mf) {{
        state.active = mf;
        setTimeout(() => {{
          try {{ mf.focus({{ preventScroll: true }}); }} catch (_) {{ try {{ mf.focus(); }} catch (_) {{}} }}
        }}, 0);
      }}
    }};
  }};

  addButton.onclick = () => {{
    if (singleQuestion) return;
    captureLocal();
    if (state.payload.latex.length >= 20) {{ status.textContent = 'Maximum 20 working steps.'; return; }}
    state.payload.latex.push('');
    state.payload.ascii.push('');
    renderRows();
    emit();
    const editors = rows.querySelectorAll('math-field');
    const mf = editors[editors.length - 1];
    if (mf) {{ state.active = mf; mf.focus(); }}
  }};

  keyboardPanel.hidden = !state.keyboardOpen;
  keyboardButton.textContent = state.keyboardOpen ? '⌨ Hide keyboard' : '⌨ Math keyboard';
  keyboardButton.setAttribute('aria-expanded', state.keyboardOpen ? 'true' : 'false');
  renderRows();

  // setStateValue() causes the Streamlit component to rerender after typing.
  // Restore focus without toggling the keyboard so the panel stays open.
  if (state.keyboardOpen) {{
    const mf = currentField();
    if (mf) {{
      state.active = mf;
      setTimeout(() => {{
        try {{ mf.focus({{ preventScroll: true }}); }} catch (_) {{ try {{ mf.focus(); }} catch (_) {{}} }}
      }}, 0);
    }}
  }}
}}
"""

try:
    _equation_editor_component = st.components.v2.component(
        "omt_math_question_editor_v3",
        html=_EQUATION_EDITOR_HTML,
        css=_EQUATION_EDITOR_CSS,
        js=_EQUATION_EDITOR_JS,
        isolate_styles=False,
    )
except Exception:
    try:
        _equation_editor_component = st.components.v2.component(
            "omt_math_question_editor_v3_fallback",
            html=_EQUATION_EDITOR_HTML,
            css=_EQUATION_EDITOR_CSS,
            js=_EQUATION_EDITOR_JS,
            isolate_styles=True,
        )
    except Exception:
        _equation_editor_component = None


_MATHIO_DISPLAY_HTML = """
<div class="omt-mathio-display" aria-live="polite"></div>
"""

_MATHIO_DISPLAY_CSS = """
.omt-mathio-display { width: 100%; overflow-x: auto; padding: .15rem 0 .25rem 0; }
.omt-mathio-display math-field {
  display: inline-block;
  width: auto;
  min-width: 0;
  border: 0 !important;
  outline: 0 !important;
  background: transparent !important;
  color: var(--st-text-color, #222);
  padding: 0;
  margin: 0;
  font-size: 1.08rem;
  pointer-events: none;
  --caret-color: transparent;
  --selection-background-color: transparent;
}
@media (pointer: coarse) { .omt-mathio-display math-field { font-size: 1.16rem; } }
"""

_MATHIO_DISPLAY_JS = f"""
const MATHLIVE_URL = 'https://cdn.jsdelivr.net/npm/mathlive@{MATHLIVE_VERSION}/+esm';

async function ensureMathLiveForDisplay() {{
  if (!customElements.get('math-field')) {{
    if (!globalThis.__omtMathLivePromise) {{
      globalThis.__omtMathLivePromise = import(MATHLIVE_URL);
    }}
    await globalThis.__omtMathLivePromise;
  }}
}}

export default async function(component) {{
  const {{ parentElement, data }} = component;
  const root = parentElement.querySelector('.omt-mathio-display');
  root.replaceChildren();
  try {{
    await ensureMathLiveForDisplay();
    const mf = document.createElement('math-field');
    mf.value = String(data?.math || '');
    mf.readOnly = true;
    mf.setAttribute('read-only', '');
    mf.setAttribute('virtual-keyboard-mode', 'off');
    mf.setAttribute('aria-label', 'Rendered mathematics');
    mf.tabIndex = -1;
    root.appendChild(mf);
  }} catch (err) {{
    const msg = document.createElement('span');
    msg.textContent = 'Equation view could not load. Reload the page to restore the maths display.';
    msg.style.opacity = '.72';
    root.appendChild(msg);
  }}
}}
"""

try:
    _mathio_display_component = st.components.v2.component(
        "omt_mathio_display",
        html=_MATHIO_DISPLAY_HTML,
        css=_MATHIO_DISPLAY_CSS,
        js=_MATHIO_DISPLAY_JS,
        isolate_styles=False,
    )
except Exception:
    _mathio_display_component = None


_MATHIO_RICH_HTML = """
<div class="omt-rich-math" aria-live="polite"></div>
"""

_MATHIO_RICH_CSS = """
.omt-rich-math {
  width: 100%; color: var(--st-text-color, #172033); line-height: 1.68;
  font-size: .98rem; overflow-wrap: anywhere;
}
.omt-rich-math .omt-rich-paragraph { margin: .15rem 0 .45rem; }
.omt-rich-math strong { font-weight: 720; }
.omt-rich-math math-field {
  display: inline-block; width: auto; min-width: 0; border: 0 !important;
  outline: 0 !important; background: transparent !important; padding: 0 .03rem;
  margin: 0 .03rem; color: var(--st-text-color, #172033); font-size: 1.03em;
  pointer-events: none; vertical-align: -0.12em; --caret-color: transparent;
  --selection-background-color: transparent;
}
.omt-rich-math math-field.omt-display-math {
  display: block; width: fit-content; max-width: 100%; margin: .42rem 0 .55rem;
  font-size: 1.12em; overflow-x: auto; vertical-align: baseline;
}
@media (pointer: coarse) { .omt-rich-math { font-size: 1rem; } }
"""

_MATHIO_RICH_JS = f"""
const MATHLIVE_URL = 'https://cdn.jsdelivr.net/npm/mathlive@{MATHLIVE_VERSION}/+esm';

async function ensureMathLiveForRich() {{
  if (!customElements.get('math-field')) {{
    if (!globalThis.__omtMathLivePromise) globalThis.__omtMathLivePromise = import(MATHLIVE_URL);
    await globalThis.__omtMathLivePromise;
  }}
}}

function appendTextWithBold(root, text) {{
  const bits = String(text || '').split(/(\\*\\*[^*]+\\*\\*)/g);
  for (const bit of bits) {{
    if (!bit) continue;
    if (bit.startsWith('**') && bit.endsWith('**')) {{
      const strong = document.createElement('strong');
      strong.textContent = bit.slice(2, -2);
      root.appendChild(strong);
    }} else {{
      const lines = bit.split('\\n');
      lines.forEach((line, index) => {{
        if (index) root.appendChild(document.createElement('br'));
        root.appendChild(document.createTextNode(line));
      }});
    }}
  }}
}}

function unwrapMath(token) {{
  if (token.startsWith('\\\\[') && token.endsWith('\\\\]')) return [token.slice(2,-2), true];
  if (token.startsWith('\\\\(') && token.endsWith('\\\\)')) return [token.slice(2,-2), false];
  if (token.startsWith('$$') && token.endsWith('$$')) return [token.slice(2,-2), true];
  if (token.startsWith('$') && token.endsWith('$')) return [token.slice(1,-1), false];
  return [token, false];
}}

export default async function(component) {{
  const {{ parentElement, data }} = component;
  const root = parentElement.querySelector('.omt-rich-math');
  root.replaceChildren();
  try {{
    await ensureMathLiveForRich();
    const raw = String(data?.text || '');
    const pattern = /(\\\\\\[[\\s\\S]*?\\\\\\]|\\\\\\([\\s\\S]*?\\\\\\)|\\$\\$[\\s\\S]*?\\$\\$|\\$[^$\\n]+?\\$)/g;
    let last = 0;
    for (const match of raw.matchAll(pattern)) {{
      appendTextWithBold(root, raw.slice(last, match.index));
      const [latex, display] = unwrapMath(match[0]);
      const mf = document.createElement('math-field');
      mf.value = latex.trim();
      mf.readOnly = true;
      mf.setAttribute('read-only', '');
      mf.setAttribute('virtual-keyboard-mode', 'off');
      mf.tabIndex = -1;
      if (display) mf.classList.add('omt-display-math');
      root.appendChild(mf);
      last = match.index + match[0].length;
    }}
    appendTextWithBold(root, raw.slice(last));
  }} catch (err) {{
    const msg = document.createElement('span');
    msg.textContent = 'Rich equation view could not load. Reload the page to restore the maths display.';
    msg.style.opacity = '.72'; root.appendChild(msg);
  }}
}}
"""

try:
    _mathio_rich_component = st.components.v2.component(
        "omt_rich_math_text",
        html=_MATHIO_RICH_HTML,
        css=_MATHIO_RICH_CSS,
        js=_MATHIO_RICH_JS,
        isolate_styles=False,
    )
except Exception:
    _mathio_rich_component = None


def equation_working_editor(label: str, *, key: str, single_question: bool = False) -> tuple[list[str], list[str]]:
    """Render MathLive working and persist the last saved payload across Streamlit reruns."""
    if _equation_editor_component is None:
        fallback = st.text_area(
            label,
            key=f"{key}_fallback",
            height=120 if single_question else 150,
            placeholder=(
                "Fallback: type the mathematical expression for your question, e.g. x^2 + 3x - 4 = 0"
                if single_question
                else "Fallback: type one mathematical step per line, e.g. m=(4-1)/(-2-7)"
            ),
        )
        lines = [line.strip() for line in fallback.splitlines() if line.strip()]
        if single_question:
            first = lines[0] if lines else ""
            return [first], [first]
        return lines, lines

    # IMPORTANT: component state and Python widget state are not the same thing.
    # Keep a separate shadow payload so a component-triggered rerun cannot feed an
    # older/default value back into MathLive and erase what the student just saved.
    shadow_key = f"{key}__saved_payload"
    payload = st.session_state.get(shadow_key, {"latex": [""], "ascii": [""]})
    if not isinstance(payload, dict):
        payload = {"latex": [""], "ascii": [""]}
    payload = {
        "latex": [str(x) for x in (payload.get("latex") or [""])],
        "ascii": [str(x) for x in (payload.get("ascii") or [""])],
    }
    if not payload["latex"]:
        payload["latex"] = [""]
    while len(payload["ascii"]) < len(payload["latex"]):
        payload["ascii"].append("")

    if single_question:
        payload["latex"] = [payload["latex"][0] if payload.get("latex") else ""]
        payload["ascii"] = [payload["ascii"][0] if payload.get("ascii") else ""]

    result = _equation_editor_component(
        data={"label": label, "payload": payload, "singleQuestion": bool(single_question)},
        default={"payload": payload},
        key=key,
        on_payload_change=lambda: None,
        width="stretch",
        height="content",
    )

    returned = getattr(result, "payload", None)
    if isinstance(returned, dict):
        saved = {
            "latex": [str(x) for x in (returned.get("latex") or [""])],
            "ascii": [str(x) for x in (returned.get("ascii") or [""])],
        }
        if not saved["latex"]:
            saved["latex"] = [""]
        while len(saved["ascii"]) < len(saved["latex"]):
            saved["ascii"].append("")
        # Persist the payload that caused this rerun. On the next render it becomes
        # the authoritative input instead of the old/default component value.
        st.session_state[shadow_key] = saved
        payload = saved

    latex = [str(x).strip() for x in payload.get("latex", [])]
    ascii_values = [str(x).strip() for x in payload.get("ascii", [])]
    if single_question:
        latex = [latex[0] if latex else ""]
        ascii_values = [ascii_values[0] if ascii_values else ""]
    while len(ascii_values) < len(latex):
        ascii_values.append("")
    return latex, ascii_values




def question_math_editor(label: str, *, key: str) -> str:
    """Single MathIO entry field for mathematics that belongs to the question."""
    if _equation_editor_component is None:
        st.warning(
            "The Math equation keyboard could not load in this browser session. "
            "Reload the page once; if it still fails, the fallback input below remains usable."
        )
    latex, _ascii = equation_working_editor(
        label,
        key=key,
        single_question=True,
    )
    return str(latex[0] if latex else "").strip()


def working_input(
    label: str,
    *,
    text_key: str,
    format_key: str,
    height: int = 170,
    plain_placeholder: str = "Show the important reasoning steps, not only the final answer.",
) -> tuple[str, str, str]:
    mode = st.radio(
        "Working input method",
        ["Equation editor", "Text working"],
        horizontal=True,
        key=format_key,
        help="Equation editor gives a visual maths keyboard; Text working is useful for sentences and explanations.",
    )

    if mode == "Equation editor":
        latex_lines, ascii_lines = equation_working_editor(label, key=f"{text_key}_equation")
        explanation = st.text_area(
            "Optional explanation in words",
            key=f"{text_key}_explanation",
            height=90,
            placeholder="Example: I expanded the bracket first, then collected like terms.",
        )
        used_latex = [line for line in latex_lines if line]
        used_ascii = [line for line in ascii_lines if line]
        working_lines = [f"Step {i}: \\({line}\\)" for i, line in enumerate(used_latex, 1)]
        if explanation.strip():
            working_lines.append(f"Student explanation: {explanation.strip()}")
        working_for_gemini = "\n".join(working_lines)
        offline_text = "\n".join(used_ascii)
        st.caption("The equation editor stores the mathematical structure automatically; students do not need to type equation code.")
        return working_for_gemini, mode, offline_text

    value = st.text_area(label, key=text_key, height=height, placeholder=plain_placeholder)
    return value, mode, value



def question_input_with_math_keyboard(*, key_base: str = "ai_question") -> str:
    """Question entry with either plain text or prose + structured maths boxes."""
    mode = st.radio(
        "Question input method",
        ["Text", "Text + math keyboard"],
        horizontal=True,
        key=f"{key_base}_input_mode",
        help=(
            "Use Text for ordinary typing. Use Text + math keyboard when the question "
            "contains fractions, powers, roots, trigonometry, vectors or other mathematical notation."
        ),
    )

    if mode == "Text":
        return st.text_area(
            "Question text",
            key=f"{key_base}_text",
            height=132,
            placeholder="Type the question here, or leave blank if it is visible in the upload.",
            label_visibility="collapsed",
        )

    prose = st.text_area(
        "Question wording",
        key=f"{key_base}_prose",
        height=100,
        placeholder="Example: Solve the equation, giving your answer to 3 significant figures.",
    )
    latex_lines, _ = equation_working_editor(
        "Question mathematics",
        key=f"{key_base}_math",
    )
    used = [line.strip() for line in latex_lines if str(line).strip()]
    if used:
        rendered = "\n".join(rf"\({line}\)" for line in used)
        return (prose.strip() + "\n\n" + rendered).strip()
    return prose.strip()


def _active_question_text_for_tools() -> str:
    """Resolve the currently entered/analyzed question text without relying on local scope."""
    for key in (
        "ai_question_text",
        "ai_question_prose",
        "current_question_text",
        "question_text",
        "last_question_text",
    ):
        value = st.session_state.get(key, "")
        if str(value or "").strip():
            return str(value).strip()
    return ""



# ---------------------------------------------------------------------------
# Animated ruler / compass / protractor constructions
# ---------------------------------------------------------------------------
def _is_construction_question(text: str) -> bool:
    """Show the simulation only for explicit construction-tool questions."""
    source = re.sub(r"\s+", " ", str(text or "")).strip().lower()
    if not source:
        return False

    explicit_tool_request = bool(
        re.search(r"\b(compass|protractor|straightedge|ruler)\b", source)
        and re.search(r"\b(construct|construction|draw|measure)\b", source)
    )
    classical_construction = bool(
        re.search(
            r"\bconstruct\b.*\b(perpendicular|parallel|angle bisector|perpendicular bisector|triangle|quadrilateral|angle)\b",
            source,
        )
    )
    return explicit_tool_request or classical_construction



def _construction_animation_spec(text: str) -> dict:
    """Build a small deterministic animation spec from construction wording.

    The first specialised template covers the common quadrilateral construction shown
    in the user's reference question. Other construction questions receive a generic
    ruler/angle/arc sequence so the tool motion remains useful rather than decorative.
    """
    raw = re.sub(r"\s+", " ", str(text or "")).strip()
    low = raw.lower()

    # Specialised quadrilateral construction:
    # PQ=8, QR=2, PS=6, angle PQR=90, angle QPS=60, parallel through S.
    pq = re.search(r"\bPQ\s*=\s*([0-9]+(?:\.[0-9]+)?)\s*cm", raw, flags=re.I)
    qr = re.search(r"\bQR\s*=\s*([0-9]+(?:\.[0-9]+)?)\s*cm", raw, flags=re.I)
    ps = re.search(r"\bPS\s*=\s*([0-9]+(?:\.[0-9]+)?)\s*cm", raw, flags=re.I)
    pqr = re.search(r"(?:angle\s*)?PQR\s*=\s*([0-9]+(?:\.[0-9]+)?)", raw, flags=re.I)
    qps = re.search(r"(?:angle\s*)?QPS\s*=\s*([0-9]+(?:\.[0-9]+)?)", raw, flags=re.I)

    if pq and qr and ps and pqr and qps:
        return {
            "type": "quadrilateral_pqrs",
            "pq": float(pq.group(1)),
            "qr": float(qr.group(1)),
            "ps": float(ps.group(1)),
            "pqr": float(pqr.group(1)),
            "qps": float(qps.group(1)),
            "parallel": bool(re.search(r"parallel\s+to\s+PQ", raw, flags=re.I)),
        }

    return {
        "type": "generic",
        "parallel": bool(re.search(r"\bparallel\b", low)),
        "perpendicular": bool(re.search(r"\bperpendicular\b|\b90\s*(?:degrees|°)", low)),
        "bisector": bool(re.search(r"\bbisect|\bbisector\b", low)),
    }


def _construction_animation_html(spec: dict) -> str:
    import json
    payload = json.dumps(spec)

    return f"""
<!doctype html>
<html>
<head>
<meta charset="utf-8">
<style>
  :root {{
    --ink:#20242c;
    --tool:#4263eb;
    --accent:#e8590c;
    --arc:#7048e8;
    --paper:#ffffff;
    --grid:#eef1f5;
  }}
  * {{ box-sizing:border-box; }}
  body {{ margin:0; font-family:Arial,Helvetica,sans-serif; background:transparent; color:var(--ink); }}
  .wrap {{ border:1px solid #d8dde7; border-radius:14px; background:#fff; padding:12px; }}
  .toolbar {{ display:flex; gap:8px; flex-wrap:wrap; margin-bottom:10px; align-items:center; }}
  button {{
    border:1px solid #ccd3df; background:#fff; border-radius:9px; padding:7px 12px;
    font-weight:600; cursor:pointer;
  }}
  button.primary {{ background:#4c6ef5; color:#fff; border-color:#4c6ef5; }}
  .status {{ margin-left:auto; font-size:13px; color:#596273; }}
  svg {{ width:100%; height:auto; display:block; background:var(--paper); border-radius:10px; }}
  .grid line {{ stroke:var(--grid); stroke-width:1; }}
  .construction {{ stroke:var(--ink); stroke-width:3; fill:none; stroke-linecap:round; }}
  .guide {{ stroke:#8b93a5; stroke-width:2; fill:none; stroke-dasharray:7 7; }}
  .arc {{ stroke:var(--arc); stroke-width:2.5; fill:none; stroke-dasharray:8 6; }}
  .point {{ fill:var(--ink); }}
  .label {{ font-size:17px; font-weight:700; fill:var(--ink); }}
  .measure {{ font-size:14px; fill:#4f5868; }}
  .tool-label {{ font-size:13px; font-weight:700; fill:#fff; }}
  .stepText {{ font-size:15px; fill:#313845; }}
  #toolRuler {{ transition:transform .55s ease, opacity .35s ease; }}
  #toolProtractor {{ transition:transform .55s ease, opacity .35s ease; }}
  #toolCompass {{ transition:transform .55s ease, opacity .35s ease; }}
  .hidden {{ opacity:0; pointer-events:none; }}
  .draw {{
    stroke-dasharray:1000;
    stroke-dashoffset:1000;
    animation:drawLine .8s ease forwards;
  }}
  @keyframes drawLine {{ to {{ stroke-dashoffset:0; }} }}
</style>
</head>
<body>
<div class="wrap">
  <div class="toolbar">
    <button class="primary" id="play">▶ Play</button>
    <button id="pause">Ⅱ Pause</button>
    <button id="step">Step</button>
    <button id="restart">↺ Restart</button>
    <span class="status" id="status">Ready</span>
  </div>

  <svg id="board" viewBox="0 0 900 560" role="img" aria-label="Animated compass and protractor construction">
    <g class="grid">
      <line x1="0" y1="70" x2="900" y2="70"/><line x1="0" y1="140" x2="900" y2="140"/>
      <line x1="0" y1="210" x2="900" y2="210"/><line x1="0" y1="280" x2="900" y2="280"/>
      <line x1="0" y1="350" x2="900" y2="350"/><line x1="0" y1="420" x2="900" y2="420"/>
      <line x1="0" y1="490" x2="900" y2="490"/>
      <line x1="90" y1="0" x2="90" y2="560"/><line x1="180" y1="0" x2="180" y2="560"/>
      <line x1="270" y1="0" x2="270" y2="560"/><line x1="360" y1="0" x2="360" y2="560"/>
      <line x1="450" y1="0" x2="450" y2="560"/><line x1="540" y1="0" x2="540" y2="560"/>
      <line x1="630" y1="0" x2="630" y2="560"/><line x1="720" y1="0" x2="720" y2="560"/>
      <line x1="810" y1="0" x2="810" y2="560"/>
    </g>

    <text id="stepText" class="stepText" x="28" y="34">Press Play to begin.</text>
    <g id="constructionLayer"></g>

    <!-- straightedge -->
    <g id="toolRuler" class="hidden">
      <rect x="-8" y="-18" width="280" height="36" rx="4" fill="#f2c94c" opacity=".88" stroke="#9a7b00"/>
      <g stroke="#7a6200">
        <line x1="10" y1="-18" x2="10" y2="-7"/><line x1="30" y1="-18" x2="30" y2="-10"/>
        <line x1="50" y1="-18" x2="50" y2="-7"/><line x1="70" y1="-18" x2="70" y2="-10"/>
        <line x1="90" y1="-18" x2="90" y2="-7"/><line x1="110" y1="-18" x2="110" y2="-10"/>
        <line x1="130" y1="-18" x2="130" y2="-7"/><line x1="150" y1="-18" x2="150" y2="-10"/>
        <line x1="170" y1="-18" x2="170" y2="-7"/><line x1="190" y1="-18" x2="190" y2="-10"/>
        <line x1="210" y1="-18" x2="210" y2="-7"/><line x1="230" y1="-18" x2="230" y2="-10"/>
      </g>
      <text x="110" y="6" font-size="13" font-weight="700" fill="#665100">straightedge</text>
    </g>

    <!-- semicircular protractor -->
    <g id="toolProtractor" class="hidden">
      <path d="M -105 0 A 105 105 0 0 1 105 0 L 0 0 Z" fill="#8ecae6" opacity=".42" stroke="#1971c2" stroke-width="3"/>
      <line x1="-105" y1="0" x2="105" y2="0" stroke="#1971c2" stroke-width="3"/>
      <line x1="0" y1="0" x2="0" y2="-96" stroke="#1971c2" stroke-width="2"/>
      <line x1="0" y1="0" x2="83" y2="-48" stroke="#1971c2" stroke-width="2"/>
      <text x="-15" y="-58" font-size="13" font-weight="700" fill="#145c93">90°</text>
      <text x="55" y="-28" font-size="13" font-weight="700" fill="#145c93">60°</text>
    </g>

    <!-- compass -->
    <g id="toolCompass" class="hidden">
      <circle cx="0" cy="0" r="9" fill="#e8590c"/>
      <line x1="0" y1="5" x2="-58" y2="120" stroke="#e8590c" stroke-width="8" stroke-linecap="round"/>
      <line x1="0" y1="5" x2="58" y2="120" stroke="#e8590c" stroke-width="8" stroke-linecap="round"/>
      <circle cx="-58" cy="120" r="5" fill="#343a40"/>
      <circle cx="58" cy="120" r="5" fill="#343a40"/>
      <text x="-29" y="145" font-size="13" font-weight="700" fill="#e8590c">compass</text>
    </g>
  </svg>
</div>

<script>
const spec = {payload};
const layer = document.getElementById("constructionLayer");
const ruler = document.getElementById("toolRuler");
const protractor = document.getElementById("toolProtractor");
const compass = document.getElementById("toolCompass");
const status = document.getElementById("status");
const stepText = document.getElementById("stepText");
let stepIndex = 0;
let playing = false;
let timer = null;

const NS = "http://www.w3.org/2000/svg";
function el(name, attrs={{}}) {{
  const node = document.createElementNS(NS,name);
  Object.entries(attrs).forEach(([k,v])=>node.setAttribute(k,v));
  layer.appendChild(node);
  return node;
}}
function line(x1,y1,x2,y2, cls="construction") {{
  return el("line",{{x1,y1,x2,y2,class:cls+" draw"}});
}}
function circle(cx,cy,r,cls="arc") {{
  return el("circle",{{cx,cy,r,class:cls+" draw"}});
}}
function arcPath(cx,cy,r,a1,a2,cls="arc") {{
  const p1=[cx+r*Math.cos(a1*Math.PI/180),cy-r*Math.sin(a1*Math.PI/180)];
  const p2=[cx+r*Math.cos(a2*Math.PI/180),cy-r*Math.sin(a2*Math.PI/180)];
  const large=Math.abs(a2-a1)>180?1:0;
  return el("path",{{d:`M ${{p1[0]}} ${{p1[1]}} A ${{r}} ${{r}} 0 ${{large}} 0 ${{p2[0]}} ${{p2[1]}}`,class:cls+" draw"}});
}}
function point(x,y,labelTxt) {{
  el("circle",{{cx:x,cy:y,r:5,class:"point"}});
  const t=el("text",{{x:x+8,y:y-8,class:"label"}}); t.textContent=labelTxt;
}}
function text(x,y,txt,cls="measure") {{
  const t=el("text",{{x,y,class:cls}}); t.textContent=txt; return t;
}}
function showTool(tool, transform) {{
  [ruler,protractor,compass].forEach(t=>t.classList.add("hidden"));
  tool.classList.remove("hidden");
  tool.setAttribute("transform",transform);
}}
function hideTools() {{ [ruler,protractor,compass].forEach(t=>t.classList.add("hidden")); }}

function buildSteps() {{
  if (spec.type === "quadrilateral_pqrs") {{
    const P=[170,430], Q=[650,430];
    const qrPx=120;
    const R=[650,430-qrPx];
    const psPx=360;
    const angle=spec.qps*Math.PI/180;
    const S=[P[0]+psPx*Math.cos(angle),P[1]-psPx*Math.sin(angle)];
    return [
      ()=>{{ stepText.textContent=`1. Draw PQ = ${{spec.pq}} cm with a straightedge.`; showTool(ruler,`translate(${{P[0]}},${{P[1]-10}})`); line(P[0],P[1],Q[0],Q[1]); point(P[0],P[1],"P"); point(Q[0],Q[1],"Q"); text(380,455,`PQ = ${{spec.pq}} cm`); }},
      ()=>{{ stepText.textContent=`2. Place the protractor at Q and mark ∠PQR = ${{spec.pqr}}°.`; showTool(protractor,`translate(${{Q[0]}},${{Q[1]}}) rotate(180)`); arcPath(Q[0],Q[1],58,90,180); text(Q[0]-76,Q[1]-62,`${{spec.pqr}}°`); }},
      ()=>{{ stepText.textContent=`3. Draw the perpendicular ray at Q.`; showTool(ruler,`translate(${{Q[0]+10}},${{Q[1]}}) rotate(-90)`); line(Q[0],Q[1],Q[0],R[1]-70,"guide"); }},
      ()=>{{ stepText.textContent=`4. Set the compass to QR = ${{spec.qr}} cm and mark R.`; showTool(compass,`translate(${{Q[0]}},${{Q[1]-115}}) scale(.7)`); arcPath(Q[0],Q[1],qrPx,75,105); point(R[0],R[1],"R"); line(Q[0],Q[1],R[0],R[1]); text(R[0]+18,(R[1]+Q[1])/2,`QR = ${{spec.qr}} cm`); }},
      ()=>{{ stepText.textContent=`5. Place the protractor at P and mark ∠QPS = ${{spec.qps}}°.`; showTool(protractor,`translate(${{P[0]}},${{P[1]}})`); arcPath(P[0],P[1],72,0,spec.qps); text(P[0]+62,P[1]-38,`${{spec.qps}}°`); }},
      ()=>{{ stepText.textContent=`6. Draw the ray from P through the ${{spec.qps}}° mark.`; showTool(ruler,`translate(${{P[0]}},${{P[1]}}) rotate(${{-spec.qps}})`); line(P[0],P[1],S[0]+65,S[1]-110,"guide"); }},
      ()=>{{ stepText.textContent=`7. Set the compass to PS = ${{spec.ps}} cm and locate S.`; showTool(compass,`translate(${{P[0]+180}},${{P[1]-135}}) rotate(${{-spec.qps/2}}) scale(.8)`); arcPath(P[0],P[1],psPx,spec.qps-10,spec.qps+10); point(S[0],S[1],"S"); line(P[0],P[1],S[0],S[1]); text(P[0]+150,P[1]-185,`PS = ${{spec.ps}} cm`); }},
      ()=>{{ stepText.textContent="8. Join S to R to complete quadrilateral PQRS."; showTool(ruler,`translate(${{S[0]}},${{S[1]}}) rotate(12)`); line(S[0],S[1],R[0],R[1]); }},
      ()=>{{ stepText.textContent="9. Through S, construct a line parallel to PQ."; showTool(ruler,`translate(${{S[0]-120}},${{S[1]-2}})`); line(S[0]-155,S[1],820,S[1],"guide"); }},
      ()=>{{ stepText.textContent="10. Extend QR to meet the parallel at U."; showTool(ruler,`translate(${{Q[0]}},${{R[1]-170}}) rotate(90)`); line(Q[0],R[1]-165,Q[0],Q[1],"guide"); const U=[Q[0],S[1]]; point(U[0],U[1],"U"); hideTools(); status.textContent="Construction complete"; }},
    ];
  }}

  // Generic fallback animation.
  const A=[180,420], B=[650,420], C=[500,190];
  return [
    ()=>{{ stepText.textContent="1. Draw the given base segment with a straightedge."; showTool(ruler,`translate(${{A[0]}},${{A[1]-10}})`); line(A[0],A[1],B[0],B[1]); point(A[0],A[1],"A"); point(B[0],B[1],"B"); }},
    ()=>{{ stepText.textContent="2. Place the protractor at the required vertex and mark the given angle."; showTool(protractor,`translate(${{A[0]}},${{A[1]}})`); arcPath(A[0],A[1],75,0,55); }},
    ()=>{{ stepText.textContent="3. Draw the construction ray through the angle mark."; showTool(ruler,`translate(${{A[0]}},${{A[1]}}) rotate(-55)`); line(A[0],A[1],C[0],C[1],"guide"); }},
    ()=>{{ stepText.textContent="4. Set the compass to the required length and draw an arc."; showTool(compass,`translate(${{A[0]+170}},${{A[1]-145}})`); arcPath(A[0],A[1],330,45,65); }},
    ()=>{{ stepText.textContent="5. Mark the intersection and join the required points."; point(C[0],C[1],"C"); line(A[0],A[1],C[0],C[1]); line(C[0],C[1],B[0],B[1]); hideTools(); status.textContent="Construction complete"; }},
  ];
}}

const steps=buildSteps();

function reset() {{
  playing=false;
  if(timer) clearTimeout(timer);
  stepIndex=0;
  layer.innerHTML="";
  hideTools();
  stepText.textContent="Press Play to begin.";
  status.textContent="Ready";
}}
function doStep() {{
  if(stepIndex>=steps.length) {{ playing=false; status.textContent="Construction complete"; return; }}
  status.textContent=`Step ${{stepIndex+1}} of ${{steps.length}}`;
  steps[stepIndex++]();
}}
function playLoop() {{
  if(!playing) return;
  doStep();
  if(stepIndex<steps.length) timer=setTimeout(playLoop,1500);
  else playing=false;
}}
document.getElementById("play").onclick=()=>{{ if(stepIndex>=steps.length) reset(); playing=true; playLoop(); }};
document.getElementById("pause").onclick=()=>{{ playing=false; if(timer) clearTimeout(timer); status.textContent="Paused"; }};
document.getElementById("step").onclick=()=>{{ playing=false; if(timer) clearTimeout(timer); doStep(); }};
document.getElementById("restart").onclick=reset;
reset();
</script>
</body>
</html>
"""


def show_construction_animation(text: str, *, key_base: str) -> None:
    """Show an animated compass/protractor/straightedge construction when relevant."""
    if not _is_construction_question(text):
        return

    spec = _construction_animation_spec(text)
    with st.expander("▶ Construction animation — compass, protractor and straightedge", expanded=False):
        st.caption(
            "Watch the tools move through the construction. "
            "Use Step when demonstrating the construction to a class."
        )
        st_components_v1.html(
            _construction_animation_html(spec),
            height=650,
            scrolling=False,
        )


def geogebra_external_tools(*, question_text: str = "", key_base: str = "geogebra_tools") -> None:
    """Offer external GeoGebra graphing/geometry workspaces without sending uploaded files."""
    text = str(question_text or "").lower()
    graph_hint = bool(re.search(r"\b(graph|function|curve|sin|cos|tan|coordinate|plot|sketch)\b", text))
    geometry_hint = bool(re.search(r"\b(circle|triangle|quadrilateral|polygon|tangent|chord|angle|geometry|vector)\b", text))

    with st.expander("↗ Open GeoGebra to illustrate the question", expanded=False):
        st.caption(
            "GeoGebra opens in a separate tab. Math Advisor does not send the uploaded question image "
            "to GeoGebra automatically; you choose what to construct there."
        )
        cols = st.columns(2)
        with cols[0]:
            st.link_button(
                "Open GeoGebra Graphing",
                "https://www.geogebra.org/graphing",
                use_container_width=True,
                type="primary" if graph_hint else "secondary",
            )
        with cols[1]:
            st.link_button(
                "Open GeoGebra Geometry",
                "https://www.geogebra.org/geometry",
                use_container_width=True,
                type="primary" if geometry_hint else "secondary",
            )
        if graph_hint:
            st.caption("Graph/function language detected — Graphing is likely the most useful workspace.")
        elif geometry_hint:
            st.caption("Geometry language detected — Geometry is likely the most useful workspace.")




_HANDWRITING_HTML = """
<div class="omt-handwriting-pad">
  <div class="omt-handwriting-help">Write with Apple Pencil, stylus, or finger. Nothing is sent to Streamlit while you are writing, so the pad will not refresh after every stroke.</div>
  <div class="omt-handwriting-toolbar">
    <button type="button" class="omt-undo-pad">Undo</button>
    <button type="button" class="omt-clear-pad">Clear</button>
    <button type="button" class="omt-save-pad omt-primary-pad">Save handwriting</button>
  </div>
  <canvas class="omt-handwriting-canvas" aria-label="Handwritten mathematics working area"></canvas>
  <div class="omt-handwriting-status" aria-live="polite">Write first, then tap Save handwriting before checking your answer.</div>
</div>
"""

_HANDWRITING_CSS = """
.omt-handwriting-pad { width:100%; font-family:var(--st-font,sans-serif); overscroll-behavior:contain; }
.omt-handwriting-help { opacity:.78; font-size:.9rem; margin:0 0 .55rem 0; }
.omt-handwriting-toolbar { display:flex; justify-content:flex-end; gap:.45rem; flex-wrap:wrap; margin-bottom:.5rem; }
.omt-handwriting-toolbar button { min-height:44px; padding:.5rem .85rem; border:1px solid rgba(128,128,128,.42); border-radius:.55rem; background:transparent; color:var(--st-text-color,#222); font-weight:600; }
.omt-handwriting-toolbar .omt-primary-pad { background:#ff4b4b; color:#fff; border-color:#ff4b4b; }
.omt-handwriting-canvas { width:100%; height:430px; display:block; background:#fff; border:1px solid rgba(128,128,128,.48); border-radius:.7rem; touch-action:none; user-select:none; -webkit-user-select:none; -webkit-touch-callout:none; overscroll-behavior:contain; box-sizing:border-box; }
.omt-handwriting-status { min-height:1.2rem; margin-top:.4rem; font-size:.8rem; opacity:.76; }
@media (max-width:900px) { .omt-handwriting-canvas { height:390px; } }
@media (pointer:coarse) { .omt-handwriting-canvas { height:48vh; min-height:340px; max-height:600px; } .omt-handwriting-toolbar button { min-height:48px; font-size:1rem; } }
"""

_HANDWRITING_JS = r"""
function validPng(value) {
  return typeof value === 'string' && value.startsWith('data:image/png;base64,') && value.length > 100;
}

export default function(component) {
  const { parentElement, data, setStateValue } = component;
  const canvas = parentElement.querySelector('.omt-handwriting-canvas');
  const undoButton = parentElement.querySelector('.omt-undo-pad');
  const clearButton = parentElement.querySelector('.omt-clear-pad');
  const saveButton = parentElement.querySelector('.omt-save-pad');
  const status = parentElement.querySelector('.omt-handwriting-status');
  const ctx = canvas.getContext('2d', { alpha:false, desynchronized:true });

  // Important: setStateValue() causes a Streamlit rerun. Therefore it is only
  // called when the student taps Save handwriting, never at pointer-up.
  let drawing = false;
  let hasInk = false;
  let dirty = false;
  let lastX = 0;
  let lastY = 0;
  let history = [];
  const restoreData = validPng(data?.image_data_url) ? data.image_data_url : '';

  const cssSize = () => {
    const r = canvas.getBoundingClientRect();
    return { w: Math.max(1, r.width), h: Math.max(1, r.height) };
  };

  const paintWhite = () => {
    const { w, h } = cssSize();
    ctx.save();
    ctx.setTransform(1,0,0,1,0,0);
    ctx.fillStyle = '#fff';
    ctx.fillRect(0,0,canvas.width,canvas.height);
    ctx.restore();
    ctx.fillStyle = '#fff';
    ctx.fillRect(0,0,w,h);
  };

  const configureCanvasOnce = () => {
    const { w, h } = cssSize();
    const ratio = Math.max(1, Math.min(window.devicePixelRatio || 1, 2));
    canvas.width = Math.max(1, Math.round(w * ratio));
    canvas.height = Math.max(1, Math.round(h * ratio));
    ctx.setTransform(ratio,0,0,ratio,0,0);
    ctx.lineCap = 'round';
    ctx.lineJoin = 'round';
    ctx.strokeStyle = '#111';
    ctx.lineWidth = 2.6;
    ctx.fillStyle = '#fff';
    ctx.fillRect(0,0,w,h);

    if (restoreData) {
      const img = new Image();
      img.onload = () => {
        ctx.drawImage(img,0,0,w,h);
        hasInk = true;
        dirty = false;
        status.textContent = 'Saved handwriting restored. Continue writing or tap Save handwriting again after changes.';
      };
      img.src = restoreData;
    }
  };

  const snapshot = () => {
    try {
      history.push(canvas.toDataURL('image/png'));
      if (history.length > 24) history.shift();
    } catch (_) {}
  };

  const restoreSnapshot = (url) => {
    if (!validPng(url)) return;
    const img = new Image();
    img.onload = () => {
      const { w, h } = cssSize();
      ctx.fillStyle='#fff'; ctx.fillRect(0,0,w,h);
      ctx.drawImage(img,0,0,w,h);
      hasInk = true; dirty = true;
      status.textContent = 'Undo applied. Tap Save handwriting when finished.';
    };
    img.src = url;
  };

  const point = (ev) => {
    const r = canvas.getBoundingClientRect();
    return [ev.clientX-r.left, ev.clientY-r.top];
  };

  const drawEvent = (ev) => {
    const events = ev.getCoalescedEvents ? ev.getCoalescedEvents() : [ev];
    for (const e of events) {
      const [x,y] = point(e);
      const pressure = e.pressure && e.pressure > 0 ? e.pressure : 0.5;
      ctx.lineWidth = 1.9 + pressure * 2.7;
      ctx.lineTo(x,y);
      ctx.stroke();
      lastX=x; lastY=y;
    }
  };

  const start = (ev) => {
    if (ev.pointerType === 'touch' && ev.isPrimary === false) return;
    ev.preventDefault();
    snapshot();
    drawing = true; hasInk = true; dirty = true;
    canvas.setPointerCapture?.(ev.pointerId);
    [lastX,lastY] = point(ev);
    ctx.beginPath(); ctx.moveTo(lastX,lastY);
    status.textContent = 'Writing… tap Save handwriting when the page is complete.';
  };

  const move = (ev) => {
    if (!drawing) return;
    ev.preventDefault();
    drawEvent(ev);
  };

  const end = (ev) => {
    if (!drawing) return;
    ev.preventDefault();
    drawing = false;
    try { canvas.releasePointerCapture?.(ev.pointerId); } catch (_) {}
    ctx.closePath();
    status.textContent = 'Unsaved handwriting. Tap Save handwriting before checking your answer.';
  };

  canvas.onpointerdown=start;
  canvas.onpointermove=move;
  canvas.onpointerup=end;
  canvas.onpointercancel=end;

  undoButton.onclick = () => {
    const previous = history.pop();
    if (previous) restoreSnapshot(previous);
  };

  clearButton.onclick = () => {
    snapshot();
    const { w,h }=cssSize();
    ctx.fillStyle='#fff'; ctx.fillRect(0,0,w,h);
    hasInk=false; dirty=true;
    status.textContent='Canvas cleared locally. Tap Save handwriting to save the blank page, or Undo to restore.';
  };

  saveButton.onclick = () => {
    // This is the only normal path that sends canvas state to Python and reruns Streamlit.
    const url = hasInk ? canvas.toDataURL('image/png') : '';
    dirty=false;
    status.textContent = hasInk ? 'Saving handwriting…' : 'Saving blank canvas…';
    setStateValue('image_data_url', url);
  };

  configureCanvasOnce();

  // Prevent Safari gestures/scrolling from stealing Pencil strokes while inside the pad.
  canvas.addEventListener('touchstart', e => e.preventDefault(), { passive:false });
  canvas.addEventListener('touchmove', e => e.preventDefault(), { passive:false });
}
"""

try:
    _handwriting_component = st.components.v2.component(
        "omt_handwriting_pad",
        html=_HANDWRITING_HTML,
        css=_HANDWRITING_CSS,
        js=_HANDWRITING_JS,
        isolate_styles=False,
    )
except Exception:
    _handwriting_component = None


def handwriting_pad(*, key: str) -> UploadedAsset | None:
    """Return a PNG UploadedAsset from a touch/Pencil handwriting canvas."""
    if _handwriting_component is None:
        st.info("The on-screen handwriting pad is unavailable in this browser. Use camera or file upload below.")
        return None
    prior = st.session_state.get(key, {})
    prior_url = prior.get("image_data_url", "") if isinstance(prior, dict) else ""
    result = _handwriting_component(
        data={"image_data_url": prior_url},
        default={"image_data_url": prior_url},
        key=key,
        on_image_data_url_change=lambda: None,
        width="stretch",
        height="content",
    )
    data_url = getattr(result, "image_data_url", "") or prior_url
    if not isinstance(data_url, str) or not data_url.startswith("data:image/png;base64,"):
        return None
    try:
        raw = base64.b64decode(data_url.split(",", 1)[1], validate=True)
    except Exception:
        return None
    if len(raw) < 200:
        return None
    return UploadedAsset(name="ipad-handwritten-working.png", mime_type="image/png", data=raw)






def student_scientific_calculator(*, key_base: str = "student_scientific_calculator") -> None:
    """Mobile-friendly scientific calculator for student solving/working only.

    The interface is original Math Advisor code, inspired by the feature set of
    conventional scientific calculators. It does not copy calculator.net JavaScript.
    """
    with st.expander("🧮 Scientific calculator", expanded=False):
        st.caption(
            "For calculations while solving. You can click the keypad or type into the display. "
            "Choose DEG or RAD before using trigonometric functions."
        )

        calc_id = re.sub(r"[^A-Za-z0-9_-]", "_", key_base)
        html = rf"""
<div id="{calc_id}" class="ma-calc">
  <div class="ma-calc-top">
    <div class="ma-mode">
      <button type="button" data-mode="DEG" class="mode active">DEG</button>
      <button type="button" data-mode="RAD" class="mode">RAD</button>
    </div>
    <div class="ma-memory" id="{calc_id}_mem">M: 0</div>
  </div>

  <input
    id="{calc_id}_display"
    class="ma-display"
    type="text"
    inputmode="text"
    autocomplete="off"
    spellcheck="false"
    placeholder="0"
    aria-label="Scientific calculator expression"
  />
  <div id="{calc_id}_result" class="ma-result">0</div>

  <div class="ma-grid">
    <button data-action="func" data-value="sin(">sin</button>
    <button data-action="func" data-value="cos(">cos</button>
    <button data-action="func" data-value="tan(">tan</button>
    <button data-action="func" data-value="asin(">sin⁻¹</button>
    <button data-action="func" data-value="acos(">cos⁻¹</button>

    <button data-action="func" data-value="atan(">tan⁻¹</button>
    <button data-action="insert" data-value="pi">π</button>
    <button data-action="insert" data-value="e">e</button>
    <button data-action="power">xʸ</button>
    <button data-action="square">x²</button>

    <button data-action="cube">x³</button>
    <button data-action="func" data-value="sqrt(">√x</button>
    <button data-action="func" data-value="cbrt(">∛x</button>
    <button data-action="func" data-value="ln(">ln</button>
    <button data-action="func" data-value="log10(">log</button>

    <button data-action="reciprocal">1/x</button>
    <button data-action="percent">%</button>
    <button data-action="factorial">n!</button>
    <button data-action="insert" data-value="(">(</button>
    <button data-action="insert" data-value=")">)</button>

    <button data-action="insert" data-value="7">7</button>
    <button data-action="insert" data-value="8">8</button>
    <button data-action="insert" data-value="9">9</button>
    <button data-action="insert" data-value="/">÷</button>
    <button data-action="back">⌫</button>

    <button data-action="insert" data-value="4">4</button>
    <button data-action="insert" data-value="5">5</button>
    <button data-action="insert" data-value="6">6</button>
    <button data-action="insert" data-value="*">×</button>
    <button data-action="ans">Ans</button>

    <button data-action="insert" data-value="1">1</button>
    <button data-action="insert" data-value="2">2</button>
    <button data-action="insert" data-value="3">3</button>
    <button data-action="insert" data-value="-">−</button>
    <button data-action="mplus">M+</button>

    <button data-action="insert" data-value="0">0</button>
    <button data-action="insert" data-value=".">.</button>
    <button data-action="exp">EXP</button>
    <button data-action="insert" data-value="+">+</button>
    <button data-action="mminus">M−</button>

    <button data-action="sign">±</button>
    <button data-action="round">RND</button>
    <button data-action="mr">MR</button>
    <button data-action="clear" class="clear">AC</button>
    <button data-action="equals" class="equals">=</button>
  </div>

  <div class="ma-calc-help">
    Examples: <code>sin(30)</code> in DEG mode, <code>sqrt(45)</code>,
    <code>3.2e5</code>, <code>log10(100)</code>.
  </div>
</div>

<style>
#{calc_id}.ma-calc {{
  max-width: 590px;
  margin: 0 auto;
  padding: 8px;
  border: 1px solid #d8dce6;
  border-radius: 14px;
  background: #f7f8fb;
  box-sizing: border-box;
  font-family: Arial, Helvetica, sans-serif;
}}
#{calc_id} .ma-calc-top {{
  display: flex;
  justify-content: space-between;
  align-items: center;
  margin-bottom: 5px;
  gap: 8px;
}}
#{calc_id} .ma-mode {{
  display: flex;
  gap: 6px;
}}
#{calc_id} .ma-mode button {{
  min-height: 30px;
  padding: 3px 10px;
  border-radius: 8px;
  border: 1px solid #b8bfcc;
  background: white;
  font-weight: 700;
}}
#{calc_id} .ma-mode button.active {{
  background: #334155;
  color: white;
}}
#{calc_id} .ma-memory {{
  font-size: 13px;
  color: #64748b;
}}
#{calc_id} .ma-display {{
  width: 100%;
  min-height: 42px;
  box-sizing: border-box;
  border-radius: 9px;
  border: 1px solid #9ca3af;
  background: white;
  padding: 6px 10px;
  font-size: 19px;
  text-align: right;
}}
#{calc_id} .ma-result {{
  min-height: 28px;
  margin: 3px 2px 6px;
  padding: 3px 6px;
  text-align: right;
  font-size: 17px;
  color: #334155;
  overflow-wrap: anywhere;
}}
#{calc_id} .ma-grid {{
  display: grid;
  grid-template-columns: repeat(5, minmax(0, 1fr));
  gap: 4px;
}}
#{calc_id} .ma-grid button {{
  min-height: 36px;
  border: 1px solid #c6cad3;
  border-radius: 9px;
  background: white;
  font-size: 14px;
  font-weight: 600;
  cursor: pointer;
  touch-action: manipulation;
}}
#{calc_id} .ma-grid button:active {{
  transform: translateY(1px);
  background: #eef2ff;
}}
#{calc_id} .ma-grid .equals {{
  background: #3156d9;
  color: white;
  border-color: #3156d9;
}}
#{calc_id} .ma-grid .clear {{
  background: #fee2e2;
  color: #991b1b;
}}
#{calc_id} .ma-calc-help {{
  margin-top: 6px;
  font-size: 11px;
  color: #64748b;
  line-height: 1.45;
}}
@media (max-width: 520px) {{
  #{calc_id}.ma-calc {{ padding: 8px; }}
  #{calc_id} .ma-grid {{ gap: 5px; }}
  #{calc_id} .ma-grid button {{ min-height: 34px; font-size: 13px; }}
  #{calc_id} .ma-display {{ font-size: 21px; }}
}}
</style>

<script>
(() => {{
  const root = document.getElementById("{calc_id}");
  const display = document.getElementById("{calc_id}_display");
  const result = document.getElementById("{calc_id}_result");
  const memLabel = document.getElementById("{calc_id}_mem");

  let angleMode = "DEG";
  let ans = 0;
  let memory = 0;

  function degToRad(x) {{ return x * Math.PI / 180; }}
  function radToDeg(x) {{ return x * 180 / Math.PI; }}

  function factorial(n) {{
    if (!Number.isFinite(n) || n < 0 || Math.floor(n) !== n || n > 170) {{
      throw new Error("Factorial requires a non-negative integer up to 170");
    }}
    let v = 1;
    for (let i = 2; i <= n; i++) v *= i;
    return v;
  }}

  function scopeFunctions() {{
    return {{
      sin: x => Math.sin(angleMode === "DEG" ? degToRad(x) : x),
      cos: x => Math.cos(angleMode === "DEG" ? degToRad(x) : x),
      tan: x => Math.tan(angleMode === "DEG" ? degToRad(x) : x),
      asin: x => angleMode === "DEG" ? radToDeg(Math.asin(x)) : Math.asin(x),
      acos: x => angleMode === "DEG" ? radToDeg(Math.acos(x)) : Math.acos(x),
      atan: x => angleMode === "DEG" ? radToDeg(Math.atan(x)) : Math.atan(x),
      sqrt: Math.sqrt,
      cbrt: Math.cbrt,
      ln: Math.log,
      log10: Math.log10,
      abs: Math.abs,
      pi: Math.PI,
      e: Math.E
    }};
  }}

  function transformExpression(source) {{
    let s = String(source || "").trim();
    if (!s) return "0";

    s = s
      .replace(/π/g, "pi")
      .replace(/×/g, "*")
      .replace(/÷/g, "/")
      .replace(/−/g, "-")
      .replace(/\^/g, "**");

    // Scientific notation entered through EXP becomes ordinary e notation.
    s = s.replace(/(\d(?:\.\d+)?)\s*EXP\s*([+\-]?\d+)/gi, "$1e$2");

    // Restrict to calculator-style tokens before evaluation.
    if (!/^[0-9A-Za-z_+\-*/().,\s*]+$/.test(s)) {{
      throw new Error("Unsupported character");
    }}
    return s;
  }}

  function evaluate(source) {{
    const expr = transformExpression(source);
    const f = scopeFunctions();

    const names = [
      "sin","cos","tan","asin","acos","atan","sqrt","cbrt",
      "ln","log10","abs","pi","e","factorial","ans"
    ];
    const values = [
      f.sin,f.cos,f.tan,f.asin,f.acos,f.atan,f.sqrt,f.cbrt,
      f.ln,f.log10,f.abs,f.pi,f.e,factorial,ans
    ];

    // Expression is limited to the calculator token whitelist above and does
    // not receive access to DOM/window objects.
    const fn = Function(...names, `"use strict"; return (${{expr}});`);
    const value = fn(...values);

    if (typeof value !== "number" || !Number.isFinite(value)) {{
      throw new Error("Result is not finite");
    }}
    return value;
  }}

  function formatNumber(x) {{
    if (!Number.isFinite(x)) return "Error";
    const ax = Math.abs(x);
    if ((ax !== 0 && ax < 1e-9) || ax >= 1e12) return x.toExponential(10).replace(/0+e/, "e");
    return Number(x.toPrecision(12)).toString();
  }}

  function insertText(text) {{
    const start = display.selectionStart ?? display.value.length;
    const end = display.selectionEnd ?? display.value.length;
    display.value = display.value.slice(0, start) + text + display.value.slice(end);
    const pos = start + text.length;
    display.focus();
    display.setSelectionRange(pos, pos);
  }}

  function wrapSelection(prefix, suffix = ")") {{
    const start = display.selectionStart ?? display.value.length;
    const end = display.selectionEnd ?? display.value.length;
    const selected = display.value.slice(start, end);
    const replacement = selected ? `${{prefix}}${{selected}}${{suffix}}` : prefix;
    display.value = display.value.slice(0, start) + replacement + display.value.slice(end);
    display.focus();
    const pos = start + replacement.length;
    display.setSelectionRange(pos, pos);
  }}

  function currentValue() {{
    return evaluate(display.value || String(ans));
  }}

  function updateMemory() {{
    memLabel.textContent = "M: " + formatNumber(memory);
  }}

  root.querySelectorAll(".mode").forEach(btn => {{
    btn.addEventListener("click", () => {{
      angleMode = btn.dataset.mode;
      root.querySelectorAll(".mode").forEach(x => x.classList.toggle("active", x === btn));
    }});
  }});

  root.querySelectorAll(".ma-grid button").forEach(btn => {{
    btn.addEventListener("click", () => {{
      const action = btn.dataset.action;
      const value = btn.dataset.value || "";

      try {{
        if (action === "insert" || action === "func") {{
          insertText(value);
        }} else if (action === "back") {{
          const start = display.selectionStart ?? display.value.length;
          const end = display.selectionEnd ?? display.value.length;
          if (start !== end) {{
            display.value = display.value.slice(0,start) + display.value.slice(end);
            display.setSelectionRange(start,start);
          }} else if (start > 0) {{
            display.value = display.value.slice(0,start-1) + display.value.slice(start);
            display.setSelectionRange(start-1,start-1);
          }}
          display.focus();
        }} else if (action === "clear") {{
          display.value = "";
          result.textContent = "0";
        }} else if (action === "ans") {{
          insertText("ans");
        }} else if (action === "power") {{
          insertText("^");
        }} else if (action === "square") {{
          wrapSelection("(", ")^2");
        }} else if (action === "cube") {{
          wrapSelection("(", ")^3");
        }} else if (action === "reciprocal") {{
          wrapSelection("1/(", ")");
        }} else if (action === "percent") {{
          wrapSelection("(", ")/100");
        }} else if (action === "factorial") {{
          const v = currentValue();
          const out = factorial(v);
          ans = out;
          display.value = formatNumber(out);
          result.textContent = formatNumber(out);
        }} else if (action === "exp") {{
          insertText("e");
        }} else if (action === "sign") {{
          if (display.value.trim()) wrapSelection("-(", ")");
          else insertText("-");
        }} else if (action === "round") {{
          const out = Math.round(currentValue());
          ans = out;
          display.value = formatNumber(out);
          result.textContent = formatNumber(out);
        }} else if (action === "mplus") {{
          memory += currentValue();
          updateMemory();
        }} else if (action === "mminus") {{
          memory -= currentValue();
          updateMemory();
        }} else if (action === "mr") {{
          insertText(formatNumber(memory));
        }} else if (action === "equals") {{
          const out = currentValue();
          ans = out;
          result.textContent = formatNumber(out);
          display.value = formatNumber(out);
        }}
      }} catch (err) {{
        result.textContent = "Error";
      }}
    }});
  }});

  display.addEventListener("keydown", evt => {{
    if (evt.key === "Enter") {{
      evt.preventDefault();
      try {{
        const out = currentValue();
        ans = out;
        result.textContent = formatNumber(out);
        display.value = formatNumber(out);
      }} catch (err) {{
        result.textContent = "Error";
      }}
    }} else if (evt.key === "Escape") {{
      display.value = "";
      result.textContent = "0";
    }}
  }});

  updateMemory();
}})();
</script>
"""

        try:
                        st_components_v1.html(html, height=505, scrolling=False)
        except Exception:
            st.warning("The scientific calculator could not be loaded in this browser session.")



def _student_table_tool(*, key_base: str) -> str:
    """Optional fillable working table whose contents are submitted with the attempt."""
    st.markdown("#### Working tools")
    use_table = st.toggle(
        "Insert a table",
        key=f"{key_base}_use_table",
        help="Useful for value tables, frequency tables, coordinates, sequences, and organised calculations.",
    )
    if not use_table:
        return ""

    c1, c2 = st.columns([1, 2])
    with c1:
        rows = int(st.number_input("Starting rows", min_value=1, max_value=15, value=4, step=1, key=f"{key_base}_table_rows"))
    with c2:
        headers_text = st.text_input(
            "Column headings",
            value="x, y",
            key=f"{key_base}_table_headers",
            help="Separate headings with commas, e.g. x, y or Class interval, Frequency.",
        )
    headers = [h.strip() for h in headers_text.split(",") if h.strip()][:8]
    if not headers:
        headers = ["Column 1", "Column 2"]

    seed_key = f"{key_base}_table_seed_{'|'.join(headers)}_{rows}"
    if seed_key not in st.session_state:
        st.session_state[seed_key] = pd.DataFrame([[""] * len(headers) for _ in range(rows)], columns=headers)

    edited = st.data_editor(
        st.session_state[seed_key],
        key=f"{key_base}_table_editor_{hashlib.sha1(seed_key.encode()).hexdigest()[:8]}",
        num_rows="dynamic",
        use_container_width=True,
        hide_index=True,
        height=min(460, 78 + max(rows, 3) * 38),
    )
    st.caption("You can add or delete rows directly in the table. Filled cells are included when your reasoning is checked.")

    if edited is None or len(edited) == 0:
        return ""
    clean = edited.fillna("").astype(str)
    nonempty = clean.apply(lambda row: any(v.strip() for v in row.tolist()), axis=1)
    clean = clean[nonempty]
    if clean.empty:
        return ""
    lines = ["Student working table:", "\t".join(headers)]
    for _, row in clean.iterrows():
        lines.append("\t".join(str(row[h]).strip() for h in headers))
    return "\n".join(lines)


def _looks_like_function_or_graph_question(pq: TargetedPracticeQuestion | None) -> bool:
    if pq is None:
        return False
    diagram = getattr(pq, "diagram_2d", None)
    if diagram is not None and bool(getattr(diagram, "show_axes", False)):
        return True
    text = " ".join([
        str(getattr(pq, "question", "") or ""),
        str(getattr(pq, "focus_prompt", "") or ""),
        str(getattr(pq, "target_skill", "") or ""),
    ]).lower()
    return bool(re.search(r"\b(function|graph|plot|curve|coordinate|coordinates|intercept|gradient|turning point|quadratic|linear graph)\b|f\s*\(|y\s*=", text))


def _normalise_function_expression(expr: str) -> str:
    value = str(expr or "").strip()
    value = _strip_math_transport_delimiters(value)
    value = value.replace("−", "-").replace("×", "*").replace("÷", "/").replace("^", "**")
    value = re.sub(r"^\s*(?:y|f\s*\(\s*x\s*\))\s*=\s*", "", value, flags=re.I)
    # Common implicit multiplication used by students: 2x, 3(x+1), x(x-2).
    value = re.sub(r"(?<=\d)(?=x\b)", "*", value, flags=re.I)
    value = re.sub(r"(?<=\d)(?=\()", "*", value)
    value = re.sub(r"(?<=x)(?=\()", "*", value, flags=re.I)
    value = re.sub(r"(?<=\))(?=(?:x|\d|\())", "*", value, flags=re.I)
    return value


_ALLOWED_FUNCS = {
    "sin": math.sin, "cos": math.cos, "tan": math.tan,
    "sqrt": math.sqrt, "exp": math.exp, "log": math.log,
    "ln": math.log, "abs": abs,
}
_ALLOWED_CONSTS = {"pi": math.pi, "e": math.e}


def _safe_eval_function(expr: str, x_value: float) -> float:
    """Evaluate a student-entered f(x) using a tiny arithmetic AST, never Python eval()."""
    tree = ast.parse(_normalise_function_expression(expr), mode="eval")

    def walk(node):
        if isinstance(node, ast.Expression):
            return walk(node.body)
        if isinstance(node, ast.Constant) and isinstance(node.value, (int, float)):
            return float(node.value)
        if isinstance(node, ast.Name):
            if node.id.lower() == "x":
                return float(x_value)
            if node.id.lower() in _ALLOWED_CONSTS:
                return float(_ALLOWED_CONSTS[node.id.lower()])
            raise ValueError("Only x and standard constants are allowed.")
        if isinstance(node, ast.UnaryOp) and isinstance(node.op, (ast.UAdd, ast.USub)):
            val = walk(node.operand)
            return val if isinstance(node.op, ast.UAdd) else -val
        if isinstance(node, ast.BinOp):
            a, b = walk(node.left), walk(node.right)
            if isinstance(node.op, ast.Add): return a + b
            if isinstance(node.op, ast.Sub): return a - b
            if isinstance(node.op, ast.Mult): return a * b
            if isinstance(node.op, ast.Div): return a / b
            if isinstance(node.op, ast.Pow): return a ** b
            raise ValueError("Unsupported operator.")
        if isinstance(node, ast.Call) and isinstance(node.func, ast.Name):
            name = node.func.id.lower()
            if name not in _ALLOWED_FUNCS or len(node.args) != 1:
                raise ValueError("Use sin, cos, tan, sqrt, exp, log/ln, or abs with one argument.")
            return float(_ALLOWED_FUNCS[name](walk(node.args[0])))
        raise ValueError("Unsupported function expression.")

    result = float(walk(tree))
    if not math.isfinite(result):
        raise ValueError("Function is not finite here.")
    return result


def _sample_function_scene(expr: str, *, x_min: float, x_max: float) -> dict[str, Any]:
    samples: list[list[float]] = []
    segments: list[list[list[float]]] = []
    n = 420
    dx = (x_max - x_min) / max(n - 1, 1)
    prior_y: float | None = None
    current: list[list[float]] = []
    finite_ys: list[float] = []
    for i in range(n):
        x = x_min + i * dx
        try:
            y = _safe_eval_function(expr, x)
            # Break the curve across vertical asymptotes / extreme jumps.
            if abs(y) > 1e5 or (prior_y is not None and abs(y - prior_y) > 80):
                raise ValueError
            current.append([round(x, 6), round(y, 6)])
            finite_ys.append(y)
            prior_y = y
        except Exception:
            if len(current) >= 2:
                segments.append(current)
            current = []
            prior_y = None
    if len(current) >= 2:
        segments.append(current)
    if not segments:
        raise ValueError("No plottable points were found in this x-range.")

    if finite_ys:
        finite_sorted = sorted(finite_ys)
        lo = finite_sorted[max(0, int(len(finite_sorted) * .03) - 1)]
        hi = finite_sorted[min(len(finite_sorted)-1, int(len(finite_sorted) * .97))]
        span = max(4.0, hi - lo)
        y_min, y_max = lo - .18 * span, hi + .18 * span
        y_min, y_max = max(-50.0, y_min), min(50.0, y_max)
        if y_max - y_min < 4:
            mid=(y_min+y_max)/2; y_min=mid-2; y_max=mid+2
    else:
        y_min, y_max = -10.0, 10.0

    polylines = [
        {"id": f"student_function_{idx}", "points": pts, "label": "", "dashed": False}
        for idx, pts in enumerate(segments)
    ]
    return {
        "x_min": float(x_min), "x_max": float(x_max),
        "y_min": float(y_min), "y_max": float(y_max),
        "show_axes": True, "keep_aspect": False,
        "points": [], "segments": [], "polylines": polylines, "circles": [], "angles": [],
    }


def _render_function_graph_tool(pq: TargetedPracticeQuestion | None, *, key_base: str) -> None:
    if not _looks_like_function_or_graph_question(pq):
        return
    st.markdown("#### Function / graph tool")
    show_graph = st.toggle(
        "Show graph of a function",
        key=f"{key_base}_show_function_graph",
        help="Enter a function of x, then explore it with the same plotting and geometry tools.",
    )
    if not show_graph:
        return
    c1, c2, c3 = st.columns([2.2, .8, .8])
    with c1:
        expr = st.text_input(
            "Function",
            key=f"{key_base}_function_expr",
            placeholder="e.g. y = x^2 - 4x + 3",
            help="Use x, +, −, ×, ÷, powers (^), brackets, and sin/cos/tan/sqrt/log if needed.",
        )
    with c2:
        x_min = float(st.number_input("x min", value=-10.0, step=1.0, key=f"{key_base}_function_xmin"))
    with c3:
        x_max = float(st.number_input("x max", value=10.0, step=1.0, key=f"{key_base}_function_xmax"))
    if not expr.strip():
        st.caption("Enter the function you want to plot.")
        return
    if x_max <= x_min:
        st.warning("x max must be greater than x min.")
        return
    try:
        scene = _sample_function_scene(expr, x_min=x_min, x_max=x_max)
    except Exception as exc:
        st.warning(f"The function could not be plotted: {exc}")
        return
    if _practice_diagram_component is None:
        st.info("The interactive graph component is unavailable in this browser session.")
        return
    graph_grid = st.toggle("Show gridlines", value=True, key=f"{key_base}_function_grid")
    _practice_diagram_component(
        data={
            "scene": scene,
            "show_grid": graph_grid,
            "step": {"highlight_ids": [], "dim_ids": [], "animate_ids": []},
            "visible_ids": [], "animate_ids": [], "reveal_mode": False, "animation_nonce": 0,
        },
        default={},
        key=f"{key_base}_function_graph",
        width="stretch",
        height="content",
    )
    st.caption("Use Point, Line, Segment, Angle, Distance and the other graph tools to explore the function. This graph is a working aid and does not reveal the answer automatically.")


def targeted_practice_input(
    label: str,
    *,
    key_base: str,
    height: int = 150,
    practice_question: TargetedPracticeQuestion | None = None,
) -> tuple[str, str, str, list[UploadedAsset]]:
    """Collect targeted-practice working from equation editor, text, or iPad handwriting."""
    _render_function_graph_tool(practice_question, key_base=key_base)
    table_text = _student_table_tool(key_base=key_base)
    mode = st.radio(
        "Working input method",
        ["Equation editor", "Handwrite on iPad", "Text working"],
        horizontal=True,
        key=f"{key_base}_mode",
        help=(
            "Equation editor is best for typed mathematics. Handwrite on iPad lets a student write with Apple Pencil, "
            "stylus, or finger, take a camera photo, or upload an image/PDF."
        ),
    )

    if mode == "Equation editor":
        latex_lines, ascii_lines = equation_working_editor(label, key=f"{key_base}_equation")
        explanation = st.text_area(
            "Optional explanation in words",
            key=f"{key_base}_explanation",
            height=80,
            placeholder="Example: I used the gradient formula first.",
        )
        used_latex = [line for line in latex_lines if line]
        used_ascii = [line for line in ascii_lines if line]
        working_lines = [f"Step {i}: \\({line}\\)" for i, line in enumerate(used_latex, 1)]
        if explanation.strip():
            working_lines.append(f"Student explanation: {explanation.strip()}")
        main_text = "\n".join(working_lines)
        if table_text:
            main_text = (main_text + "\n\n" + table_text).strip()
        offline = "\n".join(used_ascii)
        if table_text:
            offline = (offline + "\n\n" + table_text).strip()
        return main_text, mode, offline, []

    if mode == "Text working":
        value = st.text_area(
            label,
            key=f"{key_base}_text",
            height=height,
            placeholder="Show all parts and important reasoning steps, not only the final answer.",
        )
        main_text = value.strip()
        if table_text:
            main_text = (main_text + "\n\n" + table_text).strip()
        return main_text, mode, main_text, []

    st.caption(
        "On iPad, write directly with Apple Pencil/finger. The pad no longer submits after every stroke. "
        "When the page is complete, tap **Save handwriting** once before checking the answer. "
        "You can also use the camera/upload controls. For multi-part questions, label (a), (b), (c)."
    )
    canvas_asset = handwriting_pad(key=f"{key_base}_handwriting")
    if canvas_asset is None:
        st.info("If you are using the handwriting pad, finish the page and tap **Save handwriting** before pressing the marking button.")
    camera_file = st.camera_input(
        "Take a photo of handwritten working",
        key=f"{key_base}_camera",
        help="Allow camera access in Safari/Chrome when prompted.",
    )
    upload_files = st.file_uploader(
        "Or upload handwritten page(s)",
        type=["png", "jpg", "jpeg", "webp", "pdf"],
        accept_multiple_files=True,
        key=f"{key_base}_uploads",
        help="Useful for multiple pages or an existing photo/PDF from the iPad Files/Photos library.",
    )
    explanation = st.text_area(
        "Optional note to the tutor",
        key=f"{key_base}_hand_note",
        height=70,
        placeholder="Example: My working for (b) continues on the second page.",
    )

    browser_files: list[Any] = list(upload_files or [])
    if camera_file is not None:
        browser_files.insert(0, camera_file)
    try:
        assets = uploaded_assets(browser_files)
    except GeminiTutorError as exc:
        st.error(str(exc))
        assets = []
    if canvas_asset is not None:
        assets.insert(0, canvas_asset)
    total = sum(len(asset.data) for asset in assets)
    if total > MAX_TOTAL_BYTES:
        st.error("Handwritten working exceeds the app's 30 MB total upload limit.")
        assets = []

    text = explanation.strip()
    if table_text:
        text = (text + "\n\n" + table_text).strip()
    return text, "Handwritten working", table_text, assets


# Interactive visual explanations. The model supplies only declarative primitives;
# JavaScript rendering is owned by the app so uploaded/model content cannot execute code.
JSXGRAPH_VERSION = "2.4.0"
THREE_VERSION = "0.185.0"  # three.js r185

_VISUAL_2D_HTML = """
<div class="omt-visual2d-shell">
  <div class="omt-gg-toolbar" role="toolbar" aria-label="Interactive geometry tools">
    <button type="button" data-tool="move" class="active">Move</button>
    <button type="button" data-tool="point">Point</button>
    <button type="button" data-tool="line">Line</button>
    <button type="button" data-tool="segment">Segment</button>
    <button type="button" data-tool="ray">Ray</button>
    <button type="button" data-tool="vector">Vector</button>
    <button type="button" data-tool="circle">Circle</button>
    <button type="button" data-tool="polygon">Polygon</button>
    <button type="button" data-tool="finish" class="secondary">Finish</button>
    <button type="button" data-tool="midpoint">Midpoint</button>
    <button type="button" data-tool="perpendicular">Perpendicular</button>
    <button type="button" data-tool="parallel">Parallel</button>
    <button type="button" data-tool="angle">Measure angle</button>
    <button type="button" data-tool="distance">Distance</button>
    <button type="button" data-tool="delete">Delete</button>
    <button type="button" data-tool="undo" class="secondary">Undo</button>
    <button type="button" data-tool="clear" class="secondary">Clear</button>
    <button type="button" data-tool="snap" class="secondary active">Snap 0.5</button>
  </div>
  <div class="omt-gg-status">Use Move to pan/zoom, or select a construction tool.</div>
  <div class="omt-visual2d-board"></div>
  <div class="omt-visual-help">GeoGebra-style tools are for exploration. The tutor's construction remains separate from your added objects.</div>
</div>
"""

_VISUAL_2D_CSS = """
.omt-visual2d-shell { width: 100%; }
.omt-gg-toolbar { display:flex; gap:.38rem; overflow-x:auto; padding:.1rem 0 .48rem; scrollbar-width:thin; -webkit-overflow-scrolling:touch; }
.omt-gg-toolbar button { flex:0 0 auto; border:1px solid #cbd5e1; background:#fff; color:#334155; border-radius:.62rem; padding:.5rem .68rem; min-height:38px; font:650 .78rem/1 system-ui,sans-serif; cursor:pointer; }
.omt-gg-toolbar button.active { background:#eaf2ff; border-color:#60a5fa; color:#1d4ed8; }
.omt-gg-toolbar button.secondary { background:#f8fafc; }
.omt-gg-status { font-size:.78rem; color:#64748b; margin:0 0 .42rem; min-height:1.1rem; }
.omt-visual2d-board { width: 100%; height: min(62vw, 520px); min-height: 360px; border: 1px solid rgba(128,128,128,.28); border-radius: .75rem; overflow: hidden; background: #ffffff; touch-action:none; }
.omt-visual-help { margin-top: .35rem; font-size: .78rem; opacity: .68; }
@media (max-width: 640px) { .omt-visual2d-board { height: 420px; min-height: 340px; } .omt-gg-toolbar button { min-height:44px; padding:.62rem .76rem; } }
"""

_VISUAL_2D_JS = r"""
const JXG_URL = 'https://cdn.jsdelivr.net/npm/jsxgraph@1.12.2/distrib/jsxgraphcore.mjs';

async function loadJXG() {
  if (!globalThis.__omtJXGPromise) globalThis.__omtJXGPromise = import(JXG_URL);
  const mod = await globalThis.__omtJXGPromise;
  return mod.default || mod.JXG || mod;
}

function installGeoTools(board, toolbar, status, JXG) {
  if (!toolbar || !status) return () => {};
  let tool='move', picks=[], polygonPts=[], snap=true;
  const groups=[];
  const studentObjects=new Set();
  const pointStyle={name:'',fixed:false,size:4,strokeColor:'#dc2626',fillColor:'#fff1f2',strokeWidth:2,highlight:true};
  const lineStyle={fixed:false,strokeColor:'#dc2626',strokeWidth:2.6,highlight:true};
  const addGroup=(objs)=>{ const arr=objs.filter(Boolean); arr.forEach(o=>studentObjects.add(o)); groups.push(arr); board.update(); };
  const removeObjects=(objs)=>{ for(const o of [...objs].reverse()){ try{studentObjects.delete(o);board.removeObject(o);}catch(_){ } } board.update(); };
  const undo=()=>{ if(polygonPts.length){removeObjects([polygonPts.pop()]); return;} const g=groups.pop(); if(g) removeObjects(g); };
  const clear=()=>{ polygonPts=[]; while(groups.length) removeObjects(groups.pop()); };
  const roundSnap=(v)=>snap?Math.round(v*2)/2:v;
  const coords=(ev)=>{ const c=new JXG.Coords(JXG.COORDS_BY_SCREEN,[ev.offsetX,ev.offsetY],board); return [roundSnap(c.usrCoords[1]),roundSnap(c.usrCoords[2])]; };
  const mkPoint=(xy)=>board.create('point',xy,{...pointStyle});
  const length=(a,b)=>Math.hypot(a.X()-b.X(),a.Y()-b.Y());
  const midpointXY=(a,b)=>[(a.X()+b.X())/2,(a.Y()+b.Y())/2];
  const statusMap={
    move:'Drag to pan. Pinch or use the wheel to zoom.', point:'Tap to plot a point. Drag it to adjust.', line:'Tap two positions to draw an infinite line.', segment:'Tap two positions to draw a segment.', ray:'Tap the endpoint, then a second point for the ray direction.', vector:'Tap the start and end points of the vector.', circle:'Tap the centre, then a point on the circle.', polygon:'Tap polygon vertices, then press Finish.', midpoint:'Tap two positions to construct their midpoint.', perpendicular:'Tap two points for a reference line, then tap the point the perpendicular passes through.', parallel:'Tap two points for a reference line, then tap the point the parallel passes through.', angle:'Tap arm point 1, then the vertex, then arm point 2.', distance:'Tap two positions to measure their distance.', delete:'Tap one of your red constructions to delete it.'
  };
  const setTool=(name)=>{
    if(name!=='polygon' && polygonPts.length){ status.textContent='Finish or Undo the current polygon before switching tools.'; return; }
    tool=name; picks=[];
    toolbar.querySelectorAll('button[data-tool]').forEach(b=>{ if(!['snap','finish','undo','clear'].includes(b.dataset.tool)) b.classList.toggle('active',b.dataset.tool===name); });
    status.textContent=statusMap[name]||'Select a construction tool.';
  };
  const finishPolygon=()=>{
    if(polygonPts.length<3){status.textContent='A polygon needs at least 3 vertices.';return;}
    const poly=board.create('polygon',polygonPts,{withLines:true,fillColor:'#fecaca',fillOpacity:.12,borders:{strokeColor:'#dc2626',strokeWidth:2.4},vertices:{visible:true}});
    addGroup([poly,...polygonPts]); polygonPts=[]; status.textContent='Polygon added. Choose another tool or start another polygon.';
  };
  const deleteUnder=(ev)=>{
    let hits=[]; try{hits=board.getAllUnderMouse(ev)||[];}catch(_){ }
    const hit=hits.find(o=>studentObjects.has(o));
    if(hit){ const idx=groups.findIndex(g=>g.includes(hit)); if(idx>=0){const [g]=groups.splice(idx,1);removeObjects(g);} else removeObjects([hit]); status.textContent='Construction deleted.'; }
    else { status.textContent='Tap a red construction to delete it. If selection is difficult, use Undo.'; }
  };
  const addAngleMeasure=(a,v,b)=>{
    const ang=board.create('angle',[a,v,b],{...lineStyle,radius:.7,fillColor:'#fee2e2',fillOpacity:.18,name:'',withLabel:false});
    const txt=board.create('text',[()=>v.X()+0.55,()=>v.Y()+0.55,()=>`${(ang.Value()*180/Math.PI).toFixed(1)}°`],{fixed:true,fontSize:13,color:'#b91c1c'});
    addGroup([a,v,b,ang,txt]);
  };
  const addDistance=(a,b)=>{
    const seg=board.create('segment',[a,b],{...lineStyle,dash:2});
    const txt=board.create('text',[()=> (a.X()+b.X())/2,()=> (a.Y()+b.Y())/2,()=> length(a,b).toFixed(2)],{fixed:true,fontSize:13,color:'#b91c1c'});
    addGroup([a,b,seg,txt]);
  };
  const handlePointTool=(xy)=>{ const p=mkPoint(xy); addGroup([p]); };
  const handleMulti=(xy)=>{
    const p=mkPoint(xy); picks.push(p);
    const need=(['perpendicular','parallel','angle'].includes(tool)?3:2);
    status.textContent=`${picks.length}/${need} point${need===1?'':'s'} selected.`;
    if(picks.length<need) return;
    const [a,b,c]=picks; picks=[];
    if(tool==='line'){const o=board.create('line',[a,b],{...lineStyle,straightFirst:true,straightLast:true});addGroup([a,b,o]);}
    else if(tool==='segment'){const o=board.create('segment',[a,b],lineStyle);addGroup([a,b,o]);}
    else if(tool==='ray'){const o=board.create('line',[a,b],{...lineStyle,straightFirst:false,straightLast:true});addGroup([a,b,o]);}
    else if(tool==='vector'){const o=board.create('arrow',[a,b],lineStyle);addGroup([a,b,o]);}
    else if(tool==='circle'){const o=board.create('circle',[a,b],{...lineStyle,fillOpacity:0});addGroup([a,b,o]);}
    else if(tool==='midpoint'){const m=board.create('midpoint',[a,b],{...pointStyle,fillColor:'#fef3c7',strokeColor:'#d97706'});addGroup([a,b,m]);}
    else if(tool==='distance') addDistance(a,b);
    else if(tool==='angle') addAngleMeasure(a,b,c);
    else if(tool==='perpendicular'){
      const base=board.create('line',[a,b],{...lineStyle,strokeColor:'#94a3b8',strokeWidth:1.6,dash:2});
      const perp=board.create('perpendicular',[base,c],{...lineStyle}); addGroup([a,b,c,base,perp]);
    }
    else if(tool==='parallel'){
      const base=board.create('line',[a,b],{...lineStyle,strokeColor:'#94a3b8',strokeWidth:1.6,dash:2});
      const para=board.create('parallel',[base,c],{...lineStyle}); addGroup([a,b,c,base,para]);
    }
    status.textContent=(statusMap[tool]||'Construction added.')+' Construction added.';
  };
  const clickHandler=(ev)=>{
    const b=ev.target.closest('button[data-tool]'); if(!b)return;
    const name=b.dataset.tool;
    if(name==='clear'){clear();status.textContent='Your constructions were cleared.';return;}
    if(name==='undo'){undo();status.textContent='Last construction removed.';return;}
    if(name==='finish'){finishPolygon();return;}
    if(name==='snap'){snap=!snap;b.classList.toggle('active',snap);b.textContent=snap?'Snap 0.5':'Snap off';status.textContent=snap?'Coordinate snapping is on (0.5 units).':'Coordinate snapping is off.';return;}
    setTool(name);
  };
  toolbar.addEventListener('click',clickHandler);
  const downHandler=(ev)=>{
    if(tool==='move')return;
    if(tool==='delete'){deleteUnder(ev);return;}
    const xy=coords(ev);
    if(tool==='point'){handlePointTool(xy);return;}
    if(tool==='polygon'){const p=mkPoint(xy);polygonPts.push(p);status.textContent=`Polygon: ${polygonPts.length} vertices. Add more or press Finish.`;board.update();return;}
    handleMulti(xy);
    board.update();
  };
  board.on('down',downHandler);
  setTool('move');
  return ()=>{ try{toolbar.removeEventListener('click',clickHandler);}catch(_){ } };
}

function styleFor(id, highlight, dim, kind='line') {
  const hi = highlight.has(id);
  const low = dim.has(id);
  if (kind === 'point') {
    return { strokeColor: hi ? '#dc2626' : (low ? '#cbd5e1' : '#0f172a'), fillColor: hi ? '#dc2626' : (low ? '#e2e8f0' : '#0f172a'), opacity: low ? 0.35 : 1, size: hi ? 5 : 3.5 };
  }
  return { strokeColor: hi ? '#dc2626' : (low ? '#cbd5e1' : '#475569'), strokeWidth: hi ? 4 : 2.2, opacity: low ? 0.28 : 0.95 };
}

function pulsePoint(board, point, targetSize, targetOpacity) {
  const start = performance.now();
  const duration = 620;
  const tick = (now) => {
    const t = Math.min(1, (now - start) / duration);
    const ease = 1 - Math.pow(1 - t, 3);
    const size = 0.7 + (targetSize * 1.25 - 0.7) * ease;
    point.setAttribute({ size, opacity: Math.max(0.05, targetOpacity * ease) });
    board.update();
    if (t < 1) requestAnimationFrame(tick);
    else { point.setAttribute({ size: targetSize, opacity: targetOpacity }); board.update(); }
  };
  requestAnimationFrame(tick);
}

export default async function(component) {
  const { parentElement, data } = component;
  const stage = parentElement.querySelector('.omt-visual2d-board');
  const toolbar = parentElement.querySelector('.omt-gg-toolbar');
  const status = parentElement.querySelector('.omt-gg-status');
  const scene = data?.scene || {};
  const step = data?.step || {};
  const highlight = new Set(step.highlight_ids || []);
  const dim = new Set(step.dim_ids || []);
  const animate = new Set([...(step.highlight_ids || []), ...(data?.animate_ids || step.animate_ids || [])]);
  const revealMode = Boolean(data?.reveal_mode);
  const visible = new Set(data?.visible_ids || []);
  const isVisible = (id) => !revealMode || visible.has(id) || highlight.has(id) || animate.has(id);
  // Replay/Next must always produce visible motion. If this step contains no
  // explicit construction action, replay the currently visible construction.
  if (animate.size === 0 && Number(data?.animation_nonce || 0) > 0) {
    for (const group of [scene.points || [], scene.segments || [], scene.polylines || [], scene.circles || [], scene.angles || []]) {
      for (const item of group) if (isVisible(item.id)) animate.add(item.id);
    }
  }
  let JXG;
  try { JXG = await loadJXG(); } catch (err) { console.error('JSXGraph load failed', err); stage.textContent = 'Interactive 2D visual could not load. Reload once; if this persists, the JSXGraph library may be blocked by the network.'; return; }

  try { if (parentElement.__omtBoard) JXG.JSXGraph.freeBoard(parentElement.__omtBoard); } catch (_) {}
  stage.replaceChildren();
  stage.id = `omt-jxg-${Math.random().toString(36).slice(2)}`;
  const xMin = Number(scene.x_min ?? -5), xMax = Number(scene.x_max ?? 5);
  const yMin = Number(scene.y_min ?? -5), yMax = Number(scene.y_max ?? 5);
  const board = JXG.JSXGraph.initBoard(stage.id, {
    boundingbox: [xMin, yMax, xMax, yMin],
    axis: Boolean(scene.show_axes),
    grid: Boolean(data?.show_grid),
    keepaspectratio: scene.keep_aspect !== false,
    showNavigation: false,
    showCopyright: false,
    pan: { enabled: true, needShift: false },
    zoom: { wheel: true, needShift: false, factorX: 1.2, factorY: 1.2 },
  });
  parentElement.__omtBoard = board;

  // If a visible construction depends on endpoints/angle arms, show those points too.
  const neededPoints = new Set();
  for (const seg of (scene.segments || [])) if (isVisible(seg.id)) { neededPoints.add(seg.start); neededPoints.add(seg.end); }
  for (const ang of (scene.angles || [])) if (isVisible(ang.id)) { neededPoints.add(ang.arm1); neededPoints.add(ang.vertex); neededPoints.add(ang.arm2); }

  const pts = new Map();
  for (const p of (scene.points || [])) {
    const st = styleFor(p.id, highlight, dim, 'point');
    const show = isVisible(p.id) || neededPoints.has(p.id);
    const obj = board.create('point', [Number(p.x), Number(p.y)], {
      name: p.label || '', fixed: true, highlight: false, withLabel: show && Boolean(p.label), visible: show,
      strokeColor: st.strokeColor, fillColor: st.fillColor, opacity: animate.has(p.id) ? 0.04 : st.opacity,
      size: animate.has(p.id) ? 0.7 : st.size,
      label: { fontSize: 14, offset: [7, 7] },
    });
    pts.set(p.id, obj);
    if (show && animate.has(p.id)) pulsePoint(board, obj, st.size, st.opacity);
  }

  for (const seg of (scene.segments || [])) {
    if (!isVisible(seg.id)) continue;
    const a = pts.get(seg.start), b = pts.get(seg.end); if (!a || !b) continue;
    const st = styleFor(seg.id, highlight, dim);
    const attrs = {
      name: seg.label || '', withLabel: Boolean(seg.label) && !animate.has(seg.id), fixed: true, highlight: false,
      strokeColor: st.strokeColor, strokeWidth: st.strokeWidth, opacity: st.opacity,
      dash: seg.dashed ? 2 : 0, label: { fontSize: 13 },
    };
    if (animate.has(seg.id)) {
      const mover = board.create('point', [a.X(), a.Y()], { visible:false, fixed:true, name:'' });
      board.create('segment', [a, mover], attrs);
      setTimeout(() => mover.moveTo([b.X(), b.Y()], 900), 90);
      if (seg.label) setTimeout(() => board.create('text', [(a.X()+b.X())/2, (a.Y()+b.Y())/2, seg.label], {fixed:true, fontSize:13, color:st.strokeColor}), 980);
    } else {
      board.create('segment', [a,b], attrs);
    }
  }

  for (const poly of (scene.polylines || [])) {
    if (!isVisible(poly.id)) continue;
    const samples = Array.isArray(poly.points) ? poly.points.filter(v => Array.isArray(v) && v.length >= 2) : [];
    if (samples.length < 2) continue;
    const st = styleFor(poly.id, highlight, dim);
    if (animate.has(poly.id)) {
      const delay = Math.max(18, Math.min(90, 850 / Math.max(1, samples.length - 1)));
      samples.slice(0, -1).forEach((a, i) => {
        const b = samples[i + 1];
        setTimeout(() => board.create('segment', [[Number(a[0]),Number(a[1])],[Number(b[0]),Number(b[1])]], {
          fixed:true, highlight:false, strokeColor:st.strokeColor, strokeWidth:st.strokeWidth, opacity:st.opacity, dash:poly.dashed ? 2 : 0,
        }), 80 + i * delay);
      });
      if (poly.label) {
        const m = samples[Math.floor(samples.length/2)];
        setTimeout(() => board.create('text', [Number(m[0]), Number(m[1]), poly.label], { fixed:true, fontSize:13, color: st.strokeColor }), 120 + samples.length * delay);
      }
    } else {
      const xs = samples.map(v => Number(v[0])), ys = samples.map(v => Number(v[1]));
      board.create('curve', [xs, ys], { fixed:true, highlight:false, strokeColor:st.strokeColor, strokeWidth:st.strokeWidth, opacity:st.opacity, dash:poly.dashed ? 2 : 0 });
      if (poly.label) {
        const m = samples[Math.floor(samples.length/2)];
        board.create('text', [Number(m[0]), Number(m[1]), poly.label], { fixed:true, fontSize:13, color: st.strokeColor });
      }
    }
  }

  for (const c of (scene.circles || [])) {
    if (!isVisible(c.id)) continue;
    const st = styleFor(c.id, highlight, dim);
    board.create('circle', [[Number(c.center_x), Number(c.center_y)], Number(c.radius)], {
      fixed:true, highlight:false, strokeColor:st.strokeColor, strokeWidth:st.strokeWidth, opacity:animate.has(c.id) ? 0.25 : st.opacity,
      fillOpacity:0, name:c.label || '', withLabel:Boolean(c.label),
    });
  }

  for (const ang of (scene.angles || [])) {
    if (!isVisible(ang.id)) continue;
    const a = pts.get(ang.arm1), v = pts.get(ang.vertex), c = pts.get(ang.arm2); if (!a || !v || !c) continue;
    const st = styleFor(ang.id, highlight, dim);
    board.create('angle', [a,v,c], {
      name:ang.label || '', withLabel:Boolean(ang.label), fixed:true, highlight:false,
      strokeColor:st.strokeColor, fillColor:st.strokeColor, fillOpacity:highlight.has(ang.id) ? 0.16 : 0.06,
      strokeWidth:st.strokeWidth, radius:0.55, label:{fontSize:13},
    });
  }
  const removeGeoTools = installGeoTools(board, toolbar, status, JXG);
  parentElement.__omtGeoToolsCleanup = removeGeoTools;
  board.update();
}
"""

_VISUAL_3D_HTML = """
<div class="omt-visual3d-shell">
  <div class="omt-visual3d-toolbar" role="toolbar" aria-label="3D view controls">
    <button type="button" data-action="rotate" class="is-active">Rotate</button>
    <button type="button" data-action="pan">Pan</button>
    <span class="omt-visual3d-divider"></span>
    <button type="button" data-action="home">Fit</button>
    <button type="button" data-action="source">Question view</button>
    <button type="button" data-action="iso">Explore iso</button>
    <button type="button" data-action="front">Front</button>
    <button type="button" data-action="top">Top</button>
    <button type="button" data-action="side">Side</button>
  </div>
  <div class="omt-visual3d-stage"></div>
  <div class="omt-visual-help">Rotate mode explores the solid; Pan moves it. For top/front/side questions, use Front, Top and Side to compare the reconstruction with the source views.</div>
</div>
"""

_VISUAL_3D_CSS = """
.omt-visual3d-shell { width: 100%; }
.omt-visual3d-toolbar { display:flex; flex-wrap:wrap; gap:.4rem; align-items:center; margin:0 0 .45rem 0; }
.omt-visual3d-toolbar button { appearance:none; border:1px solid rgba(100,116,139,.35); background:#fff; color:#334155; border-radius:.55rem; padding:.46rem .72rem; font:600 .78rem/1 system-ui,sans-serif; cursor:pointer; min-height:36px; }
.omt-visual3d-toolbar button:hover { background:#f1f5f9; }
.omt-visual3d-toolbar button.is-active { background:#e0f2fe; border-color:#38bdf8; color:#075985; }
.omt-visual3d-divider { width:1px; height:26px; background:rgba(100,116,139,.25); margin:0 .1rem; }
.omt-visual3d-stage { width: 100%; height: min(66vw, 590px); min-height: 430px; border: 1px solid rgba(100,116,139,.32); border-radius: .75rem; overflow: hidden; background:#f8fafc; touch-action: none; position: relative; }
.omt-visual3d-stage canvas { display:block; width:100%; height:100%; }
.omt-visual-help { margin-top:.38rem; font-size:.78rem; color:#64748b; }
@media (max-width: 640px) {
  .omt-visual3d-stage { height: 520px; min-height: 430px; }
  .omt-visual3d-toolbar button { min-height:42px; padding:.55rem .78rem; font-size:.82rem; }
}
"""

_VISUAL_3D_JS = r"""
const THREE_URL = 'https://esm.sh/three@0.185.0';
const ORBIT_URL = 'https://esm.sh/three@0.185.0/examples/jsm/controls/OrbitControls.js';

async function loadThree() {
  if (!globalThis.__omtThreePromise) globalThis.__omtThreePromise = Promise.all([import(THREE_URL), import(ORBIT_URL)]);
  const [THREE, controls] = await globalThis.__omtThreePromise;
  return { THREE, OrbitControls: controls.OrbitControls };
}

function textSprite(THREE, text, color='#0f172a', scale=1.0) {
  if (!text) return null;
  const canvas=document.createElement('canvas'); canvas.width=640; canvas.height=160;
  const ctx=canvas.getContext('2d'); ctx.clearRect(0,0,640,160);
  ctx.font='600 48px system-ui, sans-serif';
  const w=Math.min(600, Math.max(120, ctx.measureText(text).width+42));
  ctx.fillStyle='rgba(255,255,255,.88)'; ctx.beginPath(); ctx.roundRect((640-w)/2,34,w,92,20); ctx.fill();
  ctx.strokeStyle='rgba(100,116,139,.25)'; ctx.stroke();
  ctx.fillStyle=color; ctx.textAlign='center'; ctx.textBaseline='middle'; ctx.fillText(text,320,80);
  const texture=new THREE.CanvasTexture(canvas); texture.minFilter=THREE.LinearFilter;
  const material=new THREE.SpriteMaterial({map:texture,transparent:true,depthTest:false});
  const sprite=new THREE.Sprite(material); sprite.scale.set(2.1*scale,.52*scale,1); sprite.userData.__texture=texture; return sprite;
}

function edgeMaterial(THREE,id,highlight,dim,dashed=false){
  const hi=highlight.has(id), low=dim.has(id); const color=hi?0xdc2626:(low?0x94a3b8:0x334155); const opacity=low?0.28:1;
  if(dashed) return new THREE.LineDashedMaterial({color,transparent:true,opacity,dashSize:.18,gapSize:.1});
  return new THREE.LineBasicMaterial({color,transparent:true,opacity});
}

function hashColor(id){
  const palette=[0x60a5fa,0x34d399,0xf59e0b,0xa78bfa,0x22d3ee,0xfb7185,0x84cc16];
  let h=0; for(const ch of String(id||'')) h=(h*31+ch.charCodeAt(0))>>>0;
  return palette[h%palette.length];
}

function solidMaterial(THREE,id,highlight,dim){
  const hi=highlight.has(id), low=dim.has(id);
  return new THREE.MeshStandardMaterial({
    color: hi?0xf97316:hashColor(id),
    roughness:.72, metalness:.02,
    transparent: low,
    opacity: low?.42:1,
    side: THREE.DoubleSide,
    depthWrite: !low,
  });
}

function orientAxis(mesh,axis){
  if(axis==='x') mesh.rotation.z=-Math.PI/2;
  else if(axis==='z') mesh.rotation.x=Math.PI/2;
}

function addSolidEdges(THREE,mesh,id,highlight,dim,scene){
  const geom=new THREE.EdgesGeometry(mesh.geometry,20);
  const mat=new THREE.LineBasicMaterial({color:highlight.has(id)?0xc2410c:(dim.has(id)?0x94a3b8:0x334155),transparent:true,opacity:dim.has(id)?.42:.96});
  const lines=new THREE.LineSegments(geom,mat); lines.position.copy(mesh.position); lines.rotation.copy(mesh.rotation); lines.scale.copy(mesh.scale); scene.add(lines);
  return lines;
}

function addSolidLabel(THREE,mesh,label,id,highlight,scene){
  if(!label) return;
  const box=new THREE.Box3().setFromObject(mesh); const top=box.getCenter(new THREE.Vector3()); top.y=box.max.y+.22;
  const sp=textSprite(THREE,label,highlight.has(id)?'#c2410c':'#334155',.78); if(sp){sp.position.copy(top);scene.add(sp);}
}

export default async function(component) {
  const {parentElement,data}=component;
  const stage=parentElement.querySelector('.omt-visual3d-stage');
  const toolbar=parentElement.querySelector('.omt-visual3d-toolbar');
  const sceneData=data?.scene||{}, step=data?.step||{};
  const sourceView=sceneData.source_view||null;
  const sourceProjection=String(sourceView?.projection||'unknown').toLowerCase();
  const isOrthographicSet=sourceProjection==='orthographic_set';
  const useOrthographic=['isometric','orthographic','orthographic_set','oblique'].includes(sourceProjection);
  const sourceButton=toolbar?.querySelector('button[data-action="source"]');
  const help=parentElement.querySelector('.omt-visual-help');
  if(isOrthographicSet){
    if(sourceButton) sourceButton.textContent='Reconstruction view';
    if(help) help.textContent='This question provides top/front/side orthographic views, not a single isometric source view. Use Front, Top and Side to compare the reconstructed solid with the question; Rotate/Pan are for exploration.';
  }
  const highlight=new Set(step.highlight_ids||[]), dim=new Set(step.dim_ids||[]);
  const explicitAnimate=new Set(data?.animate_ids||step.animate_ids||[]);
  const animate=new Set([...(step.highlight_ids||[]), ...explicitAnimate]);
  const newlyRevealed=new Set(step.reveal_ids||[]);
  const revealMode=Boolean(data?.reveal_mode), visible=new Set(data?.visible_ids||[]);
  const solidGroups=[sceneData.boxes||[],sceneData.cylinders||[],sceneData.cones||[],sceneData.spheres||[],sceneData.extrusions||[]];
  const solidIds=new Set(solidGroups.flat().map(x=>x.id));
  // Physical solids remain visible throughout. revealMode is for construction geometry,
  // not for hiding the object the student is trying to understand.
  const isVisible=(id)=>solidIds.has(id)||!revealMode||visible.has(id)||highlight.has(id)||animate.has(id);
  if(animate.size===0 && Number(data?.animation_nonce||0)>0){
    for(const group of [sceneData.vertices||[],sceneData.edges||[],sceneData.faces||[],sceneData.angles||[],...solidGroups]){
      for(const item of group) if(isVisible(item.id)) animate.add(item.id);
    }
  }
  if(parentElement.__omtThreeCleanup) parentElement.__omtThreeCleanup();
  stage.replaceChildren();

  let THREE,OrbitControls;
  try{({THREE,OrbitControls}=await loadThree());}catch(err){stage.textContent='Interactive 3D visual could not load.';return;}

  const scene=new THREE.Scene();
  // Printed isometric/orthographic exam diagrams use parallel projection. Using a
  // perspective camera changes edge directions and makes a correct solid look unlike
  // the source. Source-calibrated views therefore use OrthographicCamera by default.
  const camera=useOrthographic
    ? new THREE.OrthographicCamera(-5,5,5,-5,.01,2000)
    : new THREE.PerspectiveCamera(38,1,.01,2000);
  if(Array.isArray(sourceView?.camera_up)&&sourceView.camera_up.length===3){
    camera.up.set(Number(sourceView.camera_up[0]),Number(sourceView.camera_up[1]),Number(sourceView.camera_up[2])).normalize();
  }
  const renderer=new THREE.WebGLRenderer({antialias:true,alpha:true});
  renderer.setPixelRatio(Math.min(window.devicePixelRatio||1,2)); renderer.setClearColor(0xf8fafc,1);
  renderer.outputColorSpace=THREE.SRGBColorSpace; stage.appendChild(renderer.domElement);
  const controls=new OrbitControls(camera,renderer.domElement);
  controls.enableDamping=true; controls.dampingFactor=.075; controls.enablePan=true; controls.enableRotate=true;
  controls.screenSpacePanning=true; controls.rotateSpeed=.62; controls.panSpeed=.72; controls.zoomSpeed=.82;
  controls.mouseButtons={LEFT:THREE.MOUSE.ROTATE,MIDDLE:THREE.MOUSE.DOLLY,RIGHT:THREE.MOUSE.PAN};
  controls.touches={ONE:THREE.TOUCH.ROTATE,TWO:THREE.TOUCH.DOLLY_PAN};
  scene.add(new THREE.HemisphereLight(0xffffff,0x94a3b8,2.6));
  const keyLight=new THREE.DirectionalLight(0xffffff,2.2); keyLight.position.set(7,10,8); scene.add(keyLight);
  const fillLight=new THREE.DirectionalLight(0xffffff,1.0); fillLight.position.set(-7,4,-6); scene.add(fillLight);

  const vertices=new Map(), vertexCoords=[], solidCoords=[];
  for(const v of (sceneData.vertices||[])){const pos=new THREE.Vector3(Number(v.x),Number(v.y),Number(v.z));vertexCoords.push(pos);vertices.set(v.id,pos);}
  const neededVertices=new Set();
  for(const e of (sceneData.edges||[])) if(isVisible(e.id)){neededVertices.add(e.start);neededVertices.add(e.end);}
  for(const f of (sceneData.faces||[])) if(isVisible(f.id)) for(const id of (f.vertices||[])) neededVertices.add(id);
  for(const a of (sceneData.angles||[])) if(isVisible(a.id)){neededVertices.add(a.arm1);neededVertices.add(a.vertex);neededVertices.add(a.arm2);}

  const solidTweens=[];
  const registerSolid=(mesh,id,label)=>{
    const baseScale=mesh.scale.clone(); const doAnimate=animate.has(id);
    if(doAnimate){const startScale=newlyRevealed.has(id)?.08:.86;mesh.scale.multiplyScalar(startScale); solidTweens.push({mesh,baseScale,startScale,start:performance.now()+60,duration:720});}
    scene.add(mesh); const edge=addSolidEdges(THREE,mesh,id,highlight,dim,scene); mesh.userData.__edge=edge;
    addSolidLabel(THREE,mesh,label,id,highlight,scene);
    const box=new THREE.Box3().setFromObject(mesh); if(Number.isFinite(box.min.x)){solidCoords.push(box.min.clone(),box.max.clone());}
  };

  for(const b of (sceneData.boxes||[])){
    if(!isVisible(b.id)) continue;
    const geom=new THREE.BoxGeometry(Number(b.width),Number(b.height),Number(b.depth));
    const mesh=new THREE.Mesh(geom,solidMaterial(THREE,b.id,highlight,dim));
    mesh.position.set(...(b.center||[0,0,0])); const r=b.rotation||[0,0,0]; mesh.rotation.set(Number(r[0]||0),Number(r[1]||0),Number(r[2]||0));
    registerSolid(mesh,b.id,b.label||'');
  }

  for(const c of (sceneData.cylinders||[])){
    if(!isVisible(c.id)) continue;
    const geom=new THREE.CylinderGeometry(Number(c.radius),Number(c.radius),Number(c.height),48,1,false);
    const mesh=new THREE.Mesh(geom,solidMaterial(THREE,c.id,highlight,dim)); mesh.position.set(...(c.center||[0,0,0])); orientAxis(mesh,c.axis||'y'); registerSolid(mesh,c.id,c.label||'');
  }

  for(const c of (sceneData.cones||[])){
    if(!isVisible(c.id)) continue;
    const geom=new THREE.ConeGeometry(Number(c.radius),Number(c.height),48,1,false);
    const mesh=new THREE.Mesh(geom,solidMaterial(THREE,c.id,highlight,dim)); mesh.position.set(...(c.center||[0,0,0])); orientAxis(mesh,c.axis||'y'); registerSolid(mesh,c.id,c.label||'');
  }

  for(const sp of (sceneData.spheres||[])){
    if(!isVisible(sp.id)) continue;
    const geom=new THREE.SphereGeometry(Number(sp.radius),40,24);
    const mesh=new THREE.Mesh(geom,solidMaterial(THREE,sp.id,highlight,dim)); mesh.position.set(...(sp.center||[0,0,0])); registerSolid(mesh,sp.id,sp.label||'');
  }

  for(const ex of (sceneData.extrusions||[])){
    if(!isVisible(ex.id)||!(ex.profile||[]).length) continue;
    const profile=ex.profile; const shape=new THREE.Shape(); shape.moveTo(Number(profile[0][0]),Number(profile[0][1]));
    for(let i=1;i<profile.length;i++) shape.lineTo(Number(profile[i][0]),Number(profile[i][1])); shape.closePath();
    const geom=new THREE.ExtrudeGeometry(shape,{depth:Number(ex.depth),bevelEnabled:false,steps:1}); geom.center();
    if((ex.axis||'z')==='x') geom.rotateY(Math.PI/2); else if((ex.axis||'z')==='y') geom.rotateX(-Math.PI/2);
    const mesh=new THREE.Mesh(geom,solidMaterial(THREE,ex.id,highlight,dim)); mesh.position.set(...(ex.center||[0,0,0])); registerSolid(mesh,ex.id,ex.label||'');
  }

  const focusVertices=new Set();
  for(const e of (sceneData.edges||[])) if(highlight.has(e.id)){focusVertices.add(e.start);focusVertices.add(e.end);}
  for(const a of (sceneData.angles||[])) if(highlight.has(a.id)){focusVertices.add(a.arm1);focusVertices.add(a.vertex);focusVertices.add(a.arm2);}
  for(const v of (sceneData.vertices||[])){
    if(!isVisible(v.id)&&!neededVertices.has(v.id)) continue;
    const pos=vertices.get(v.id), hi=highlight.has(v.id), low=dim.has(v.id);
    const geom=new THREE.SphereGeometry(hi?.075:.038,16,10); const mat=new THREE.MeshBasicMaterial({color:hi?0xdc2626:(low?0xcbd5e1:0x334155),transparent:true,opacity:low?.22:(hi?1:.48)});
    const mesh=new THREE.Mesh(geom,mat);mesh.position.copy(pos);scene.add(mesh);
    // Avoid the unreadable cloud of labels seen in complex composite solids.
    // Labels appear only when the current teaching step actually uses the vertex.
    if(hi||focusVertices.has(v.id)){
      const label=textSprite(THREE,v.label||v.id,hi?'#dc2626':'#334155',.64);if(label){label.position.copy(pos).add(new THREE.Vector3(.11,.15,.07));scene.add(label);}
    }
  }

  const edgeTweens=[];
  for(const e of (sceneData.edges||[])){
    if(!isVisible(e.id)) continue;
    const a=vertices.get(e.start),b=vertices.get(e.end);if(!a||!b)continue;
    const mat=edgeMaterial(THREE,e.id,highlight,dim,Boolean(e.dashed));
    if(animate.has(e.id)){
      const arr=new Float32Array([a.x,a.y,a.z,a.x,a.y,a.z]);
      const geom=new THREE.BufferGeometry();geom.setAttribute('position',new THREE.BufferAttribute(arr,3));
      const line=new THREE.Line(geom,mat);scene.add(line);edgeTweens.push({attr:geom.getAttribute('position'),a:a.clone(),b:b.clone(),start:performance.now()+80,duration:900,line});
    }else{
      const geom=new THREE.BufferGeometry().setFromPoints([a,b]);const line=new THREE.Line(geom,mat);if(e.dashed)line.computeLineDistances();scene.add(line);
    }
    if(e.label){const m=a.clone().add(b).multiplyScalar(.5);const sp=textSprite(THREE,e.label,highlight.has(e.id)?'#dc2626':'#334155',.66);if(sp){sp.position.copy(m).add(new THREE.Vector3(.08,.08,.08));scene.add(sp);}}
  }

  for(const f of (sceneData.faces||[])){
    if(!isVisible(f.id))continue;const vv=(f.vertices||[]).map(id=>vertices.get(id)).filter(Boolean);if(vv.length<3)continue;
    const arr=[];for(let i=1;i<vv.length-1;i++){for(const p of[vv[0],vv[i],vv[i+1]])arr.push(p.x,p.y,p.z);}const geom=new THREE.BufferGeometry();geom.setAttribute('position',new THREE.Float32BufferAttribute(arr,3));geom.computeVertexNormals();
    const hi=highlight.has(f.id),low=dim.has(f.id);const mat=new THREE.MeshStandardMaterial({color:hi?0xf97316:0x94a3b8,transparent:true,opacity:low?.06:(hi?.32:.18),side:THREE.DoubleSide,depthWrite:false,roughness:.85});scene.add(new THREE.Mesh(geom,mat));
  }

  for(const aDef of (sceneData.angles||[])){
    if(!isVisible(aDef.id))continue;const pa=vertices.get(aDef.arm1),pv=vertices.get(aDef.vertex),pc=vertices.get(aDef.arm2);if(!pa||!pv||!pc)continue;
    const u=pa.clone().sub(pv).normalize(),w=pc.clone().sub(pv).normalize(),dot=THREE.MathUtils.clamp(u.dot(w),-1,1),theta=Math.acos(dot),sin=Math.sin(theta);if(theta<.02||Math.abs(sin)<1e-4)continue;
    const tangent=w.clone().sub(u.clone().multiplyScalar(dot)).normalize(),radius=.34,samples=[];for(let i=0;i<=28;i++){const t=theta*i/28;samples.push(pv.clone().add(u.clone().multiplyScalar(Math.cos(t)*radius)).add(tangent.clone().multiplyScalar(Math.sin(t)*radius)));}
    const geom=new THREE.BufferGeometry().setFromPoints(samples),mat=edgeMaterial(THREE,aDef.id,highlight,dim,false);scene.add(new THREE.Line(geom,mat));
    if(aDef.label){const mid=samples[Math.floor(samples.length/2)],sp=textSprite(THREE,aDef.label,highlight.has(aDef.id)?'#dc2626':'#334155',.66);if(sp){sp.position.copy(mid);scene.add(sp);}}
  }

  // Frame the PHYSICAL SOLID first. Scattered named vertices must never make the
  // camera zoom so far out that the object disappears.
  const fitCoords=solidCoords.length?solidCoords:vertexCoords;
  let center=new THREE.Vector3(0,0,0),radius=2,fitBox=null;
  if(fitCoords.length){fitBox=new THREE.Box3().setFromPoints(fitCoords);center=fitBox.getCenter(new THREE.Vector3());const sphere=fitBox.getBoundingSphere(new THREE.Sphere());radius=Math.max(sphere.radius,1.0);}
  const fovRad=THREE.MathUtils.degToRad(useOrthographic?38:camera.fov); const fitDistance=Math.max(radius/Math.sin(fovRad/2)*1.18,radius*2.25);
  camera.near=Math.max(.005,radius/1000); camera.far=Math.max(150,radius*80);
  let orthoHalfHeight=Math.max(radius*1.32,1.4);
  if(useOrthographic){
    controls.minZoom=.35; controls.maxZoom=5.5; camera.zoom=1;
  }else{
    controls.minDistance=Math.max(radius*.45,.25); controls.maxDistance=Math.max(radius*12,12);
  }
  camera.updateProjectionMatrix();
  const groundY=fitBox?fitBox.min.y-.05:center.y-radius*.55;
  const grid=new THREE.GridHelper(Math.max(radius*3.2,3),10,0xcbd5e1,0xe2e8f0); grid.position.set(center.x,groundY,center.z); grid.material.transparent=true; grid.material.opacity=.16; scene.add(grid);
  const sourceCp=Array.isArray(sourceView?.camera_position)&&sourceView.camera_position.length===3?sourceView.camera_position:null;
  const sourceCt=Array.isArray(sourceView?.camera_target)&&sourceView.camera_target.length===3?sourceView.camera_target:null;
  const cp=Array.isArray(step.camera_position)&&step.camera_position.length===3?step.camera_position:sourceCp;
  const ct=Array.isArray(step.camera_target)&&step.camera_target.length===3?step.camera_target:sourceCt;
  let targetLook=new THREE.Vector3(...(ct?ct:[center.x,center.y,center.z]));
  // Reject model-supplied camera targets/positions that are wildly outside the reconstructed solid.
  if(targetLook.distanceTo(center)>radius*3.5) targetLook=center.clone();
  let targetPos=cp?new THREE.Vector3(...cp):new THREE.Vector3(center.x+fitDistance*.72,center.y+fitDistance*.48,center.z+fitDistance*.72);
  const suppliedDistance=targetPos.distanceTo(targetLook);
  if(!Number.isFinite(suppliedDistance)||suppliedDistance<radius*.7||suppliedDistance>radius*7){targetPos.set(center.x+fitDistance*.72,center.y+fitDistance*.48,center.z+fitDistance*.72);}
  const pcp=Array.isArray(data?.previous_camera_position)&&data.previous_camera_position.length===3?new THREE.Vector3(...data.previous_camera_position):targetPos.clone();
  const pct=Array.isArray(data?.previous_camera_target)&&data.previous_camera_target.length===3?new THREE.Vector3(...data.previous_camera_target):targetLook.clone();
  camera.position.copy(pcp);controls.target.copy(pct);camera.lookAt(controls.target);controls.update();
  const camStart=performance.now(),camDuration=850;
  let cameraTweenActive=true;
  controls.addEventListener('start',()=>{cameraTweenActive=false;});

  const setInteractionMode=(mode)=>{
    const rotate=mode!=='pan';
    controls.mouseButtons.LEFT=rotate?THREE.MOUSE.ROTATE:THREE.MOUSE.PAN;
    controls.touches.ONE=rotate?THREE.TOUCH.ROTATE:THREE.TOUCH.PAN;
    for(const btn of toolbar?.querySelectorAll('button[data-action="rotate"],button[data-action="pan"]')||[]){
      btn.classList.toggle('is-active',btn.dataset.action===mode);
    }
  };
  const resetZoom=()=>{ if(useOrthographic){camera.zoom=1;camera.updateProjectionMatrix();} };
  const moveCamera=(position,look=center,up=null)=>{
    cameraTweenActive=false;
    if(up&&up.length===3) camera.up.set(Number(up[0]),Number(up[1]),Number(up[2])).normalize();
    camera.position.copy(position); controls.target.copy(look); camera.lookAt(look); resetZoom(); controls.update();
  };
  const sourceViewPosition=()=>{
    const p=Array.isArray(sourceView?.camera_position)&&sourceView.camera_position.length===3
      ? new THREE.Vector3(...sourceView.camera_position)
      : new THREE.Vector3(center.x+fitDistance*.72,center.y+fitDistance*.48,center.z+fitDistance*.72);
    const t=Array.isArray(sourceView?.camera_target)&&sourceView.camera_target.length===3
      ? new THREE.Vector3(...sourceView.camera_target):center.clone();
    return {p,t,up:Array.isArray(sourceView?.camera_up)?sourceView.camera_up:null};
  };
  const standardView=(name)=>{
    const d=fitDistance;
    if(name==='source'){
      if(isOrthographicSet) moveCamera(new THREE.Vector3(center.x+d*.72,center.y+d*.48,center.z+d*.72),center,[0,1,0]);
      else {const sv=sourceViewPosition();moveCamera(sv.p,sv.t,sv.up);}
    }
    else if(name==='front') moveCamera(new THREE.Vector3(center.x,center.y,center.z+d),center,[0,1,0]);
    else if(name==='top') moveCamera(new THREE.Vector3(center.x,center.y+d,center.z+.001),center,[0,0,-1]);
    else if(name==='side') moveCamera(new THREE.Vector3(center.x+d,center.y,center.z),center,[0,1,0]);
    else moveCamera(new THREE.Vector3(center.x+d*.72,center.y+d*.48,center.z+d*.72),center,[0,1,0]);
  };
  const toolbarHandler=(event)=>{
    const btn=event.target.closest('button[data-action]'); if(!btn)return;
    event.preventDefault();event.stopPropagation();const action=btn.dataset.action;
    if(action==='rotate'||action==='pan') setInteractionMode(action);
    else if(action==='home') standardView(sourceView?'source':'iso');
    else standardView(action);
  };
  toolbar?.addEventListener('click',toolbarHandler);
  setInteractionMode('rotate');

  const resize=()=>{
    const w=Math.max(stage.clientWidth,320),h=Math.max(stage.clientHeight,390);renderer.setSize(w,h,false);
    const aspect=w/h;
    if(useOrthographic){camera.left=-orthoHalfHeight*aspect;camera.right=orthoHalfHeight*aspect;camera.top=orthoHalfHeight;camera.bottom=-orthoHalfHeight;}
    else camera.aspect=aspect;
    camera.updateProjectionMatrix();
  };resize();const ro=new ResizeObserver(resize);ro.observe(stage);
  let raf=0;const animateFrame=(now)=>{raf=requestAnimationFrame(animateFrame);
    if(cameraTweenActive){
      const ctween=Math.min(1,(now-camStart)/camDuration);const ce=1-Math.pow(1-ctween,3);camera.position.lerpVectors(pcp,targetPos,ce);controls.target.lerpVectors(pct,targetLook,ce);
      if(ctween>=1) cameraTweenActive=false;
    }
    for(const tw of edgeTweens){const t=Math.max(0,Math.min(1,(now-tw.start)/tw.duration)),e=1-Math.pow(1-t,3),cur=tw.a.clone().lerp(tw.b,e);tw.attr.setXYZ(1,cur.x,cur.y,cur.z);tw.attr.needsUpdate=true;if(tw.line.material.isLineDashedMaterial)tw.line.computeLineDistances();}
    for(const tw of solidTweens){const t=Math.max(0,Math.min(1,(now-tw.start)/tw.duration)),e=1-Math.pow(1-t,3),k=tw.startScale+(1-tw.startScale)*e;tw.mesh.scale.copy(tw.baseScale).multiplyScalar(k); if(tw.mesh.userData.__edge) tw.mesh.userData.__edge.scale.copy(tw.mesh.scale);}
    controls.update();renderer.render(scene,camera);
  };raf=requestAnimationFrame(animateFrame);
  parentElement.__omtThreeCleanup=()=>{cancelAnimationFrame(raf);ro.disconnect();toolbar?.removeEventListener('click',toolbarHandler);controls.dispose();scene.traverse(o=>{o.geometry?.dispose?.();if(o.material){const mats=Array.isArray(o.material)?o.material:[o.material];for(const m of mats){m.map?.dispose?.();m.dispose?.();}}});renderer.dispose();renderer.domElement.remove();};
}

"""

try:
    _visual_2d_component = st.components.v2.component(
        "omt_visual_explanation_2d",
        html=_VISUAL_2D_HTML,
        css=_VISUAL_2D_CSS,
        js=_VISUAL_2D_JS,
        isolate_styles=False,
    )
except Exception:
    _visual_2d_component = None

# Interactive 2D workspace used inside targeted practice. Geometry questions get a
# clean schematic; graph/coordinate questions additionally get a GeoGebra-like
# student workspace with draggable points and segment construction tools. All
# interaction remains in the browser, so plotting does not rerun Streamlit.
_PRACTICE_DIAGRAM_HTML = """
<div class="omt-practice-workspace">
  <div class="omt-gg-toolbar" role="toolbar" aria-label="GeoGebra-style construction tools">
    <button type="button" data-tool="move" class="active">Move</button>
    <button type="button" data-tool="point">Point</button>
    <button type="button" data-tool="line">Line</button>
    <button type="button" data-tool="segment">Segment</button>
    <button type="button" data-tool="ray">Ray</button>
    <button type="button" data-tool="vector">Vector</button>
    <button type="button" data-tool="circle">Circle</button>
    <button type="button" data-tool="polygon">Polygon</button>
    <button type="button" data-tool="finish" class="secondary">Finish</button>
    <button type="button" data-tool="midpoint">Midpoint</button>
    <button type="button" data-tool="perpendicular">Perpendicular</button>
    <button type="button" data-tool="parallel">Parallel</button>
    <button type="button" data-tool="angle">Measure angle</button>
    <button type="button" data-tool="distance">Distance</button>
    <button type="button" data-tool="delete">Delete</button>
    <button type="button" data-tool="undo" class="secondary">Undo</button>
    <button type="button" data-tool="clear" class="secondary">Clear</button>
    <button type="button" data-tool="snap" class="secondary active">Snap 0.5</button>
  </div>
  <div class="omt-gg-status">Use Move to pan/zoom, or select a construction tool.</div>
  <div class="omt-visual2d-board"></div>
</div>
"""

_PRACTICE_DIAGRAM_CSS = """
.omt-practice-workspace { width:100%; }
.omt-gg-toolbar { display:flex; gap:.38rem; overflow-x:auto; padding:.1rem 0 .48rem; scrollbar-width:thin; -webkit-overflow-scrolling:touch; }
.omt-gg-toolbar button { flex:0 0 auto; border:1px solid #cbd5e1; background:#fff; color:#334155; border-radius:.62rem; padding:.5rem .68rem; min-height:38px; font:650 .78rem/1 system-ui,sans-serif; cursor:pointer; }
.omt-gg-toolbar button.active { background:#eaf2ff; border-color:#60a5fa; color:#1d4ed8; }
.omt-gg-toolbar button.secondary { background:#f8fafc; }
.omt-gg-status { font-size:.78rem; color:#64748b; margin:0 0 .42rem; min-height:1.1rem; }
.omt-practice-workspace .omt-visual2d-board { width:100%; height:390px; min-height:320px; border:1px solid rgba(128,128,128,.28); border-radius:.9rem; overflow:hidden; background:#fff; touch-action:none; }
@media (max-width:640px) { .omt-practice-workspace .omt-visual2d-board { height:360px; min-height:310px; } .omt-gg-toolbar button { min-height:44px; padding:.62rem .76rem; } }
"""

_PRACTICE_DIAGRAM_JS = r"""
const JXG_URL = 'https://cdn.jsdelivr.net/npm/jsxgraph@1.12.2/distrib/jsxgraphcore.mjs';
async function loadJXG(){
  if(!globalThis.__omtPracticeJXGPromise) globalThis.__omtPracticeJXGPromise=import(JXG_URL);
  const mod=await globalThis.__omtPracticeJXGPromise; return mod.default||mod.JXG||mod;
}

function installGeoTools(board, toolbar, status, JXG) {
  if (!toolbar || !status) return () => {};
  let tool='move', picks=[], polygonPts=[], snap=true;
  const groups=[];
  const studentObjects=new Set();
  const pointStyle={name:'',fixed:false,size:4,strokeColor:'#dc2626',fillColor:'#fff1f2',strokeWidth:2,highlight:true};
  const lineStyle={fixed:false,strokeColor:'#dc2626',strokeWidth:2.6,highlight:true};
  const addGroup=(objs)=>{ const arr=objs.filter(Boolean); arr.forEach(o=>studentObjects.add(o)); groups.push(arr); board.update(); };
  const removeObjects=(objs)=>{ for(const o of [...objs].reverse()){ try{studentObjects.delete(o);board.removeObject(o);}catch(_){ } } board.update(); };
  const undo=()=>{ if(polygonPts.length){removeObjects([polygonPts.pop()]); return;} const g=groups.pop(); if(g) removeObjects(g); };
  const clear=()=>{ polygonPts=[]; while(groups.length) removeObjects(groups.pop()); };
  const roundSnap=(v)=>snap?Math.round(v*2)/2:v;
  const coords=(ev)=>{ const c=new JXG.Coords(JXG.COORDS_BY_SCREEN,[ev.offsetX,ev.offsetY],board); return [roundSnap(c.usrCoords[1]),roundSnap(c.usrCoords[2])]; };
  const mkPoint=(xy)=>board.create('point',xy,{...pointStyle});
  const length=(a,b)=>Math.hypot(a.X()-b.X(),a.Y()-b.Y());
  const midpointXY=(a,b)=>[(a.X()+b.X())/2,(a.Y()+b.Y())/2];
  const statusMap={
    move:'Drag to pan. Pinch or use the wheel to zoom.', point:'Tap to plot a point. Drag it to adjust.', line:'Tap two positions to draw an infinite line.', segment:'Tap two positions to draw a segment.', ray:'Tap the endpoint, then a second point for the ray direction.', vector:'Tap the start and end points of the vector.', circle:'Tap the centre, then a point on the circle.', polygon:'Tap polygon vertices, then press Finish.', midpoint:'Tap two positions to construct their midpoint.', perpendicular:'Tap two points for a reference line, then tap the point the perpendicular passes through.', parallel:'Tap two points for a reference line, then tap the point the parallel passes through.', angle:'Tap arm point 1, then the vertex, then arm point 2.', distance:'Tap two positions to measure their distance.', delete:'Tap one of your red constructions to delete it.'
  };
  const setTool=(name)=>{
    if(name!=='polygon' && polygonPts.length){ status.textContent='Finish or Undo the current polygon before switching tools.'; return; }
    tool=name; picks=[];
    toolbar.querySelectorAll('button[data-tool]').forEach(b=>{ if(!['snap','finish','undo','clear'].includes(b.dataset.tool)) b.classList.toggle('active',b.dataset.tool===name); });
    status.textContent=statusMap[name]||'Select a construction tool.';
  };
  const finishPolygon=()=>{
    if(polygonPts.length<3){status.textContent='A polygon needs at least 3 vertices.';return;}
    const poly=board.create('polygon',polygonPts,{withLines:true,fillColor:'#fecaca',fillOpacity:.12,borders:{strokeColor:'#dc2626',strokeWidth:2.4},vertices:{visible:true}});
    addGroup([poly,...polygonPts]); polygonPts=[]; status.textContent='Polygon added. Choose another tool or start another polygon.';
  };
  const deleteUnder=(ev)=>{
    let hits=[]; try{hits=board.getAllUnderMouse(ev)||[];}catch(_){ }
    const hit=hits.find(o=>studentObjects.has(o));
    if(hit){ const idx=groups.findIndex(g=>g.includes(hit)); if(idx>=0){const [g]=groups.splice(idx,1);removeObjects(g);} else removeObjects([hit]); status.textContent='Construction deleted.'; }
    else { status.textContent='Tap a red construction to delete it. If selection is difficult, use Undo.'; }
  };
  const addAngleMeasure=(a,v,b)=>{
    const ang=board.create('angle',[a,v,b],{...lineStyle,radius:.7,fillColor:'#fee2e2',fillOpacity:.18,name:'',withLabel:false});
    const txt=board.create('text',[()=>v.X()+0.55,()=>v.Y()+0.55,()=>`${(ang.Value()*180/Math.PI).toFixed(1)}°`],{fixed:true,fontSize:13,color:'#b91c1c'});
    addGroup([a,v,b,ang,txt]);
  };
  const addDistance=(a,b)=>{
    const seg=board.create('segment',[a,b],{...lineStyle,dash:2});
    const txt=board.create('text',[()=> (a.X()+b.X())/2,()=> (a.Y()+b.Y())/2,()=> length(a,b).toFixed(2)],{fixed:true,fontSize:13,color:'#b91c1c'});
    addGroup([a,b,seg,txt]);
  };
  const handlePointTool=(xy)=>{ const p=mkPoint(xy); addGroup([p]); };
  const handleMulti=(xy)=>{
    const p=mkPoint(xy); picks.push(p);
    const need=(['perpendicular','parallel','angle'].includes(tool)?3:2);
    status.textContent=`${picks.length}/${need} point${need===1?'':'s'} selected.`;
    if(picks.length<need) return;
    const [a,b,c]=picks; picks=[];
    if(tool==='line'){const o=board.create('line',[a,b],{...lineStyle,straightFirst:true,straightLast:true});addGroup([a,b,o]);}
    else if(tool==='segment'){const o=board.create('segment',[a,b],lineStyle);addGroup([a,b,o]);}
    else if(tool==='ray'){const o=board.create('line',[a,b],{...lineStyle,straightFirst:false,straightLast:true});addGroup([a,b,o]);}
    else if(tool==='vector'){const o=board.create('arrow',[a,b],lineStyle);addGroup([a,b,o]);}
    else if(tool==='circle'){const o=board.create('circle',[a,b],{...lineStyle,fillOpacity:0});addGroup([a,b,o]);}
    else if(tool==='midpoint'){const m=board.create('midpoint',[a,b],{...pointStyle,fillColor:'#fef3c7',strokeColor:'#d97706'});addGroup([a,b,m]);}
    else if(tool==='distance') addDistance(a,b);
    else if(tool==='angle') addAngleMeasure(a,b,c);
    else if(tool==='perpendicular'){
      const base=board.create('line',[a,b],{...lineStyle,strokeColor:'#94a3b8',strokeWidth:1.6,dash:2});
      const perp=board.create('perpendicular',[base,c],{...lineStyle}); addGroup([a,b,c,base,perp]);
    }
    else if(tool==='parallel'){
      const base=board.create('line',[a,b],{...lineStyle,strokeColor:'#94a3b8',strokeWidth:1.6,dash:2});
      const para=board.create('parallel',[base,c],{...lineStyle}); addGroup([a,b,c,base,para]);
    }
    status.textContent=(statusMap[tool]||'Construction added.')+' Construction added.';
  };
  const clickHandler=(ev)=>{
    const b=ev.target.closest('button[data-tool]'); if(!b)return;
    const name=b.dataset.tool;
    if(name==='clear'){clear();status.textContent='Your constructions were cleared.';return;}
    if(name==='undo'){undo();status.textContent='Last construction removed.';return;}
    if(name==='finish'){finishPolygon();return;}
    if(name==='snap'){snap=!snap;b.classList.toggle('active',snap);b.textContent=snap?'Snap 0.5':'Snap off';status.textContent=snap?'Coordinate snapping is on (0.5 units).':'Coordinate snapping is off.';return;}
    setTool(name);
  };
  toolbar.addEventListener('click',clickHandler);
  const downHandler=(ev)=>{
    if(tool==='move')return;
    if(tool==='delete'){deleteUnder(ev);return;}
    const xy=coords(ev);
    if(tool==='point'){handlePointTool(xy);return;}
    if(tool==='polygon'){const p=mkPoint(xy);polygonPts.push(p);status.textContent=`Polygon: ${polygonPts.length} vertices. Add more or press Finish.`;board.update();return;}
    handleMulti(xy);
    board.update();
  };
  board.on('down',downHandler);
  setTool('move');
  return ()=>{ try{toolbar.removeEventListener('click',clickHandler);}catch(_){ } };
}

export default async function(component){
  const {parentElement,data}=component;
  const stage=parentElement.querySelector('.omt-visual2d-board');
  const toolbar=parentElement.querySelector('.omt-gg-toolbar');
  const status=parentElement.querySelector('.omt-gg-status');
  const scene=data?.scene||{};
  let JXG;
  try{JXG=await loadJXG();}catch(err){console.error(err);stage.textContent='Interactive maths workspace could not load.';return;}
  try{if(parentElement.__omtPracticeBoard)JXG.JSXGraph.freeBoard(parentElement.__omtPracticeBoard);}catch(_){ }
  stage.replaceChildren(); stage.id=`omt-practice-${Math.random().toString(36).slice(2)}`;
  const xMin=Number(scene.x_min??-5),xMax=Number(scene.x_max??5),yMin=Number(scene.y_min??-5),yMax=Number(scene.y_max??5);
  const graphMode=Boolean(scene.show_axes);
  const board=JXG.JSXGraph.initBoard(stage.id,{boundingbox:[xMin,yMax,xMax,yMin],axis:graphMode,grid:Boolean(data?.show_grid),keepaspectratio:scene.keep_aspect!==false,showNavigation:false,showCopyright:false,pan:{enabled:true,needShift:false},zoom:{wheel:true,needShift:false,factorX:1.18,factorY:1.18}});
  parentElement.__omtPracticeBoard=board;
  const pts=new Map();
  for(const p of(scene.points||[])){
    const obj=board.create('point',[Number(p.x),Number(p.y)],{name:p.label||'',fixed:true,highlight:false,size:3.8,strokeColor:'#0f172a',fillColor:'#0f172a',label:{fontSize:14,offset:[7,7]}});pts.set(p.id,obj);
  }
  for(const seg of(scene.segments||[])){
    const a=pts.get(seg.start),b=pts.get(seg.end);if(!a||!b)continue;
    board.create('segment',[a,b],{name:seg.label||'',withLabel:Boolean(seg.label),fixed:true,highlight:false,strokeColor:'#475569',strokeWidth:2.3,dash:seg.dashed?2:0,label:{fontSize:13}});
  }
  for(const poly of(scene.polylines||[])){
    const arr=Array.isArray(poly.points)?poly.points:[];if(arr.length<2)continue;
    const xs=arr.map(v=>Number(v[0])),ys=arr.map(v=>Number(v[1]));
    board.create('curve',[xs,ys],{fixed:true,highlight:false,strokeColor:'#475569',strokeWidth:2.3,dash:poly.dashed?2:0});
  }
  for(const c of(scene.circles||[])){
    const center=pts.get(c.center);if(center)board.create('circle',[center,Number(c.radius)],{fixed:true,highlight:false,strokeColor:'#475569',strokeWidth:2.2,fillOpacity:0});
    else if(Number.isFinite(Number(c.center_x))&&Number.isFinite(Number(c.center_y)))board.create('circle',[[Number(c.center_x),Number(c.center_y)],Number(c.radius)],{fixed:true,highlight:false,strokeColor:'#475569',strokeWidth:2.2,fillOpacity:0});
  }
  for(const a of(scene.angles||[])){
    const p1=pts.get(a.arm1),v=pts.get(a.vertex),p2=pts.get(a.arm2);if(!p1||!v||!p2)continue;
    board.create('angle',[p1,v,p2],{name:a.label||'',withLabel:Boolean(a.label),fixed:true,highlight:false,radius:Number(a.radius||.7),strokeColor:'#2563eb',fillColor:'#dbeafe',fillOpacity:.35,label:{fontSize:13}});
  }
  installGeoTools(board,toolbar,status,JXG);
  status.textContent=graphMode?'Plot points, draw lines, measure angles/distances, or pan/zoom the coordinate plane.':'Use the construction tools to explore the geometry. The given diagram remains fixed.';
  board.update();
}
"""

try:
    _practice_diagram_component = st.components.v2.component(
        "omt_targeted_practice_diagram",
        html=_PRACTICE_DIAGRAM_HTML,
        css=_PRACTICE_DIAGRAM_CSS,
        js=_PRACTICE_DIAGRAM_JS,
        isolate_styles=False,
    )
except Exception:
    _practice_diagram_component = None

try:
    _visual_3d_component = st.components.v2.component(
        "omt_visual_explanation_3d",
        html=_VISUAL_3D_HTML,
        css=_VISUAL_3D_CSS,
        js=_VISUAL_3D_JS,
        isolate_styles=False,
    )
except Exception:
    _visual_3d_component = None


def _visual_plan_is_recommended(analysis: VisualExplanationResult | GeminiAnalysis, question_text: str = "") -> bool:
    """Show simulations only when a diagram/graph/spatial view materially supports the maths."""
    if isinstance(analysis, VisualExplanationResult):
        return analysis.mode in {"geometry2d", "graph2d", "geometry3d"}

    topic = str(getattr(analysis, "likely_syllabus_topic", "") or "").lower()
    interpreted = str(getattr(analysis, "interpreted_question", "") or "").lower()
    raw_question = str(question_text or "").lower()
    haystack = f"{topic} {interpreted} {raw_question}"

    # Strong visual cues: these topics normally benefit from a diagram, graph, coordinate plane,
    # construction, or spatial model. Keep this list intentionally narrower than the old filter.
    visual_keywords = (
        "geometry", "coordinate geometry", "coordinate", "graph", "plot", "sketch",
        "straight-line graph", "straight line graph", "gradient of the line", "intercept",
        "triangle", "quadrilateral", "polygon", "circle", "angle", "bearing",
        "trigonometry", "trigonometric", "angle of elevation", "angle of depression",
        "similar triangles", "congruent", "transformation", "reflection", "rotation",
        "translation", "enlargement", "locus", "construction", "scale drawing",
        "mensuration", "perimeter", "area of",
        "cuboid", "prism", "pyramid", "cone", "cylinder", "sphere", "3d",
        "three-dimensional", "isometric", "orthographic", "top view", "front view", "side view",
        "diagram", "figure",
    )

    # Explicitly non-visual topics should not get a simulation merely because a generic word such
    # as "line" or "gradient" appears in explanatory prose.
    nonvisual_topic_keywords = (
        "standard form", "indices", "surds", "algebraic manipulation", "factorisation",
        "factorization", "equations and inequalities", "linear equation", "quadratic equation",
        "simultaneous equation", "number", "ratio", "percentage", "proportion", "sets",
        "probability", "statistics", "mean", "median", "mode", "arithmetic", "sequence",
    )

    has_visual_cue = any(token in haystack for token in visual_keywords)
    if not has_visual_cue:
        return False

    # If the syllabus topic is clearly non-visual, require an explicit visual cue in the actual
    # question itself (e.g. "plot the graph" or "in the diagram") before allowing a simulation.
    topic_is_nonvisual = any(token in topic for token in nonvisual_topic_keywords)
    question_has_explicit_visual_cue = any(token in raw_question or token in interpreted for token in (
        "diagram", "graph", "plot", "sketch", "coordinate", "triangle", "circle", "bearing",
        "elevation", "depression", "cuboid", "prism", "pyramid", "cone", "cylinder", "sphere",
        "isometric", "orthographic", "top view", "front view", "side view",
    ))
    if topic_is_nonvisual and not question_has_explicit_visual_cue:
        return False

    return True


def _render_source_3d_reference(plan: VisualExplanationResult, question_files: list[Any] | None) -> None:
    """Show the exact source isometric/orthographic drawing used to calibrate the 3D model."""
    if plan.mode != "geometry3d" or plan.scene_3d is None or not question_files:
        return
    source_view = getattr(plan.scene_3d, "source_view", None)
    if source_view is None:
        return
    source_index = int(getattr(source_view, "source_index", 1) or 1)
    page_number = int(getattr(source_view, "page_number", 1) or 1)
    if not (1 <= source_index <= len(question_files)):
        return
    image = _question_source_image(question_files[source_index - 1], page_number)
    if image is None:
        return
    box = list(getattr(source_view, "diagram_box_2d", []) or [])
    if len(box) == 4:
        px = _normalized_box_to_pixels(box, image.width, image.height)
        if px is not None:
            x1, y1, x2, y2 = px
            pad_x = max(8, int((x2 - x1) * 0.04))
            pad_y = max(8, int((y2 - y1) * 0.04))
            image = image.crop((max(0, x1-pad_x), max(0, y1-pad_y), min(image.width, x2+pad_x), min(image.height, y2+pad_y)))
    projection_raw = str(getattr(source_view, "projection", "unknown"))
    is_orthographic_set = projection_raw == "orthographic_set"
    title = "Compare reconstruction with the question's top/front/side views" if is_orthographic_set else "Compare with the question's original 3D/isometric view"
    caption = "Orthographic source views used to reconstruct the 3D object" if is_orthographic_set else "Source diagram used to calibrate the 3D model"
    with st.expander(title, expanded=True):
        st.image(image, caption=caption, use_container_width=True)
        projection = projection_raw.replace("_", " ").title()
        confidence = str(getattr(source_view, "match_confidence", "medium")).title()
        if is_orthographic_set:
            st.caption(f"Source type: {projection} · Projection-consistency confidence: {confidence}")
            st.info("This question does not provide a single isometric drawing. The 3D model is reconstructed by combining the top, front and side projections. Use the Top, Front and Side buttons in the 3D viewer to verify the reconstruction.")
        else:
            st.caption(f"Source-view projection: {projection} · Match confidence: {confidence}")
        checks = list(getattr(source_view, "view_consistency_checks", []) or [])
        if checks:
            st.markdown("**Projection checks**")
            for check in checks:
                st.markdown(f"- {check}")
        if is_orthographic_set:
            components = list(getattr(plan.scene_3d, "orthographic_components", []) or [])
            if components:
                st.markdown("**How the 3D form was inferred from the three views**")
                for item in sorted(components, key=lambda x: int(getattr(x, "vertical_order", 0))):
                    kind = str(getattr(item, "inferred_kind", "component")).replace("_", " ").title()
                    relation = str(getattr(item, "stacking_relation", "")).strip()
                    st.markdown(f"**{kind}**" + (f" — {relation}" if relation else ""))
                    cols = st.columns(3)
                    cols[0].caption("Top: " + str(getattr(item, "top_view_evidence", "")))
                    cols[1].caption("Front: " + str(getattr(item, "front_view_evidence", "")))
                    cols[2].caption("Side: " + str(getattr(item, "side_view_evidence", "")))
        note = str(getattr(source_view, "match_note", "")).strip()
        if note:
            st.caption(note)


def render_visual_explanation(plan: VisualExplanationResult, question_files: list[Any] | None = None) -> None:
    if plan.mode == "none":
        if plan.reconstruction_note:
            st.info("A reliable interactive reconstruction was not generated: " + plan.reconstruction_note)
        return
    if not plan.steps:
        return

    st.markdown("### Visual step-by-step simulation")
    st.caption(
        f"Reconstruction confidence: {plan.reconstruction_confidence.title()}. {plan.reconstruction_note}"
    )
    show_visual_grid = False
    if plan.mode in {"geometry2d", "graph2d"}:
        show_visual_grid = st.toggle(
            "Show gridlines",
            value=(plan.mode == "graph2d"),
            key="ai_visual_show_grid",
            help="Turn gridlines on or off without changing the mathematical construction.",
        )
    if plan.mode == "geometry3d" and getattr(plan, "reconstructed_parts", None):
        st.markdown("**3D form identified from the question:** " + " · ".join(plan.reconstructed_parts))
    if plan.mode == "geometry3d":
        _render_source_3d_reference(plan, question_files)
    max_index = len(plan.steps) - 1
    idx = max(0, min(int(st.session_state.get("ai_visual_step", 0)), max_index))
    st.session_state.ai_visual_step = idx

    def _go_previous() -> None:
        current = int(st.session_state.get("ai_visual_step", 0))
        st.session_state.ai_visual_step = max(0, current - 1)
        st.session_state.ai_visual_replay_nonce = int(st.session_state.get("ai_visual_replay_nonce", 0)) + 1

    def _go_next() -> None:
        current = int(st.session_state.get("ai_visual_step", 0))
        st.session_state.ai_visual_step = min(max_index, current + 1)
        st.session_state.ai_visual_replay_nonce = int(st.session_state.get("ai_visual_replay_nonce", 0)) + 1

    def _replay_current() -> None:
        # The nonce is sent through component data. Streamlit Components v2 calls
        # the frontend renderer again whenever data changes, so this reliably
        # restarts the current construction without changing the selected step.
        st.session_state.ai_visual_replay_nonce = int(st.session_state.get("ai_visual_replay_nonce", 0)) + 1

    b1, mid, b2, replay = st.columns([1, 1.7, 1, 1])
    b1.button(
        "← Previous",
        disabled=idx <= 0,
        use_container_width=True,
        key="ai_visual_prev",
        on_click=_go_previous,
    )
    mid.markdown(f"<div style='text-align:center;padding:.55rem'><strong>Step {idx + 1} of {len(plan.steps)}</strong></div>", unsafe_allow_html=True)
    b2.button(
        "Next →",
        disabled=idx >= max_index,
        use_container_width=True,
        key="ai_visual_next",
        on_click=_go_next,
    )
    replay.button(
        "↻ Replay",
        use_container_width=True,
        key="ai_visual_replay",
        on_click=_replay_current,
    )

    step = plan.steps[idx]
    # New plans progressively reveal the construction. Old saved plans without reveal/animate
    # fields continue to show the complete scene for backward compatibility.
    reveal_mode = any(bool(getattr(item, "reveal_ids", [])) or bool(getattr(item, "animate_ids", [])) for item in plan.steps)
    visible_ids: set[str] = set()
    if reveal_mode:
        for earlier in plan.steps[: idx + 1]:
            visible_ids.update(getattr(earlier, "reveal_ids", []) or [])
            visible_ids.update(getattr(earlier, "animate_ids", []) or [])
        visible_ids.update(getattr(step, "highlight_ids", []) or [])
    animate_ids = list(getattr(step, "animate_ids", []) or [])
    replay_nonce = int(st.session_state.get("ai_visual_replay_nonce", 0))

    previous_step = plan.steps[idx - 1] if idx > 0 else None
    previous_camera_position = list(getattr(previous_step, "camera_position", []) or []) if previous_step else []
    previous_camera_target = list(getattr(previous_step, "camera_target", []) or []) if previous_step else []

    scene_payload: dict[str, Any] | None = None
    component_data = {
        "step": step.model_dump(),
        "visible_ids": sorted(visible_ids),
        "animate_ids": animate_ids,
        "reveal_mode": reveal_mode,
        "animation_nonce": replay_nonce,
        "previous_camera_position": previous_camera_position,
        "previous_camera_target": previous_camera_target,
    }
    if plan.mode in {"geometry2d", "graph2d"} and plan.scene_2d is not None:
        scene_payload = plan.scene_2d.model_dump()
        if _visual_2d_component is not None:
            _visual_2d_component(
                data={"scene": scene_payload, "show_grid": show_visual_grid, **component_data},
                default={},
                key="ai_visual2d",
                width="stretch",
                height="content",
            )
        else:
            st.info("The interactive 2D renderer is unavailable in this browser session.")
    elif plan.mode == "geometry3d" and plan.scene_3d is not None:
        scene_payload = plan.scene_3d.model_dump()
        if _visual_3d_component is not None:
            _visual_3d_component(
                data={"scene": scene_payload, **component_data},
                default={},
                key="ai_visual3d",
                width="stretch",
                height="content",
            )
        else:
            st.info("The interactive 3D renderer is unavailable in this browser session.")

    st.markdown(f"#### {step.title}")
    st.caption(f"Matches corrected solution step {getattr(step, 'source_step_index', idx + 1)}")
    if getattr(step, "simulation_note", ""):
        st.info("Simulation: " + step.simulation_note)
    if not (getattr(step, "animate_ids", []) or getattr(step, "highlight_ids", []) or getattr(step, "reveal_ids", [])):
        st.caption("This corrected step is mainly algebraic, so no new diagram object is introduced. Replay redraws the current construction for orientation.")
    st.markdown("**Matching corrected step**")
    for formula in step.math:
        render_mathio(formula)
    st.markdown("**Why this simulation matches the step**")
    render_mathio_mixed(step.explanation)

    if plan.mode == "geometry3d":
        st.caption("iPad: drag with one finger to rotate the solid and pinch with two fingers to zoom. Use Replay to watch the current construction/camera movement again.")


def _clean_practice_display_text(text: str) -> str:
    """Keep generated practice wording compact and prevent model Markdown from taking over the UI."""
    value = str(text or "").strip()
    value = value.replace("**", "")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def _render_practice_key_information(items: list[str]) -> None:
    for item in (items or [])[:6]:
        if not str(item).strip():
            continue
        with st.container(border=True):
            render_mathio_mixed(_clean_practice_display_text(str(item)))


def _compact_task_prompt(focus_prompt: str, full_question: str) -> str:
    """Prefer a one-line action prompt even when the model repeats the whole story."""
    candidate = _clean_practice_display_text(focus_prompt)
    if candidate and len(candidate) <= 180 and candidate.count("\n") <= 1:
        return candidate
    source = _clean_practice_display_text(full_question)
    # Choose the last sentence/clause that contains an exam-style command.
    pieces = re.split(r"(?<=[.!?])\s+|\n+", source)
    commands = re.compile(r"\b(calculate|find|determine|solve|show|state|express|sketch|work out|give|write down)\b", re.I)
    matches = [piece.strip() for piece in pieces if commands.search(piece)]
    if matches:
        return matches[-1][:260].strip()
    return (candidate or source)[:260].strip()


def render_targeted_practice_focus(pq: TargetedPracticeQuestion, *, key: str) -> None:
    """Present the task as a compact student card, with visual information before story text."""
    full_question = _clean_practice_display_text(getattr(pq, "question", "") or "")
    focus_prompt = _compact_task_prompt(getattr(pq, "focus_prompt", "") or "", full_question)
    key_information = list(getattr(pq, "key_information", []) or [])
    diagram = getattr(pq, "diagram_2d", None)
    diagram_note = str(getattr(pq, "diagram_note", "") or "").strip()

    with st.container(border=True):
        st.markdown('<div class="omt-focus-title">Your task</div>', unsafe_allow_html=True)
        render_mathio_mixed(focus_prompt or full_question)

        if diagram is not None:
            visual_col, info_col = st.columns([1.25, .85], gap="large", vertical_alignment="top")
            with visual_col:
                st.caption("Interactive graph workspace" if bool(getattr(diagram, "show_axes", False)) else "Question diagram")
                practice_grid = st.toggle(
                    "Show gridlines",
                    value=bool(getattr(diagram, "show_axes", False)),
                    key=f"practice_grid_{key}",
                )
                if _practice_diagram_component is not None:
                    _practice_diagram_component(
                        data={
                            "scene": diagram.model_dump(),
                            "show_grid": practice_grid,
                            "step": {"highlight_ids": [], "dim_ids": [], "animate_ids": []},
                            "visible_ids": [],
                            "animate_ids": [],
                            "reveal_mode": False,
                            "animation_nonce": 0,
                        },
                        default={},
                        key=f"practice_diagram_{key}",
                        width="stretch",
                        height="content",
                    )
                else:
                    st.info("The practice diagram could not load in this browser session.")
                st.caption(diagram_note or ("Plot points or draw segments to explore the graph. Your red constructions are for working only." if bool(getattr(diagram, "show_axes", False)) else "Schematic only · not drawn to scale"))
            with info_col:
                if key_information:
                    st.caption("Given")
                    _render_practice_key_information(key_information)
        elif key_information:
            st.caption("Given")
            info_cols = st.columns(2) if len(key_information) > 1 else [st.container()]
            if len(key_information) > 1:
                for idx, item in enumerate(key_information[:6]):
                    with info_cols[idx % 2]:
                        with st.container(border=True):
                            render_mathio_mixed(_clean_practice_display_text(str(item)))
            else:
                _render_practice_key_information(key_information)

        with st.expander("Full wording", expanded=False):
            render_mathio_mixed(full_question)


def init_state() -> None:
    defaults: dict[str, Any] = {
        "session_id": secrets.token_hex(8),
        "question": None,
        "attempt_result": None,
        "history": [],
        "hint_level": 0,
        "reveal_solution": False,
        "seed_counter": 1,
        "ai_analysis": None,
        "ai_error": "",
        "ai_visual_explanation": None,
        "ai_visual_error": "",
        "ai_visual_step": 0,
        "ai_visual_replay_nonce": 0,
        "ai_fallback_result": None,
        "ai_question_detection": None,
        "ai_question_detection_error": "",
        "ai_question_file_signature": "",
        "ai_selected_question_index": 0,
        "ai_question_feasibility": None,
        "ai_question_feasibility_error": "",
        "ai_question_feasibility_signature": "",
        "ai_practice_stage": 0,
        "ai_practice_current_question": None,
        "ai_practice_evaluation": None,
        "ai_practice_last_working": "",
        "ai_practice_misses": {"Near transfer": 0, "Varied context": 0, "Stretch": 0},
        "ai_practice_consecutive_correct": {"Near transfer": 0, "Varied context": 0, "Stretch": 0},
        "ai_practice_completed": {"Near transfer": False, "Varied context": False, "Stretch": False},
        "ai_practice_ready_to_advance": False,
        "ai_practice_finished": False,
        "ai_practice_question_version": 0,
        "ai_cached_verification": None,
        "ai_cached_verification_signature": "",
        "batch_results": [],
        "batch_error": "",
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


init_state()

PRACTICE_STAGES = ["Near transfer", "Varied context", "Stretch"]


def clear_ai_practice_state() -> None:
    st.session_state.ai_practice_stage = 0
    st.session_state.ai_practice_current_question = None
    st.session_state.ai_practice_evaluation = None
    st.session_state.ai_practice_last_working = ""
    st.session_state.ai_practice_misses = {kind: 0 for kind in PRACTICE_STAGES}
    st.session_state.ai_practice_consecutive_correct = {kind: 0 for kind in PRACTICE_STAGES}
    st.session_state.ai_practice_completed = {kind: False for kind in PRACTICE_STAGES}
    st.session_state.ai_practice_ready_to_advance = False
    st.session_state.ai_practice_finished = False
    st.session_state.ai_practice_question_version = 0


def initialize_ai_practice(analysis: GeminiAnalysis) -> None:
    clear_ai_practice_state()
    by_kind = {q.kind: q for q in analysis.practice_questions}
    st.session_state.ai_practice_current_question = by_kind["Near transfer"]


def practice_attempt_is_secure(result: PracticeEvaluation) -> bool:
    # Older Streamlit sessions may still contain evaluation objects from the pre-multipart schema.
    # Treat those as non-secure rather than allowing an accidental category advance.
    return (
        bool(getattr(result, "is_correct", False))
        and bool(getattr(result, "all_required_parts_complete", False))
        and not list(getattr(result, "missing_or_incorrect_parts", []) or [])
        and not list(getattr(result, "presentation_errors", []) or [])
        and int(getattr(result, "answer_score", 0) or 0) >= 80
        and int(getattr(result, "reasoning_score", 0) or 0) >= 80
        and getattr(result, "mastery", "") in {"Secure", "Strong"}
    )


def initial_practice_question(analysis: GeminiAnalysis, kind: str) -> TargetedPracticeQuestion:
    for question in analysis.practice_questions:
        if question.kind == kind:
            return question
    raise ValueError(f"No initial practice question found for {kind}")

# Streamlit Community Cloud stores app secrets in st.secrets.
# Copy the Gemini key into the process environment so the service layer can read it.
try:
    if "GEMINI_API_KEY" in st.secrets and not os.getenv("GEMINI_API_KEY"):
        os.environ["GEMINI_API_KEY"] = str(st.secrets["GEMINI_API_KEY"])
except Exception:
    pass


def track_code(label: str) -> str:
    return selected_track_info(label)["engine_code"]


def reset_current_question():
    """Reset per-question UI state without mutating already-instantiated widget keys.

    Streamlit raises StreamlitAPIException when code assigns to a session-state key
    belonging to a widget after that widget has been created in the current run.
    For widget-bound keys, delete the key instead and let the widget recreate its
    default value on the next rerun.
    """
    # Non-widget state can be assigned normally.
    st.session_state.hint_level = 0

    # Widget-bound/per-question keys must be removed rather than reassigned.
    reset_keys = [
        "reveal_solution",
        "student_answer",
        "student_working",
        "working_format",
        "offline_answer",
        "offline_working",
        "offline_working_format",
        "show_hint",
    ]

    # Also clear generated editor/table/calculator state tied to the previous question.
    prefixes = (
        "offline_",
        "student_working_",
        "practice_",
        "near_transfer_",
    )

    for key in list(st.session_state.keys()):
        if key in reset_keys or any(key.startswith(prefix) for prefix in prefixes):
            try:
                del st.session_state[key]
            except Exception:
                pass



def make_new_question(track: str, topic: str, difficulty: str) -> None:
    seed = int(datetime.now().timestamp() * 1000) + st.session_state.seed_counter
    st.session_state.seed_counter += 1
    previous = st.session_state.get("question")
    candidate = generate_question(track, topic, difficulty, seed=seed)
    # Avoid showing the same template family repeatedly when alternatives exist.
    for offset in range(1, 7):
        if not (
            previous is not None
            and getattr(previous, "track", None) == track
            and getattr(previous, "topic_code", None) == topic
            and getattr(previous, "family", "")
            and getattr(candidate, "family", "") == getattr(previous, "family", "")
        ):
            break
        candidate = generate_question(track, topic, difficulty, seed=seed + 7919 * offset)
    st.session_state.question = candidate
    reset_current_question()


def record_history(question: Question, result: AttemptResult) -> None:
    st.session_state.history.append(
        {
            "time_utc": datetime.now(timezone.utc).isoformat(timespec="seconds"),
            "mode": "offline generated",
            "track": question.track,
            "topic_code": official_topic_code(question.track, question.topic_code),
            "topic": question.topic_name,
            "difficulty": question.difficulty,
            "question": question.prompt,
            "correct": result.is_correct,
            "answer_score": result.answer_score,
            "reasoning_score": result.reasoning_score,
            "mastery": result.mastery,
        }
    )


def record_ai_practice_history(track: str, q: TargetedPracticeQuestion, result: PracticeEvaluation) -> None:
    st.session_state.history.append(
        {
            "time_utc": datetime.now(timezone.utc).isoformat(timespec="seconds"),
            "mode": "Gemini targeted practice",
            "track": track,
            "topic_code": "AI",
            "topic": q.target_skill,
            "difficulty": q.kind,
            "question": q.question,
            "correct": result.is_correct,
            "answer_score": result.answer_score,
            "reasoning_score": result.reasoning_score,
            "mastery": result.mastery,
        }
    )


def render_attempt(result: AttemptResult) -> None:
    st.markdown("### Feedback")
    c1, c2, c3 = st.columns(3)
    c1.metric("Answer", f"{result.answer_score}%")
    c2.metric("Reasoning", f"{result.reasoning_score}%")
    c3.metric("Current mastery", result.mastery)
    if result.is_correct and result.first_logic_break is None:
        st.success(result.summary)
    else:
        st.info(result.summary)
    if result.first_logic_break is not None:
        st.warning(
            f"First detected logic break: line {result.first_logic_break}. "
            f"{result.first_logic_break_explanation}"
        )
    if result.step_feedback:
        st.markdown("#### Step-by-step check")
        icon = {"correct": "✅", "incorrect": "❌", "unparsed": "🔎", "checked": "•"}
        for item in result.step_feedback:
            with st.expander(f"{icon.get(item.status, '•')} Line {item.line_number}"):
                st.markdown("**Student step**")
                render_mathio(item.line)
                render_mathio_mixed(item.feedback)
    if result.strengths:
        st.markdown("**What is working**")
        for item in result.strengths:
            st.write(f"• {item}")
    if result.gaps:
        st.markdown("**What to improve**")
        for item in result.gaps:
            st.write(f"• {item}")
    st.markdown("**Next hint:**")
    render_mathio_mixed(result.next_hint)



def call_analyze_submission_compat(**kwargs):
    """Avoid transient deploy mismatches between app.py and gemini_service.py.

    Streamlit Cloud can briefly serve a new app.py while a dependency module is still
    being reloaded. Only pass keyword arguments supported by the currently imported
    analyze_submission() signature.
    """
    try:
        supported = inspect.signature(analyze_submission).parameters
        safe_kwargs = {key: value for key, value in kwargs.items() if key in supported}
        return analyze_submission(**safe_kwargs)
    except (TypeError, ValueError):
        # Fallback for unusual wrapped callables: retry once without the newest optional argument.
        kwargs.pop("verification", None)
        return analyze_submission(**kwargs)




st.markdown("""<style>
/* Guided-solving readability */
.omt-guidance-item {
    font-size: 1rem;
    line-height: 1.65;
    margin: 0.35rem 0 0.55rem 0;
}
.omt-guidance-item p {
    margin: 0;
}

.compact-guided-title {
    margin-bottom: 0.45rem !important;
}
div[data-testid="stVerticalBlock"] > div:has(.omt-guidance-item) {
    gap: 0.15rem !important;
}
.omt-guidance-item {
    margin: 0.08rem 0 0.18rem 0 !important;
    line-height: 1.45 !important;
}
.omt-guidance-item p {
    margin: 0 !important;
}
</style>""", unsafe_allow_html=True)

def clean_guidance_text(value: str) -> str:
    """Sanitize model-generated guidance without turning prose into mathematics."""
    text = str(value or "").strip()
    # Normalize scalable delimiter commands before any other cleanup.
    # This prevents "\\right)" from ever degrading into visible "ight)".
    text = re.sub(r"\\left\s*([\(\[\{\|])", r"\1", text)
    text = re.sub(r"\\right\s*([\)\]\}\|])", r"\1", text)
    text = re.sub(r"\\left\b", "", text)
    text = re.sub(r"\\right\b", "", text)
    text = re.sub(r"(?<!\\)\bpi\b", r"\\pi", text, flags=re.IGNORECASE)
    text = re.sub(r"(?<!\\)\btheta\b", r"\\theta", text, flags=re.IGNORECASE)


    # Remove one or more literal backslashes before bullet commands. This handles
    # \textbullet, \\textbullet and similar escaped variants returned by models.
    text = re.sub(r"\\+(?:textbullet|bullet)\b\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^\s*[•●▪◦*-]+\s*", "", text)

    # Unwrap LaTeX text-formatting commands when they were incorrectly used for prose.
    # Repeat a few times to handle nested/simple escaped values.
    for _ in range(3):
        new_text = re.sub(r"\\+(?:text|mathrm|mathbf|operatorname)\{([^{}]*)\}", r"\1", text)
        if new_text == text:
            break
        text = new_text

    # Fix common prose-without-spaces artifacts caused by MathIO/LaTeX generation.
    text = text.replace("90degrees", "90°")
    text = re.sub(r"\s{2,}", " ", text).strip()

    # Suppress empty/punctuation-only guidance items.
    if not re.search(r"[A-Za-z0-9]", text):
        return ""
    return text


def _plainify_embedded_math(text: str) -> str:
    """Make small math fragments readable when embedded inside prose.

    Standalone equations still use MathIO. This function is only for a sentence that
    is mostly English prose but contains a few raw LaTeX commands.
    """
    value = text

    replacements = {
        r"\pi": "π",
        r"\theta": "θ",
        r"\alpha": "α",
        r"\beta": "β",
        r"\gamma": "γ",
        r"\times": "×",
        r"\cdot": "·",
        r"\div": "÷",
        r"\pm": "±",
        r"\leq": "≤",
        r"\geq": "≥",
        r"\neq": "≠",
        r"\circ": "°",
    }
    for src, dst in replacements.items():
        value = value.replace(src, dst)

    # Simple fractions frequently appearing inside explanatory sentences.
    # Leave complex/nested fractions for standalone MathIO rendering.
    value = re.sub(
        r"\\frac\{([^{}]+)\}\{([^{}]+)\}",
        lambda m: f"({m.group(1)})/({m.group(2)})",
        value,
    )
    value = re.sub(r"\^\{([^{}]+)\}", r"^(\1)", value)
    value = re.sub(r"_\{([^{}]+)\}", r"_(\1)", value)
    value = value.replace(r"\,", " ").replace(r"\;", " ")
    return re.sub(r"\s{2,}", " ", value).strip()


def _looks_like_standalone_math(text: str) -> bool:
    """True when a value is primarily an equation/expression rather than prose."""
    value = clean_guidance_text(text)
    if not value:
        return False

    words = re.findall(r"[A-Za-z]{3,}", value)
    math_signals = len(re.findall(
        r"[=+\-×÷*/^]|\\(?:frac|sqrt|angle|pi|theta|sin|cos|tan|arcsin|arccos|arctan|log|ln)\b",
        value,
    ))
    # A short expression with strong mathematical syntax is safe for MathIO.
    return math_signals >= 1 and len(words) <= 4



_INLINE_MATH_FRAGMENT_RE = re.compile(
    r"""
    (?:
        [A-Za-z][A-Za-z0-9_]*\s*=\s*[^,.;:]+
        |
        \([^()\n]{1,80}\)\^\{?\d+\}?
        |
        \([^()\n]{1,80}\)\^\d+
        |
        [A-Za-z0-9]+\^\{?\d+\}?
        |
        \\angle\s*[A-Z]{2,4}\s*=\s*[^,.;:]*
        |
        \\(?:frac|sqrt|pi|theta|sin|cos|tan|arcsin|arccos|arctan|log|ln)\b[^,.;:]*
    )
    """,
    re.VERBOSE,
)



_MATH_TO_PROSE_CONNECTOR_RE = re.compile(
    r"(?i)\s+(?=(?:for|where|when|with|over|from|on|using|given\s+that|such\s+that|"
    r"and\s+then|so\s+that|in\s+which|whose|representing|corresponding\s+to)\b)"
)


def _split_equation_from_prose_tail(fragment: str) -> tuple[str, str]:
    """Separate a real mathematical expression from trailing English prose."""
    value = str(fragment or "").strip()
    if not value:
        return "", ""
    if not re.search(r"[=<>≤≥]|\\(?:sin|cos|tan|frac|sqrt|log|ln)\b", value):
        return value, ""
    m = _MATH_TO_PROSE_CONNECTOR_RE.search(value)
    if not m:
        return value, ""
    return value[:m.start()].rstrip(), value[m.start():].strip()


def render_guidance_mixed_mathio(value: str) -> None:
    """Render prose normally and actual inline mathematics with MathIO.

    This is used for Goal / hints / prose-heavy guidance. It avoids sending the
    whole English sentence to MathIO while still rendering expressions such as
    (1+y)^7, x^7 and (1+x+x^2)^7 properly.
    """
    text = _normalize_integral_source(clean_guidance_text(value))
    if not text:
        return

    # Remove ellipsis LaTeX commands; they are presentation shorthand and should not
    # appear in the student's worked solution.
    text = re.sub(r"\\+(?:dots|ldots|cdots)\b", "", text)
    text = re.sub(r"\.{3,}", "", text)
    text = re.sub(r"\s{2,}", " ", text).strip()

    if _looks_like_standalone_math(text):
        render_mathio(text)
        return

    parts = []
    cursor = 0
    for match in _INLINE_MATH_FRAGMENT_RE.finditer(text):
        if match.start() > cursor:
            parts.append(("text", text[cursor:match.start()]))
        math_fragment, prose_tail = _split_equation_from_prose_tail(match.group(0))
        if math_fragment:
            parts.append(("math", math_fragment))
        if prose_tail:
            parts.append(("text", prose_tail))
        cursor = match.end()
    if cursor < len(text):
        parts.append(("text", text[cursor:]))

    if not any(kind == "math" for kind, _ in parts):
        st.markdown(_plainify_embedded_math(text))
        return

    # Render in a compact flow. MathIO equations get their own line for reliability.
    for kind, chunk in parts:
        chunk = chunk.strip()
        if not chunk:
            continue
        if kind == "math":
            render_mathio(chunk)
        else:
            st.markdown(chunk)


def render_guidance_content(value: str) -> None:
    """Render guidance as readable prose, using MathIO only for real equations."""
    text = clean_guidance_text(value)
    text = re.sub(r"\{(cm|mm|km|kg|m|g|s|h|ml|l)\}", r"\1", text, flags=re.IGNORECASE)
    text = re.sub(r"(?<!\\)\bpi\b", r"\\pi", text, flags=re.IGNORECASE)
    text = re.sub(r"(?<!\\)\btheta\b", r"\\theta", text, flags=re.IGNORECASE)
    if not text:
        return

    if _looks_like_standalone_math(text):
        render_mathio(text)
        return

    # Do not feed a whole English sentence into MathIO merely because it contains
    # \pi, \frac, etc. Convert small embedded fragments to readable inline symbols.
    readable = _plainify_embedded_math(text)
    st.markdown(readable)


def guidance_item(value: str) -> None:
    """Render one compact bullet only when meaningful text exists."""
    text = clean_guidance_text(value)
    if not text:
        return
    readable = _plainify_embedded_math(text)
    st.markdown(f"- {readable}")


def render_guidance_step(step_number: int, value) -> None:
    """Keep prose in text mode and equations in MathIO mode."""
    if hasattr(value, "explanation"):
        explanation = clean_guidance_text(getattr(value, "explanation", ""))
        equations = [
            str(eq).strip()
            for eq in (getattr(value, "equations", []) or [])
            if str(eq).strip()
        ]
    else:
        explanation = clean_guidance_text(str(value or ""))
        equations = []

    if not explanation and not equations:
        return

    with st.container(border=True):
        st.markdown(f"#### Step {step_number}")

        # Explanation may occasionally contain a mathematical statement because
        # model output is not perfectly channel-separated. The mixed renderer keeps
        # ordinary English as text but automatically routes equations to MathIO.
        if explanation:
            render_guidance_mixed_mathio(explanation)

        # equations are always MathIO
        for equation in equations:
            equation = re.sub(r"\\+(?:dots|ldots|cdots)\b", "", equation)
            equation = re.sub(r"\.{3,}", "", equation)
            equation = re.sub(r"\s{2,}", " ", equation).strip()
            if equation:
                render_mathio(equation)



def _switch_guided_to_full() -> None:
    """Widget callback: safe place to change the radio-controlled session key."""
    st.session_state.guided_support_mode = "Full solution"


def _switch_guided_to_hints() -> None:
    """Widget callback: safe place to change the radio-controlled session key."""
    st.session_state.guided_support_mode = "Hints only"


def _show_all_guided_steps(total_steps: int) -> None:
    st.session_state.guided_reveal_step = total_steps


def _show_next_guided_step(total_steps: int) -> None:
    current = int(st.session_state.get("guided_reveal_step", 0))
    st.session_state.guided_reveal_step = min(current + 1, total_steps)


def _show_next_guided_hint(total_hints: int) -> None:
    current = int(st.session_state.get("guided_hint_count", 0))
    st.session_state.guided_hint_count = min(current + 1, total_hints)


def render_guided_solution(g: GuidedSolution) -> None:
    st.markdown('<div class="omt-section-kicker">Guided solving</div>', unsafe_allow_html=True)
    st.markdown('<div class="omt-section-title compact-guided-title">Work through the question</div>', unsafe_allow_html=True)

    with st.container(border=True):
        st.markdown("### 🎯 Goal")
        render_guidance_mixed_mathio(g.interpreted_goal)

        known_items = [clean_guidance_text(item) for item in g.known_information]
        known_items = [item for item in known_items if item]
        if known_items:
            st.markdown("#### What is given")
            for item in known_items:
                guidance_item(item)

        concept_items = [clean_guidance_text(item) for item in g.concepts_to_use]
        concept_items = [item for item in concept_items if item]
        if concept_items:
            with st.expander("Useful concepts", expanded=False):
                for item in concept_items:
                    guidance_item(item)

    support_mode = st.radio(
        "Choose how much help you want",
        ["Hints only", "Full solution"],
        horizontal=True,
        key="guided_support_mode",
        help=(
            "Hints only reveals progressively stronger hints without showing worked steps. "
            "Full solution lets you reveal the verified working step by step or all at once."
        ),
    )

    if support_mode == "Hints only":
        with st.container(border=True):
            st.markdown("### 💡 Hints")
            render_guidance_content(g.first_question_for_student)

            hint_count = int(st.session_state.get("guided_hint_count", 0))
            for i, hint in enumerate(g.hint_ladder[:hint_count], 1):
                with st.container(border=True):
                    st.markdown(f"**Hint {i}**")
                    render_guidance_content(hint)

            if hint_count < len(g.hint_ladder):
                st.button(
                    "Show next hint",
                    key="guided_show_hint",
                    use_container_width=True,
                    type="primary" if hint_count == 0 else "secondary",
                    on_click=_show_next_guided_hint,
                    args=(len(g.hint_ladder),),
                )
            else:
                st.success("All hints are shown. Try the question before revealing the full solution.")

            st.button(
                "Switch to full solution",
                key="guided_switch_full",
                use_container_width=True,
                on_click=_switch_guided_to_full,
            )

    else:
        total_steps = len(g.guided_steps)
        reveal = int(st.session_state.get("guided_reveal_step", 0))

        if total_steps == 0:
            st.warning("No worked steps were returned for this guided solution. Regenerate the guided solution.")
            return

        # Make Full solution immediately useful: if the student has not revealed anything yet,
        # show Step 1 automatically. This does not mutate a widget-controlled key.
        if reveal == 0:
            reveal = 1
            st.session_state.guided_reveal_step = 1

        st.markdown("### Step-by-step solution")

        c1, c2, c3 = st.columns([1, 1, 1])
        with c1:
            st.button(
                "← Hints",
                key="guided_back_hints",
                use_container_width=True,
                on_click=_switch_guided_to_hints,
            )
        with c2:
            if reveal < total_steps:
                st.button(
                    "Reveal next step",
                    key="guided_reveal_next",
                    type="primary",
                    use_container_width=True,
                    on_click=_show_next_guided_step,
                    args=(total_steps,),
                )
        with c3:
            if reveal < total_steps:
                st.button(
                    "Show full solution",
                    key="guided_reveal_all",
                    use_container_width=True,
                    on_click=_show_all_guided_steps,
                    args=(total_steps,),
                )

        for i, step in enumerate(g.guided_steps, 1):
            if i <= reveal:
                render_guidance_step(i, step)

        if reveal >= total_steps:
            with st.container(border=True):
                st.markdown("### ✅ Verified final answer")
                if str(g.final_answer_mathio or "").strip():
                    render_mathio(str(g.final_answer_mathio).strip())
                else:
                    st.warning("The verifier did not return a final answer for this question.")

                if g.common_pitfalls:
                    with st.expander("Common mistakes to avoid", expanded=False):
                        for item in g.common_pitfalls:
                            guidance_item(item)


def render_ai_analysis(a: GeminiAnalysis) -> None:
    st.markdown('<div class="omt-section-kicker">Diagnosis</div>', unsafe_allow_html=True)
    st.markdown('<div class="omt-section-title">What the student understands — and where the reasoning breaks</div>', unsafe_allow_html=True)

    with st.container(border=True):
        st.caption(a.likely_syllabus_topic)
        st.markdown("**Question understood as**")
        render_math_text(a.interpreted_question)
        st.markdown("**Method shown in the working**")
        render_math_text(a.student_method)

    if a.first_logic_break_step > 0:
        st.markdown(
            f'<div class="omt-logic-break"><strong>First material logic break · Step {a.first_logic_break_step}</strong><br><span style="color:#7c4a10">Use the advice below for this point before continuing.</span></div>',
            unsafe_allow_html=True,
        )
        render_math_text(a.first_logic_break_explanation)
    else:
        st.markdown('<div class="omt-success-card"><strong>No material logic break found</strong><br>The visible method is mathematically coherent.</div>', unsafe_allow_html=True)
        if a.first_logic_break_explanation:
            render_math_text(a.first_logic_break_explanation)

    if a.steps:
        st.markdown("### Working, step by step")
        icons = {
            "correct": "✓",
            "partly_correct": "◐",
            "incorrect": "×",
            "unclear": "?",
            "unsupported": "•",
        }
        labels = {
            "correct": "Correct",
            "partly_correct": "Partly correct",
            "incorrect": "Needs advice",
            "unclear": "Unclear",
            "unsupported": "Needs support",
        }
        for step in a.steps:
            presentation_flag = bool(getattr(step, "presentation_error", False))
            icon = "!" if presentation_flag else icons.get(step.status, "•")
            status_label = "Presentation issue" if presentation_flag else labels.get(step.status, step.status.replace("_", " ").title())
            with st.expander(f"{icon}  Step {step.line_number} · {status_label}", expanded=(step.line_number == a.first_logic_break_step)):
                st.caption("Student wrote")
                render_mathio(step.student_step)
                if presentation_flag:
                    st.warning("This line is not written as a complete, unambiguous mathematical statement.")
                    presentation_explanation = getattr(step, "presentation_error_explanation", "")
                    if presentation_explanation:
                        render_math_text(presentation_explanation)
                detail_left, detail_right = st.columns([1.1, .9], gap="large")
                with detail_left:
                    st.markdown("**What this step is trying to do**")
                    render_mathio_mixed(step.logic_inferred)
                with detail_right:
                    st.markdown("**Tutor feedback**")
                    render_mathio_mixed(step.feedback)
                    st.caption(f"Issue type · {step.issue_type.replace('_', ' ').title()}")
                for formula in list(getattr(step, "supporting_math", []) or []):
                    render_mathio(formula)

    c1, c2 = st.columns(2, gap="large")
    with c1:
        with st.container(border=True):
            st.markdown("### ✓ What is working")
            if a.strengths:
                for item in a.strengths:
                    render_math_text(f"• {item}")
            else:
                st.caption("No specific strength was identified from the visible work.")
    with c2:
        with st.container(border=True):
            st.markdown("### → Main advice focus")
            render_math_text(a.misconception_or_gap)
            st.markdown("**Check your thinking**")
            render_math_text(a.diagnostic_question)

    st.markdown("### Guided advice")
    st.caption("Reveal only as much help as the student needs.")
    hint_cols = st.columns(3)
    for i, hint in enumerate(a.hint_ladder[:3], 1):
        with hint_cols[i - 1]:
            with st.expander(f"Hint {i}"):
                render_mathio_mixed(hint)
    with st.expander("Show corrected path and final answer"):
        for i, line in enumerate(a.corrected_path, 1):
            st.caption(f"Corrected step {i}")
            render_mathio(line)
        st.markdown("**Final answer**")
        render_mathio(a.final_answer)


def render_practice_evaluation(e: PracticeEvaluation) -> None:
    st.markdown("### Attempt feedback")
    c1, c2, c3 = st.columns(3)
    c1.metric("Answer", f"{e.answer_score}%")
    c2.metric("Reasoning", f"{e.reasoning_score}%")
    c3.metric("Mastery", e.mastery)

    with st.container(border=True):
        render_math_text(e.summary)

    if e.first_logic_break_step > 0:
        st.markdown(
            f'<div class="omt-logic-break"><strong>First reasoning break · Step {e.first_logic_break_step}</strong></div>',
            unsafe_allow_html=True,
        )
        render_math_text(e.first_logic_break_explanation)

    left, right = st.columns(2, gap="large")
    with left:
        with st.container(border=True):
            st.markdown("**✓ Strengths**")
            if e.strengths:
                for item in e.strengths:
                    render_math_text(f"• {item}")
            else:
                st.caption("No secure strength identified yet.")
    with right:
        with st.container(border=True):
            st.markdown("**→ Advice for next step**")
            if e.missing_or_incorrect_parts:
                st.warning("Complete: " + ", ".join(e.missing_or_incorrect_parts))
            presentation_errors = list(getattr(e, "presentation_errors", []) or [])
            for item in presentation_errors:
                render_mathio_mixed(item)
            if e.gaps:
                for item in e.gaps:
                    render_math_text(f"• {item}")

    with st.expander("Next hint", expanded=False):
        render_mathio_mixed(e.next_hint)
    with st.expander("Show corrected next step", expanded=False):
        render_mathio(e.corrected_next_step)



FULL_PAPER_MAX_BYTES = 30 * 1024 * 1024


def extract_docx_exam(file_obj: Any) -> tuple[str, list[UploadedAsset]]:
    """Extract readable text, tables and embedded images from a .docx exam paper."""
    data = file_obj.getvalue()
    if len(data) > FULL_PAPER_MAX_BYTES:
        raise GeminiTutorError("The Word paper is larger than the 30 MB full-paper limit.", category="input")

    document = Document(BytesIO(data))
    chunks: list[str] = []
    table_index = 0
    for child in document.element.body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            texts = [node.text for node in child.iter() if node.tag.rsplit("}", 1)[-1] == "t" and node.text]
            text = "".join(texts).strip()
            if text:
                chunks.append(text)
        elif tag == "tbl":
            table_index += 1
            chunks.append(f"[Table {table_index}]")
            for tr in child.iter():
                if tr.tag.rsplit("}", 1)[-1] != "tr":
                    continue
                cells = []
                for tc in tr:
                    if tc.tag.rsplit("}", 1)[-1] != "tc":
                        continue
                    texts = [node.text for node in tc.iter() if node.tag.rsplit("}", 1)[-1] == "t" and node.text]
                    cells.append("".join(texts).strip())
                if cells:
                    chunks.append(" | ".join(cells))

    assets: list[UploadedAsset] = []
    try:
        with zipfile.ZipFile(BytesIO(data)) as archive:
            media_names = [
                name for name in archive.namelist()
                if name.startswith("word/media/")
            ]
            for index, name in enumerate(media_names, 1):
                blob = archive.read(name)
                suffix = os.path.splitext(name)[1].lower()
                mime = {
                    ".png": "image/png",
                    ".jpg": "image/jpeg",
                    ".jpeg": "image/jpeg",
                    ".webp": "image/webp",
                }.get(suffix)
                if mime:
                    assets.append(
                        UploadedAsset(
                            name=f"{os.path.splitext(file_obj.name)[0]}_image_{index}{suffix}",
                            mime_type=mime,
                            data=blob,
                        )
                    )
    except zipfile.BadZipFile as exc:
        raise GeminiTutorError("The Word file could not be read as a valid .docx file.", category="input") from exc

    return "\n".join(chunks), assets


def full_paper_input(file_obj: Any) -> tuple[str, list[UploadedAsset]]:
    """Convert a PDF or DOCX paper into text/assets suitable for Gemini."""
    if file_obj is None:
        return "", []

    data = file_obj.getvalue()
    if len(data) > FULL_PAPER_MAX_BYTES:
        raise GeminiTutorError("The exam paper is larger than the 30 MB full-paper limit.", category="input")

    suffix = os.path.splitext(file_obj.name)[1].lower()
    if suffix == ".pdf":
        return "", [
            UploadedAsset(
                name=file_obj.name,
                mime_type="application/pdf",
                data=data,
            )
        ]
    if suffix == ".docx":
        return extract_docx_exam(file_obj)
    if suffix == ".doc":
        raise GeminiTutorError(
            "Legacy .doc files are not reliably readable in Streamlit Cloud. Save the paper as .docx or PDF and upload it again.",
            category="input",
        )
    raise GeminiTutorError("Upload a PDF or Word (.docx) exam paper.", category="input")


def scope_pdf_asset_to_pages(asset: UploadedAsset, pages: list[int]) -> UploadedAsset:
    """Create a smaller PDF containing only the pages for one detected question."""
    if asset.mime_type != "application/pdf" or not pages:
        return asset
    try:
        reader = PdfReader(BytesIO(asset.data))
        writer = PdfWriter()
        valid_pages = sorted({p for p in pages if 1 <= p <= len(reader.pages)})
        if not valid_pages:
            return asset
        for page_number in valid_pages:
            writer.add_page(reader.pages[page_number - 1])
        output = BytesIO()
        writer.write(output)
        return UploadedAsset(
            name=f"{os.path.splitext(asset.name)[0]}_pages_{'-'.join(map(str, valid_pages))}.pdf",
            mime_type="application/pdf",
            data=output.getvalue(),
        )
    except Exception:
        return asset


def paper_question_text_context(detected_question: Any, paper_text: str, next_question_number: str | None = None) -> str:
    pieces=[f"Question {detected_question.question_number}", detected_question.question_text or ""]
    if detected_question.subparts:
        pieces.append("Detected subparts:")
        pieces.extend(f"{p.label}: {p.question_text}" for p in detected_question.subparts)
    if paper_text:
        qn=re.escape(str(detected_question.question_number))
        hits=[re.search(rf"(?im)^\s*Question\s+{qn}\b",paper_text),
              re.search(rf"(?im)^\s*{qn}\s*[\.\)]\s+",paper_text)]
        hit=next((m for m in hits if m),None); start=hit.start() if hit else 0
        finish=min(len(paper_text),start+18000)
        if next_question_number:
            nn=re.escape(str(next_question_number))
            nhits=[re.search(rf"(?im)^\s*Question\s+{nn}\b",paper_text[start+1:]),
                   re.search(rf"(?im)^\s*{nn}\s*[\.\)]\s+",paper_text[start+1:])]
            nh=next((m for m in nhits if m),None)
            if nh: finish=start+1+nh.start()
        pieces+=["AUTHORITATIVE EXTRACTED QUESTION BLOCK:",paper_text[start:finish]]
    return "\n".join(x for x in pieces if x).strip()



def scoped_assets_for_paper_question(
    assets: list[UploadedAsset],
    page_numbers: list[int],
) -> list[UploadedAsset]:
    if not page_numbers:
        return assets
    scoped: list[UploadedAsset] = []
    for asset in assets:
        if asset.mime_type == "application/pdf":
            scoped.append(scope_pdf_asset_to_pages(asset, page_numbers))
        else:
            scoped.append(asset)
    return scoped


def paper_solution_markdown(
    *,
    track_label: str,
    paper_title: str,
    solutions: list[PaperQuestionSolution],
) -> str:
    lines = [
        f"# {paper_title or 'Full Paper Worked Solutions'}",
        "",
        f"**Track:** {track_label}",
        "",
        "> Suggested marking guides are AI-generated and are not official SEAB/MOE marking schemes.",
        "",
    ]
    for question in solutions:
        lines.append(f"## Question {question.question_number} — {question.topic}")
        lines.append("")
        for part in question.parts:
            lines.append(f"### {part.label} ({part.marks_available} marks; {part.mark_source})")
            lines.append("")
            lines.append(part.question_text)
            lines.append("")
            for idx, step in enumerate(part.worked_steps, 1):
                if step.explanation.strip():
                    lines.append(f"{idx}. {step.explanation.strip()}")
                for equation in step.equations:
                    if str(equation).strip():
                        lines.append(f"   `{str(equation).strip()}`")
            if part.final_answer_mathio.strip():
                lines.extend(["", f"**Final answer:** `{part.final_answer_mathio.strip()}`", ""])
            if part.marking_points:
                lines.append("**Suggested marking guide**")
                for point in part.marking_points:
                    ft = " (follow-through allowed)" if point.allow_follow_through else ""
                    lines.append(f"- {point.code} [{point.marks}]: {point.description}{ft}")
                lines.append("")
        lines.append(f"**Question total:** {question.total_marks} marks")
        lines.append("")
    return "\n".join(lines)


def build_paper_solution_docx(
    *,
    track_label: str,
    paper_title: str,
    solutions: list[PaperQuestionSolution],
) -> bytes:
    """Create a teacher-friendly Word export. Equations are kept as readable source text."""
    document = Document()
    document.add_heading(paper_title or "Full Paper Worked Solutions", level=0)
    document.add_paragraph(f"Track: {track_label}")
    p = document.add_paragraph()
    run = p.add_run("Important: ")
    run.bold = True
    p.add_run("The marking guides in this document are AI-generated suggestions, not official SEAB/MOE marking schemes.")

    for question in solutions:
        document.add_heading(f"Question {question.question_number} — {question.topic}", level=1)
        if question.page_numbers:
            document.add_paragraph("Source page(s): " + ", ".join(map(str, question.page_numbers)))
        for part in question.parts:
            document.add_heading(
                f"{part.label} — {part.marks_available} marks ({part.mark_source})",
                level=2,
            )
            if part.question_text.strip():
                document.add_paragraph(part.question_text.strip())

            for index, step in enumerate(part.worked_steps, 1):
                paragraph = document.add_paragraph(style="List Number")
                paragraph.add_run(step.explanation.strip() or f"Step {index}")
                for equation in step.equations:
                    if str(equation).strip():
                        document.add_paragraph(str(equation).strip())

            if part.final_answer_mathio.strip():
                p = document.add_paragraph()
                r = p.add_run("Final answer: ")
                r.bold = True
                p.add_run(part.final_answer_mathio.strip())

            if part.marking_points:
                document.add_heading("Suggested marking guide", level=3)
                for point in part.marking_points:
                    p = document.add_paragraph(style="List Bullet")
                    p.add_run(f"{point.code} [{point.marks}] ").bold = True
                    p.add_run(point.description)
                    if point.allow_follow_through:
                        p.add_run(" — follow-through allowed")

            if part.common_errors:
                document.add_heading("Common errors", level=3)
                for error in part.common_errors:
                    document.add_paragraph(error, style="List Bullet")

        document.add_paragraph(f"Question total: {question.total_marks} marks")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_full_paper_worked_step(step_number: int, step) -> None:
    """Render Full-paper worked solutions entirely through the MathIO-aware path.

    Explanatory prose stays readable, but any mathematical fragments embedded in it
    are detected and rendered using MathIO. Standalone equations always use MathIO.
    """
    explanation = clean_guidance_text(getattr(step, "explanation", "") or "")
    equations = [
        str(eq).strip()
        for eq in (getattr(step, "equations", []) or [])
        if str(eq).strip()
    ]

    # Remove presentation-only ellipsis commands.
    explanation = re.sub(r"\\+(?:dots|ldots|cdots)\b", "", explanation)
    explanation = re.sub(r"\.{3,}", "", explanation)
    explanation = re.sub(r"\s{2,}", " ", explanation).strip()

    cleaned_equations = []
    for equation in equations:
        equation = re.sub(r"\\+(?:dots|ldots|cdots)\b", "", equation)
        equation = re.sub(r"\.{3,}", "", equation)
        equation = re.sub(r"\s{2,}", " ", equation).strip()
        if equation:
            cleaned_equations.append(equation)

    if not explanation and not cleaned_equations:
        return

    with st.container(border=True):
        st.markdown(f"#### Step {step_number}")

        # This is the key difference from the old Full-paper renderer:
        # mixed prose/mathematics uses the MathIO-aware guidance renderer.
        if explanation:
            render_guidance_mixed_mathio(explanation)

        for equation in cleaned_equations:
            render_mathio(equation)


def render_paper_question_solution(solution: PaperQuestionSolution) -> None:
    title = f"Question {solution.question_number} · {solution.topic}"
    with st.expander(title, expanded=False):
        if solution.page_numbers:
            st.caption("Paper page(s): " + ", ".join(map(str, solution.page_numbers)))
        st.caption(f"Verification confidence: {solution.confidence.title()}")

        if solution.verification_note.strip():
            st.info(clean_guidance_text(solution.verification_note))

        if getattr(solution, "diagram_scene_3d", None) is not None:
            show_scene3d(
                solution.diagram_scene_3d,
                caption="3D diagram used in the worked solution",
            )
        elif getattr(solution, "diagram_scene_2d", None) is not None:
            show_scene2d(
                solution.diagram_scene_2d,
                caption="Diagram used in the worked solution",
            )

        for part in solution.parts:
            st.markdown(f"### {part.label}")

            # Question wording can itself contain equations/functions.
            if part.question_text.strip():
                render_guidance_mixed_mathio(part.question_text)

            if part.worked_steps:
                st.markdown("#### Worked solution")
                for idx, step in enumerate(part.worked_steps, 1):
                    render_full_paper_worked_step(idx, step)

            if part.final_answer_mathio.strip():
                st.markdown("#### ✅ Final answer")
                render_mathio(part.final_answer_mathio.strip())

            st.markdown(
                f"#### Suggested marking guide · {part.marks_available} marks"
                + (" · printed allocation" if part.mark_source == "printed" else " · AI-suggested allocation")
            )

            if part.marking_points:
                # Keep the compact teacher table, but sanitize raw LaTeX commands so
                # they are not exposed in a dataframe cell.
                table = pd.DataFrame(
                    [
                        {
                            "Code": point.code,
                            "Marks": point.marks,
                            "Criterion": _plainify_embedded_math(
                                clean_guidance_text(point.description)
                            ),
                            "Follow-through": "Yes" if point.allow_follow_through else "",
                        }
                        for point in part.marking_points
                    ]
                )
                st.dataframe(table, hide_index=True, use_container_width=True)

                # Optional MathIO view of marking criteria that contain mathematics.
                math_criteria = [
                    point
                    for point in part.marking_points
                    if _contains_raw_math_source(point.description)
                    or re.search(r"[A-Za-z0-9]\^\{?\d+\}?|=", point.description)
                ]
                if math_criteria:
                    with st.expander("MathIO view of mathematical marking criteria", expanded=False):
                        for point in math_criteria:
                            st.markdown(f"**{point.code} · {point.marks} mark(s)**")
                            render_guidance_mixed_mathio(point.description)

            else:
                st.caption("No marking points were generated for this part.")

            if part.common_errors:
                with st.expander("Common errors to watch for", expanded=False):
                    for error in part.common_errors:
                        render_guidance_mixed_mathio(f"• {error}")

        st.markdown(f"**Question total: {solution.total_marks} marks**")



def _omml_run(text: str):
    run = OxmlElement("m:r")
    t = OxmlElement("m:t")
    t.text = text
    run.append(t)
    return run


def _latex_display_text(value: str) -> str:
    """Normalize MathIO/LaTeX source before conversion to native Word equations."""
    text = str(value or "").strip()
    text = re.sub(r"\\+(?:dots|ldots|cdots)\b", "", text)
    text = re.sub(r"\\(?:left|right)", "", text)
    replacements = {
        r"\times": "×", r"\div": "÷", r"\pi": "π", r"\theta": "θ",
        r"\alpha": "α", r"\beta": "β", r"\gamma": "γ", r"\delta": "δ",
        r"\leq": "≤", r"\le": "≤", r"\geq": "≥", r"\ge": "≥",
        r"\neq": "≠", r"\pm": "±", r"\circ": "°", r"\cdot": "·",
        r"\infty": "∞", r"\approx": "≈", r"\therefore": "∴",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    text = re.sub(r"\\(?:text|mathrm|mathbf|operatorname)\{([^{}]*)\}", r"\1", text)
    return re.sub(r"\s{2,}", " ", text).strip()


def _balanced_group(src: str, start: int) -> tuple[str, int] | None:
    """Read a {...} group starting at start. Returns (inside, next_index)."""
    if start >= len(src) or src[start] != "{":
        return None
    depth = 0
    for i in range(start, len(src)):
        if src[i] == "{":
            depth += 1
        elif src[i] == "}":
            depth -= 1
            if depth == 0:
                return src[start + 1:i], i + 1
    return None



def _normalize_word_math_source(source: str) -> str:
    """Normalize generated math source before converting it to Word Equation Editor."""
    text = str(source or "").strip()
    text = _normalise_math_variable_italics(text)
    text = re.sub(r"\\left\s*([\(\[\{\|])", r"\1", text)
    text = re.sub(r"\\right\s*([\)\]\}\|])", r"\1", text)
    text = re.sub(r"\\left\b", "", text)
    text = re.sub(r"\\right\b", "", text)
    text = re.sub(r"(?<!\\)\bpi\b", r"\\pi", text, flags=re.IGNORECASE)
    text = re.sub(r"(?<!\\)\btheta\b", r"\\theta", text, flags=re.IGNORECASE)


    # Repair common missing-leading-slash model output.
    text = re.sub(r"(?<!\\)\boverrightarrow\s*\{", r"\\overrightarrow{", text)
    text = re.sub(r"(?<!\\)\bvec\s*\{", r"\\vec{", text)
    text = re.sub(r"(?<!\\)\bquad\b", r"\\quad", text)
    text = re.sub(r"(?<!\\)\bqquad\b", r"\\qquad", text)
    text = re.sub(r"(?<!\\)\bcdot\b", r"\\cdot", text)
    text = re.sub(r"(?<!\\)\btimes\b", r"\\times", text)

    # Layout commands should become spacing, never visible command names.
    text = text.replace(r"\qquad", "    ")
    text = text.replace(r"\quad", "  ")
    text = text.replace(r"\;", " ")
    text = text.replace(r"\:", " ")
    text = text.replace(r"\,", " ")
    text = text.replace(r"\!", "")

    replacements = {
        r"\cdot": "·",
        r"\times": "×",
        r"\pm": "±",
        r"\leq": "≤",
        r"\geq": "≥",
        r"\neq": "≠",
        r"\parallel": "∥",
        r"\perp": "⊥",
        r"\infty": "∞",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    return re.sub(r"[ \t]{3,}", "  ", text).strip()


def _append_omml_expression(parent, source: str) -> None:
    """Append common school mathematics as structured, editable OMML.

    Supports nested fractions, roots, superscripts/subscripts and ordinary symbols.
    Unknown constructs remain editable Word-math text rather than raw body text.
    """
    src = _normalize_word_math_source(_latex_display_text(source))
    i = 0
    plain = []

    def flush_plain():
        if plain:
            parent.append(_omml_run("".join(plain)))
            plain.clear()

    while i < len(src):
        if src.startswith(r"\frac", i):
            g1 = _balanced_group(src, i + 5)
            if g1:
                g2 = _balanced_group(src, g1[1])
                if g2:
                    flush_plain()
                    frac = OxmlElement("m:f")
                    num = OxmlElement("m:num")
                    den = OxmlElement("m:den")
                    _append_omml_expression(num, g1[0])
                    _append_omml_expression(den, g2[0])
                    frac.extend([num, den])
                    parent.append(frac)
                    i = g2[1]
                    continue

        # Vector arrow / over-arrow as a native Word math accent.
        vector_cmd = None
        cmd_len = 0
        if src.startswith(r"\overrightarrow", i):
            vector_cmd = r"\overrightarrow"
            cmd_len = len(vector_cmd)
        elif src.startswith(r"\vec", i):
            vector_cmd = r"\vec"
            cmd_len = len(vector_cmd)

        if vector_cmd:
            g = _balanced_group(src, i + cmd_len)
            if g:
                flush_plain()
                accent = OxmlElement("m:acc")
                acc_pr = OxmlElement("m:accPr")
                char = OxmlElement("m:chr")
                char.set(qn("m:val"), "→")
                acc_pr.append(char)
                elem = OxmlElement("m:e")
                _append_omml_expression(elem, g[0])
                accent.extend([acc_pr, elem])
                parent.append(accent)
                i = g[1]
                continue

        if src.startswith(r"\sqrt", i):
            g = _balanced_group(src, i + 5)
            if g:
                flush_plain()
                rad = OxmlElement("m:rad")
                rad_pr = OxmlElement("m:radPr")
                deg_hide = OxmlElement("m:degHide")
                deg_hide.set(qn("m:val"), "1")
                rad_pr.append(deg_hide)
                deg = OxmlElement("m:deg")
                elem = OxmlElement("m:e")
                _append_omml_expression(elem, g[0])
                rad.extend([rad_pr, deg, elem])
                parent.append(rad)
                i = g[1]
                continue

        # Superscript or subscript applying to the immediately preceding textual base.
        if src[i] in "^_" and i + 1 < len(src):
            op = src[i]
            if src[i + 1] == "{":
                g = _balanced_group(src, i + 1)
                if g:
                    exponent, next_i = g
                else:
                    exponent, next_i = src[i + 1], i + 2
            else:
                exponent, next_i = src[i + 1], i + 2

            if plain:
                base_char = plain.pop()
                flush_plain()
                node = OxmlElement("m:sSup" if op == "^" else "m:sSub")
                base = OxmlElement("m:e"); base.append(_omml_run(base_char))
                script = OxmlElement("m:sup" if op == "^" else "m:sub")
                _append_omml_expression(script, exponent)
                node.extend([base, script])
                parent.append(node)
                i = next_i
                continue

        # Common functions are retained as native math text inside the OMML zone.
        command_match = re.match(r"\\(sin|cos|tan|log|ln)\\?", src[i:])
        if command_match:
            plain.append(command_match.group(1))
            i += command_match.end()
            continue

        # Remaining commands: convert known symbols and drop layout-only commands.
        # Never expose source words such as "quad" in the generated paper.
        if src[i] == "\\":
            command = re.match(r"\\([A-Za-z]+)", src[i:])
            if command:
                name = command.group(1)
                symbol_map = {
                    "theta": "θ", "alpha": "α", "beta": "β", "gamma": "γ",
                    "delta": "δ", "pi": "π", "angle": "∠",
                    "rightarrow": "→", "leftarrow": "←",
                }
                presentation_only = {
                    "quad", "qquad", "textstyle", "displaystyle",
                    "left", "right", "mathrm", "mathbf", "mathit",
                }
                if name in symbol_map:
                    plain.append(symbol_map[name])
                elif name in presentation_only:
                    pass
                else:
                    plain.append(name)
                i += command.end()
                continue

        plain.append(src[i])
        i += 1

    flush_plain()



def _normalize_integral_source(source: str) -> str:
    text=str(source or "")
    text=re.sub(r"\\?int_([A-Za-z0-9.+-]+)\^\(([^)]+)\)",r"\\int_{\1}^{\2}",text)
    text=re.sub(r"\\?int_([A-Za-z0-9.+-]+)\^([A-Za-z0-9.+-]+)",r"\\int_{\1}^{\2}",text)
    text=re.sub(r"(?<!\\)\bint(?=_|\s)",r"\\int",text)
    return text

def _append_native_integral(paragraph, source: str) -> bool:
    src=_normalize_integral_source(source).strip()
    m=re.match(r"\\int(?:_\{([^{}]+)\})?(?:\^\{([^{}]+)\})?\s*(.*)$",src,re.S)
    if not m:
        return False
    lower,upper,body=m.group(1),m.group(2),m.group(3)
    math=OxmlElement("m:oMath")
    nary=OxmlElement("m:nary")
    pr=OxmlElement("m:naryPr")
    ch=OxmlElement("m:chr"); ch.set(qn("m:val"),"∫")
    lim=OxmlElement("m:limLoc"); lim.set(qn("m:val"),"subSup")
    pr.extend([ch,lim]); nary.append(pr)
    sub=OxmlElement("m:sub")
    if lower: _append_omml_expression(sub,lower)
    nary.append(sub)
    sup=OxmlElement("m:sup")
    if upper: _append_omml_expression(sup,upper)
    nary.append(sup)
    elem=OxmlElement("m:e"); _append_omml_expression(elem,body); nary.append(elem)
    math.append(nary); paragraph._p.append(math)
    return True

def append_word_math(paragraph, latex: str) -> None:
    """Insert an editable Microsoft Word Equation Editor object. DOCX stores these internally as OMML."""
    source = _normalize_word_math_source(_normalize_integral_source(str(latex or "").strip()))
    if not source:
        return
    if source.startswith(r"\int") and _append_native_integral(paragraph, source):
        return
    math = OxmlElement("m:oMath")
    _append_omml_expression(math, source)
    paragraph._p.append(math)


_WORD_MATH_FRAGMENT_RE = re.compile(
    r"""
    (?:
        (?<!\w)[A-Za-z]\s*=\s*\\frac\{[^{}]+\}\{[^{}]+\}
        |
        (?<!\w)[A-Za-z]\s*=\s*\\sqrt\{[^{}]+\}
        |
        \\frac\{[^{}]+\}\{[^{}]+\}
        |
        \\sqrt\{[^{}]+\}
        |
        \\(?:pi|theta|alpha|beta|gamma|delta)\b
        |
        \\?(?:overrightarrow|vec)\s*\{[A-Za-z0-9]+\}
        (?:\s*(?:=|≤|≥|<|>)\s*[^,.;\n]+)?
        |
        \\(?:sin|cos|tan|log|ln)\b(?:\s*\([^)]*\)|\s*\{[^{}]*\})?
        |
        (?<!\w)(?:\d+(?:\.\d+)?\s*)?[A-Za-z](?:\^\{?[-+]?\d+\}?|_\{?[A-Za-z0-9]+\}?)?
        (?:\s*[+\-×÷*/]\s*(?:\d+(?:\.\d+)?\s*)?[A-Za-z0-9](?:\^\{?[-+]?\d+\}?)?)*
        \s*(?:=|≤|≥|<|>)\s*
        [-+]?\d*(?:\.\d+)?[A-Za-z0-9]*(?:\^\{?[-+]?\d+\}?)?
        (?:\s*[+\-×÷*/]\s*[-+]?\d*(?:\.\d+)?[A-Za-z0-9]*(?:\^\{?[-+]?\d+\}?)?)*
        |
        \\?int(?:_\{?[^ \t\n{}()]+\}?|_\([^)]*\))?(?:\^\{?[^ \t\n{}()]+\}?|\^\([^)]*\))?\s*[^.;\n]+?\s*d[xyt]\b
        |
        \([^()\n]{1,100}\)\^\{?[-+]?\d+\}?
        |
        (?<!\w)[A-Za-z0-9]+\^\{?[-+]?\d+\}?
        |
        (?<!\w)[A-Za-z][A-Za-z0-9]*_\{?[A-Za-z0-9]+\}?
        |
        \([-+]?\d+(?:\.\d+)?\s*,\s*[-+]?\d+(?:\.\d+)?\)
        |
        (?<!\w)-?\d+(?:\.\d+)?\s*(?:cm|mm|m|km|g|kg|s|h|°|%)(?:\^2|\^3)?\b
    )
    """,
    re.VERBOSE,
)


def append_word_inline_math_linear(paragraph, latex: str) -> None:
    """Reliable inline native Word math for complex fragments inside prose.

    Word/LibreOffice interoperability is better when inline fractions/radicals are
    represented in linear mathematical notation. Standalone equation fields still
    use the fully structured OMML builder above.
    """
    text = _latex_display_text(latex)
    # Convert simple fraction/radical source to readable linear math inside an OMML zone.
    for _ in range(4):
        new = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", text)
        if new == text:
            break
        text = new
    text = re.sub(r"\\sqrt\{([^{}]+)\}", r"√(\1)", text)
    text = re.sub(r"\^\{([^{}]+)\}", r"^\1", text)
    text = re.sub(r"_\{([^{}]+)\}", r"_\1", text)
    math = OxmlElement("m:oMath")
    math.append(_omml_run(text))
    paragraph._p.append(math)



def _add_word_text_with_variable_italics(paragraph, text: str, *, bold: bool = False):
    """Add ordinary Word text while italicising standalone mathematical variables."""
    value = str(text or "")
    if not value:
        return

    # Variables are single Latin/Greek symbols used mathematically.
    token_re = re.compile(
        r"(?<![A-Za-z])("
        r"[A-Za-z]|"
        r"θ|π|α|β|γ|δ|λ|μ|σ|φ|ω"
        r")(?![A-Za-z])"
    )

    cursor = 0
    for m in token_re.finditer(value):
        if m.start() > cursor:
            run = paragraph.add_run(value[cursor:m.start()])
            run.bold = bold
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

        token = m.group(1)

        # Do not italicise common English article/pronoun tokens in prose.
        # A/I are only treated as variables when surrounded by mathematical punctuation.
        left = value[m.start()-1] if m.start() > 0 else ""
        right = value[m.end()] if m.end() < len(value) else ""
        math_context = (
            left in "=+-*/(^,<>≤≥"
            or right in "=+-*/)^,<>≤≥"
            or token not in {"A", "I", "a"}
        )

        run = paragraph.add_run(token)
        run.bold = bold
        run.italic = bool(math_context)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        cursor = m.end()

    if cursor < len(value):
        run = paragraph.add_run(value[cursor:])
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)


def append_word_mixed_math(paragraph, value: str, *, bold_prefix: str = "") -> None:
    """Write prose normally but place every detected mathematical fragment in OMML.

    This is the Word equivalent of MathIO mixed rendering. MathIO itself is browser-only;
    native OMML is the editable Word representation.
    """
    text = clean_guidance_text(value)
    text = re.sub(r"\\+(?:dots|ldots|cdots)\b", "", text)
    text = re.sub(r"\.{3,}", "", text)
    text = re.sub(r"\s{2,}", " ", text).strip()
    if bold_prefix:
        r = paragraph.add_run(bold_prefix)
        r.bold = True
    if not text:
        return

    def append_plain(segment: str) -> None:
        if not segment:
            return
        leading = " " if segment[:1].isspace() else ""
        trailing = " " if segment[-1:].isspace() else ""
        core = _plainify_embedded_math(segment.strip()) if segment.strip() else ""
        paragraph.add_run(leading + core + trailing)

    cursor = 0
    for match in _WORD_MATH_FRAGMENT_RE.finditer(text):
        if match.start() > cursor:
            append_plain(text[cursor:match.start()])
        fragment = match.group(0).strip()
        if r"\frac" in fragment or r"\sqrt" in fragment:
            append_word_inline_math_linear(paragraph, fragment)
        else:
            append_word_math(paragraph, fragment)
        cursor = match.end()
    if cursor < len(text):
        append_plain(text[cursor:])


def _scene_items(scene, name: str):
    return list(getattr(scene, name, []) or [])



def _diagram_font(size: int = 16):
    """Use Times New Roman for diagram labels; fall back to a metrically similar serif font."""
    candidates = [
        "/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf",
        "/usr/share/fonts/truetype/msttcorefonts/times.ttf",
        "/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman_Bold.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSerif-Regular.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSerif-Regular.ttf",
    ]
    for path in candidates:
        try:
            if Path(path).exists():
                return ImageFont.truetype(path, size=size)
        except Exception:
            pass
    return ImageFont.load_default()


def _diagram_label_font_for_canvas(width: int) -> object:
    # Approximate Times New Roman 11 pt on the generated raster.
    # Larger Word-export canvas gets proportionally larger raster text.
    px = 17 if width <= 980 else 20
    return _diagram_font(px)


def render_scene2d_png(scene, *, width: int = 960, height: int = 560, padding: int = 54) -> bytes:
    """Render a structured 2D maths scene to a clean PNG for app/Word use."""
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font = _diagram_label_font_for_canvas(width)

    x_min = float(getattr(scene, "x_min", -5) or -5)
    x_max = float(getattr(scene, "x_max", 5) or 5)
    y_min = float(getattr(scene, "y_min", -5) or -5)
    y_max = float(getattr(scene, "y_max", 5) or 5)
    if x_max <= x_min:
        x_min, x_max = -5, 5
    if y_max <= y_min:
        y_min, y_max = -5, 5

    plot_w = max(1, width - 2 * padding)
    plot_h = max(1, height - 2 * padding)

    x_span = max(1e-9, x_max - x_min)
    y_span = max(1e-9, y_max - y_min)
    has_circles = bool(_scene_items(scene, "circles"))
    preserve_aspect = has_circles or not bool(getattr(scene, "show_axes", False))

    if preserve_aspect:
        scale = min(plot_w / x_span, plot_h / y_span)
        used_w = x_span * scale
        used_h = y_span * scale
        left = padding + (plot_w - used_w) / 2
        top = padding + (plot_h - used_h) / 2
        sx = sy = scale

        def xy(x, y):
            px = left + (float(x) - x_min) * scale
            py = top + used_h - (float(y) - y_min) * scale
            return (int(round(px)), int(round(py)))

        frame = [int(left), int(top), int(left + used_w), int(top + used_h)]
    else:
        sx = plot_w / x_span
        sy = plot_h / y_span

        def xy(x, y):
            px = padding + (float(x) - x_min) / x_span * plot_w
            py = height - padding - (float(y) - y_min) / y_span * plot_h
            return (int(round(px)), int(round(py)))

        frame = [padding, padding, width-padding, height-padding]

    draw.rectangle(frame, outline=(205,205,205), width=1)

    if bool(getattr(scene, "show_axes", False)):
        # Grid at sensible integer positions when range is moderate.
        x0 = max(math.ceil(x_min), -30)
        x1 = min(math.floor(x_max), 30)
        y0 = max(math.ceil(y_min), -30)
        y1 = min(math.floor(y_max), 30)
        if x1-x0 <= 30 and y1-y0 <= 30:
            for xv in range(x0, x1+1):
                p1=xy(xv,y_min); p2=xy(xv,y_max)
                draw.line([p1,p2], fill=(236,236,236), width=1)
            for yv in range(y0, y1+1):
                p1=xy(x_min,yv); p2=xy(x_max,yv)
                draw.line([p1,p2], fill=(236,236,236), width=1)
        if x_min <= 0 <= x_max:
            draw.line([xy(0,y_min),xy(0,y_max)], fill=(70,70,70), width=2)
        if y_min <= 0 <= y_max:
            draw.line([xy(x_min,0),xy(x_max,0)], fill=(70,70,70), width=2)

    points = {str(getattr(p,"id","")): p for p in _scene_items(scene,"points")}

    # Curves / graphs.
    for poly in _scene_items(scene,"polylines"):
        samples = list(getattr(poly,"points",[]) or [])
        coords=[]
        for v in samples:
            if isinstance(v,(list,tuple)) and len(v)>=2:
                coords.append(xy(v[0],v[1]))
        if len(coords)>=2:
            draw.line(coords, fill=(35,35,35), width=3)
            label=str(getattr(poly,"label","") or "")
            if label:
                mx,my=coords[len(coords)//2]
                draw.text((mx+6,my-16),label,fill=(25,25,25),font=font)

    # Circles.
    for c in _scene_items(scene,"circles"):
        cx=float(getattr(c,"center_x",0)); cy=float(getattr(c,"center_y",0)); r=float(getattr(c,"radius",1))
        center=xy(cx,cy)
        rx=max(2,int(round(r*sx))); ry=max(2,int(round(r*sy)))
        box=[center[0]-rx,center[1]-ry,center[0]+rx,center[1]+ry]
        draw.ellipse(box,outline=(35,35,35),width=3)
        label=str(getattr(c,"label","") or "")
        if label:
            draw.text((center[0]+rx+5,center[1]-8),label,fill=(20,20,20),font=font)

    # Segments.
    for seg in _scene_items(scene,"segments"):
        a=points.get(str(getattr(seg,"start",""))); b=points.get(str(getattr(seg,"end","")))
        if not a or not b:
            continue
        pa=xy(getattr(a,"x",0),getattr(a,"y",0)); pb=xy(getattr(b,"x",0),getattr(b,"y",0))
        draw.line([pa,pb],fill=(30,30,30),width=3)
        label=str(getattr(seg,"label","") or "")
        if label:
            mx=(pa[0]+pb[0])//2; my=(pa[1]+pb[1])//2
            draw.rectangle([mx-2,my-17,mx+max(20,6*len(label)),my-3],fill="white")
            draw.text((mx,my-16),label,fill=(20,20,20),font=font)

    # Points and labels.
    for p in points.values():
        px,py=xy(getattr(p,"x",0),getattr(p,"y",0))
        draw.ellipse([px-4,py-4,px+4,py+4],fill=(20,20,20))
        label=str(getattr(p,"label","") or "")
        if label:
            draw.text((px+7,py-14),label,fill=(20,20,20),font=font)

    # Angle labels. Arms should normally already be represented by segments.
    for ang in _scene_items(scene,"angles"):
        vertex=points.get(str(getattr(ang,"vertex","")))
        if not vertex:
            continue
        px,py=xy(getattr(vertex,"x",0),getattr(vertex,"y",0))
        label=str(getattr(ang,"label","") or "")
        if label:
            draw.text((px+12,py+8),label,fill=(20,20,20),font=font)

    buf=BytesIO()
    img.save(buf,format="PNG")
    return buf.getvalue()



def _compile_generated_function(expression: str):
    text = str(expression or "").replace("−","-").replace("×","*").replace("÷","/")
    text = text.replace(r"\left","").replace(r"\right","").replace(r"\pi","pi").replace("π","pi")
    text = re.sub(r"\\sqrt\s*\(([^()]*)\)", r"sqrt(\1)", text)
    text = re.sub(r"\\sqrt\s*\{([^{}]*)\}", r"sqrt(\1)", text)
    for fn in ("sin","cos","tan","sqrt","log","ln"):
        text = text.replace("\\"+fn,fn)
    text = text.replace("{","(").replace("}",")").replace("^","**")
    # Convert common implicit multiplication while preserving function names.
    text = re.sub(r"(?<=\d)(?=x\b)", "*", text)
    text = re.sub(r"(?<=\d)(?=(?:sin|cos|tan|sqrt|log|ln)\b)", "*", text)
    text = re.sub(r"(?<=x)(?=\()", "*", text)
    text = re.sub(r"(?<=\))(?=[A-Za-z0-9(])", "*", text)
    if not re.fullmatch(r"[0-9A-Za-z_+\-*/().,\s*]+", text):
        return None
    allowed={
        "sin":math.sin, "cos":math.cos, "tan":math.tan,
        "sqrt":math.sqrt, "log":math.log, "ln":math.log, "pi":math.pi,
    }
    def f(x):
        env=dict(allowed); env["x"]=float(x)
        return float(eval(text,{"__builtins__":{}},env))
    return f


def _question_equation_sources(question) -> list[str]:
    values = []
    values.append(str(getattr(question, "stem_text", "") or ""))
    values.extend(str(x) for x in (getattr(question, "stem_equations", []) or []))
    for part in (getattr(question, "parts", []) or []):
        values.append(str(getattr(part, "prompt_text", "") or ""))
        values.extend(str(x) for x in (getattr(part, "equations", []) or []))
    return [v for v in values if v.strip()]


def _extract_y_functions(question) -> list[tuple[str, object]]:
    """Extract exact y=f(x) definitions, including hidden graph-construction equations."""
    found = []
    seen = set()

    # graph_equations is authoritative for graph-reading questions where the
    # printed equation intentionally contains unknown parameters a,b,c,d.
    sources = list(getattr(question, "graph_equations", []) or [])
    sources.extend(_question_equation_sources(question))

    for source in sources:
        # split generously on punctuation but stop before prose clauses
        for m in re.finditer(r"\by\s*=\s*([^.;\n]+)", source, flags=re.I):
            expr = m.group(1).strip()
            # trim common prose that follows a displayed equation
            expr = re.split(r"\s+(?:and|where|for|with|intersect|meets|at)\s+", expr, maxsplit=1, flags=re.I)[0].strip()
            key = re.sub(r"\s+","",expr)
            if not expr or key in seen:
                continue
            # A general form such as a*sin(b*x+c)+d cannot be plotted until
            # numerical parameter values are known. Do not silently draw blank axes.
            symbolic_parameters = set(re.findall(r"\b[a-d]\b", expr, flags=re.I))
            if symbolic_parameters and not re.search(r"\b(?:a|b|c|d)\s*=\s*[-+]?\d", source, flags=re.I):
                continue

            fn = _compile_generated_function(expr)
            if fn is not None:
                seen.add(key)
                found.append((expr, fn))
    return found


def _geogebra_safe_expression(expr: str) -> str | None:
    """Convert a generated function expression to a conservative GeoGebra input."""
    value = str(expr or "").strip()
    if not value:
        return None

    # Convert common MathIO/LaTeX notation to GeoGebra input syntax.
    replacements = {
        r"\pi": "pi",
        "π": "pi",
        r"\theta": "theta",
        r"\cdot": "*",
        r"\times": "*",
        "×": "*",
        "÷": "/",
        r"\left": "",
        r"\right": "",
    }
    for old, new in replacements.items():
        value = value.replace(old, new)

    # A few common fraction forms. Nested fractions continue to use local fallback.
    for _ in range(4):
        new_value = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"((\1)/(\2))", value)
        if new_value == value:
            break
        value = new_value

    value = re.sub(r"\^\{([^{}]+)\}", r"^(\1)", value)
    value = re.sub(r"\bsqrt\s*\{([^{}]+)\}", r"sqrt(\1)", value, flags=re.I)
    value = re.sub(r"\\sqrt\{([^{}]+)\}", r"sqrt(\1)", value)
    value = re.sub(r"\s+", " ", value).strip()

    # Insert multiplication in common textbook forms: 3x, 2(x+1), )( .
    value = re.sub(r"(?<=\d)(?=[A-Za-z(])", "*", value)
    value = re.sub(r"(?<=[A-Za-z)])(?=\()", "*", value)

    # Whitelist only ordinary graphing syntax. Never send arbitrary LLM text as a command.
    if not re.fullmatch(r"[A-Za-z0-9_+\-*/^().,\s]+", value):
        return None

    return value


def _question_graph_spec(question) -> dict | None:
    """Return sanitized GeoGebra commands and a sensible view for y=f(x) questions."""
    funcs = _extract_y_functions(question)
    if not funcs:
        return None

    expressions = []
    functions = []
    for expr, fn in funcs:
        safe = _geogebra_safe_expression(expr)
        if not safe:
            continue
        expressions.append(safe)
        functions.append(fn)

    if not expressions:
        return None

    scene = getattr(question, "diagram_scene_2d", None)
    text = " ".join(_question_equation_sources(question)).lower()
    is_trig = any(token in text for token in ("sin", "cos", "tan", "trigonometric"))

    # Prefer question-specified scene bounds. Trig defaults show multiple periods.
    if scene is not None:
        xmin = float(getattr(scene, "x_min", -6.5 if is_trig else -5) or (-6.5 if is_trig else -5))
        xmax = float(getattr(scene, "x_max", 6.5 if is_trig else 5) or (6.5 if is_trig else 5))
    else:
        xmin, xmax = ((-2 * math.pi, 2 * math.pi) if is_trig else (-5.0, 5.0))

    if xmax <= xmin:
        xmin, xmax = (-2 * math.pi, 2 * math.pi) if is_trig else (-5.0, 5.0)

    # Sample locally only to choose a useful y-view; GeoGebra draws the actual curve.
    finite = []
    for fn in functions:
        for i in range(361):
            x = xmin + (xmax - xmin) * i / 360
            try:
                y = float(fn(x))
            except Exception:
                continue
            if math.isfinite(y) and abs(y) < 1000:
                finite.append(y)

    if scene is not None:
        raw_ymin = getattr(scene, "y_min", None)
        raw_ymax = getattr(scene, "y_max", None)
    else:
        raw_ymin = raw_ymax = None

    if finite:
        lo, hi = min(finite), max(finite)
        span = max(1.0, hi - lo)
        pad = max(1.0, 0.12 * span)
        calc_ymin = math.floor(lo - pad)
        calc_ymax = math.ceil(hi + pad)
    else:
        calc_ymin, calc_ymax = -5.0, 5.0

    ymin = float(raw_ymin) if raw_ymin is not None else float(calc_ymin)
    ymax = float(raw_ymax) if raw_ymax is not None else float(calc_ymax)
    if ymax <= ymin:
        ymin, ymax = float(calc_ymin), float(calc_ymax)

    commands = [f"f{i}(x)={expr}" for i, expr in enumerate(expressions, 1)]
    signature_source = "|".join(commands) + f"|{xmin:.6g}|{xmax:.6g}|{ymin:.6g}|{ymax:.6g}"
    signature = hashlib.sha256(signature_source.encode("utf-8")).hexdigest()[:20]

    return {
        "commands": commands,
        "expressions": expressions,
        "xmin": xmin,
        "xmax": xmax,
        "ymin": ymin,
        "ymax": ymax,
        "grid": True,
        "height": 450,
        "signature": signature,
    }


def _geogebra_graph_store() -> dict[str, bytes]:
    store = st.session_state.get("setter_geogebra_graphs")
    if not isinstance(store, dict):
        store = {}
        st.session_state.setter_geogebra_graphs = store
    return store


def render_geogebra_question_graph(question, *, figure_caption: str = "") -> bytes | None:
    """Mount GeoGebra, capture its PNG, and cache it for Word export."""
    spec = _question_graph_spec(question)
    if spec is None:
        return None

    qnum = str(getattr(question, "question_number", "graph"))
    store = _geogebra_graph_store()
    cache_key = f"{qnum}:{spec['signature']}"

    # Existing capture survives Streamlit reruns and is immediately reusable.
    if cache_key in store:
        png = store[cache_key]
        if png:
            st.image(png, caption=figure_caption or None, use_container_width=True)
            st.caption("Graph source: GeoGebra (captured and ready for Word export).")
            return png

    if not _GEOGEBRA_COMPONENT_AVAILABLE or _GEOGEBRA_COMPONENT is None:
        st.caption("GeoGebra component is unavailable in this Streamlit version; using the local graph renderer.")
        return None

    st.caption("Interactive graph source: GeoGebra. The exact function below is also captured for the Word paper.")
    try:
        result = _GEOGEBRA_COMPONENT(
            data=spec,
            default={"capture": None},
            on_capture_change=lambda: None,
            key=f"setter_geogebra_{qnum}_{spec['signature']}",
        )
        capture = getattr(result, "capture", None)
        if isinstance(capture, dict) and capture.get("ok") and capture.get("png_base64"):
            try:
                png = base64.b64decode(capture["png_base64"])
                if png:
                    store[cache_key] = png
                    st.session_state.setter_geogebra_graphs = store
                    st.caption("GeoGebra graph captured successfully for download.")
                    return png
            except Exception:
                pass
    except Exception as exc:
        st.caption(f"GeoGebra component unavailable for this graph; local fallback will be used. ({type(exc).__name__})")
    return None


def _captured_geogebra_png(question) -> bytes | None:
    """Return the cached GeoGebra PNG for the current question/function signature."""
    spec = _question_graph_spec(question)
    if spec is None:
        return None
    qnum = str(getattr(question, "question_number", "graph"))
    return _geogebra_graph_store().get(f"{qnum}:{spec['signature']}")


def add_png_to_word(doc: Document, png: bytes, *, caption: str = "") -> None:
    """Insert a captured graph image into a Word paper with exam-paper styling."""
    if not png:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(BytesIO(png), width=Cm(13.5))
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = cp.runs[0]
        rr.italic = True
        rr.font.name = "Times New Roman"
        rr.font.size = Pt(11)
        rr._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        rr._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rr._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")




# ---------------------------------------------------------------------------
# Real-life context images (Wikimedia Commons)
# ---------------------------------------------------------------------------
_REAL_LIFE_CONTEXT_TERMS = [
    (r"\bjugs?\b", "water jug"),
    (r"\bbottles?\b", "water bottle"),
    (r"\bcups?\b", "drinking cup"),
    (r"\bglasses?\b", "drinking glass"),
    (r"\bcontainers?\b", "container"),
    (r"\bbuckets?\b", "bucket"),
    (r"\btanks?\b", "water tank"),
    (r"\bcars?\b", "car"),
    (r"\bbuses?\b", "bus"),
    (r"\btrains?\b", "train"),
    (r"\bbicycles?\b|\bbikes?\b", "bicycle"),
    (r"\bboats?\b", "boat"),
    (r"\bships?\b", "ship"),
    (r"\baircraft\b|\baeroplanes?\b|\bairplanes?\b", "airplane"),
    (r"\bbuildings?\b", "building"),
    (r"\btowers?\b", "tower"),
    (r"\btrees?\b", "tree"),
    (r"\bladders?\b", "ladder"),
    (r"\bflags?\b", "flag pole"),
    (r"\bclocks?\b", "clock"),
    (r"\bthermometers?\b", "thermometer"),
    (r"\bbooks?\b", "book"),
    (r"\bbox(?:es)?\b", "cardboard box"),
    (r"\bparcels?\b", "parcel"),
    (r"\bballs?\b", "ball"),
    (r"\bcones?\s+of\s+(?:ice cream|popcorn)\b", "ice cream cone"),
]


def _real_life_context_query(text: str) -> str | None:
    """Return a Commons query only when the physical object itself is contextually useful.

    Do not add decorative photos to questions whose primary information source is a
    table, graph, chart, frequency distribution or statistical dataset.
    """
    source = re.sub(r"\s+", " ", str(text or "")).strip().lower()
    if not source:
        return None

    # Structured-data questions should display their mathematical/statistical
    # representation, not an unrelated decorative photograph.
    suppress_patterns = [
        r"\bthe table (?:below )?shows\b",
        r"\btable of\b",
        r"\bfrequency table\b",
        r"\bfrequency distribution\b",
        r"\bcumulative frequency\b",
        r"\bhistogram\b",
        r"\bbox(?:-and-whisker)? plot\b",
        r"\bscatter (?:plot|diagram|graph)\b",
        r"\bbar (?:chart|graph)\b",
        r"\bpie chart\b",
        r"\bline graph\b",
        r"\bthe graph (?:below )?shows\b",
        r"\bthe chart (?:below )?shows\b",
        r"\bdata set\b",
        r"\bdistribution of\b",
    ]
    if any(re.search(pattern, source, flags=re.IGNORECASE) for pattern in suppress_patterns):
        return None

    for pattern, query in _REAL_LIFE_CONTEXT_TERMS:
        if re.search(pattern, source, flags=re.IGNORECASE):
            return query
    return None



def _fetch_commons_context_image(query: str) -> dict | None:
    """Best-effort reusable context image lookup from Wikimedia Commons."""
    query = str(query or "").strip()
    if not query:
        return None

    params = {
        "action": "query",
        "format": "json",
        "generator": "search",
        "gsrsearch": f"File:{query}",
        "gsrnamespace": "6",
        "gsrlimit": "8",
        "prop": "imageinfo",
        "iiprop": "url|mime|extmetadata",
        "iiurlwidth": "900",
        "origin": "*",
    }
    api_url = "https://commons.wikimedia.org/w/api.php?" + urlencode(params)

    try:
        request = Request(
            api_url,
            headers={"User-Agent": "Math-Advisor-Education-App/1.0"},
        )
        with urlopen(request, timeout=6) as response:
            payload = json.loads(response.read().decode("utf-8"))

        pages = list((payload.get("query", {}) or {}).get("pages", {}).values())
        pages.sort(key=lambda item: int(item.get("index", 9999)))

        for page in pages:
            infos = page.get("imageinfo") or []
            if not infos:
                continue
            info = infos[0]
            mime = str(info.get("mime", "") or "")
            if mime not in {"image/jpeg", "image/png", "image/webp"}:
                continue

            image_url = info.get("thumburl") or info.get("url")
            if not image_url:
                continue

            image_request = Request(
                image_url,
                headers={"User-Agent": "Math-Advisor-Education-App/1.0"},
            )
            with urlopen(image_request, timeout=7) as image_response:
                content = image_response.read(7_000_000)
            if not content or len(content) >= 7_000_000:
                continue

            metadata = info.get("extmetadata") or {}
            def meta(name: str) -> str:
                raw = metadata.get(name) or {}
                return re.sub(r"<[^>]+>", "", str(raw.get("value", "") or "")).strip()

            return {
                "bytes": content,
                "title": str(page.get("title", "Wikimedia Commons image")).replace("File:", ""),
                "license": meta("LicenseShortName") or meta("UsageTerms") or "Wikimedia Commons",
                "artist": meta("Artist"),
                "page_url": info.get("descriptionurl") or "",
                "query": query,
            }
    except Exception:
        return None

    return None


def _context_image_for_text(text: str) -> dict | None:
    query = _real_life_context_query(text)
    if not query:
        return None
    return _fetch_commons_context_image(query)


def _question_context_text(question) -> str:
    parts = [str(getattr(question, "stem_text", "") or "")]
    parts.extend(str(v) for v in (getattr(question, "stem_equations", []) or []))
    for part in list(getattr(question, "parts", []) or []):
        parts.append(str(getattr(part, "prompt_text", "") or ""))
    return " ".join(parts)


def _question_context_image(question) -> dict | None:
    return _context_image_for_text(_question_context_text(question))


def _context_image_caption(image: dict, *, figure_label: str = "") -> str:
    source = "Wikimedia Commons"
    license_name = str(image.get("license", "") or "").strip()
    title = str(image.get("title", "") or "").strip()
    prefix = (figure_label + ". ") if figure_label else ""
    details = title or str(image.get("query", "context image"))
    if license_name:
        return f"{prefix}Context image: {details} — {source}, {license_name}"
    return f"{prefix}Context image: {details} — {source}"


def show_context_image_for_text(text: str) -> None:
    image = _context_image_for_text(text)
    if not image:
        return
    st.image(image["bytes"], width=320)
    st.caption(_context_image_caption(image))


def add_context_image_to_word(doc: Document, image: dict, *, caption: str) -> None:
    if not image or not image.get("bytes"):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        run.add_picture(BytesIO(image["bytes"]), width=Cm(7.5))
    except Exception:
        return
    cp = doc.add_paragraph(_context_image_caption(image, figure_label=caption))
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cp.runs:
        rr = cp.runs[0]
        rr.italic = True
        rr.font.name = "Times New Roman"
        rr.font.size = Pt(9)


# ---------------------------------------------------------------------------
# Deterministic statistics graph engine
# ---------------------------------------------------------------------------
def _stats_graph_spec(question):
    return getattr(question, "statistics_graph", None)


def _stats_graph_values(spec, name: str) -> list:
    value = getattr(spec, name, None)
    return list(value or [])


def validate_statistics_graph_spec(spec) -> list[str]:
    """Structural validation before a statistics graph may enter a paper."""
    if spec is None:
        return []
    issues = []
    graph_type = str(getattr(spec, "graph_type", "") or "")

    xs = _stats_graph_values(spec, "x_values")
    ys = _stats_graph_values(spec, "y_values")
    boundaries = _stats_graph_values(spec, "class_boundaries")
    frequencies = _stats_graph_values(spec, "frequencies")
    cumulatives = _stats_graph_values(spec, "cumulative_frequencies")
    five = _stats_graph_values(spec, "five_number_summary")

    if graph_type == "cumulative_frequency":
        if len(boundaries) < 2:
            issues.append("cumulative-frequency graph needs at least two class boundaries")
        if frequencies and len(frequencies) != max(0, len(boundaries) - 1):
            issues.append("frequency count must be one fewer than boundary count")
        if not cumulatives and frequencies:
            running = 0.0
            cumulatives = [0.0]
            for f in frequencies:
                running += float(f)
                cumulatives.append(running)
        if cumulatives:
            if len(cumulatives) != len(boundaries):
                issues.append("cumulative-frequency values must align with class boundaries")
            if any(float(b) < float(a) for a, b in zip(cumulatives, cumulatives[1:])):
                issues.append("cumulative frequencies must not decrease")
            if frequencies and abs(float(cumulatives[-1]) - sum(float(f) for f in frequencies)) > 1e-6:
                issues.append("final cumulative frequency must equal the total frequency")

    elif graph_type == "histogram":
        if len(boundaries) < 2 or len(frequencies) != max(0, len(boundaries) - 1):
            issues.append("histogram needs class boundaries and one frequency per class")

    elif graph_type == "frequency_polygon":
        if len(xs) < 2 or len(xs) != len(ys):
            issues.append("frequency polygon requires paired x/y values")

    elif graph_type in {"scatter", "line_graph"}:
        if len(xs) < 2 or len(xs) != len(ys):
            issues.append(f"{graph_type} requires paired x/y values")

    elif graph_type == "bar_chart":
        labels = _stats_graph_values(spec, "labels")
        if not ys or (labels and len(labels) != len(ys)):
            issues.append("bar chart requires values and matching labels")

    elif graph_type == "box_plot":
        if len(five) != 5:
            issues.append("box plot requires [minimum, Q1, median, Q3, maximum]")
        elif any(float(b) < float(a) for a, b in zip(five, five[1:])):
            issues.append("five-number summary must be non-decreasing")

    else:
        issues.append("unsupported statistics graph type")

    return issues


def _statistics_cf_points(spec) -> tuple[list[float], list[float]]:
    boundaries = [float(v) for v in _stats_graph_values(spec, "class_boundaries")]
    cumulatives = [float(v) for v in _stats_graph_values(spec, "cumulative_frequencies")]
    frequencies = [float(v) for v in _stats_graph_values(spec, "frequencies")]
    if not cumulatives and frequencies and boundaries:
        cumulatives = [0.0]
        running = 0.0
        for f in frequencies:
            running += f
            cumulatives.append(running)
    return boundaries, cumulatives


def _nice_ticks(lo: float, hi: float, count: int = 8) -> list[float]:
    if not math.isfinite(lo) or not math.isfinite(hi) or hi <= lo:
        return [lo]
    span = hi - lo
    raw = span / max(2, count)
    power = 10 ** math.floor(math.log10(raw)) if raw > 0 else 1
    fraction = raw / power
    step = (1 if fraction <= 1 else 2 if fraction <= 2 else 5 if fraction <= 5 else 10) * power
    start = math.floor(lo / step) * step
    values = []
    v = start
    for _ in range(100):
        if v >= lo - 1e-9 and v <= hi + 1e-9:
            values.append(v)
        v += step
        if v > hi + step:
            break
    return values or [lo, hi]



def _monotone_cubic_curve_points(xs, ys, *, samples_per_interval: int = 28):
    """Return smooth monotone cubic Hermite samples through cumulative-frequency points.

    Uses the Fritsch-Carlson slope limiter. This preserves monotonicity and prevents
    the overshoot that ordinary cubic splines can introduce into an ogive.
    """
    xs = [float(v) for v in xs]
    ys = [float(v) for v in ys]
    n = min(len(xs), len(ys))
    if n < 2:
        return list(zip(xs, ys))

    xs = xs[:n]
    ys = ys[:n]
    h = [xs[i+1] - xs[i] for i in range(n-1)]
    if any(v <= 0 for v in h):
        return list(zip(xs, ys))

    delta = [(ys[i+1] - ys[i]) / h[i] for i in range(n-1)]
    m = [0.0] * n
    m[0] = delta[0]
    m[-1] = delta[-1]

    for i in range(1, n-1):
        if delta[i-1] * delta[i] <= 0:
            m[i] = 0.0
        else:
            w1 = 2*h[i] + h[i-1]
            w2 = h[i] + 2*h[i-1]
            m[i] = (w1 + w2) / (w1/delta[i-1] + w2/delta[i])

    # Fritsch-Carlson monotonicity limiter.
    for i in range(n-1):
        if abs(delta[i]) < 1e-12:
            m[i] = 0.0
            m[i+1] = 0.0
            continue
        a = m[i] / delta[i]
        b = m[i+1] / delta[i]
        s = a*a + b*b
        if s > 9.0:
            tau = 3.0 / (s ** 0.5)
            m[i] = tau * a * delta[i]
            m[i+1] = tau * b * delta[i]

    points = []
    spi = max(8, int(samples_per_interval))
    for i in range(n-1):
        x0, x1 = xs[i], xs[i+1]
        y0, y1 = ys[i], ys[i+1]
        hi = h[i]
        for j in range(spi):
            t = j / spi
            h00 = 2*t**3 - 3*t**2 + 1
            h10 = t**3 - 2*t**2 + t
            h01 = -2*t**3 + 3*t**2
            h11 = t**3 - t**2
            x = x0 + t*hi
            y = h00*y0 + h10*hi*m[i] + h01*y1 + h11*hi*m[i+1]
            points.append((x, y))
    points.append((xs[-1], ys[-1]))
    return points


def render_statistics_graph_png(spec, *, width: int = 1100, height: int = 650, completed: bool = True) -> bytes:
    """Render exam-quality statistics graphs as a PNG using PIL only."""
    issues = validate_statistics_graph_spec(spec)
    if issues:
        raise ValueError("; ".join(issues))

    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    margin_l, margin_r, margin_t, margin_b = 105, 50, 50, 85
    plot_l, plot_r = margin_l, width - margin_r
    plot_t, plot_b = margin_t, height - margin_b

    try:
        font = ImageFont.truetype("DejaVuSerif.ttf", 24)
        small = ImageFont.truetype("DejaVuSerif.ttf", 19)
        title_font = ImageFont.truetype("DejaVuSerif.ttf", 24)
    except Exception:
        font = small = title_font = ImageFont.load_default()

    graph_type = str(getattr(spec, "graph_type", ""))
    x_label = str(getattr(spec, "x_label", "") or "")
    y_label = str(getattr(spec, "y_label", "") or "")
    title = str(getattr(spec, "title", "") or "")
    show_grid = bool(getattr(spec, "show_grid", True))

    xs, ys = [], []
    if graph_type == "cumulative_frequency":
        xs, ys = _statistics_cf_points(spec)
    elif graph_type in {"frequency_polygon", "scatter", "line_graph"}:
        xs = [float(v) for v in _stats_graph_values(spec, "x_values")]
        ys = [float(v) for v in _stats_graph_values(spec, "y_values")]
    elif graph_type == "histogram":
        xs = [float(v) for v in _stats_graph_values(spec, "class_boundaries")]
        ys = [float(v) for v in _stats_graph_values(spec, "frequencies")]
    elif graph_type == "bar_chart":
        ys = [float(v) for v in _stats_graph_values(spec, "y_values")]
        xs = list(range(len(ys)))
    elif graph_type == "box_plot":
        five = [float(v) for v in _stats_graph_values(spec, "five_number_summary")]
        xs = five
        ys = [1.0] * len(five)

    if graph_type == "bar_chart":
        x_min, x_max = -0.75, max(0.75, len(xs)-0.25)
    else:
        x_min = min(xs) if xs else 0.0
        x_max = max(xs) if xs else 1.0
        if x_max <= x_min:
            x_max = x_min + 1.0
        x_pad = 0.04 * (x_max - x_min)
        x_min -= x_pad
        x_max += x_pad

    if graph_type == "box_plot":
        y_min, y_max = 0.0, 2.0
    else:
        y_min = min(0.0, min(ys) if ys else 0.0)
        y_max = max(1.0, max(ys) if ys else 1.0)
        y_max += max(1.0, 0.08 * (y_max-y_min))

    def px(x):
        return plot_l + (float(x)-x_min)/(x_max-x_min)*(plot_r-plot_l)
    def py(y):
        return plot_b - (float(y)-y_min)/(y_max-y_min)*(plot_b-plot_t)

    # Grid and ticks.
    x_ticks = _nice_ticks(x_min, x_max, 9)
    y_ticks = _nice_ticks(y_min, y_max, 8)
    if show_grid:
        for x in x_ticks:
            xx = px(x)
            draw.line((xx, plot_t, xx, plot_b), fill=(226,226,226), width=1)
        for y in y_ticks:
            yy = py(y)
            draw.line((plot_l, yy, plot_r, yy), fill=(226,226,226), width=1)

    draw.line((plot_l, plot_b, plot_r, plot_b), fill="black", width=2)
    draw.line((plot_l, plot_t, plot_l, plot_b), fill="black", width=2)

    # Tick labels.
    for x in x_ticks:
        label = f"{x:g}"
        bb = draw.textbbox((0,0), label, font=small)
        draw.text((px(x)-(bb[2]-bb[0])/2, plot_b+10), label, fill="black", font=small)
    for y in y_ticks:
        label = f"{y:g}"
        bb = draw.textbbox((0,0), label, font=small)
        draw.text((plot_l-12-(bb[2]-bb[0]), py(y)-(bb[3]-bb[1])/2), label, fill="black", font=small)

    if completed:
        if graph_type == "cumulative_frequency":
            pts = [(px(x), py(y)) for x,y in zip(xs,ys)]
            if len(pts) >= 2:
                smooth_data = _monotone_cubic_curve_points(
                    xs,
                    ys,
                    samples_per_interval=32,
                )
                smooth = [(px(x), py(y)) for x, y in smooth_data]
                if len(smooth) >= 2:
                    draw.line(smooth, fill="black", width=4, joint="curve")
                for p in pts:
                    draw.ellipse((p[0]-4,p[1]-4,p[0]+4,p[1]+4), fill="black")

        elif graph_type == "histogram":
            boundaries = [float(v) for v in _stats_graph_values(spec, "class_boundaries")]
            freqs = [float(v) for v in _stats_graph_values(spec, "frequencies")]
            for left,right,f in zip(boundaries,boundaries[1:],freqs):
                draw.rectangle((px(left),py(f),px(right),py(0)), outline="black", width=3)

        elif graph_type == "frequency_polygon":
            pts = [(px(x),py(y)) for x,y in zip(xs,ys)]
            draw.line(pts, fill="black", width=3)
            for p in pts:
                draw.ellipse((p[0]-4,p[1]-4,p[0]+4,p[1]+4), fill="black")

        elif graph_type == "scatter":
            for x,y in zip(xs,ys):
                p=(px(x),py(y))
                draw.ellipse((p[0]-5,p[1]-5,p[0]+5,p[1]+5), fill="black")

        elif graph_type == "line_graph":
            pts=[(px(x),py(y)) for x,y in zip(xs,ys)]
            draw.line(pts,fill="black",width=3)
            for p in pts:
                draw.ellipse((p[0]-4,p[1]-4,p[0]+4,p[1]+4),fill="black")

        elif graph_type == "bar_chart":
            labels = [str(v) for v in _stats_graph_values(spec, "labels")]
            for i,y in enumerate(ys):
                xl=px(i-0.32); xr=px(i+0.32)
                draw.rectangle((xl,py(y),xr,py(0)), outline="black", width=3)
                if i < len(labels):
                    bb=draw.textbbox((0,0),labels[i],font=small)
                    draw.text(((xl+xr)/2-(bb[2]-bb[0])/2,plot_b+35),labels[i],fill="black",font=small)

        elif graph_type == "box_plot":
            mn,q1,med,q3,mx=[float(v) for v in _stats_graph_values(spec,"five_number_summary")]
            yy=py(1)
            draw.line((px(mn),yy,px(mx),yy),fill="black",width=3)
            draw.rectangle((px(q1),yy-45,px(q3),yy+45),outline="black",width=3)
            draw.line((px(med),yy-45,px(med),yy+45),fill="black",width=3)
            draw.line((px(mn),yy-22,px(mn),yy+22),fill="black",width=3)
            draw.line((px(mx),yy-22,px(mx),yy+22),fill="black",width=3)

    if title:
        bb=draw.textbbox((0,0),title,font=title_font)
        draw.text(((width-(bb[2]-bb[0]))/2,10),title,fill="black",font=title_font)
    if x_label:
        bb=draw.textbbox((0,0),x_label,font=font)
        draw.text(((plot_l+plot_r-(bb[2]-bb[0]))/2,height-42),x_label,fill="black",font=font)
    if y_label:
        # PIL text rotation for y label.
        tmp=Image.new("RGBA",(300,50),(255,255,255,0))
        td=ImageDraw.Draw(tmp)
        td.text((5,5),y_label,fill="black",font=font)
        tmp=tmp.rotate(90,expand=True)
        image.paste(tmp,(12,int((height-tmp.height)/2)),tmp)

    buf=BytesIO()
    image.save(buf,format="PNG",dpi=(300,300))
    return buf.getvalue()


def show_statistics_graph(spec, *, caption: str = "", completed: bool = True) -> None:
    png=render_statistics_graph_png(spec,completed=completed)
    st.image(png,caption=caption or None,use_container_width=True)


def add_statistics_graph_to_word(doc: Document, spec, *, caption: str = "", completed: bool = True) -> None:
    png=render_statistics_graph_png(spec,width=1100,height=650,completed=completed)
    add_png_to_word(doc,png,caption=caption)


def _question_requests_student_draw_graph(question) -> bool:
    text=" ".join(_question_equation_sources(question)).lower()
    return bool(re.search(
        r"\b(draw|sketch|plot|construct|complete)\b.{0,45}\b(graph|curve|histogram|polygon|box plot|scatter)",
        text,
        flags=re.I,
    ))


def _question_claims_graph_is_shown(question) -> bool:
    text=" ".join(_question_equation_sources(question)).lower()
    return bool(re.search(
        r"\b(graph|curve|diagram)\b.{0,35}\b(shows|shown|below|given)\b|"
        r"\b(shows|shown)\b.{0,35}\b(graph|curve)\b|"
        r"\bfrom the graph\b",
        text,
        flags=re.I,
    ))


def validate_function_graph_readiness(question) -> list[str]:
    """A graph-reading question may never pass with only unknown parameters."""
    if not _question_claims_graph_is_shown(question):
        return []
    if _question_requests_student_draw_graph(question):
        return []

    hidden = list(getattr(question,"graph_equations",[]) or [])
    functions = _extract_y_functions(question)
    issues=[]
    if not hidden:
        issues.append("graph-reading question has no hidden numeric graph equation")
    if not functions:
        issues.append("no plot-ready numerical function is available")
    return issues


def build_function_graph_scene(question):
    """Build a complete 2D function scene directly from the exact question equations.

    This is the guaranteed fallback for generated papers. It does not require
    Gemini to have supplied diagram_scene_2d.
    """
    functions = _extract_y_functions(question)
    if not functions:
        return None

    original = getattr(question, "diagram_scene_2d", None)
    spec = _question_graph_spec(question)
    if spec is None:
        return None

    # Use the existing scene as a style/bounds source when available, but never
    # depend on it for the actual function curve.
    if original is not None:
        try:
            scene = original.model_copy(deep=True)
        except Exception:
            scene = copy.deepcopy(original)
    else:
        scene = SimpleNamespace(
            x_min=float(spec["xmin"]),
            x_max=float(spec["xmax"]),
            y_min=float(spec["ymin"]),
            y_max=float(spec["ymax"]),
            show_axes=True,
            points=[],
            segments=[],
            polylines=[],
            circles=[],
            arcs=[],
            polygons=[],
            texts=[],
        )

    scene.show_axes = True
    scene.x_min = float(spec["xmin"])
    scene.x_max = float(spec["xmax"])
    scene.y_min = float(spec["ymin"])
    scene.y_max = float(spec["ymax"])

    all_chunks = []
    finite_y = []
    xmin = float(scene.x_min)
    xmax = float(scene.x_max)
    visible_span = max(1.0, float(scene.y_max) - float(scene.y_min))

    for function_index, (expr, fn) in enumerate(functions, 1):
        current = []
        previous_y = None

        # Dense sampling keeps trig graphs smooth in the Word/PDF export.
        for i in range(1201):
            x = xmin + (xmax - xmin) * i / 1200
            try:
                y = float(fn(x))
            except Exception:
                y = float("nan")

            if not math.isfinite(y) or abs(y) > 1e5:
                if len(current) >= 2:
                    all_chunks.append(
                        SimpleNamespace(
                            id=f"paper_curve_{function_index}_{len(all_chunks)+1}",
                            points=current,
                            label=(f"y = {expr}" if not any(
                                str(getattr(c, "label", "")) == f"y = {expr}" for c in all_chunks
                            ) else ""),
                        )
                    )
                current = []
                previous_y = None
                continue

            # Split discontinuities instead of drawing a vertical line through an
            # asymptote (important for tangent/rational functions).
            if previous_y is not None and abs(y - previous_y) > max(30.0, 4.0 * visible_span):
                if len(current) >= 2:
                    all_chunks.append(
                        SimpleNamespace(
                            id=f"paper_curve_{function_index}_{len(all_chunks)+1}",
                            points=current,
                            label=(f"y = {expr}" if not any(
                                str(getattr(c, "label", "")) == f"y = {expr}" for c in all_chunks
                            ) else ""),
                        )
                    )
                current = []

            current.append([x, y])
            previous_y = y
            if abs(y) < 1e4:
                finite_y.append(y)

        if len(current) >= 2:
            all_chunks.append(
                SimpleNamespace(
                    id=f"paper_curve_{function_index}_{len(all_chunks)+1}",
                    points=current,
                    label=(f"y = {expr}" if not any(
                        str(getattr(c, "label", "")) == f"y = {expr}" for c in all_chunks
                    ) else ""),
                )
            )

    if not all_chunks:
        return None

    # Replace any placeholder/model polylines with the authoritative curves.
    scene.polylines = all_chunks
    return scene


def render_function_graph_png(question, *, width: int = 1100, height: int = 650) -> bytes | None:
    """Render an exact function graph PNG for paper export."""
    scene = build_function_graph_scene(question)
    if scene is None:
        return None
    return render_scene2d_png(scene, width=width, height=height)


def ensure_question_function_curve(question):
    """Return an authoritative function scene built from the question equations."""
    exact_scene = build_function_graph_scene(question)
    if exact_scene is not None:
        return exact_scene
    return getattr(question, "diagram_scene_2d", None)



def _segment_distance_to_point(a, b, p) -> float:
    ax,ay=a; bx,by=b; px,py=p
    dx,dy=bx-ax,by-ay
    denom=dx*dx+dy*dy
    if denom <= 1e-12:
        return math.hypot(px-ax,py-ay)
    t=max(0.0,min(1.0,((px-ax)*dx+(py-ay)*dy)/denom))
    qx,qy=ax+t*dx,ay+t*dy
    return math.hypot(px-qx,py-qy)


def validate_question_scene_2d(question, scene) -> list[str]:
    """Heuristic structural validation so diagrams match the wording."""
    if scene is None:
        return []
    text=" ".join(_question_equation_sources(question))
    lower=text.lower()
    issues=[]
    pts={str(getattr(p,"label","") or getattr(p,"id","")).strip():p for p in _scene_items(scene,"points")}
    circles=list(_scene_items(scene,"circles"))
    segs=list(_scene_items(scene,"segments"))

    if "circle" in lower and not circles:
        issues.append("question mentions a circle but the diagram has no circle")

    # Explicit point lists: Points D, E, and F ... / point B
    named=set()
    for m in re.finditer(r"\bpoints?\s+([A-Z](?:\s*,\s*[A-Z])*(?:\s*,?\s*and\s+[A-Z])?)", text):
        named.update(re.findall(r"\b[A-Z]\b",m.group(1)))
    for m in re.finditer(r"\bpoint\s+([A-Z])\b", text):
        named.add(m.group(1))
    for m in re.finditer(r"\b(?:chord|line|segment)\s+([A-Z]{2})\b", text):
        named.update(m.group(1))
    for m in re.finditer(r"\bintersect(?:s|ed)?\s+at\s+([A-Z])\b", text, flags=re.I):
        named.add(m.group(1).upper())
    missing=sorted(x for x in named if x not in pts)
    if missing:
        issues.append("missing labelled point(s): "+", ".join(missing))

    # Chords XY must actually be drawn between X and Y.
    def has_segment(u,v):
        for s in segs:
            a=str(getattr(s,"start","")); b=str(getattr(s,"end",""))
            if {a,b}=={u,v}:
                return True
            pa=pts.get(u); pb=pts.get(v)
            sa=pts.get(a); sb=pts.get(b)
            if pa and pb and sa and sb:
                if {str(getattr(sa,"label","")),str(getattr(sb,"label",""))}=={u,v}:
                    return True
        return False

    for m in re.finditer(r"\bchords?\s+([A-Z]{2})(?:\s+and\s+([A-Z]{2}))?", text):
        for chord in [m.group(1),m.group(2)]:
            if chord and not has_segment(chord[0],chord[1]):
                issues.append(f"missing chord {chord}")

    # Tangent ABC at B should contain points A-B-C and B should lie on circle.
    tang = re.search(r"\b([A-Z]{3})\s+is\s+a\s+tangent\b.*?\bat\s+(?:the\s+)?point\s+([A-Z])", text, flags=re.I)
    if tang:
        letters=tang.group(1)
        touch=tang.group(2).upper()
        for ch in letters:
            if ch not in pts:
                issues.append(f"tangent point {ch} is missing")
        if touch in pts and circles:
            p=pts[touch]
            ok=False
            for c in circles:
                d=math.hypot(float(p.x)-float(c.center_x),float(p.y)-float(c.center_y))
                if abs(d-float(c.radius)) <= max(0.08,0.04*float(c.radius)):
                    ok=True
            if not ok:
                issues.append(f"tangent contact {touch} is not on the circle")

    # Intersections: if chords XY and UV intersect at G, G should lie on both.
    inter=re.search(r"\bchords?\s+([A-Z]{2})\s+and\s+([A-Z]{2})\s+intersect\s+at\s+([A-Z])", text, flags=re.I)
    if inter:
        c1,c2,g=inter.group(1),inter.group(2),inter.group(3).upper()
        if all(ch in pts for ch in c1+c2+g):
            gp=(float(pts[g].x),float(pts[g].y))
            scale=max(1.0, float(getattr(scene,"x_max",5))-float(getattr(scene,"x_min",-5)))
            tol=0.035*scale
            for chord in (c1,c2):
                a=(float(pts[chord[0]].x),float(pts[chord[0]].y))
                b=(float(pts[chord[1]].x),float(pts[chord[1]].y))
                if _segment_distance_to_point(a,b,gp)>tol:
                    issues.append(f"intersection point {g} is not on chord {chord}")
    return issues



def show_scene2d(scene, *, caption: str = "") -> None:
    if scene is None:
        return
    try:
        png = render_scene2d_png(scene)
        st.image(png, caption=caption or None, use_container_width=True)
    except Exception as exc:
        st.info("The diagram could not be rendered reliably. " + str(exc))


def add_scene2d_to_word(doc: Document, scene, *, caption: str = "") -> None:
    if scene is None:
        return
    try:
        png = render_scene2d_png(scene, width=1100, height=650)
        p=doc.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run()
        run.add_picture(BytesIO(png), width=Cm(13.5))
        if caption:
            cp=doc.add_paragraph(caption)
            cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for rr in cp.runs:
                rr.italic = True
                rr.font.name = "Times New Roman"
                rr.font.size = Pt(11)
                rr._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
                rr._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
                rr._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    except Exception:
        # Do not break paper generation if a malformed scene slips through.
        pass



def _project3d_point(x: float, y: float, z: float, *, scale: float = 1.0) -> tuple[float,float]:
    """Stable exam-style isometric projection."""
    # Horizontal x, depth z at 30 degrees, vertical y.
    u = float(x) + 0.55 * float(z)
    v = float(y) - 0.32 * float(z)
    return (u*scale, v*scale)


def _scene3d_wire_segments(scene):
    """Return projected labelled line segments for vertices/edges and solid primitives."""
    segs=[]
    verts={str(getattr(v,"id","")):v for v in list(getattr(scene,"vertices",[]) or [])}

    # Explicit edges are authoritative.
    for e in list(getattr(scene,"edges",[]) or []):
        a=verts.get(str(getattr(e,"start",""))); b=verts.get(str(getattr(e,"end","")))
        if a and b:
            segs.append((
                _project3d_point(a.x,a.y,a.z),
                _project3d_point(b.x,b.y,b.z),
                str(getattr(e,"label","") or ""),
                bool(getattr(e,"dashed",False)),
            ))

    # Boxes/cuboids.
    for box in list(getattr(scene,"boxes",[]) or []):
        cx,cy,cz=[float(v) for v in box.center]
        hx=float(box.width)/2; hy=float(box.height)/2; hz=float(box.depth)/2
        pts={}
        for ix,sx in enumerate((-1,1)):
            for iy,sy in enumerate((-1,1)):
                for iz,sz in enumerate((-1,1)):
                    pts[(ix,iy,iz)]=_project3d_point(cx+sx*hx,cy+sy*hy,cz+sz*hz)
        for iy in (0,1):
            for iz in (0,1): segs.append((pts[(0,iy,iz)],pts[(1,iy,iz)],"",False))
        for ix in (0,1):
            for iz in (0,1): segs.append((pts[(ix,0,iz)],pts[(ix,1,iz)],"",False))
        for ix in (0,1):
            for iy in (0,1): segs.append((pts[(ix,iy,0)],pts[(ix,iy,1)],"",False))

    # Extruded prisms.
    for ex in list(getattr(scene,"extrusions",[]) or []):
        prof=list(getattr(ex,"profile",[]) or [])
        if len(prof)>=3:
            cx,cy,cz=[float(v) for v in ex.center]
            d=float(ex.depth)/2
            front=[]; back=[]
            for uv in prof:
                if len(uv)<2: continue
                u,v=float(uv[0]),float(uv[1])
                front.append(_project3d_point(cx+u,cy+v,cz-d))
                back.append(_project3d_point(cx+u,cy+v,cz+d))
            for ring in (front,back):
                for i in range(len(ring)): segs.append((ring[i],ring[(i+1)%len(ring)],"",False))
            for a,b in zip(front,back): segs.append((a,b,"",False))

    # Cylinders
    for cyl in list(getattr(scene,"cylinders",[]) or []):
        cx,cy,cz=[float(v) for v in cyl.center]; r=float(cyl.radius); h=float(cyl.height)
        axis=str(getattr(cyl,"axis","y")); rings=[]
        for sign in (-1,1):
            ring=[]
            for k in range(40):
                t=2*math.pi*k/40
                if axis=="y": p=(cx+r*math.cos(t),cy+sign*h/2,cz+r*math.sin(t))
                elif axis=="x": p=(cx+sign*h/2,cy+r*math.cos(t),cz+r*math.sin(t))
                else: p=(cx+r*math.cos(t),cy+r*math.sin(t),cz+sign*h/2)
                ring.append(_project3d_point(*p))
            rings.append(ring)
            for k in range(40): segs.append((ring[k],ring[(k+1)%40],"",False))
        for k in (0,10,20,30): segs.append((rings[0][k],rings[1][k],"",False))

    # Cones
    for cone in list(getattr(scene,"cones",[]) or []):
        cx,cy,cz=[float(v) for v in cone.center]; r=float(cone.radius); h=float(cone.height)
        axis=str(getattr(cone,"axis","y")); direction=str(getattr(cone,"direction","positive"))
        vs=1 if direction!="negative" else -1; bs=-vs
        rim=[]
        for k in range(48):
            t=2*math.pi*k/48
            if axis=="y":
                p=(cx+r*math.cos(t),cy+bs*h/2,cz+r*math.sin(t)); apex=(cx,cy+vs*h/2,cz)
            elif axis=="x":
                p=(cx+bs*h/2,cy+r*math.cos(t),cz+r*math.sin(t)); apex=(cx+vs*h/2,cy,cz)
            else:
                p=(cx+r*math.cos(t),cy+r*math.sin(t),cz+bs*h/2); apex=(cx,cy,cz+vs*h/2)
            rim.append(_project3d_point(*p))
        ap=_project3d_point(*apex)
        for k in range(48): segs.append((rim[k],rim[(k+1)%48],"",False))
        for k in (0,12,24,36): segs.append((rim[k],ap,"",False))

    # Spheres
    for sph in list(getattr(scene,"spheres",[]) or []):
        cx,cy,cz=[float(v) for v in sph.center]; r=float(sph.radius)
        for plane in ("xy","xz","yz"):
            ring=[]
            for k in range(48):
                t=2*math.pi*k/48
                if plane=="xy": p=(cx+r*math.cos(t),cy+r*math.sin(t),cz)
                elif plane=="xz": p=(cx+r*math.cos(t),cy,cz+r*math.sin(t))
                else: p=(cx,cy+r*math.cos(t),cz+r*math.sin(t))
                ring.append(_project3d_point(*p))
            for k in range(48): segs.append((ring[k],ring[(k+1)%48],"",False))

    return segs, verts



def render_exam_3d_wireframe_png(scene, *, width: int = 1100, height: int = 700) -> bytes:
    """Exam-style 3D wireframe PNG following the supplied cuboid/prism reference."""
    image=Image.new("RGB",(width,height),"white")
    draw=ImageDraw.Draw(image)
    try:
        label_font=ImageFont.truetype("DejaVuSerif.ttf",22)
        dim_font=ImageFont.truetype("DejaVuSerif.ttf",20)
    except Exception:
        label_font=dim_font=ImageFont.load_default()

    cx,cy=width*.5,height*.56
    scale=min(width,height)*.055
    def proj(pt):
        x,y,z=[float(v) for v in pt]
        return (cx+scale*(x-.68*y), cy-scale*(z-.34*y))

    vertices={}
    for p in list(getattr(scene,"points",[]) or []):
        pid=str(getattr(p,"id","") or getattr(p,"label","") or "")
        xyz=getattr(p,"xyz",None) or getattr(p,"coords",None) or getattr(p,"point",None)
        if xyz and len(xyz)>=3: vertices[pid]=tuple(xyz[:3])

    edges=[]
    for e in list(getattr(scene,"segments",[]) or [])+list(getattr(scene,"edges",[]) or []):
        a=str(getattr(e,"a","") or getattr(e,"start","") or getattr(e,"from_id","") or "")
        b=str(getattr(e,"b","") or getattr(e,"end","") or getattr(e,"to_id","") or "")
        if a in vertices and b in vertices:
            edges.append((a,b,bool(getattr(e,"hidden",False))))

    if vertices and edges:
        for a,b,hidden in edges:
            pa,pb=proj(vertices[a]),proj(vertices[b])
            if hidden:
                for k in range(0,16,2):
                    t1=k/16; t2=(k+1)/16
                    draw.line((pa[0]+(pb[0]-pa[0])*t1,pa[1]+(pb[1]-pa[1])*t1,
                               pa[0]+(pb[0]-pa[0])*t2,pa[1]+(pb[1]-pa[1])*t2),fill="black",width=2)
            else:
                draw.line((*pa,*pb),fill="black",width=3)
        for name,xyz in vertices.items():
            p=proj(xyz)
            draw.ellipse((p[0]-3,p[1]-3,p[0]+3,p[1]+3),fill="black")
            if name: draw.text((p[0]+6,p[1]-24),name,fill="black",font=label_font)
        for txt in list(getattr(scene,"texts",[]) or []):
            label=str(getattr(txt,"text","") or getattr(txt,"label","") or "")
            xyz=getattr(txt,"xyz",None) or getattr(txt,"point",None)
            if label and xyz and len(xyz)>=3:
                p=proj(xyz); draw.text((p[0]+5,p[1]+5),label,fill="black",font=dim_font)
    else:
        kind=str(getattr(scene,"kind","") or getattr(scene,"solid_type","") or "").lower()
        dims=getattr(scene,"dimensions",{}) or {}
        if "cuboid" in kind or "rectangular prism" in kind:
            L=float(dims.get("length",8)); W=float(dims.get("width",5)); H=float(dims.get("height",4))
            pts={"A":(0,0,0),"B":(L,0,0),"C":(L,W,0),"D":(0,W,0),
                 "E":(0,0,H),"F":(L,0,H),"G":(L,W,H),"H":(0,W,H)}
            ed=[("A","B"),("B","C"),("C","D"),("D","A"),("E","F"),("F","G"),("G","H"),("H","E"),
                ("A","E"),("B","F"),("C","G"),("D","H")]
            for a,b in ed: draw.line((*proj(pts[a]),*proj(pts[b])),fill="black",width=3)
            for name,xyz in pts.items():
                p=proj(xyz); draw.text((p[0]+5,p[1]-22),name,fill="black",font=label_font)
        else:
            draw.text((70,70),"3D solid requires structured point/edge data.",fill="black",font=label_font)

    buf=BytesIO(); image.save(buf,format="PNG",dpi=(300,300)); return buf.getvalue()


def render_exam_3d_wireframe_png(scene, *, width: int=960, height: int=620, padding: int=60) -> bytes:
    """Render a clear isometric 3D exam schematic to PNG."""
    img=Image.new("RGB",(width,height),"white")
    draw=ImageDraw.Draw(img)
    font=_diagram_label_font_for_canvas(width)
    segs,verts=_scene3d_wire_segments(scene)

    projected=[]
    for a,b,label,dashed in segs: projected.extend([a,b])
    for v in verts.values(): projected.append(_project3d_point(v.x,v.y,v.z))
    if not projected:
        buf=BytesIO(); img.save(buf,format="PNG"); return buf.getvalue()

    xs=[p[0] for p in projected]; ys=[p[1] for p in projected]
    xmin,xmax=min(xs),max(xs); ymin,ymax=min(ys),max(ys)
    dx=max(1e-6,xmax-xmin); dy=max(1e-6,ymax-ymin)
    scale=min((width-2*padding)/dx,(height-2*padding)/dy)

    def screen(p):
        x=padding+(p[0]-xmin)*scale
        y=height-padding-(p[1]-ymin)*scale
        return (int(round(x)),int(round(y)))

    for a,b,label,dashed in segs:
        pa,pb=screen(a),screen(b)
        if dashed:
            steps=12
            for i in range(0,steps,2):
                t1=i/steps; t2=min(1,(i+1)/steps)
                aa=(int(pa[0]+(pb[0]-pa[0])*t1),int(pa[1]+(pb[1]-pa[1])*t1))
                bb=(int(pa[0]+(pb[0]-pa[0])*t2),int(pa[1]+(pb[1]-pa[1])*t2))
                draw.line([aa,bb],fill=(75,75,75),width=2)
        else:
            draw.line([pa,pb],fill=(30,30,30),width=3)
        if label:
            mx=(pa[0]+pb[0])//2; my=(pa[1]+pb[1])//2
            draw.text((mx+5,my-15),label,fill=(20,20,20),font=font)

    for v in verts.values():
        p=screen(_project3d_point(v.x,v.y,v.z))
        draw.ellipse([p[0]-3,p[1]-3,p[0]+3,p[1]+3],fill=(20,20,20))
        if getattr(v,"label",""):
            draw.text((p[0]+6,p[1]-13),str(v.label),fill=(20,20,20),font=font)

    buf=BytesIO(); img.save(buf,format="PNG"); return buf.getvalue()



# ---------------------------------------------------------------------------
# Deterministic dimensioned 3D solid renderer
# ---------------------------------------------------------------------------

_SOLID3D_TAG_RE = re.compile(r"(?is)\[SOLID3D\](.*?)\[/SOLID3D\]")

_SOLID3D_GENERATION_REQUIREMENTS = """
IMPORTANT DIMENSIONED 3D GEOMETRY REQUIREMENTS:
- For questions involving cuboids, cubes, cylinders, cones, spheres, hemispheres,
  triangular prisms or composite solids, provide a deterministic diagram specification.
- Put the specification in diagram_spec using this exact transport form:

  [SOLID3D]
  type=composite_cylinder_cuboid
  cuboid_length=12
  cuboid_depth=8
  cuboid_height=5
  cylinder_radius=3
  cylinder_height=10
  unit=cm
  [/SOLID3D]

- Supported type values:
  cuboid, cube, cylinder, cone, sphere, hemisphere, triangular_prism,
  composite_cylinder_cuboid, composite_cone_cylinder,
  composite_prism_cuboid, composite_cylinder_prism,
  composite_cylinder_cuboid_prism.
- Use only positive numerical dimensions.
- Every dimension needed by the question must appear in the SOLID3D block.
- The wording and the SOLID3D values must agree exactly.
- For triangular prisms use:
  prism_base, prism_triangle_height, prism_length.
- For composite solids made from a cylinder, cuboid and triangular prism use:
  cuboid_length, cuboid_depth, cuboid_height,
  cylinder_radius, cylinder_height,
  prism_base, prism_triangle_height, prism_length,
  and placement values where needed.
- Composite parts must touch along a clearly defined face or circular base.
- Do not allow impossible overlaps or a mounted solid whose footprint is larger than
  the supporting face unless the wording explicitly describes an overhang.
- Questions may test volume, total surface area, exposed surface area, missing dimensions,
  comparison of solids, reverse problems, or the volume/material removed or added.
"""

def _solid_num(value, default=None):
    try:
        x = float(str(value).strip())
        if x > 0 and math.isfinite(x):
            return x
    except Exception:
        pass
    return default


def _parse_solid3d_tag(text: str) -> dict | None:
    source = str(text or "")
    m = _SOLID3D_TAG_RE.search(source)
    if not m:
        return None
    payload = {}
    for raw in m.group(1).splitlines():
        line = raw.strip()
        if not line or "=" not in line:
            continue
        key, value = line.split("=", 1)
        payload[key.strip().lower()] = value.strip()
    if "type" not in payload:
        return None
    return payload


def _dimensioned_solid_spec(question) -> dict | None:
    """Return a conservative, deterministic solid specification."""
    diagram_text = str(getattr(question, "diagram_spec", "") or "")
    tagged = _parse_solid3d_tag(diagram_text)
    if tagged:
        spec = dict(tagged)
        spec["type"] = str(spec.get("type", "")).strip().lower()
        spec["unit"] = str(spec.get("unit", "cm") or "cm").strip()
        for key in list(spec):
            if key not in {"type", "unit"}:
                num = _solid_num(spec[key])
                if num is not None:
                    spec[key] = num
        return spec

    # Fallback parser for natural-language generated questions.
    text = " ".join([
        str(getattr(question, "stem_text", "") or ""),
        str(getattr(question, "diagram_spec", "") or ""),
        " ".join(str(getattr(p, "prompt_text", "") or "") for p in (getattr(question, "parts", []) or [])),
    ])
    low = text.lower()
    unit_match = re.search(r"\b(mm|cm|m)\b", low)
    unit = unit_match.group(1) if unit_match else "cm"

    def first(patterns):
        for pattern in patterns:
            m = re.search(pattern, low, flags=re.I)
            if m:
                return _solid_num(m.group(1))
        return None

    # Composite cylinder on cuboid.
    if "cylinder" in low and "cuboid" in low:
        cuboid_dims = re.search(
            r"cuboid(?:\s+\w+){0,5}?\s*(\d+(?:\.\d+)?)\s*(?:cm|m|mm)?\s*(?:by|×|x)\s*"
            r"(\d+(?:\.\d+)?)\s*(?:cm|m|mm)?\s*(?:by|×|x)\s*(\d+(?:\.\d+)?)",
            low,
        )
        radius = first([
            r"cylinder(?:\s+\w+){0,5}?\s+radius\s*(?:is|=|of)?\s*(\d+(?:\.\d+)?)",
            r"radius\s*(?:is|=|of)?\s*(\d+(?:\.\d+)?)",
        ])
        cyl_height = first([
            r"cylinder(?:\s+\w+){0,6}?\s+height\s*(?:is|=|of)?\s*(\d+(?:\.\d+)?)",
            r"height\s*(?:is|=|of)?\s*(\d+(?:\.\d+)?)",
        ])
        if cuboid_dims and radius and cyl_height:
            return {
                "type": "composite_cylinder_cuboid",
                "cuboid_length": float(cuboid_dims.group(1)),
                "cuboid_depth": float(cuboid_dims.group(2)),
                "cuboid_height": float(cuboid_dims.group(3)),
                "cylinder_radius": radius,
                "cylinder_height": cyl_height,
                "unit": unit,
            }

    # Cuboid / cube.
    dims = re.search(
        r"(?:cuboid|block)(?:\s+\w+){0,6}?\s*(\d+(?:\.\d+)?)\s*(?:cm|m|mm)?\s*(?:by|×|x)\s*"
        r"(\d+(?:\.\d+)?)\s*(?:cm|m|mm)?\s*(?:by|×|x)\s*(\d+(?:\.\d+)?)",
        low,
    )
    if dims:
        return {
            "type": "cuboid",
            "length": float(dims.group(1)),
            "depth": float(dims.group(2)),
            "height": float(dims.group(3)),
            "unit": unit,
        }

    if "cube" in low:
        side = first([r"(?:side|edge)(?:\s+length)?\s*(?:is|=|of)?\s*(\d+(?:\.\d+)?)"])
        if side:
            return {"type": "cube", "side": side, "unit": unit}

    # Cylinder / cone / hemisphere / sphere.
    if "cylinder" in low:
        r = first([r"radius\s*(?:is|=|of)?\s*(\d+(?:\.\d+)?)"])
        h = first([r"height\s*(?:is|=|of)?\s*(\d+(?:\.\d+)?)"])
        if r and h:
            return {"type": "cylinder", "radius": r, "height": h, "unit": unit}

    if "triangular prism" in low or "prism" in low:
        base = first([
            r"(?:triangular\s+)?prism(?:\s+\w+){0,5}?\s+base\s*(?:is|=|of)?\s*(\d+(?:\.\d+)?)",
            r"base\s*(?:is|=|of)?\s*(\d+(?:\.\d+)?)",
        ])
        tri_h = first([
            r"triangle(?:\s+\w+){0,3}?\s+height\s*(?:is|=|of)?\s*(\d+(?:\.\d+)?)",
            r"perpendicular\s+height\s*(?:is|=|of)?\s*(\d+(?:\.\d+)?)",
        ])
        length = first([
            r"prism(?:\s+\w+){0,5}?\s+length\s*(?:is|=|of)?\s*(\d+(?:\.\d+)?)",
            r"length\s*(?:is|=|of)?\s*(\d+(?:\.\d+)?)",
        ])
        if base and tri_h and length:
            return {
                "type": "triangular_prism",
                "base": base,
                "triangle_height": tri_h,
                "length": length,
                "unit": unit,
            }

    if "cone" in low:
        r = first([r"radius\s*(?:is|=|of)?\s*(\d+(?:\.\d+)?)"])
        h = first([r"height\s*(?:is|=|of)?\s*(\d+(?:\.\d+)?)"])
        slant = first([r"slant(?:\s+height)?\s*(?:is|=|of)?\s*(\d+(?:\.\d+)?)"])
        if r and (h or slant):
            spec = {"type": "cone", "radius": r, "unit": unit}
            if h: spec["height"] = h
            if slant: spec["slant_height"] = slant
            return spec

    if "hemisphere" in low or "hemi-sphere" in low:
        r = first([r"radius\s*(?:is|=|of)?\s*(\d+(?:\.\d+)?)"])
        if r:
            return {"type": "hemisphere", "radius": r, "unit": unit}

    if "sphere" in low:
        r = first([r"radius\s*(?:is|=|of)?\s*(\d+(?:\.\d+)?)"])
        if r:
            return {"type": "sphere", "radius": r, "unit": unit}

    return None


def _solid_label(value, unit="cm") -> str:
    x = float(value)
    text = str(int(x)) if abs(x - round(x)) < 1e-9 else f"{x:g}"
    return f"{text} {unit}"


def _draw_dimension(draw, p1, p2, label, font, *, offset=(0,0), fill=(80,80,80)):
    x1,y1 = p1; x2,y2 = p2
    ox,oy = offset
    x1 += ox; y1 += oy; x2 += ox; y2 += oy
    draw.line((x1,y1,x2,y2), fill=fill, width=2)
    # arrowheads
    angle = math.atan2(y2-y1, x2-x1)
    for x,y,a in [(x1,y1,angle),(x2,y2,angle+math.pi)]:
        size=7
        pts=[
            (x, y),
            (x+size*math.cos(a+0.55), y+size*math.sin(a+0.55)),
            (x+size*math.cos(a-0.55), y+size*math.sin(a-0.55)),
        ]
        draw.line(pts+[pts[0]], fill=fill, width=2)
    mx=(x1+x2)/2; my=(y1+y2)/2
    bbox=draw.textbbox((0,0), label, font=font)
    tw=bbox[2]-bbox[0]; th=bbox[3]-bbox[1]
    draw.rectangle((mx-tw/2-4,my-th/2-3,mx+tw/2+4,my+th/2+3), fill="white")
    draw.text((mx-tw/2,my-th/2), label, fill=(40,40,40), font=font)


def _draw_cuboid_iso(draw, x, y, w, h, d, *, front=(217,233,249), side=(198,220,244), top=(232,242,252)):
    dx=d; dy=-int(d*0.62)
    draw.polygon([(x,y),(x+w,y),(x+w+dx,y+dy),(x+dx,y+dy)], fill=top, outline=(24,95,165))
    draw.polygon([(x+w,y),(x+w,y+h),(x+w+dx,y+h+dy),(x+w+dx,y+dy)], fill=side, outline=(24,95,165))
    draw.rectangle((x,y,x+w,y+h), fill=front, outline=(24,95,165))
    return {"front_tl":(x,y),"front_tr":(x+w,y),"front_br":(x+w,y+h),"front_bl":(x,y+h),
            "back_tr":(x+w+dx,y+dy),"back_br":(x+w+dx,y+h+dy)}


def _draw_cylinder(draw, cx, top_y, radius_px, height_px):
    ry=max(8,int(radius_px*0.34))
    left=cx-radius_px; right=cx+radius_px
    bottom_y=top_y+height_px
    draw.rectangle((left,top_y,right,bottom_y), fill=(220,245,237), outline=(15,110,86))
    draw.ellipse((left,top_y-ry,right,top_y+ry), fill=(225,248,240), outline=(15,110,86), width=2)
    draw.arc((left,bottom_y-ry,right,bottom_y+ry), 0, 180, fill=(15,110,86), width=2)
    draw.line((cx,top_y,right,top_y), fill=(15,110,86), width=2)
    draw.ellipse((cx-2,top_y-2,cx+2,top_y+2), fill=(15,110,86))
    return {"left":left,"right":right,"top_y":top_y,"bottom_y":bottom_y,"ry":ry,"cx":cx}


def render_dimensioned_solid_png(spec: dict, *, width=900, height=620) -> bytes:
    """Render a clean exam-style dimensioned 3D solid PNG."""
    img=Image.new("RGB",(width,height),"white")
    draw=ImageDraw.Draw(img)
    font=ImageFont.load_default()
    t=str(spec.get("type","")).lower()
    unit=str(spec.get("unit","cm"))

    if t == "composite_cylinder_cuboid":
        L=_solid_num(spec.get("cuboid_length"),12)
        D=_solid_num(spec.get("cuboid_depth"),8)
        H=_solid_num(spec.get("cuboid_height"),5)
        R=_solid_num(spec.get("cylinder_radius"),3)
        CH=_solid_num(spec.get("cylinder_height"),10)

        # Geometry closely follows the supplied SVG logic.
        cub=_draw_cuboid_iso(draw,220,330,300,130,80)
        cyl=_draw_cylinder(draw,370,145,65,180)

        _draw_dimension(draw,(220,485),(520,485),_solid_label(L,unit),font)
        _draw_dimension(draw,(205,330),(205,460),_solid_label(H,unit),font)
        _draw_dimension(draw,(535,475),(615,425),_solid_label(D,unit),font)
        _draw_dimension(draw,(455,145),(455,325),_solid_label(CH,unit),font)
        # radius leader
        draw.line((370,145,435,145),fill=(15,110,86),width=2)
        draw.line((410,145,610,105),fill=(125,125,125),width=1)
        draw.text((620,98),f"r = {_solid_label(R,unit)}",fill=(55,55,55),font=font)

    elif t in {"cuboid","cube"}:
        if t=="cube":
            L=D=H=_solid_num(spec.get("side"),5)
        else:
            L=_solid_num(spec.get("length"),10)
            D=_solid_num(spec.get("depth"),6)
            H=_solid_num(spec.get("height"),5)
        cub=_draw_cuboid_iso(draw,230,240,300,170,90)
        _draw_dimension(draw,(230,440),(530,440),_solid_label(L,unit),font)
        _draw_dimension(draw,(210,240),(210,410),_solid_label(H,unit),font)
        _draw_dimension(draw,(545,425),(635,370),_solid_label(D,unit),font)

    elif t == "cylinder":
        R=_solid_num(spec.get("radius"),3); H=_solid_num(spec.get("height"),8)
        c=_draw_cylinder(draw,390,150,90,250)
        _draw_dimension(draw,(500,150),(500,400),_solid_label(H,unit),font)
        draw.line((390,150,480,150),fill=(15,110,86),width=2)
        draw.text((410,120),f"r = {_solid_label(R,unit)}",fill=(55,55,55),font=font)

    elif t == "cone":
        R=_solid_num(spec.get("radius"),3)
        H=_solid_num(spec.get("height"),8)
        S=_solid_num(spec.get("slant_height"),None)
        cx=390; base_y=390; rx=100; ry=30; apex_y=120
        draw.ellipse((cx-rx,base_y-ry,cx+rx,base_y+ry),fill=(224,244,250),outline=(18,120,145),width=2)
        draw.polygon([(cx,apex_y),(cx-rx,base_y),(cx+rx,base_y)],fill=(244,218,120),outline=(155,105,0))
        draw.arc((cx-rx,base_y-ry,cx+rx,base_y+ry),0,180,fill=(18,120,145),width=2)
        draw.line((cx,base_y,cx+rx,base_y),fill=(18,120,145),width=2)
        draw.text((cx+30,base_y-28),f"r = {_solid_label(R,unit)}",fill=(55,55,55),font=font)
        if H:
            _draw_dimension(draw,(520,apex_y),(520,base_y),_solid_label(H,unit),font)
        if S:
            draw.text((470,250),f"l = {_solid_label(S,unit)}",fill=(55,55,55),font=font)

    elif t in {"sphere","hemisphere"}:
        R=_solid_num(spec.get("radius"),4)
        cx=390; cy=285; rpx=145
        if t=="sphere":
            draw.ellipse((cx-rpx,cy-rpx,cx+rpx,cy+rpx),fill=(230,243,251),outline=(26,112,168),width=2)
            draw.ellipse((cx-rpx,cy-40,cx+rpx,cy+40),outline=(110,160,195),width=1)
        else:
            draw.pieslice((cx-rpx,cy-rpx,cx+rpx,cy+rpx),180,360,fill=(218,246,244),outline=(25,150,150))
            draw.ellipse((cx-rpx,cy-35,cx+rpx,cy+35),fill=(230,251,249),outline=(25,150,150),width=2)
        draw.line((cx,cy,cx+rpx,cy),fill=(40,110,150),width=2)
        draw.ellipse((cx-2,cy-2,cx+2,cy+2),fill=(40,110,150))
        draw.text((cx+35,cy-25),f"r = {_solid_label(R,unit)}",fill=(55,55,55),font=font)

    elif t == "triangular_prism":
        base=_solid_num(spec.get("base"),6)
        tri_h=_solid_num(spec.get("triangle_height"),4)
        length=_solid_num(spec.get("length"),10)
        # Front and rear triangular faces.
        A=(240,390); B=(390,390); C=(315,230); dx=180; dy=-70
        A2=(A[0]+dx,A[1]+dy); B2=(B[0]+dx,B[1]+dy); C2=(C[0]+dx,C[1]+dy)
        for poly in [(A,B,C),(A2,B2,C2)]:
            draw.polygon(poly,fill=(239,246,252),outline=(30,80,120))
        for p,q in [(A,A2),(B,B2),(C,C2)]:
            draw.line((p,q),fill=(30,80,120),width=2)
        _draw_dimension(draw,(A[0],420),(B[0],420),_solid_label(base,unit),font)
        _draw_dimension(draw,(210,A[1]),(210,C[1]),_solid_label(tri_h,unit),font)
        _draw_dimension(draw,(B[0]+10,B[1]+25),(B2[0]+10,B2[1]+25),_solid_label(length,unit),font)

    elif t == "composite_prism_cuboid":
        L=_solid_num(spec.get("cuboid_length"),12)
        D=_solid_num(spec.get("cuboid_depth"),8)
        H=_solid_num(spec.get("cuboid_height"),5)
        PB=_solid_num(spec.get("prism_base"),6)
        PH=_solid_num(spec.get("prism_triangle_height"),4)
        PL=_solid_num(spec.get("prism_length"),8)

        # Cuboid base
        cub=_draw_cuboid_iso(draw,190,330,330,125,85)

        # Triangular prism mounted centrally on cuboid.
        A=(270,300); B=(430,300); C=(350,210); dx=105; dy=-48
        A2=(A[0]+dx,A[1]+dy); B2=(B[0]+dx,B[1]+dy); C2=(C[0]+dx,C[1]+dy)

        # visible faces
        draw.polygon([A,B,C],fill=(245,239,231),outline=(132,83,54))
        draw.polygon([B,B2,C2,C],fill=(230,218,205),outline=(132,83,54))
        draw.polygon([A,A2,C2,C],fill=(250,244,238),outline=(132,83,54))
        for p,q in [(A,A2),(B,B2),(C,C2)]:
            draw.line((p,q),fill=(132,83,54),width=2)

        _draw_dimension(draw,(190,485),(520,485),_solid_label(L,unit),font)
        _draw_dimension(draw,(175,330),(175,455),_solid_label(H,unit),font)
        _draw_dimension(draw,(535,470),(620,420),_solid_label(D,unit),font)

        _draw_dimension(draw,(A[0],315),(B[0],315),_solid_label(PB,unit),font)
        _draw_dimension(draw,(250,A[1]),(250,C[1]),_solid_label(PH,unit),font)
        _draw_dimension(draw,(B[0]+5,B[1]-5),(B2[0]+5,B2[1]-5),_solid_label(PL,unit),font)

    elif t == "composite_cylinder_prism":
        R=_solid_num(spec.get("cylinder_radius"),3)
        CH=_solid_num(spec.get("cylinder_height"),8)
        PB=_solid_num(spec.get("prism_base"),6)
        PH=_solid_num(spec.get("prism_triangle_height"),4)
        PL=_solid_num(spec.get("prism_length"),10)

        # Triangular prism base.
        A=(220,390); B=(430,390); C=(325,255); dx=150; dy=-65
        A2=(A[0]+dx,A[1]+dy); B2=(B[0]+dx,B[1]+dy); C2=(C[0]+dx,C[1]+dy)

        draw.polygon([A,B,C],fill=(245,239,231),outline=(132,83,54))
        draw.polygon([B,B2,C2,C],fill=(230,218,205),outline=(132,83,54))
        draw.polygon([A,A2,C2,C],fill=(250,244,238),outline=(132,83,54))
        for p,q in [(A,A2),(B,B2),(C,C2)]:
            draw.line((p,q),fill=(132,83,54),width=2)

        # Cylinder mounted on top rectangular face of prism.
        cyl=_draw_cylinder(draw,405,145,62,150)

        _draw_dimension(draw,(A[0],425),(B[0],425),_solid_label(PB,unit),font)
        _draw_dimension(draw,(195,A[1]),(195,C[1]),_solid_label(PH,unit),font)
        _draw_dimension(draw,(B[0]+8,B[1]+20),(B2[0]+8,B2[1]+20),_solid_label(PL,unit),font)
        _draw_dimension(draw,(485,145),(485,295),_solid_label(CH,unit),font)
        draw.line((405,145,467,145),fill=(15,110,86),width=2)
        draw.line((440,145,620,105),fill=(125,125,125),width=1)
        draw.text((625,98),f"r = {_solid_label(R,unit)}",fill=(55,55,55),font=font)

    elif t == "composite_cylinder_cuboid_prism":
        L=_solid_num(spec.get("cuboid_length"),14)
        D=_solid_num(spec.get("cuboid_depth"),9)
        H=_solid_num(spec.get("cuboid_height"),5)
        R=_solid_num(spec.get("cylinder_radius"),2.5)
        CH=_solid_num(spec.get("cylinder_height"),7)
        PB=_solid_num(spec.get("prism_base"),5)
        PH=_solid_num(spec.get("prism_triangle_height"),4)
        PL=_solid_num(spec.get("prism_length"),7)

        # Cuboid foundation.
        cub=_draw_cuboid_iso(draw,165,345,360,120,90)

        # Triangular prism mounted on the left half.
        A=(215,325); B=(365,325); C=(290,245); dx=90; dy=-42
        A2=(A[0]+dx,A[1]+dy); B2=(B[0]+dx,B[1]+dy); C2=(C[0]+dx,C[1]+dy)
        draw.polygon([A,B,C],fill=(245,239,231),outline=(132,83,54))
        draw.polygon([B,B2,C2,C],fill=(230,218,205),outline=(132,83,54))
        draw.polygon([A,A2,C2,C],fill=(250,244,238),outline=(132,83,54))
        for p,q in [(A,A2),(B,B2),(C,C2)]:
            draw.line((p,q),fill=(132,83,54),width=2)

        # Cylinder mounted on the right half.
        cyl=_draw_cylinder(draw,470,155,55,165)

        # Cuboid dimensions
        _draw_dimension(draw,(165,492),(525,492),_solid_label(L,unit),font)
        _draw_dimension(draw,(150,345),(150,465),_solid_label(H,unit),font)
        _draw_dimension(draw,(540,480),(630,425),_solid_label(D,unit),font)

        # Prism dimensions
        _draw_dimension(draw,(A[0],338),(B[0],338),_solid_label(PB,unit),font)
        _draw_dimension(draw,(195,A[1]),(195,C[1]),_solid_label(PH,unit),font)
        _draw_dimension(draw,(B[0]+3,B[1]-3),(B2[0]+3,B2[1]-3),_solid_label(PL,unit),font)

        # Cylinder dimensions
        _draw_dimension(draw,(540,155),(540,320),_solid_label(CH,unit),font)
        draw.line((470,155,525,155),fill=(15,110,86),width=2)
        draw.line((500,155,650,115),fill=(125,125,125),width=1)
        draw.text((655,108),f"r = {_solid_label(R,unit)}",fill=(55,55,55),font=font)

    elif t == "composite_cone_cylinder":
        R=_solid_num(spec.get("radius"),3)
        CH=_solid_num(spec.get("cylinder_height"),7)
        COH=_solid_num(spec.get("cone_height"),5)
        # cylinder
        c=_draw_cylinder(draw,390,230,85,210)
        # cone on top
        apex=(390,90); left=(305,230); right=(475,230)
        draw.polygon([apex,left,right],fill=(244,218,120),outline=(155,105,0))
        draw.ellipse((305,202,475,258),fill=(225,248,240),outline=(15,110,86),width=2)
        _draw_dimension(draw,(500,230),(500,440),_solid_label(CH,unit),font)
        _draw_dimension(draw,(535,90),(535,230),_solid_label(COH,unit),font)
        draw.line((390,230,475,230),fill=(15,110,86),width=2)
        draw.text((410,200),f"r = {_solid_label(R,unit)}",fill=(55,55,55),font=font)

    else:
        raise ValueError(f"Unsupported solid type: {t}")

    buf=BytesIO()
    img.save(buf,format="PNG")
    return buf.getvalue()


def show_dimensioned_solid(spec: dict, *, caption="") -> None:
    try:
        st.image(
            render_dimensioned_solid_png(spec),
            caption=caption or None,
            use_container_width=True,
        )
    except Exception as exc:
        st.warning("The dimensioned solid diagram could not be rendered: " + str(exc))


def add_dimensioned_solid_to_word(doc: Document, spec: dict, *, caption="") -> None:
    try:
        png=render_dimensioned_solid_png(spec,width=1100,height=720)
        p=doc.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(BytesIO(png),width=Cm(14.2))
        if caption:
            cp=doc.add_paragraph(caption)
            cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for rr in cp.runs:
                rr.italic=True
                rr.font.name="Times New Roman"
                rr.font.size=Pt(11)
    except Exception:
        pass


def _strip_solid3d_block(text: str) -> str:
    return _SOLID3D_TAG_RE.sub("", str(text or "")).strip()


def show_scene3d(scene, *, caption: str="") -> None:
    if scene is None: return
    try:
        st.image(render_exam_3d_wireframe_png(scene),caption=caption or None,use_container_width=True)
    except Exception as exc:
        st.info("The 3D diagram could not be rendered reliably. "+str(exc))


def add_scene3d_to_word(doc: Document, scene, *, caption: str="") -> None:
    if scene is None: return
    try:
        png=render_exam_3d_wireframe_png(scene,width=1100,height=700)
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(BytesIO(png),width=Cm(13.5))
        if caption:
            cp=doc.add_paragraph(caption); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for rr in cp.runs:
                rr.italic = True
                rr.font.name = "Times New Roman"
                rr.font.size = Pt(11)
                rr._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
                rr._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
                rr._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    except Exception:
        pass


def apply_word_tnr11(document: Document) -> None:
    """Force Times New Roman 11 pt across normal text and table content."""
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(11)
                        run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
                        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _is_worksheet_draft(draft) -> bool:
    return str(getattr(draft, "assessment_type", "") or "").strip().lower() == "worksheet"



_TABLE_TAG_RE = re.compile(
    r"(?is)\[TABLE\](.*?)\[/TABLE\]"
)

def _clean_table_cell(value: str) -> str:
    text = str(value or "").strip()
    text = re.sub(r"^\|+|\|+$", "", text).strip()
    return text


def _parse_generated_table_block(text: str) -> tuple[list[str], list[list[str]]] | None:
    """Parse markdown, pipe-delimited or [TABLE] generated question tables."""
    source = str(text or "").strip()
    if not source:
        return None

    tagged = _TABLE_TAG_RE.search(source)
    if tagged:
        source = tagged.group(1).strip()

    lines = [line.strip() for line in source.splitlines() if line.strip()]
    pipe_lines = [line for line in lines if "|" in line]

    if len(pipe_lines) < 2:
        # Semicolon form:
        # Class interval | Frequency; 10<m<=20 | 5; ...
        if ";" in source and "|" in source:
            pipe_lines = [x.strip() for x in source.split(";") if "|" in x]

    if len(pipe_lines) < 2:
        return None

    rows = []
    for line in pipe_lines:
        cells = [_clean_table_cell(x) for x in line.strip("|").split("|")]
        if len(cells) < 2:
            continue
        # Skip markdown separator rows.
        if all(re.fullmatch(r":?-{2,}:?", c.replace(" ", "")) for c in cells):
            continue
        rows.append(cells)

    if len(rows) < 2:
        return None

    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]

    header = rows[0]
    body = rows[1:]

    # A real table requires meaningful headings and at least one populated row.
    if not any(header) or not body:
        return None

    return header, body


def _question_table_payload(question) -> tuple[list[str], list[list[str]]] | None:
    """Find structured table content on any supported question field."""
    candidates = []

    # Future-proof: accept table-like fields if gemini_service adds them.
    for attr in (
        "table_text", "table_data", "data_table", "frequency_table",
        "stem_table", "table_spec",
    ):
        value = getattr(question, attr, None)
        if value:
            candidates.append(value)

    candidates.extend([
        getattr(question, "stem_text", ""),
        getattr(question, "diagram_spec", ""),
    ])

    # Dataclass/dict payloads.
    for value in candidates:
        if isinstance(value, dict):
            headers = value.get("headers") or value.get("columns") or []
            rows = value.get("rows") or value.get("data") or []
            if headers and rows:
                return [str(x) for x in headers], [[str(c) for c in row] for row in rows]
        elif isinstance(value, list) and value and isinstance(value[0], dict):
            headers = list(value[0].keys())
            return headers, [[str(row.get(h, "")) for h in headers] for row in value]
        else:
            parsed = _parse_generated_table_block(str(value or ""))
            if parsed is not None:
                return parsed

    return None


def _strip_generated_table_block(text: str) -> str:
    """Remove the structured table markup from prose after extracting it."""
    source = str(text or "")
    source = _TABLE_TAG_RE.sub("", source)

    lines = source.splitlines()
    # Remove contiguous markdown/pipe table lines.
    cleaned = []
    for line in lines:
        if "|" in line and len(line.strip("|").split("|")) >= 2:
            continue
        cleaned.append(line)
    return "\n".join(cleaned).strip()


def _looks_like_missing_frequency_table(question) -> bool:
    """Catch invalid questions that refer to a frequency table but contain no data."""
    text = " ".join([
        str(getattr(question, "stem_text", "") or ""),
        str(getattr(question, "diagram_spec", "") or ""),
    ]).lower()

    if not re.search(r"\b(frequency table|table below|table shows|results are summarised in the .*table)\b", text):
        return False

    if _question_table_payload(question) is not None:
        return False

    # Class intervals without frequencies are a common malformed generation.
    equations = [str(x) for x in (getattr(question, "stem_equations", []) or [])]
    interval_count = sum(
        1 for x in equations
        if re.search(r"(?:<|\\le|≤).*(?:<|\\le|≤)", x)
    )
    return interval_count >= 2


def _render_question_table_preview(question) -> bool:
    payload = _question_table_payload(question)
    if payload is None:
        return False

    headers, rows = payload
    df = pd.DataFrame(rows, columns=headers)
    st.dataframe(
        df,
        hide_index=True,
        use_container_width=False,
    )
    return True


def _add_question_table_to_word(doc: Document, question) -> bool:
    payload = _question_table_payload(question)
    if payload is None:
        return False

    headers, rows = payload
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    for cell, header in zip(table.rows[0].cells, headers):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(header))
        run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                append_word_mixed_math(p, str(value))
            except Exception:
                p.add_run(str(value))

    doc.add_paragraph()
    return True


_TABLE_GENERATION_REQUIREMENTS = """
IMPORTANT TABLE REQUIREMENTS FOR GENERATED ASSESSMENTS:
- If a question refers to a table, frequency table, grouped-frequency table, value table,
  probability table or data table, the complete table MUST be supplied.
- Never place table rows as separate stem_equations.
- Encode the table inside stem_text using this exact transport form:

  [TABLE]
  Class interval | Frequency
  10 < m <= 20 | 5
  20 < m <= 30 | 8
  ...
  [/TABLE]

- Include every value required to answer the question.
- For grouped-frequency questions, frequencies must be positive whole numbers and their
  total must agree exactly with the stated sample size.
- A question saying "the table below" without a complete [TABLE] block is invalid.
"""


def build_setter_question_paper_docx(draft: ExamPaperDraft) -> bytes:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.5); sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(1.8); sec.right_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    for heading_name, heading_size in [("Title", 15), ("Heading 1", 14), ("Heading 2", 12)]:
        if heading_name in styles:
            styles[heading_name].font.name = "Times New Roman"
            styles[heading_name].font.size = Pt(heading_size)
            styles[heading_name]._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
            styles[heading_name]._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
            styles[heading_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(draft.school_name or "School Mathematics Department"); r.bold = True; r.font.size = Pt(12)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(draft.paper_title); r.bold = True; r.font.size = Pt(15)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if _is_worksheet_draft(draft):
        p.add_run(f"{draft.track_label}    |    Worksheet")
    else:
        p.add_run(f"{draft.track_label}    |    {draft.duration_minutes} minutes    |    {draft.total_marks} marks")
    note = doc.add_paragraph(); note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = note.add_run("Mathematical expressions are editable using Microsoft Word Equation Editor.")
    rr.italic = True; rr.font.size = Pt(8.5)

    if draft.instructions:
        doc.add_heading("Instructions", level=2)
        for item in draft.instructions:
            ip = doc.add_paragraph(style="List Bullet")
            append_word_mixed_math(ip, item)

    doc.add_paragraph()
    figure_number = 0
    for question_index, q in enumerate(draft.questions):
        if question_index > 0:
            doc.add_page_break()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        r = p.add_run(f"{q.question_number}. "); r.bold = True
        append_word_mixed_math(p, _strip_solid3d_block(_strip_generated_table_block(q.stem_text)))
        word_table_rendered = _add_question_table_to_word(doc, q)
        if not word_table_rendered and _looks_like_missing_frequency_table(q):
            warning = doc.add_paragraph()
            wr = warning.add_run(
                "[INCOMPLETE GENERATED QUESTION: referenced frequency table data is missing]"
            )
            wr.bold = True
        for eq in q.stem_equations:
            if (
                word_table_rendered
                and re.search(r"(?:<|\\le|≤).*(?:<|\\le|≤)", str(eq or ""))
            ):
                continue
            append_word_math(doc.add_paragraph(), eq)

        context_image = _question_context_image(q)
        if context_image is not None:
            figure_number += 1
            add_context_image_to_word(
                doc,
                context_image,
                caption=f"Figure {figure_number}",
            )

        solid_spec = _dimensioned_solid_spec(q)
        if solid_spec is not None:
            figure_number += 1
            add_dimensioned_solid_to_word(
                doc,
                solid_spec,
                caption=f"Figure {figure_number}",
            )
        else:
            stats_spec = _stats_graph_spec(q)
        if solid_spec is None and stats_spec is not None:
            figure_number += 1
            stats_issues = validate_statistics_graph_spec(stats_spec)
            if not stats_issues:
                completed = bool(getattr(stats_spec, "show_completed_graph_in_question", True))
                add_statistics_graph_to_word(
                    doc,
                    stats_spec,
                    caption=f"Figure {figure_number}",
                    completed=completed,
                )
            else:
                note=doc.add_paragraph()
                rr=note.add_run("Statistics graph omitted pending correction: " + "; ".join(stats_issues))
                rr.italic=True
                rr.font.name="Times New Roman"
                rr.font.size=Pt(11)
        elif solid_spec is None:
            graph_ready_issues = validate_function_graph_readiness(q)
            graph_spec = _question_graph_spec(q)
            if graph_ready_issues:
                note=doc.add_paragraph()
                rr=note.add_run(
                    "Graph question requires regeneration: " + "; ".join(graph_ready_issues)
                )
                rr.italic=True
                rr.font.name="Times New Roman"
                rr.font.size=Pt(11)
            elif graph_spec is not None:
                figure_number += 1
                ggb_png = _captured_geogebra_png(q)
                if ggb_png:
                    add_png_to_word(
                        doc,
                        ggb_png,
                        caption=f"Figure {figure_number}",
                    )
                else:
                    # IMPORTANT: downloading must never wait for browser-side GeoGebra.
                    # Build the same curve directly from the exact equation.
                    fallback_png = render_function_graph_png(q, width=1100, height=650)
                    if fallback_png:
                        add_png_to_word(
                            doc,
                            fallback_png,
                            caption=f"Figure {figure_number}",
                        )
                    else:
                        note = doc.add_paragraph()
                        rr = note.add_run(
                            "Function graph omitted because neither GeoGebra nor the deterministic "
                            "equation renderer could construct the graph reliably."
                        )
                        rr.italic = True
                        rr.font.name = "Times New Roman"
                        rr.font.size = Pt(11)
            elif getattr(q, "diagram_scene_3d", None) is not None:
                figure_number += 1
                add_scene3d_to_word(
                    doc,
                    q.diagram_scene_3d,
                    caption=f"Figure {figure_number}",
                )
            elif getattr(q, "diagram_scene_2d", None) is not None:
                figure_number += 1
                effective_scene_2d = ensure_question_function_curve(q)
                scene_issues = validate_question_scene_2d(q, effective_scene_2d)
                if not scene_issues:
                    add_scene2d_to_word(
                        doc,
                        effective_scene_2d,
                        caption=f"Figure {figure_number}",
                    )
                else:
                    note = doc.add_paragraph()
                    rr = note.add_run("Diagram omitted pending correction: " + "; ".join(scene_issues))
                    rr.italic = True
                    rr.font.name = "Times New Roman"
                    rr.font.size = Pt(11)
            elif q.diagram_spec:
                box = doc.add_paragraph()
                rr = box.add_run("Diagram / figure specification: ")
                rr.italic = True
                append_word_mixed_math(box, q.diagram_spec)

        for part in q.parts:
            pp = doc.add_paragraph()
            pp.paragraph_format.left_indent = Cm(0.5)
            if part.label:
                rr = pp.add_run(part.label + " "); rr.bold = True
            append_word_mixed_math(pp, part.prompt_text)
            for eq in part.equations:
                ep = doc.add_paragraph(); ep.paragraph_format.left_indent = Cm(1.0)
                append_word_math(ep, eq)
            if not _is_worksheet_draft(draft):
                markp = doc.add_paragraph(); markp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                markp.add_run(f"[{part.marks}]")
            for _ in range(max(1, min(part.answer_space_lines, 12))):
                doc.add_paragraph(" ")

    buf = BytesIO(); doc.save(buf); return buf.getvalue()


def _set_cell_shading(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def build_setter_marking_scheme_docx(draft: ExamPaperDraft) -> bytes:
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"; doc.styles["Normal"].font.size = Pt(9.5)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run((draft.school_name or "School Mathematics Department") + "\n"); r.bold = True
    guide_suffix = " - Answer Guide" if _is_worksheet_draft(draft) else " - Marking Scheme"
    r = p.add_run(draft.paper_title + guide_suffix); r.bold = True; r.font.size = Pt(14)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if _is_worksheet_draft(draft):
        p.add_run(f"{draft.track_label} | Worksheet")
    else:
        p.add_run(f"{draft.track_label} | Total: {draft.total_marks} marks")
    eqnote = doc.add_paragraph(); eqnote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = eqnote.add_run("Mathematical expressions are editable using Microsoft Word Equation Editor.")
    rr.italic = True; rr.font.size = Pt(8.5)
    doc.add_paragraph("AI-generated teacher draft. Recheck against departmental/official marking conventions before formal use.")

    figure_number = 0
    for question_index, q in enumerate(draft.questions):
        if question_index > 0:
            doc.add_page_break()
        doc.add_heading(f"Question {q.question_number}", level=2)
        if getattr(q, "diagram_scene_3d", None) is not None:
            figure_number += 1
            add_scene3d_to_word(doc, q.diagram_scene_3d, caption=f"Figure {figure_number}")
        elif getattr(q, "diagram_scene_2d", None) is not None:
            figure_number += 1
            add_scene2d_to_word(doc, q.diagram_scene_2d, caption=f"Figure {figure_number}")
        if _is_worksheet_draft(draft):
            for part in q.parts:
                label = part.label or "Answer"
                pp = doc.add_paragraph()
                rr = pp.add_run(label + " "); rr.bold = True
                if part.final_answer_mathio:
                    append_word_math(pp, part.final_answer_mathio)
                for step_no, step in enumerate(part.solution_steps, 1):
                    sp = doc.add_paragraph()
                    sr = sp.add_run(f"Step {step_no}: "); sr.bold = True
                    append_word_mixed_math(sp, step.explanation)
                    for eq in step.equations:
                        append_word_math(doc.add_paragraph(), eq)
            continue

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ["Answer", "Marks", "Partial Marks", "Guidance"]
        for cell, text in zip(table.rows[0].cells, headers):
            _set_cell_shading(cell, "D9D9D9")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            pr = cell.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = pr.add_run(text); rr.bold = True

        for part in q.parts:
            rows = part.marking_points or [PaperMarkPoint(code="", marks=part.marks, description="Correct complete solution", allow_follow_through=False)]
            for i, mp in enumerate(rows):
                cells = table.add_row().cells
                cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if i == 0:
                    if part.final_answer_mathio:
                        append_word_math(cells[0].paragraphs[0], part.final_answer_mathio)
                    else:
                        cells[0].paragraphs[0].add_run(part.label or "Answer")
                cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells[1].paragraphs[0].add_run(str(mp.marks))
                cells[2].paragraphs[0].add_run(mp.code)
                guidance = mp.description + ("; ft allowed" if mp.allow_follow_through else "")
                append_word_mixed_math(cells[3].paragraphs[0], guidance)

    buf = BytesIO(); doc.save(buf); return buf.getvalue()


def audit_generated_solids(draft) -> list[str]:
    issues = []
    for q in list(getattr(draft, "questions", []) or []):
        text = " ".join([
            str(getattr(q, "stem_text", "") or ""),
            str(getattr(q, "diagram_spec", "") or ""),
        ]).lower()
        if re.search(r"\b(cuboid|cube|cylinder|cone|sphere|hemisphere|triangular prism|prism|composite solid)\b", text):
            if re.search(r"\b(volume|surface area|solid|mounted|formed by|prism)\b", text):
                spec = _dimensioned_solid_spec(q)
                if spec is None:
                    issues.append(
                        f"Question {getattr(q, 'question_number', '?')}: "
                        "3D solid question has no reliable dimensioned diagram specification."
                    )
                    continue

                t = str(spec.get("type", "")).lower()

                if t in {"composite_cylinder_cuboid", "composite_cylinder_cuboid_prism"}:
                    r = _solid_num(spec.get("cylinder_radius"))
                    L = _solid_num(spec.get("cuboid_length"))
                    D = _solid_num(spec.get("cuboid_depth"))
                    if r and L and D and 2*r > min(L, D):
                        issues.append(
                            f"Question {getattr(q, 'question_number', '?')}: "
                            "cylinder diameter is larger than the supporting cuboid face."
                        )

                if t in {"composite_prism_cuboid", "composite_cylinder_cuboid_prism"}:
                    pb = _solid_num(spec.get("prism_base"))
                    pl = _solid_num(spec.get("prism_length"))
                    L = _solid_num(spec.get("cuboid_length"))
                    D = _solid_num(spec.get("cuboid_depth"))
                    if pb and pl and L and D and (pb > L or pl > max(L, D)):
                        issues.append(
                            f"Question {getattr(q, 'question_number', '?')}: "
                            "triangular prism footprint is larger than the supporting cuboid face."
                        )
    return issues



def audit_generated_tables(draft) -> list[str]:
    issues = []
    for q in list(getattr(draft, "questions", []) or []):
        if _looks_like_missing_frequency_table(q):
            issues.append(
                f"Question {getattr(q, 'question_number', '?')}: "
                "question refers to a frequency table but does not contain complete table data."
            )
    return issues


def audit_generated_graphs(draft) -> list[str]:
    issues=[]
    for q in list(getattr(draft,"questions",[]) or []):
        qn=str(getattr(q,"question_number","?"))
        for issue in validate_function_graph_readiness(q):
            issues.append(f"Question {qn}: {issue}")
        stats_spec=_stats_graph_spec(q)
        if stats_spec is not None:
            for issue in validate_statistics_graph_spec(stats_spec):
                issues.append(f"Question {qn}: {issue}")
    return issues


def render_setter_preview(draft: ExamPaperDraft) -> None:
    """Render the generated assessment paper with MathIO for all mathematics."""
    st.markdown("### Generated paper preview")
    if _is_worksheet_draft(draft):
        c1, c2 = st.columns(2)
        c1.metric("Assessment", "Worksheet")
        c2.metric("Questions", len(draft.questions))
    else:
        c1, c2, c3 = st.columns(3)
        c1.metric("Total marks", draft.total_marks)
        c2.metric("Questions", len(draft.questions))
        c3.metric("Duration", f"{draft.duration_minutes} min")

    rows = []
    for q in draft.questions:
        row = {
            "Q": q.question_number,
            "Topic": q.topic,
            "AO": q.ao,
            "Difficulty": q.difficulty,
        }
        if not _is_worksheet_draft(draft):
            row["Marks"] = q.marks
        rows.append(row)
    if rows:
        st.dataframe(pd.DataFrame(rows), hide_index=True, use_container_width=True)

    st.markdown("---")
    st.markdown("## Question paper")

    figure_number = 0
    for q in draft.questions:
        with st.container(border=True):
            if _is_worksheet_draft(draft):
                st.markdown(f"### Question {q.question_number}")
            else:
                st.markdown(
                    f"### Question {q.question_number} "
                    f"<span style='font-size:.78em;font-weight:500'>[{q.marks} marks]</span>",
                    unsafe_allow_html=True,
                )

            geogebra_external_tools(
                question_text=" ".join(
                    [_normalise_unit_braces(str(q.stem_text or ""))] +
                    [str(x) for x in (getattr(q, "stem_equations", []) or [])] +
                    [str(getattr(q, "diagram_spec", "") or "")]
                ),
                key_base=f"setter_geogebra_external_{q.question_number}",
            )

            # Stem prose can itself contain mathematical expressions, so use the
            # MathIO-aware mixed renderer rather than st.write().
            if str(q.stem_text or "").strip():
                preview_stem = _normalise_unit_braces(
                    _strip_solid3d_block(_strip_generated_table_block(str(q.stem_text or "")))
                )
                preview_stem = re.sub(r"(?<!\\)\bpi\b", r"\\pi", preview_stem, flags=re.IGNORECASE)
                preview_stem = re.sub(r"(?<!\\)\btheta\b", r"\\theta", preview_stem, flags=re.IGNORECASE)
                if preview_stem.strip():
                    render_mathio_mixed(preview_stem)

            table_rendered = _render_question_table_preview(q)
            if not table_rendered and _looks_like_missing_frequency_table(q):
                st.error(
                    "This generated question refers to a frequency table but the frequency "
                    "data is missing. Regenerate the paper before using or downloading it."
                )

            context_image = _question_context_image(q)
            if context_image is not None:
                st.image(context_image["bytes"], width=320)
                st.caption(_context_image_caption(context_image))

            # Explicit equation fields always render through MathIO.
            # If a structured table exists, do not duplicate class-interval rows here.
            for eq in q.stem_equations:
                eq_text = re.sub(r"\\+(?:dots|ldots|cdots)\b", "", str(eq or ""))
                eq_text = re.sub(r"\.{3,}", "", eq_text)
                eq_text = re.sub(r"\s{2,}", " ", eq_text).strip()
                if (
                    table_rendered
                    and re.search(r"(?:<|\\le|≤).*(?:<|\\le|≤)", eq_text)
                ):
                    continue
                if eq_text:
                    render_mathio(eq_text)

            solid_spec = _dimensioned_solid_spec(q)
            if solid_spec is not None:
                figure_number += 1
                show_dimensioned_solid(
                    solid_spec,
                    caption=f"Figure {figure_number}",
                )
            else:
                stats_spec = _stats_graph_spec(q)
            if solid_spec is None and stats_spec is not None:
                figure_number += 1
                stats_issues = validate_statistics_graph_spec(stats_spec)
                if stats_issues:
                    st.warning("Statistics graph withheld: " + "; ".join(stats_issues))
                else:
                    completed = bool(getattr(stats_spec, "show_completed_graph_in_question", True))
                    show_statistics_graph(
                        stats_spec,
                        caption=f"Figure {figure_number}",
                        completed=completed,
                    )
            elif solid_spec is None:
                graph_ready_issues = validate_function_graph_readiness(q)
                graph_spec = _question_graph_spec(q)
                if graph_ready_issues:
                    st.error(
                        "Function graph question is invalid and must be regenerated: "
                        + "; ".join(graph_ready_issues)
                    )
                elif graph_spec is not None:
                    figure_number += 1
                    geogebra_png = render_geogebra_question_graph(
                        q,
                        figure_caption=f"Figure {figure_number}",
                    )
                    if geogebra_png is None:
                        # Show the equation-driven graph immediately while GeoGebra capture
                        # is still pending. The downloaded paper uses this same fallback.
                        effective_scene_2d = build_function_graph_scene(q)
                        if effective_scene_2d is not None:
                            show_scene2d(
                                effective_scene_2d,
                                caption=f"Figure {figure_number} · deterministic graph",
                            )
                        else:
                            st.warning("The function could not be rendered by either graph engine.")
                elif getattr(q, "diagram_scene_3d", None) is not None:
                    figure_number += 1
                    show_scene3d(
                        q.diagram_scene_3d,
                        caption=f"Figure {figure_number}",
                    )
                elif getattr(q, "diagram_scene_2d", None) is not None:
                    figure_number += 1
                    effective_scene_2d = ensure_question_function_curve(q)
                    scene_issues = validate_question_scene_2d(q, effective_scene_2d)
                    if scene_issues:
                        st.warning(
                            "Diagram withheld because it does not yet match the question: "
                            + "; ".join(scene_issues)
                        )
                    else:
                        show_scene2d(
                            effective_scene_2d,
                            caption=f"Figure {figure_number}",
                        )
                elif q.diagram_spec:
                    with st.expander("Diagram / figure information", expanded=False):
                        render_mathio_mixed(_strip_solid3d_block(q.diagram_spec))

            for part in q.parts:
                label = part.label or "Question"
                if _is_worksheet_draft(draft):
                    st.markdown(f"#### {label}")
                else:
                    st.markdown(f"#### {label} [{part.marks} marks]")

                if str(part.prompt_text or "").strip():
                    preview_prompt = re.sub(r"(?<!\\)\bpi\b", r"\\pi", part.prompt_text, flags=re.IGNORECASE)
                    preview_prompt = re.sub(r"(?<!\\)\btheta\b", r"\\theta", preview_prompt, flags=re.IGNORECASE)
                    render_mathio_mixed(preview_prompt)

                for eq in part.equations:
                    eq_text = re.sub(r"\\+(?:dots|ldots|cdots)\b", "", str(eq or ""))
                    eq_text = re.sub(r"\.{3,}", "", eq_text)
                    eq_text = re.sub(r"\s{2,}", " ", eq_text).strip()
                    if eq_text:
                        render_mathio(eq_text)

                # Give the preview some visual answer space without excessive blank area.
                if getattr(part, "answer_space_lines", 0):
                    st.caption("Answer space included in downloaded paper.")


def uploaded_assets(files: list[Any] | None) -> list[UploadedAsset]:
    files = files or []
    assets: list[UploadedAsset] = []
    total = 0
    for f in files:
        data = f.getvalue()
        total += len(data)
        if len(data) > MAX_FILE_BYTES:
            raise GeminiTutorError(f"{f.name} is larger than the app's 12 MB per-file limit.", category="input")
        mime = f.type or "application/octet-stream"
        assets.append(UploadedAsset(name=f.name, mime_type=mime, data=data))
    if total > MAX_TOTAL_BYTES:
        raise GeminiTutorError("Uploads exceed the app's 30 MB total limit.", category="input")
    return assets


def question_file_signature(files: list[Any] | None) -> str:
    files = files or []
    return "|".join(f"{getattr(f, 'name', '')}:{getattr(f, 'size', 0)}:{getattr(f, 'type', '')}" for f in files)


def detected_question_context(detection: QuestionDetectionResult, index: int) -> str:
    if not detection.questions:
        return ""
    index = max(0, min(index, len(detection.questions) - 1))
    q = detection.questions[index]
    lines = [
        "[SELECTED QUESTION FROM UPLOADED WORKSHEET]",
        f"Main question number: {q.question_number}",
        f"Main question stem: {q.question_text}",
        f"Likely topic: {q.topic_hint}",
    ]
    if q.subparts:
        lines.append("Subparts:")
        for part in q.subparts:
            lines.append(f"- {part.label}: {part.question_text}")
    lines.append("IMPORTANT: Analyse this selected main question only. Ignore other unrelated questions visible in the uploaded source.")
    return "\n".join(lines)


def render_question_detection(detection: QuestionDetectionResult) -> int:
    subpart_count = sum(len(q.subparts) for q in detection.questions)
    st.success(
        f"Detected {detection.main_question_count} confirmed main question(s)"
        + (f" with {subpart_count} subpart(s)." if subpart_count else ".")
    )
    if detection.possible_additional_question_count:
        st.warning(
            f"Gemini also saw {detection.possible_additional_question_count} possible additional question(s) that were too cropped or unclear to confirm."
        )
    st.caption(f"Detection confidence: {detection.overall_confidence.title()}")
    for note in detection.notes:
        st.caption(f"• {note}")

    if not detection.questions:
        st.info("No confirmed main question could be extracted. Try a clearer image or smaller crop.")
        return 0

    options = list(range(len(detection.questions)))
    current = min(int(st.session_state.ai_selected_question_index), len(options) - 1)
    selected = st.selectbox(
        "Choose the main question to analyse",
        options,
        index=current,
        format_func=lambda i: (
            f"Question {detection.questions[i].question_number} · {detection.questions[i].topic_hint}"
            + (f" · {len(detection.questions[i].subparts)} subpart(s)" if detection.questions[i].subparts else "")
        ),
        key="ai_detected_question_selector",
    )
    st.session_state.ai_selected_question_index = int(selected)
    q = detection.questions[selected]
    with st.expander("Review detected question text", expanded=True):
        render_math_text(f"**Question {q.question_number}:** {q.question_text}")
        for part in q.subparts:
            render_math_text(f"**{part.label}** {part.question_text}")
        st.caption(f"Topic hint: {q.topic_hint} · Confidence: {q.confidence}")
    return int(selected)


def question_for_selected_analysis(
    typed_text: str,
    detection: QuestionDetectionResult | None,
    selected_index: int,
) -> str:
    parts = [typed_text.strip()]
    if detection is not None and detection.questions:
        parts.append(detected_question_context(detection, selected_index))
    return "\n\n".join(part for part in parts if part)


def question_feasibility_signature(question_text: str, files: list[Any] | None, selected_index: int) -> str:
    return f"{question_text.strip()}||{question_file_signature(files)}||selected={selected_index}"


def _question_source_image(file_obj: Any, page_number: int = 1) -> Image.Image | None:
    """Load an uploaded question image or rasterize one PDF page for visual callouts."""
    if file_obj is None:
        return None
    try:
        data = file_obj.getvalue() if hasattr(file_obj, "getvalue") else bytes(file_obj.read())
        name = str(getattr(file_obj, "name", "")).lower()
        mime = str(getattr(file_obj, "type", "")).lower()
        if name.endswith(".pdf") or mime == "application/pdf":
            import fitz

            doc = fitz.open(stream=data, filetype="pdf")
            if doc.page_count < 1:
                return None
            page_index = max(0, min(int(page_number) - 1, doc.page_count - 1))
            page = doc.load_page(page_index)
            pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
            image = Image.open(BytesIO(pix.tobytes("png"))).convert("RGB")
            doc.close()
            return image
        return Image.open(BytesIO(data)).convert("RGB")
    except Exception:
        return None


def _normalized_box_to_pixels(box: list[int], width: int, height: int) -> tuple[int, int, int, int] | None:
    if len(box) != 4:
        return None
    try:
        ymin, xmin, ymax, xmax = [max(0, min(1000, int(v))) for v in box]
    except (TypeError, ValueError):
        return None
    if ymax <= ymin or xmax <= xmin:
        return None
    x1 = int(xmin / 1000 * width)
    y1 = int(ymin / 1000 * height)
    x2 = int(xmax / 1000 * width)
    y2 = int(ymax / 1000 * height)
    return x1, y1, x2, y2


def _annotate_issue_regions(image: Image.Image, callouts: list[tuple[int, Any, str]]) -> Image.Image:
    """Draw numbered issue callouts on a copy of the original question diagram/page."""
    annotated = image.copy()
    draw = ImageDraw.Draw(annotated)
    font = ImageFont.load_default()
    width, height = annotated.size
    line_width = max(3, round(min(width, height) * 0.006))

    for issue_number, region, severity in callouts:
        box = _normalized_box_to_pixels(list(getattr(region, "box_2d", []) or []), width, height)
        if box is None:
            continue
        x1, y1, x2, y2 = box
        outline = (190, 30, 45) if severity == "blocking" else (175, 115, 0)
        draw.rectangle((x1, y1, x2, y2), outline=outline, width=line_width)

        badge = f"{issue_number}"
        left = max(0, x1)
        top = max(0, y1 - 24)
        try:
            bbox = draw.textbbox((left, top), badge, font=font)
            tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
        except Exception:
            tw, th = 10, 12
        pad = 4
        draw.rectangle((left, top, left + tw + 2 * pad, top + th + 2 * pad), fill=outline)
        draw.text((left + pad, top + pad), badge, fill=(255, 255, 255), font=font)
    return annotated


def render_feasibility_visual_map(result: QuestionFeasibilityResult, question_files: list[Any] | None) -> None:
    """Show the original diagram/page with numbered highlights matching feasibility issues."""
    if not question_files:
        return

    grouped: dict[tuple[int, int], list[tuple[int, Any, str]]] = {}
    for issue_number, issue in enumerate(result.issues, 1):
        for region in list(getattr(issue, "visual_regions", []) or []):
            source_index = int(getattr(region, "source_index", 0) or 0)
            page_number = int(getattr(region, "page_number", 1) or 1)
            if source_index < 1 or source_index > len(question_files):
                continue
            grouped.setdefault((source_index, page_number), []).append((issue_number, region, issue.severity))

    if not grouped:
        return

    st.markdown("#### Visual issue map")
    st.caption("Numbered highlights show the parts of the original question diagram/page that support each issue or warning below.")
    for (source_index, page_number), callouts in sorted(grouped.items()):
        file_obj = question_files[source_index - 1]
        image = _question_source_image(file_obj, page_number)
        if image is None:
            continue
        annotated = _annotate_issue_regions(image, callouts)
        name = str(getattr(file_obj, "name", f"Question source {source_index}"))
        page_note = f" · page {page_number}" if name.lower().endswith(".pdf") else ""
        st.image(annotated, caption=f"{name}{page_note}", use_container_width=True)
        labels = []
        for issue_number, region, _severity in callouts:
            label = str(getattr(region, "label", "")).strip()
            if label:
                labels.append(f"{issue_number}: {label}")
        if labels:
            st.caption(" · ".join(labels))


def render_question_feasibility(result: QuestionFeasibilityResult, question_files: list[Any] | None = None) -> None:
    labels = {
        "feasible": "Ready to analyse",
        "feasible_with_caveats": "Ready with caveats",
        "needs_clarification": "Clarification needed",
        "infeasible": "Question issue detected",
    }
    message = labels.get(result.status, result.status.replace("_", " ").title())
    if result.status == "feasible":
        st.markdown(f'<div class="omt-success-card"><strong>✓ {message}</strong><br>The question is sufficiently clear and consistent for reasoning analysis.</div>', unsafe_allow_html=True)
    elif result.status == "feasible_with_caveats":
        st.warning(f"{message} — review the note below before marking.")
    else:
        st.error(f"{message} — student-working analysis stays locked until the question is clarified or corrected.")

    with st.container(border=True):
        st.caption("Question interpreted as")
        render_math_text(result.interpreted_question)
        c1, c2, c3 = st.columns(3)
        c1.metric("Answerability", result.answerability.replace("_", " ").title())
        c2.metric("Information", "Complete" if result.required_information_present else "Needs attention")
        c3.metric("Diagram / table", "Sufficient" if result.diagram_or_table_sufficient else "Needs attention")
        st.caption(
            f"Syllabus fit · {result.syllabus_fit.replace('_', ' ').title()}    |    "
            f"Confidence · {result.confidence.title()}"
        )

    render_feasibility_visual_map(result, question_files)

    if result.issues:
        st.markdown("### Issues to check")
        for issue_number, issue in enumerate(result.issues, 1):
            label = "Blocking issue" if issue.severity == "blocking" else "Warning"
            visual_note = f" · diagram {issue_number}" if list(getattr(issue, "visual_regions", []) or []) else ""
            with st.container(border=True):
                st.markdown(f"**{label}{visual_note}**")
                render_mathio_mixed(issue.description)
                if issue.suggested_fix:
                    st.caption("Suggested fix")
                    render_math_text(issue.suggested_fix)

    if result.suspected_corrections:
        with st.expander("Possible corrections to verify"):
            for item in result.suspected_corrections:
                render_math_text(f"• {item}")

    if result.action_needed:
        st.markdown("**Next action**")
        render_math_text(result.action_needed)


def offline_evidence_for(question_text: str, working_text: str) -> tuple[str, AttemptResult | None]:
    if is_additional_math_track(track_label):
        return "No deterministic offline algebra evidence is used for this Additional Mathematics track; Gemini code execution performs independent verification.", None
    if not question_text.strip() or not working_text.strip():
        return "", None
    try:
        result = analyze_own_algebra_question(question_text, working_text)
    except ValueError:
        return "", None
    evidence = (
        f"Offline algebra checker says is_correct={result.is_correct}; "
        f"first_logic_break={result.first_logic_break}; "
        f"explanation={result.first_logic_break_explanation}; "
        f"answer_score={result.answer_score}; reasoning_score={result.reasoning_score}."
    )
    return evidence, result


# ---------- Sidebar ----------
with st.sidebar:
    st.markdown(
        """
        <div class="omt-side-brand">
          <div class="title">✦ SG Math Tutor</div>
          <div class="sub">Reasoning-first support for 2027 SEC G1/G2/G3 Mathematics and Additional Mathematics, with 2026 O/N-Level transition support.</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    track_label = st.selectbox("Exam track", list(APP_TRACKS.keys()), index=0)
    track_info = selected_track_info(track_label)
    tcode = track_code(track_label)
    if track_info["year"] == 2027:
        syllabus_name = (
            f'2027 SEC · {track_info["level"]} {track_info["subject"]} · {track_info["subject_code"]}'
        )
    else:
        syllabus_name = f'{track_info["level"]} {track_info["subject"]} · {track_info["subject_code"]}'
    st.markdown(f'<div class="omt-status-pill neutral">📘 <span>{syllabus_name}</span></div>', unsafe_allow_html=True)
    if track_info["year"] == 2027:
        st.caption(
            "Singapore-Cambridge Secondary Education Certificate (SEC) · "
            + " · ".join(track_info["strands"])
        )

    with st.expander("⚙️ Gemini connection", expanded=False):
        explicit_key = st.text_input(
            "Gemini API key (optional here)",
            type="password",
            help="Prefer Streamlit Community Cloud Secrets with the name GEMINI_API_KEY.",
        )
        has_key = bool(get_api_key(explicit_key))
        model = st.selectbox(
            "Gemini model",
            ["gemini-3.5-flash-lite", "gemini-3.1-flash-lite"],
            index=0,
            help="Free-tier availability and quotas depend on the Google account/project.",
        )
        analysis_speed = st.radio(
            "Analysis mode",
            ["Fast", "Full"],
            horizontal=True,
            help="Fast keeps mathematical verification but defers the optional visual-explanation API call until you request it. Full builds visuals automatically when appropriate.",
        )
    if has_key:
        st.markdown('<div class="omt-status-pill good">● <span>Gemini online</span></div>', unsafe_allow_html=True)
    else:
        st.markdown('<div class="omt-status-pill neutral">○ <span>Offline tools available</span></div>', unsafe_allow_html=True)

    with st.expander("Privacy & data", expanded=False):
        st.caption(
            "Online analysis sends only the selected question/work to Gemini. Remove names, NRICs and unnecessary identifiers. "
            "Offline practice do not call Gemini."
        )

    if st.button("↻ Reset learning session", use_container_width=True):
        for key in list(st.session_state.keys()):
            if key not in {"session_id"}:
                del st.session_state[key]
        st.rerun()



def inject_pwa_home_screen_metadata() -> None:
    """Inject iPad Home Screen / PWA metadata without blocking app startup."""
    try:
        st_components_v1.html(
            """
            <script>
            (function () {
              try {
                const doc = window.parent.document;
                const head = doc.head;

                function setMeta(name, content) {
                  let el = head.querySelector(`meta[name="${name}"]`);
                  if (!el) {
                    el = doc.createElement('meta');
                    el.setAttribute('name', name);
                    head.appendChild(el);
                  }
                  el.setAttribute('content', content);
                }

                function setLink(rel, href, attrs) {
                  let el = head.querySelector(`link[rel="${rel}"][href="${href}"]`);
                  if (!el) {
                    el = doc.createElement('link');
                    el.setAttribute('rel', rel);
                    el.setAttribute('href', href);
                    Object.entries(attrs || {}).forEach(([k,v]) => el.setAttribute(k,v));
                    head.appendChild(el);
                  }
                }

                setMeta('apple-mobile-web-app-capable', 'yes');
                setMeta('apple-mobile-web-app-status-bar-style', 'default');
                setMeta('apple-mobile-web-app-title', 'Math Buddy');
                setMeta('mobile-web-app-capable', 'yes');
                setMeta('theme-color', '#4b63e6');

                setLink(
                  'apple-touch-icon',
                  '/app/static/math-buddy-180.png',
                  {sizes:'180x180'}
                );
                setLink(
                  'icon',
                  '/app/static/math-buddy-192.png',
                  {sizes:'192x192', type:'image/png'}
                );
                setLink('manifest', '/app/static/manifest.webmanifest');

                doc.title = 'Math Buddy';

                const standalone =
                  (window.parent.matchMedia &&
                   window.parent.matchMedia('(display-mode: standalone)').matches) ||
                  window.parent.navigator.standalone === true;

                if (standalone) {
                  doc.documentElement.classList.add('math-buddy-standalone');
                }
              } catch (err) {
                console.debug('Math Buddy Home Screen metadata skipped:', err);
              }
            })();
            </script>
            """,
            height=0,
            width=0,
        )
    except Exception:
        # PWA metadata must never prevent Math Buddy from loading.
        return


inject_pwa_home_screen_metadata()

st.markdown(
    """
    <style>
    /* iPad / Home Screen refinements */
    @media (pointer: coarse) {
      .stButton > button,
      [data-testid="stSegmentedControl"] button,
      [data-baseweb="select"] > div {
        min-height: 44px !important;
      }

      textarea,
      input {
        font-size: 16px !important;
      }
    }

    @media (max-width: 900px) {
      .block-container {
        padding-left: 1rem !important;
        padding-right: 1rem !important;
        padding-top: max(1rem, env(safe-area-inset-top)) !important;
        padding-bottom: max(1.25rem, env(safe-area-inset-bottom)) !important;
      }

      .omt-hero {
        border-radius: 18px !important;
      }
    }

    html.math-buddy-standalone body {
      overscroll-behavior-y: none;
    }
    </style>
    """,
    unsafe_allow_html=True,
)


role_mode = st.segmented_control(
    "Choose workspace",
    ["For Teacher", "For Student"],
    default="For Student",
    key="math_advisor_role_mode",
)

workspace_heading = "Math Buddy" if role_mode == "For Student" else "Math Teacher Aid"
workspace_description = (
    "Your space for lesson notes, Mathematics questions and syllabus-based practice."
    if role_mode == "For Student"
    else "Understand the student's method, find the first reasoning break, advise the student clearly, then build mastery through adaptive practice."
)

st.markdown(
    f"""
    <section class="omt-hero">
      <div class="omt-eyebrow">Singapore secondary mathematics</div>
      <h1>{workspace_heading}</h1>
      <p>{workspace_description}</p>
      <div class="omt-chip-row">
        <span class="omt-chip">✍️ Handwriting & iPad</span>
        <span class="omt-chip">∑ MathIO equation view</span>
        <span class="omt-chip">◫ Visual geometry</span>
        <span class="omt-chip">↗ Adaptive mastery</span>
      </div>
    </section>
    """,
    unsafe_allow_html=True,
)

def _topic_offline_support(topic) -> str:
    """Support both legacy syllabus Topic objects and the new learning-outcome model."""
    legacy = getattr(topic, "offline_support", None)
    if legacy:
        return str(legacy)

    # Topics in the new engine are backed by compiled learning outcomes.
    if getattr(topic, "code", None) and getattr(topic, "name", None):
        return "Strong"
    return "Unknown"


def _topic_coverage_note(topic) -> str:
    """Return legacy notes when available, otherwise describe the outcome-backed practice."""
    note = getattr(topic, "notes", None)
    if note:
        return str(note)

    keywords = tuple(getattr(topic, "keywords", ()) or ())
    if keywords:
        return (
            "Offline practice is generated from compiled learning outcomes for this topic. "
            "Question families include: " + ", ".join(keywords[:6]) + "."
        )

    return "Offline practice is generated from the compiled learning outcomes for this topic."

if role_mode == "For Teacher":
    ai_tab, setter_tab, syllabus_tab = st.tabs(
        ['✨ Analyse', '🧑\u200d🏫 Paper setter', '📚 Syllabus']
    )
else:
    student_whiteboard_tab, student_ask_tab, student_practice_tab = st.tabs(
        ["📝 Lesson whiteboard", "💬 Ask Math Advisor", "🧠 Offline practice"]
    )


# Guided-solving session defaults
st.session_state.setdefault("ai_guided_solution", None)
st.session_state.setdefault("student_ask_guided_solution", None)
st.session_state.setdefault("student_ask_error", "")
st.session_state.setdefault("student_ask_hint_count", 0)
st.session_state.setdefault("student_ask_show_full_solution", False)
st.session_state.setdefault("student_ask_reveal_step", 0)
st.session_state.setdefault("student_ask_signature", "")
st.session_state.setdefault("ai_guided_error", "")
st.session_state.setdefault("guided_hint_count", 0)
st.session_state.setdefault("guided_reveal_step", 0)
st.session_state.setdefault("guided_support_mode", "Hints only")
st.session_state.setdefault("paper_detection", None)
st.session_state.setdefault("paper_solutions", [])
st.session_state.setdefault("paper_errors", [])
st.session_state.setdefault("paper_signature", "")
st.session_state.setdefault("setter_draft", None)
st.session_state.setdefault("setter_error", "")
st.session_state.setdefault("setter_reference_signature", "")

if role_mode == "For Teacher":
    # ---------- Combined teacher workflow ----------
    with setter_tab:
        st.caption("Build 2026-08-20 · smooth monotone ogive + context-image relevance")
        st.markdown('<div class="omt-section-kicker">Teacher assessment tools</div>', unsafe_allow_html=True)
        st.markdown('<div class="omt-section-title">Paper setter, solutions & marking scheme</div>', unsafe_allow_html=True)
        teacher_workflow_mode = st.radio(
            "What would you like to do?",
            ["Create a new assessment paper", "Upload an existing question paper"],
            horizontal=True,
            key="teacher_paper_workflow_mode",
            help=(
                "Create a fresh assessment using a reference-format paper, or upload an existing question paper "
                "to generate worked solutions and a teacher marking scheme."
            ),
        )

        if teacher_workflow_mode == "Create a new assessment paper":
            st.caption("Build 2026-08-20 · worksheet None fix + near-transfer working tools")
            st.markdown('<div class="omt-section-kicker">Teacher assessment design</div>', unsafe_allow_html=True)
            st.markdown('<div class="omt-section-title">Set a new Mathematics paper</div>', unsafe_allow_html=True)
            st.write(
                "Choose the syllabus scope, assessment type, marks and duration, then upload a past paper that defines the format. "
                "The generator follows that reference structure while writing fresh questions."
            )
            st.info(
                "The reference paper is required: format, numbering, mark placement and difficulty gradient are not guessed. "
                "Generated marking schemes are teacher drafts, not official SEAB/MOE schemes."
            )

            left, right = st.columns([1, 1], gap="large")
            with left:
                st.markdown("#### 1 · Assessment settings")
                setter_track_label = st.selectbox(
                    "Selected level / syllabus",
                    options=list(APP_TRACKS.keys()),
                    index=list(APP_TRACKS.keys()).index(track_label) if track_label in APP_TRACKS else 0,
                    key="setter_track_label",
                    help="Choose the syllabus for this assessment independently of the main tutor sidebar.",
                )
                setter_info = selected_track_info(setter_track_label)
                if setter_info.get("year") == 2027:
                    st.caption(
                        f"{setter_info.get('year')} SEC · {setter_info.get('level')} "
                        f"{setter_info.get('subject')} · {setter_info.get('subject_code')}"
                    )

                setter_assessment = st.selectbox(
                    "Assessment type",
                    ["Weighted Assessment (WA)", "End-of-Year (EOY)", "Class test", "Worksheet", "Preliminary examination"],
                    key="setter_assessment_type",
                )
                is_worksheet = setter_assessment == "Worksheet"
                if is_worksheet:
                    setter_marks = 0
                    setter_duration = 0
                    setter_questions = st.number_input(
                        "Main questions", min_value=1, max_value=60, value=12, step=1,
                        key="setter_question_count",
                    )
                    st.caption("Worksheet mode: no total marks and no duration.")
                else:
                    c1, c2 = st.columns(2)
                    setter_marks = c1.number_input("Total marks", min_value=10, max_value=200, value=50, step=5, key="setter_total_marks")
                    setter_questions = c2.number_input("Main questions", min_value=1, max_value=40, value=12, step=1, key="setter_question_count")
                    setter_duration = st.number_input("Duration (minutes)", min_value=20, max_value=240, value=75, step=5, key="setter_duration")
                setter_school = st.text_input("School name (optional - otherwise infer from reference)", key="setter_school")
                setter_title = st.text_input("Paper title (optional)", key="setter_title")

            with right:
                st.markdown("#### 2 · Syllabus scope")
                setter_track_info = selected_track_info(setter_track_label)
                available_topics = paper_setting_topics(track_code(setter_track_label))
                setter_topics = st.multiselect(
                    "Topics / chapters to test",
                    options=available_topics,
                    default=available_topics,
                    key=f"setter_topics_{track_code(setter_track_label)}",
                    help="Topics are taken from the uploaded learning-outcomes workbook for this exam track.",
                )
                setter_syllabus_notes = st.text_area(
                    "Additional syllabus notes / exclusions",
                    key="setter_syllabus_notes",
                    height=130,
                    placeholder=(
                        "Example: Algebraic manipulation; linear equations; Pythagoras; trigonometry. "
                        "List exclusions too, e.g. no quadratic formula yet."
                    ),
                    help="The topic list already comes from the uploaded syllabus workbook. Use this box only for exclusions or school-specific emphasis.",
                )
                include_scheme = st.checkbox(
                    "Generate marking scheme together with the paper",
                    value=False,
                    key="setter_include_scheme",
                )

            st.markdown("#### 3 · Reference format paper (optional)")
            setter_reference = st.file_uploader(
                "Optional: upload a past paper of the same assessment type",
                type=["pdf", "docx", "doc"],
                accept_multiple_files=False,
                key="setter_reference_upload",
                help=("Optional. If supplied, this paper guides section structure, numbering, mark placement and difficulty gradient. "      "If omitted, Math Advisor uses the selected syllabus, assessment settings and built-in Singapore paper conventions."),
            )

            reference_ready = False
            reference_text = ""
            reference_assets: list[UploadedAsset] = []
            if setter_reference is not None:
                try:
                    reference_text, reference_assets = full_paper_input(setter_reference)
                    reference_ready = bool(reference_assets or reference_text.strip())
                    st.success(
                        f"Reference loaded: {setter_reference.name}"
                        + (f" · {len(reference_text):,} characters extracted" if reference_text else "")
                    )
                except GeminiTutorError as exc:
                    st.error(str(exc))

            scope_ready = bool(setter_topics or setter_syllabus_notes.strip())
            can_generate = scope_ready
            if setter_reference is None:
                st.caption(
                    "No reference paper selected. Math Advisor will use the chosen syllabus, assessment type, "
                    "marks, duration and built-in Singapore assessment conventions."
                )

            if st.button(
                "Generate assessment paper",
                type="primary",
                use_container_width=True,
                disabled=not can_generate,
                key="setter_generate_button",
            ):
                st.session_state.setter_error = ""
                st.session_state.setter_draft = None
                st.session_state.setter_geogebra_graphs = {}
                try:
                    spinner_text = (
                        "Reading the optional reference format, setting questions and auditing mark totals..."
                        if reference_ready
                        else "Setting questions from the selected syllabus and assessment settings..."
                    )
                    with st.spinner(spinner_text):
                        source_syllabus_notes = topic_notes_for_selection(
                            track_code(setter_track_label), list(setter_topics)
                        )
                        combined_syllabus_notes = "\n\n".join(
                            x for x in [
                                source_syllabus_notes,
                                setter_syllabus_notes.strip(),
                                _TABLE_GENERATION_REQUIREMENTS,
                                _SOLID3D_GENERATION_REQUIREMENTS,
                            ]
                            if x
                        )
                        draft = generate_exam_paper_draft(
                            track_label=setter_track_label,
                            assessment_type=setter_assessment,
                            total_marks=(0 if setter_assessment == "Worksheet" else int(setter_marks)),
                            number_of_questions=int(setter_questions),
                            duration_minutes=(0 if setter_assessment == "Worksheet" else int(setter_duration)),
                            topics=list(setter_topics),
                            syllabus_notes=combined_syllabus_notes,
                            reference_text=reference_text,
                            reference_assets=reference_assets,
                            school_name=setter_school,
                            paper_title=setter_title,
                            include_marking_scheme=include_scheme,
                            api_key=explicit_key,
                            model=model,
                        )
                    st.session_state.setter_draft = draft
                    st.rerun()
                except GeminiTutorError as exc:
                    st.session_state.setter_error = str(exc)
                    st.rerun()
                except Exception as exc:
                    st.session_state.setter_error = f"{type(exc).__name__}: {str(exc)[:400]}"
                    st.rerun()

            if st.session_state.get("setter_error"):
                st.error(st.session_state.setter_error)

            setter_draft = st.session_state.get("setter_draft")
            if setter_draft is not None:
                solid_audit_issues = audit_generated_solids(setter_draft)
                if solid_audit_issues:
                    st.warning(
                        "Some 3D-solid questions have incomplete dimension-diagram data: "
                        + "; ".join(solid_audit_issues)
                    )

                table_audit_issues = audit_generated_tables(setter_draft)
                if table_audit_issues:
                    st.error(
                        "Generated paper has incomplete table-data issues and should be regenerated: "
                        + "; ".join(table_audit_issues)
                    )

                graph_audit_issues = audit_generated_graphs(setter_draft)
                if graph_audit_issues:
                    st.error(
                        "Generated paper has graph-data issues and should be regenerated before use: "
                        + "; ".join(graph_audit_issues)
                    )
                render_setter_preview(setter_draft)

                if setter_draft.reference_format_summary:
                    with st.expander("Reference-format features used", expanded=False):
                        for item in setter_draft.reference_format_summary:
                            st.markdown(f"- {item}")

                if setter_draft.verification_notes:
                    with st.expander("Paper audit notes", expanded=False):
                        for item in setter_draft.verification_notes:
                            st.markdown(f"- {item}")

                graph_questions = [q for q in setter_draft.questions if _question_graph_spec(q) is not None]
                if graph_questions:
                    captured_count = sum(1 for q in graph_questions if _captured_geogebra_png(q))
                    if captured_count == len(graph_questions):
                        st.success(
                            f"GeoGebra graph capture ready: {captured_count}/{len(graph_questions)} function graph(s) "
                            "will be inserted into the Word paper."
                        )
                    else:
                        st.info(
                            f"GeoGebra graph capture: {captured_count}/{len(graph_questions)} ready. "
                            "You may download now: any graph not yet captured by GeoGebra will be generated "
                            "directly from its exact equation by Math Advisor's deterministic renderer."
                        )

                question_docx = build_setter_question_paper_docx(setter_draft)
                downloads = st.columns(2 if include_scheme else 1)
                downloads[0].download_button(
                    "Download question paper (.docx)",
                    data=question_docx,
                    file_name="Generated_Maths_Question_Paper.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
                if include_scheme:
                    scheme_docx = build_setter_marking_scheme_docx(setter_draft)
                    downloads[1].download_button(
                        "Download marking scheme (.docx)",
                        data=scheme_docx,
                        file_name="Generated_Maths_Marking_Scheme.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )

    # ---------- Existing question paper → solutions + marking scheme ----------
    with setter_tab:
        if teacher_workflow_mode == "Upload an existing question paper":
            st.markdown("### Upload question paper → solutions + marking scheme")
            st.caption(
                "Upload a complete PDF or Word question paper. Math Advisor will detect the questions, "
                "generate worked solutions, and create a teacher marking scheme."
            )
            st.caption("Build 2026-08-17 · full-paper MathIO solutions")
            st.markdown('<div class="omt-section-kicker">Teacher / revision workflow</div>', unsafe_allow_html=True)
            st.markdown('<div class="omt-section-title">Worked solutions + marking scheme</div>', unsafe_allow_html=True)
            st.write(
                "Upload a complete PDF or Word exam paper. The tutor detects every main question and subpart, "
                "then generates verified worked solutions and an AI-suggested marking scheme."
            )
            st.warning(
                "The marking scheme is a teaching aid, not an official SEAB/MOE mark scheme. "
                "For formal grading, use the official marking scheme where available."
            )

            paper_file = st.file_uploader(
                "Upload full exam paper",
                type=["pdf", "docx", "doc"],
                accept_multiple_files=False,
                key="full_paper_upload",
                help="PDF and modern Word .docx files are supported. Legacy .doc files should be saved as .docx or PDF.",
            )
            paper_generation_mode = st.radio(
                "Full-paper generation mode",
                ["Reliable", "Faster"],
                horizontal=True,
                key="paper_generation_mode",
                help=(
                    "Reliable uses structured retries and independent verification for every question. "
                    "Faster reduces retry effort but may need more manual review."
                ),
            )
            st.caption(
                "Questions are solved one at a time. A failure on one question no longer stops or invalidates the rest of the paper."
            )

            paper_title = st.text_input(
                "Paper title (optional)",
                key="full_paper_title",
                placeholder="Example: 2027 G3 Mathematics Revision Paper 1",
            )

            if paper_file is not None:
                signature = f"visual-safe-v2:{paper_file.name}:{getattr(paper_file, 'size', 0)}"
                if st.session_state.paper_signature != signature:
                    st.session_state.paper_signature = signature
                    st.session_state.paper_detection = None
                    st.session_state.paper_solutions = []
                    st.session_state.paper_errors = []

                try:
                    paper_text, paper_assets = full_paper_input(paper_file)
                    st.success(
                        f"Loaded {paper_file.name}"
                        + (f" · extracted {len(paper_text):,} characters of Word text" if paper_text else "")
                    )

                    if st.button("1 · Detect all questions and subparts", use_container_width=True):
                        with st.spinner("Reading the full paper structure..."):
                            detection = detect_questions_in_assets(
                                track_label=track_label,
                                question_assets=paper_assets,
                                paper_text=paper_text,
                                api_key=explicit_key,
                                model=model,
                            )
                        st.session_state.paper_detection = detection
                        st.session_state.paper_solutions = []
                        st.session_state.paper_errors = []
                        st.rerun()

                    detection = st.session_state.get("paper_detection")
                    if detection is not None:
                        st.markdown("### Detected paper structure")
                        st.success(
                            f"Detected {detection.main_question_count} main question(s)"
                            + (
                                f"; {sum(len(q.subparts) for q in detection.questions)} subpart(s)."
                                if detection.questions else "."
                            )
                        )
                        if detection.notes:
                            for note in detection.notes:
                                st.caption(f"• {note}")

                        preview_rows = []
                        for q in detection.questions:
                            preview_rows.append(
                                {
                                    "Question": q.question_number,
                                    "Topic": q.topic_hint,
                                    "Subparts": ", ".join(p.label for p in q.subparts) or "Whole question",
                                    "Pages": ", ".join(map(str, q.page_numbers)) or "—",
                                    "Confidence": q.confidence.title(),
                                }
                            )
                        if preview_rows:
                            st.dataframe(pd.DataFrame(preview_rows), hide_index=True, use_container_width=True)

                        st.caption(
                            "Generation verifies each question independently. Geometry/shaded-area questions use the topology-first accuracy checks."
                        )

                        if st.button(
                            "2 · Generate worked solutions + marking scheme",
                            type="primary",
                            use_container_width=True,
                        ):
                            st.session_state.paper_solutions = []
                            st.session_state.paper_errors = []
                            progress = st.progress(0.0, text="Starting paper solution generation...")
                            total_questions = max(1, len(detection.questions))

                            solutions: list[PaperQuestionSolution] = []
                            errors: list[str] = []
                            for index, detected_question in enumerate(detection.questions, 1):
                                progress.progress(
                                    (index - 1) / total_questions,
                                    text=f"Solving Question {detected_question.question_number} ({index}/{total_questions}) · {paper_generation_mode} mode...",
                                )
                                try:
                                    next_question_number = detection.questions[index].question_number if index < len(detection.questions) else None
                                    focused_paper_text = paper_question_text_context(
                                        detected_question, paper_text, next_question_number
                                    )
                                    scoped_assets = scoped_assets_for_paper_question(
                                        paper_assets,
                                        detected_question.page_numbers,
                                    )
                                    if str(getattr(paper_file, "name", "")).lower().endswith(".docx"):
                                        question_visual_text = " ".join(
                                            [
                                                detected_question.question_text or "",
                                                *[
                                                    getattr(part, "question_text", "") or ""
                                                    for part in (getattr(detected_question, "subparts", []) or [])
                                                ],
                                            ]
                                        )
                                        needs_visual = bool(
                                            re.search(
                                                r"\b(diagram|figure|graph|table|chart|grid|shape|circle|semicircle|triangle|angle|coordinates?|plot|draw|sketch|construction|map|pie chart|histogram)\b",
                                                question_visual_text,
                                                re.IGNORECASE,
                                            )
                                        )
                                        scoped_assets = scoped_assets[:8] if needs_visual else []
                                    solution = generate_paper_question_solution(
                                        track_label=track_label,
                                        detected_question=detected_question,
                                        question_assets=scoped_assets,
                                        paper_text_context=focused_paper_text,
                                        api_key=explicit_key,
                                        model=model,
                                    )
                                    solutions.append(solution)
                                except GeminiTutorError as exc:
                                    errors.append(f"Question {detected_question.question_number}: {exc}")
                                    message = str(exc).lower()
                                    if "quota" in message or "rate limit" in message or "ratelimit" in message or "free-tier" in message:
                                        errors.append("Generation stopped because the Gemini quota/rate limit was reached. Completed questions are preserved; continue later.")
                                        break
                                except Exception as exc:
                                    errors.append(
                                        f"Question {detected_question.question_number}: {type(exc).__name__}: {str(exc)[:300]}"
                                    )

                            progress.progress(1.0, text="Paper generation complete.")
                            st.session_state.paper_solutions = solutions
                            st.session_state.paper_errors = errors
                            st.rerun()

                        solutions = st.session_state.get("paper_solutions") or []
                        errors = st.session_state.get("paper_errors") or []

                        if errors:
                            st.warning(
                                f"{len(errors)} question(s) could not be completed reliably. "
                                "They are listed below rather than guessed."
                            )
                            for error in errors:
                                st.error(error)

                        if solutions:
                            st.markdown("## Worked solutions")
                            total_marks = sum(solution.total_marks for solution in solutions)
                            c1, c2, c3 = st.columns(3)
                            c1.metric("Questions completed", len(solutions))
                            c2.metric("Suggested / printed marks", total_marks)
                            c3.metric("Questions needing review", len(errors))

                            for solution in solutions:
                                render_paper_question_solution(solution)

                            markdown_export = paper_solution_markdown(
                                track_label=track_label,
                                paper_title=paper_title or os.path.splitext(paper_file.name)[0],
                                solutions=solutions,
                            )
                            docx_export = build_paper_solution_docx(
                                track_label=track_label,
                                paper_title=paper_title or os.path.splitext(paper_file.name)[0],
                                solutions=solutions,
                            )
                            d1, d2 = st.columns(2)
                            d1.download_button(
                                "Download solutions as Markdown",
                                data=markdown_export.encode("utf-8"),
                                file_name=f"{os.path.splitext(paper_file.name)[0]}_worked_solutions.md",
                                mime="text/markdown",
                                use_container_width=True,
                            )
                            d2.download_button(
                                "Download solutions + marking scheme as Word",
                                data=docx_export,
                                file_name=f"{os.path.splitext(paper_file.name)[0]}_worked_solutions_marking_guide.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                            )
                except GeminiTutorError as exc:
                    st.error(str(exc))

if role_mode == "For Teacher":
    # ---------- Gemini online analysis ----------
    with ai_tab:
        st.markdown('<div class="omt-section-kicker">Step 1 · Submit</div>', unsafe_allow_html=True)
        st.markdown('<div class="omt-section-title">Question + student working</div>', unsafe_allow_html=True)
        st.markdown(
            "<div class='omt-section-copy'>Upload a photo/PDF or type the question, then add the student's working. The tutor keeps the question and solution separate during diagnosis.</div>",
            unsafe_allow_html=True,
        )

        input_left, input_right = st.columns([.95, 1.05], gap="large")
        with input_left:
            with st.container(border=True):
                st.markdown("#### 📄 Question")
                q_text = question_input_with_math_keyboard(key_base="ai_question")
                geogebra_external_tools(question_text=q_text, key_base="ai_geogebra")
                q_files = st.file_uploader(
                    "Upload question image/PDF",
                    type=["png", "jpg", "jpeg", "webp", "pdf"],
                    accept_multiple_files=True,
                    key="ai_question_files",
                    help="Photos, screenshots and PDFs are supported.",
                )
                submission_mode = st.radio(
                    "How is this question being used?",
                    [
                        "Separate student solution",
                        "Student solution is already on the question upload",
                        "No student solution — guide me to solve it",
                    ],
                    key="ai_submission_mode",
                    help=(
                        "Choose whether you want the tutor to mark a separate solution, read working already written "
                        "on the uploaded question, or teach the question from scratch."
                    ),
                )
                working_in_question_upload = submission_mode == "Student solution is already on the question upload"
                guided_mode = submission_mode == "No student solution — guide me to solve it"

        with input_right:
            with st.container(border=True):
                if guided_mode:
                    st.markdown("#### 🧭 Guided solving")
                    st.info(
                        "No student solution is required. After the question passes the feasibility check, "
                        "the tutor will guide the student with a diagnostic question, progressive hints, "
                        "and solution steps revealed one at a time."
                    )
                    w_text, w_input_mode, w_offline_text = "", "No student working", ""
                    w_files = []
                else:
                    st.markdown("#### ✍️ Student working")
                    if working_in_question_upload:
                        st.caption(
                            "The tutor will read the handwriting/annotations from the question upload. "
                            "You may add extra working below if needed."
                        )
                    geogebra_external_tools(
                        question_text=_active_question_text_for_tools(),
                        key_base="student_working_geogebra",
                    )
                    student_scientific_calculator(
                        key_base="student_working_calculator"
                    )
                    w_text, w_input_mode, w_offline_text = working_input(
                        "Student working",
                        text_key="ai_working_text",
                        format_key="ai_working_format",
                        height=160,
                        plain_placeholder="Type the steps, use the equation editor, or leave blank when the working is uploaded.",
                    )
                    w_files = st.file_uploader(
                        "Upload student working image/PDF",
                        type=["png", "jpg", "jpeg", "webp", "pdf"],
                        accept_multiple_files=True,
                        key="ai_working_files",
                    )

        # Clear stale detection results when the uploaded source changes.
        current_signature = question_file_signature(q_files)
        if current_signature != st.session_state.ai_question_file_signature:
            st.session_state.ai_question_file_signature = current_signature
            st.session_state.ai_question_detection = None
            st.session_state.ai_question_detection_error = ""
            st.session_state.ai_selected_question_index = 0
            st.session_state.ai_question_feasibility = None
            st.session_state.ai_question_feasibility_error = ""
            st.session_state.ai_question_feasibility_signature = ""
            st.session_state.ai_analysis = None
            st.session_state.ai_visual_explanation = None
            st.session_state.ai_visual_error = ""
            st.session_state.ai_visual_step = 0
            st.session_state.ai_cached_verification = None
            st.session_state.ai_cached_verification_signature = ""
            st.session_state.ai_guided_solution = None
            st.session_state.ai_guided_error = ""
            st.session_state.guided_hint_count = 0
            st.session_state.guided_reveal_step = 0
            st.session_state.guided_support_mode = "Hints only"
            clear_ai_practice_state()
            st.session_state.pop("ai_detected_question_selector", None)

        consent = st.checkbox(
            "Allow Gemini to analyse the selected question and working",
            key="gemini_consent",
            help="Remove names, NRICs and other unnecessary personal identifiers before sending student work.",
        )

        selected_detection_index = 0
        if q_files:
            st.markdown("### Detect questions in the upload")
            st.write(
                "Gemini can count the **main questions** in the uploaded image/PDF, keep subparts grouped under their main question, "
                "and let the student choose which question to analyse."
            )
            if st.button("Detect questions in uploaded file(s)", use_container_width=True):
                st.session_state.ai_question_detection = None
                st.session_state.ai_question_detection_error = ""
                st.session_state.ai_question_feasibility = None
                st.session_state.ai_question_feasibility_error = ""
                st.session_state.ai_question_feasibility_signature = ""
                st.session_state.pop("ai_detected_question_selector", None)
                if not consent:
                    st.session_state.ai_question_detection_error = (
                        "Confirm the Gemini data-sharing acknowledgement before detecting questions."
                    )
                else:
                    try:
                        assets_q = uploaded_assets(q_files)
                        with st.spinner("Detecting main questions and subparts in the upload..."):
                            detection = detect_questions_in_assets(
                                track_label=track_label,
                                question_assets=assets_q,
                                api_key=explicit_key,
                                model=model,
                            )
                        st.session_state.ai_question_detection = detection
                        st.rerun()
                    except GeminiTutorError as exc:
                        st.session_state.ai_question_detection_error = str(exc)
                        st.rerun()

            if st.session_state.ai_question_detection_error:
                st.error(st.session_state.ai_question_detection_error)

            detection: QuestionDetectionResult | None = st.session_state.ai_question_detection
            if detection is not None:
                selected_detection_index = render_question_detection(detection)

        detection = st.session_state.ai_question_detection
        question_for_analysis = question_for_selected_analysis(q_text, detection, selected_detection_index)
        current_feasibility_signature = question_feasibility_signature(
            question_for_analysis, q_files, selected_detection_index
        )
        if (
            st.session_state.ai_question_feasibility_signature
            and st.session_state.ai_question_feasibility_signature != current_feasibility_signature
        ):
            st.session_state.ai_question_feasibility = None
            st.session_state.ai_question_feasibility_error = ""
            st.session_state.ai_question_feasibility_signature = ""
            st.session_state.ai_analysis = None
            st.session_state.ai_visual_explanation = None
            st.session_state.ai_visual_error = ""
            st.session_state.ai_visual_step = 0
            clear_ai_practice_state()

        st.markdown("### Check the question before analysis")
        st.write(
            "Gemini can first check whether the selected question is complete, internally consistent, "
            "mathematically meaningful, and sufficiently clear before marking or guiding the student."
        )
        bypass_feasibility = st.checkbox(
            "Bypass question feasibility check",
            key="ai_bypass_feasibility",
            help=(
                "Use this only when you already trust the question. Independent mathematical verification still runs "
                "before the tutor marks or guides the solution."
            ),
        )
        if bypass_feasibility:
            st.warning(
                "Question-feasibility screening is being bypassed. The tutor will still independently verify the mathematics, "
                "but missing or contradictory information may only be discovered during verification."
            )

        if st.button(
            "Check question feasibility",
            use_container_width=True,
            disabled=bypass_feasibility,
        ):
            st.session_state.ai_question_feasibility = None
            st.session_state.ai_question_feasibility_error = ""
            st.session_state.ai_analysis = None
            st.session_state.ai_visual_explanation = None
            st.session_state.ai_visual_error = ""
            st.session_state.ai_visual_step = 0
            clear_ai_practice_state()
            if not consent:
                st.session_state.ai_question_feasibility_error = (
                    "Confirm the Gemini data-sharing acknowledgement before checking the question."
                )
            elif not question_for_analysis.strip() and not q_files:
                st.session_state.ai_question_feasibility_error = "Provide the question as text or an upload first."
            else:
                # This fallback only exists for the student-solution path. Initialise it
                # before any Gemini call so guided-mode exceptions cannot trigger a NameError.
                offline_result = None
                try:
                    assets_q = uploaded_assets(q_files)
                    with st.spinner("Checking the question for missing information, contradictions, ambiguity, and mathematical feasibility..."):
                        feasibility = assess_question_feasibility(
                            track_label=track_label,
                            question_text=question_for_analysis,
                            question_assets=assets_q,
                            api_key=explicit_key,
                            model=model,
                        )
                    st.session_state.ai_question_feasibility = feasibility
                    st.session_state.ai_question_feasibility_signature = current_feasibility_signature
                    st.rerun()
                except GeminiTutorError as exc:
                    st.session_state.ai_question_feasibility_error = str(exc)
                    st.rerun()

        if st.session_state.ai_question_feasibility_error:
            st.error(st.session_state.ai_question_feasibility_error)

        feasibility: QuestionFeasibilityResult | None = st.session_state.ai_question_feasibility
        if feasibility is not None:
            render_question_feasibility(feasibility, q_files)

        feasibility_passed = bool(
            feasibility is not None
            and feasibility.can_analyse_student_work
            and st.session_state.ai_question_feasibility_signature == current_feasibility_signature
        )
        feasibility_ready = bool(bypass_feasibility or feasibility_passed)

        # If the user chose the student-solution path but has not supplied any working,
        # automatically treat the action as guided solving rather than failing with
        # "Provide the student's working".
        has_student_work = bool(
            (w_text or "").strip()
            or w_files
            or (working_in_question_upload and q_files)
        )
        effective_guided_mode = bool(guided_mode or not has_student_work)

        if not feasibility_ready:
            st.info(
                "Run the question feasibility check, or select **Bypass question feasibility check**, "
                "before continuing."
            )
        elif not guided_mode and not has_student_work:
            st.info(
                "No student working is currently supplied. Clicking the main button will **advise how to solve the question** "
                "using guided hints and step-by-step support."
            )

        primary_action_label = "Analyse student working / Advise how to solve the question"
        if st.button(
            primary_action_label,
            type="primary",
            use_container_width=True,
            disabled=not feasibility_ready,
        ):
            st.session_state.ai_analysis = None
            st.session_state.ai_error = ""
            st.session_state.ai_fallback_result = None
            st.session_state.ai_visual_explanation = None
            st.session_state.ai_visual_error = ""
            st.session_state.ai_visual_step = 0
            st.session_state.ai_guided_solution = None
            st.session_state.ai_guided_error = ""
            st.session_state.guided_hint_count = 0
            st.session_state.guided_reveal_step = 0
            clear_ai_practice_state()
            if not consent:
                st.error("Confirm the Gemini data-sharing acknowledgement before sending the question.")
            elif not feasibility_ready:
                st.error("Run the question feasibility check or enable the bypass option before continuing.")
            else:
                try:
                    assets_q = uploaded_assets(q_files)

                    # Cache the independent verification once per question.
                    verification = st.session_state.get("ai_cached_verification")
                    if st.session_state.get("ai_cached_verification_signature") != current_feasibility_signature:
                        verification = None
                    if verification is None:
                        with st.spinner("Verifying the question mathematics once..."):
                            verification = verify_question_math(
                                track_label=track_label,
                                question_text=question_for_analysis,
                                question_assets=assets_q,
                                api_key=explicit_key,
                                model=model,
                            )
                        st.session_state.ai_cached_verification = verification
                        st.session_state.ai_cached_verification_signature = current_feasibility_signature

                    if effective_guided_mode:
                        with st.spinner("Preparing guided steps without revealing the answer immediately..."):
                            guided = generate_guided_solution(
                                track_label=track_label,
                                question_text=question_for_analysis,
                                question_assets=assets_q,
                                api_key=explicit_key,
                                model=model,
                                verification=verification,
                            )
                        st.session_state.ai_guided_solution = guided
                        st.rerun()

                    # Student-solution analysis path.
                    evidence, offline_result = offline_evidence_for(question_for_analysis, w_offline_text)
                    working_for_gemini = (
                        f"[Student working input method: {w_input_mode}]\n{w_text}" if w_text.strip() else w_text
                    )
                    if working_in_question_upload:
                        embedded_note = "[Student working is visible in the same uploaded question image/PDF. Inspect the handwritten/annotated working in that upload as the student's solution.]"
                        working_for_gemini = (working_for_gemini + "\n" + embedded_note).strip()

                    assets_w = uploaded_assets(w_files)
                    if working_in_question_upload:
                        assets_w = [*assets_w, *assets_q]

                    with st.spinner("Gemini is checking the student's reasoning..."):
                        analysis = call_analyze_submission_compat(
                            track_label=track_label,
                            question_text=question_for_analysis,
                            working_text=working_for_gemini,
                            question_assets=assets_q,
                            working_assets=assets_w,
                            offline_evidence=evidence,
                            api_key=explicit_key,
                            model=model,
                            verification=verification,
                        )
                    st.session_state.ai_analysis = analysis
                    if analysis_speed == "Full" and _visual_plan_is_recommended(analysis, question_for_analysis):
                        try:
                            with st.spinner("Building an interactive visual explanation for this geometry/graph question..."):
                                visual_plan = generate_visual_explanation(
                                    track_label=track_label,
                                    question_text=question_for_analysis,
                                    analysis=analysis,
                                    question_assets=assets_q,
                                    api_key=explicit_key,
                                    model=model,
                                )
                            st.session_state.ai_visual_explanation = visual_plan
                        except GeminiTutorError as visual_exc:
                            # Visuals are an enhancement; never lose the verified reasoning analysis if this second call fails.
                            st.session_state.ai_visual_error = str(visual_exc)
                    initialize_ai_practice(analysis)
                    st.rerun()
                except GeminiTutorError as exc:
                    if effective_guided_mode:
                        st.session_state.ai_guided_error = str(exc)
                    else:
                        st.session_state.ai_error = str(exc)
                        if offline_result is not None:
                            st.session_state.ai_fallback_result = offline_result
                    st.rerun()

        guided_result = st.session_state.get("ai_guided_solution")
        if guided_result is not None:
            render_guided_solution(guided_result)

        if st.session_state.get("ai_guided_error"):
            st.error(st.session_state.ai_guided_error)

        if st.session_state.ai_error:
            st.error(st.session_state.ai_error)
            if st.session_state.ai_fallback_result is not None:
                st.info("Gemini was unavailable, so the tutor automatically used its deterministic offline algebra fallback for this typed submission.")
                render_attempt(st.session_state.ai_fallback_result)
            else:
                st.info("If you intended to solve without student working, use **Advise how to solve the question**; offline practice and algebra check also remain available.")

        analysis: GeminiAnalysis | None = st.session_state.ai_analysis
        if analysis is not None:
            render_ai_analysis(analysis)
            visual_plan: VisualExplanationResult | None = st.session_state.ai_visual_explanation
            if visual_plan is not None:
                st.markdown("---")
                render_visual_explanation(visual_plan, q_files)
            elif st.session_state.ai_visual_error:
                st.caption("Interactive visual explanation unavailable for this attempt: " + st.session_state.ai_visual_error)
            elif _visual_plan_is_recommended(analysis, question_for_analysis):
                if st.button("Build visual explanation", key="build_visual_on_demand", use_container_width=True):
                    try:
                        assets_q = uploaded_assets(q_files)
                        with st.spinner("Building the interactive visual explanation..."):
                            st.session_state.ai_visual_explanation = generate_visual_explanation(
                                track_label=track_label, question_text=question_for_analysis, analysis=analysis,
                                question_assets=assets_q, api_key=explicit_key, model=model,
                            )
                        st.rerun()
                    except GeminiTutorError as exc:
                        st.session_state.ai_visual_error = str(exc)
                        st.rerun()
            st.markdown("---")
            st.markdown('<div class="omt-section-kicker">Adaptive practice</div>', unsafe_allow_html=True)
            st.markdown('<div class="omt-section-title">Build mastery one transfer level at a time</div>', unsafe_allow_html=True)
            st.markdown(
                '<div class="omt-section-copy">Near transfer → Varied context → Stretch. A mistake keeps the student on the same skill until the reasoning becomes secure.</div>',
                unsafe_allow_html=True,
            )

            if st.session_state.ai_practice_current_question is None and not st.session_state.ai_practice_finished:
                initialize_ai_practice(analysis)

            stage_index = int(st.session_state.ai_practice_stage)
            completed = st.session_state.ai_practice_completed
            stage_html = []
            for i, kind in enumerate(PRACTICE_STAGES):
                if completed.get(kind):
                    css = "done"; icon = "✓"; detail = "Mastered"
                elif not st.session_state.ai_practice_finished and i == stage_index:
                    css = "current"; icon = "●"; detail = "Current focus"
                else:
                    css = "locked"; icon = "◌"; detail = "Locked"
                stage_html.append(
                    f'<div class="omt-stage {css}"><div class="name">{icon} {kind.title()}</div><div class="detail">{detail}</div></div>'
                )
            st.markdown('<div class="omt-stage-row">' + ''.join(stage_html) + '</div>', unsafe_allow_html=True)

            if st.session_state.ai_practice_finished:
                st.success(
                    "Adaptive practice complete: the student demonstrated secure reasoning through Near transfer, "
                    "Varied context, and Stretch."
                )
            else:
                kind = PRACTICE_STAGES[stage_index]
                pq: TargetedPracticeQuestion = st.session_state.ai_practice_current_question
                misses = st.session_state.ai_practice_misses[kind]
                streak = st.session_state.ai_practice_consecutive_correct[kind]

                st.markdown(f'<div class="omt-section-kicker">Current focus</div><div class="omt-section-title">{kind.title()}</div>', unsafe_allow_html=True)
                if misses:
                    st.warning(
                        f"This category remains active because the student has had {misses} non-secure attempt(s). "
                        f"Current recovery streak: {streak}/2 secure attempts."
                    )
                render_targeted_practice_focus(
                    pq,
                    key=f"{stage_index}_{st.session_state.ai_practice_question_version}",
                )
                st.caption("Skill being checked")
                render_mathio_mixed(_clean_practice_display_text(pq.target_skill))
                required_parts = required_parts_for_question(pq)
                if required_parts != ["whole question"]:
                    st.caption("All parts required for mastery: " + ", ".join(required_parts))
                with st.expander("Why this question"):
                    render_mathio_mixed(_clean_practice_display_text(pq.why_this_tests_understanding))
                with st.expander("Practice hints"):
                    for i, hint in enumerate(pq.hints, 1):
                        st.markdown(f"**Hint {i}:**")
                        render_mathio_mixed(hint)

                if kind == "Near transfer":
                    st.caption("Working tools")
                    geogebra_external_tools(
                        question_text=" ".join([
                            str(getattr(pq, "question", "") or ""),
                            str(getattr(pq, "focus_prompt", "") or ""),
                        ]),
                        key_base=f"near_transfer_geogebra_{stage_index}_{st.session_state.ai_practice_question_version}",
                    )
                    student_scientific_calculator(
                        key_base=f"near_transfer_calculator_{stage_index}_{st.session_state.ai_practice_question_version}"
                    )

                working_key = f"ai_practice_working_{stage_index}_{st.session_state.ai_practice_question_version}"
                attempt, practice_input_mode, _practice_offline_text, practice_assets = targeted_practice_input(
                    f"Student working for {kind}",
                    key_base=working_key,
                    height=150,
                    practice_question=pq,
                )

                if st.button(f"Check {kind} reasoning", key=f"ai_practice_check_{stage_index}_{st.session_state.ai_practice_question_version}", type="primary"):
                    if practice_input_mode == "Handwritten working" and not attempt.strip() and not practice_assets:
                        st.warning("No saved handwriting was received. Return to the handwriting pad, tap **Save handwriting**, then check the reasoning again.")
                        st.stop()
                    try:
                        with st.spinner("Checking the practice reasoning..."):
                            evaluation = evaluate_practice_attempt(
                                track_label=track_label,
                                practice_question=pq,
                                student_working=(
                                    f"[Student working input method: {practice_input_mode}]\n{attempt}"
                                    if attempt.strip() else f"[Student working input method: {practice_input_mode}]"
                                ),
                                working_assets=practice_assets,
                                original_gap=analysis.misconception_or_gap,
                                api_key=explicit_key,
                                model=model,
                            )
                        secure = practice_attempt_is_secure(evaluation)
                        if secure:
                            st.session_state.ai_practice_consecutive_correct[kind] += 1
                        else:
                            st.session_state.ai_practice_misses[kind] += 1
                            st.session_state.ai_practice_consecutive_correct[kind] = 0

                        current_misses = st.session_state.ai_practice_misses[kind]
                        current_streak = st.session_state.ai_practice_consecutive_correct[kind]
                        st.session_state.ai_practice_ready_to_advance = bool(
                            secure and (current_misses == 0 or current_streak >= 2)
                        )
                        st.session_state.ai_practice_evaluation = evaluation
                        st.session_state.ai_practice_last_working = (
                            attempt if attempt.strip() else f"[{practice_input_mode} submitted; use the marking feedback as the diagnostic summary.]"
                        )
                        record_ai_practice_history(tcode, pq, evaluation)
                        st.rerun()
                    except GeminiTutorError as exc:
                        st.error(str(exc))

                evaluation: PracticeEvaluation | None = st.session_state.ai_practice_evaluation
                if evaluation is not None:
                    render_practice_evaluation(evaluation)
                    secure = practice_attempt_is_secure(evaluation)
                    ready = bool(st.session_state.ai_practice_ready_to_advance)

                    if ready:
                        st.success(f"{kind} is secure. The next transfer level can now be unlocked.")
                        if stage_index < len(PRACTICE_STAGES) - 1:
                            next_kind = PRACTICE_STAGES[stage_index + 1]
                            if st.button(f"Continue to {next_kind}", use_container_width=True):
                                st.session_state.ai_practice_completed[kind] = True
                                st.session_state.ai_practice_stage = stage_index + 1
                                st.session_state.ai_practice_current_question = initial_practice_question(analysis, next_kind)
                                st.session_state.ai_practice_evaluation = None
                                st.session_state.ai_practice_last_working = ""
                                st.session_state.ai_practice_ready_to_advance = False
                                st.session_state.ai_practice_question_version += 1
                                st.rerun()
                        else:
                            if st.button("Complete adaptive practice", use_container_width=True):
                                st.session_state.ai_practice_completed[kind] = True
                                st.session_state.ai_practice_finished = True
                                st.session_state.ai_practice_evaluation = None
                                st.rerun()
                    else:
                        if secure:
                            remaining = max(0, 2 - st.session_state.ai_practice_consecutive_correct[kind])
                            st.info(
                                f"Good recovery. Because there was an earlier miss in {kind}, "
                                f"complete {remaining} more secure attempt(s) in this same category before advancing."
                            )
                        else:
                            st.warning(
                                f"Stay on {kind}. The next category remains locked until the student can apply the advice securely."
                            )

                        if st.button(f"Generate another {kind} question", use_container_width=True):
                            try:
                                with st.spinner(f"Creating another {kind} question focused on the same gap..."):
                                    followup = generate_followup_practice_question(
                                        track_label=track_label,
                                        kind=kind,
                                        previous_question=pq,
                                        previous_working=st.session_state.ai_practice_last_working,
                                        evaluation=evaluation,
                                        original_gap=analysis.misconception_or_gap,
                                        api_key=explicit_key,
                                        model=model,
                                    )
                                st.session_state.ai_practice_current_question = followup
                                st.session_state.ai_practice_evaluation = None
                                st.session_state.ai_practice_last_working = ""
                                st.session_state.ai_practice_ready_to_advance = False
                                st.session_state.ai_practice_question_version += 1
                                st.rerun()
                            except GeminiTutorError as exc:
                                st.error(str(exc))

                with st.expander("Reveal reference answer and worked solution"):
                    st.markdown("**Answer**")
                    render_mathio(pq.answer)
                    st.markdown("**Worked solution**")
                    for i, line in enumerate(pq.worked_solution, 1):
                        st.caption(f"Step {i}")
                        render_mathio(line)

    # ---------- Offline generated practice ----------

    # ---------- Batch / class trend analysis ----------
    _QUESTION_COMMAND_RE = re.compile(
        r"(?i)^(Simplify|Evaluate|Calculate|Find|Solve|Expand|Factorise|Factorize|Divide|"
        r"Express|Write|State|Determine|Show that|Prove that|Given that|Hence|Complete|"
        r"Sketch|Draw|Plot|Construct|Estimate|Arrange|Compare|Convert|Factor|Substitute)\b"
    )

    # Mathematical fragments that should be rendered with MathIO rather than as prose.
    _GENERIC_MATH_TOKEN_RE = re.compile(
        r"""
        (?:
            # LaTeX/MathIO commands and symbols
            \\(?:frac|sqrt|angle|theta|alpha|beta|gamma|delta|pi|sin|cos|tan|arcsin|arccos|arctan|log|ln|times|div|leq|geq|neq|pm|parallel|perp)\b[^\s,.;:!?]*
            |
            # Explicit equations / inequalities
            [A-Za-z][A-Za-z0-9_]*\s*(?:=|≤|≥|<|>)\s*[^,.;:!?]+
            |
            # Ratios / proportions
            \d+(?:\.\d+)?\s*:\s*\d+(?:\.\d+)?(?:\s*:\s*\d+(?:\.\d+)?)*
            |
            # Coordinates
            \(\s*[-+]?\d+(?:\.\d+)?\s*,\s*[-+]?\d+(?:\.\d+)?\s*\)
            |
            # Algebraic expressions with powers/brackets/operators
            (?:[-+]?\d*(?:\.\d+)?[A-Za-z](?:\^\{?[-+]?\d+\}?|\^\(?[-+]?\d+\)?)?)
            (?:\s*[+\-×÷*/]\s*(?:[-+]?\d*(?:\.\d+)?[A-Za-z0-9](?:\^\{?[-+]?\d+\}?)?|\([^)]*\)))+
            |
            # Standalone powers / standard form
            \d+(?:\.\d+)?\s*(?:×|\\times|x)\s*10\s*\^\s*[-+]?\d+
            |
            [A-Za-z0-9]+\s*\^\s*\{?[-+]?\d+\}?
            |
            # Fractions written linearly
            \d+(?:\.\d+)?\s*/\s*\d+(?:\.\d+)?
            |
            # Percentages and measured values
            [-+]?\d+(?:\.\d+)?\s*(?:%|°|cm|mm|m|km|g|kg|s|h)(?:\^2|\^3)?
            |
            # Number sequences / comma-separated numeric data
            [-+]?\d+(?:\.\d+)?(?:\s*,\s*[-+]?\d+(?:\.\d+)?){2,}(?:\s*,?\s*(?:\.\.\.|…))?
        )
        """,
        re.VERBOSE,
    )


    def _normalise_question_source(text: str) -> str:
        """Convert common verbal maths into symbolic forms without changing ordinary prose."""
        value = re.sub(r"\s+", " ", str(text or "")).strip()
        value = re.sub(r"(?<!\\)\btheta\b", r"\\theta", value, flags=re.IGNORECASE)
        value = re.sub(r"\b(\d+(?:\.\d+)?)\s+degrees?\b", r"\1^{\\circ}", value, flags=re.IGNORECASE)
        value = re.sub(r"\b([A-Za-z])\s+squared\b", lambda m: f"{m.group(1)}^2", value, flags=re.IGNORECASE)
        value = re.sub(r"\b([A-Za-z])\s+cubed\b", lambda m: f"{m.group(1)}^3", value, flags=re.IGNORECASE)

        # e.g. y = x squared divided by (2x+1)
        value = re.sub(
            r"\b([A-Za-z])\s*=\s*([A-Za-z])\^2\s+divided\s+by\s+\(([^)]+)\)",
            lambda m: rf"{m.group(1)} = \frac{{{m.group(2)}^2}}{{{m.group(3)}}}",
            value,
            flags=re.IGNORECASE,
        )
        value = value.replace("...", r"\ldots")
        return value


    def _split_question_text_math(text: str) -> list[tuple[str, str]]:
        """Return ordered ('text'|'math', content) chunks for any question type."""
        source = _normalise_question_source(text)
        if not source:
            return []

        chunks: list[tuple[str, str]] = []
        cursor = 0

        for match in _GENERIC_MATH_TOKEN_RE.finditer(source):
            # Avoid treating a bare single variable inside normal prose as a display equation.
            frag = match.group(0).strip()
            if not frag:
                continue

            if match.start() > cursor:
                prose = source[cursor:match.start()].strip()
                if prose:
                    chunks.append(("text", prose))

            chunks.append(("math", frag))
            cursor = match.end()

        if cursor < len(source):
            prose = source[cursor:].strip()
            if prose:
                chunks.append(("text", prose))

        if not chunks:
            return [("text", source)]

        # Merge adjacent chunks of the same type.
        merged: list[tuple[str, str]] = []
        for kind, content in chunks:
            if merged and merged[-1][0] == kind:
                sep = " " if kind == "text" else r"\quad "
                merged[-1] = (kind, merged[-1][1] + sep + content)
            else:
                merged.append((kind, content))
        return merged


    def _offline_prompt_parts(prompt: str) -> tuple[str, str, str]:
        """Compatibility wrapper retained for existing callers."""
        chunks = _split_question_text_math(prompt)
        prose_before = []
        maths = []
        prose_after = []
        seen_math = False

        for kind, content in chunks:
            if kind == "math":
                seen_math = True
                maths.append(content)
            elif not seen_math:
                prose_before.append(content)
            else:
                prose_after.append(content)

        return (
            " ".join(prose_before).strip(),
            r"\quad ".join(maths).strip(),
            " ".join(prose_after).strip(),
        )



    def _question_chunk_to_inline_text(kind: str, content: str) -> str:
        """Convert short maths chunks to readable inline Unicode/Markdown-safe text."""
        if kind == "text":
            return str(content or "").strip()

        value = str(content or "").strip()
        if not value:
            return ""

        # For short question fragments, preserve horizontal sentence flow instead of
        # mounting a separate MathIO component for every number/unit.
        value = value.replace(r"\theta", "θ")
        value = value.replace(r"\pi", "π")
        value = value.replace(r"\times", "×")
        value = value.replace(r"\div", "÷")
        value = value.replace(r"\leq", "≤")
        value = value.replace(r"\geq", "≥")
        value = value.replace(r"\neq", "≠")
        value = value.replace(r"\pm", "±")
        value = value.replace(r"\ldots", "…")
        value = re.sub(r"\^\{([^{}]+)\}", r"^\1", value)
        value = re.sub(r"\^\{?\\circ\}?", "°", value)
        value = re.sub(r"(?<=\d)\s*(km|cm|mm|kg|g|m|s|h)\b", r" \1", value)
        return value


    def _is_large_standalone_math(content: str) -> bool:
        """Decide when a maths fragment deserves its own MathIO display line."""
        value = str(content or "").strip()
        if not value:
            return False

        # Display genuinely structural maths separately.
        if any(token in value for token in (r"\frac", r"\sqrt", r"\int", r"\sum", r"\begin{", r"\matrix")):
            return True
        if len(value) > 42 and re.search(r"[=+\-*/^]", value):
            return True
        if value.count("=") >= 2:
            return True
        return False


    def render_question_text_mathio(prompt: str) -> None:
        """Render a question compactly, keeping prose and short maths on the same line."""
        chunks = _split_question_text_math(prompt)
        if not chunks:
            return

        inline_parts: list[str] = []

        def flush_inline() -> None:
            if not inline_parts:
                return
            text = " ".join(part for part in inline_parts if part).strip()
            text = re.sub(r"\s+([,.;:!?])", r"\1", text)
            text = re.sub(r"\(\s+", "(", text)
            text = re.sub(r"\s+\)", ")", text)
            if text:
                st.markdown(text)
            inline_parts.clear()

        for kind, content in chunks:
            if kind == "math" and _is_large_standalone_math(content):
                flush_inline()
                render_mathio(content)
            else:
                piece = _question_chunk_to_inline_text(kind, content)
                if piece:
                    inline_parts.append(piece)

        flush_inline()




def _offline_prompt_mathio_markup(prompt: str) -> str:
    """Prepare an Offline Practice prompt for mixed prose + MathIO rendering."""
    text = str(prompt or "").strip()
    if not text:
        return ""

    # Normalise generated notation.
    text = text.replace("**", "^")
    text = re.sub(r"(?<!\\)\bpi\b", r"\\pi", text, flags=re.IGNORECASE)
    text = re.sub(r"(?<!\\)\btheta\b", r"\\theta", text, flags=re.IGNORECASE)
    text = re.sub(
        r"(\d+(?:\.\d+)?)\s*degrees\b",
        lambda m: rf"\({m.group(1)}^{{\circ}}\)",
        text,
        flags=re.IGNORECASE,
    )

    # Keep standard-form products as one mathematical expression.
    standard_form_command = re.match(
        r"(?i)^(Calculate|Evaluate)\s+(.+?)(\.\s+(?:Give|Express|State)\b.*)$",
        text,
    )
    if standard_form_command and (
        r"\times" in standard_form_command.group(2)
        or re.search(r"10\s*\^", standard_form_command.group(2))
    ):
        command = standard_form_command.group(1)
        maths = standard_form_command.group(2).strip()
        tail = standard_form_command.group(3)
        text = rf"{command} \({maths}\){tail}"

    # Common Offline Practice algebra prompts:
    # "Given p = 3 and q = -2, evaluate 2p^2 - 3q."
    # Keep prose as prose and force every assignment/expression through MathIO.
    text = re.sub(
        r"(?i)\b([A-Za-z])\s*=\s*(-?\d+(?:\.\d+)?)\b",
        lambda m: rf"\({m.group(1)}={m.group(2)}\)",
        text,
    )

    command_pattern = re.compile(
        r"(?i)\b(evaluate|simplify|calculate|expand|factorise|factorize)\s+"
        r"([^.;]+)"
    )

    def command_repl(match):
        command = match.group(1)
        fragment = match.group(2).strip()

        # Stop before ordinary-English answer instructions if they occur in the same sentence.
        tail = re.search(
            r"(?i)\s+(give|express|state|write|correct|hence|where|when|if)\b",
            fragment,
        )
        if tail:
            maths = fragment[:tail.start()].strip()
            prose_tail = fragment[tail.start():]
        else:
            maths = fragment
            prose_tail = ""

        # Do not wrap prose-only material accidentally.
        if not re.search(r"[A-Za-z0-9][+\-*/^=()]|[+\-*/^=][A-Za-z0-9]", maths):
            return match.group(0)

        maths = maths.replace("**", "^")
        maths = re.sub(r"(?<=\d)\*(?=[A-Za-z(])", "", maths)
        maths = maths.replace("*", r"\times ")
        return rf"{command} \({maths}\){prose_tail}"

    # Apply only outside already-delimited MathIO.
    rebuilt = []
    cursor = 0
    for m in _MATHIO_MIXED_PATTERN.finditer(text):
        rebuilt.append(command_pattern.sub(command_repl, text[cursor:m.start()]))
        rebuilt.append(m.group(0))
        cursor = m.end()
    rebuilt.append(command_pattern.sub(command_repl, text[cursor:]))
    text = "".join(rebuilt)

    # Natural-language logarithms -> MathIO.
    text = re.sub(
        r"log\s+base\s+([^ ]+)\s+of\s+([^ ,.;]+)\s+equals\s+([A-Za-z][A-Za-z0-9_]*)",
        lambda m: rf"\(\log_{{{m.group(1)}}}({m.group(2)}) = {m.group(3)}\)",
        text,
        flags=re.IGNORECASE,
    )

    # Function notation must become real MathIO rather than raw words like sqrt(27).
    text = re.sub(
        r"(?<!\\)sqrt\(([^()]+)\)",
        lambda m: rf"\(\sqrt{{{m.group(1)}}}\)",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"(?<!\\)\bsin\(([^()]+)\)",
        lambda m: rf"\(\sin({m.group(1)})\)",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"(?<!\\)\bcos\(([^()]+)\)",
        lambda m: rf"\(\cos({m.group(1)})\)",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"(?<!\\)\btan\(([^()]+)\)",
        lambda m: rf"\(\tan({m.group(1)})\)",
        text,
        flags=re.IGNORECASE,
    )

    # Algebraic equations that are still plain text.
    equation_pattern = re.compile(
        r"(?<![\w\\])([A-Za-z][A-Za-z0-9_]*(?:\([^)]*\))?\s*=\s*[^,.;:]+)"
    )

    def equation_repl(match):
        fragment = match.group(1).strip()
        tail = re.search(
            r"(?i)\s+(for|where|when|with|from|over|giving|correct)\b",
            fragment,
        )
        if tail:
            maths = fragment[:tail.start()].strip()
            prose = fragment[tail.start():]
        else:
            maths, prose = fragment, ""
        maths = re.sub(r"(?<=\d)\*(?=[A-Za-z(])", "", maths)
        maths = maths.replace("*", r"\times ")
        return rf"\({maths}\){prose}"

    # Apply equation conversion without re-wrapping existing MathIO fragments.
    pieces = []
    cursor = 0
    for m in _MATHIO_MIXED_PATTERN.finditer(text):
        prose = text[cursor:m.start()]
        pieces.append(equation_pattern.sub(equation_repl, prose))
        pieces.append(m.group(0))
        cursor = m.end()
    pieces.append(equation_pattern.sub(equation_repl, text[cursor:]))

    return re.sub(r"\s{2,}", " ", "".join(pieces)).strip()



def render_offline_practice_prompt(prompt: str) -> None:
    """Render offline questions as normal prose with mathematical fragments in MathIO."""
    value = _offline_prompt_mathio_markup(prompt)
    if not value:
        return
    render_mathio_mixed(value)



def render_learning_outcome_mixed_mathio(value: str) -> None:
    """Render learning-outcome prose as text and only mathematical fragments through MathIO."""
    text = str(value or "").strip()
    if not text:
        return

    # Normalize common syllabus notation first.
    text = text.replace("π", r"\pi").replace("θ", r"\theta")
    text = re.sub(r"(?<!\\)\bpi\b", r"\\pi", text, flags=re.IGNORECASE)
    text = re.sub(r"(?<!\\)\btheta\b", r"\\theta", text, flags=re.IGNORECASE)

    # Candidate maths fragments. Keep surrounding syllabus prose out of MathIO.
    math_patterns = [
        r"y\s*=\s*[^,;]+",
        r"[A-Za-z]\s*=\s*[^,;]+",
        r"\\(?:sin|cos|tan|log|ln|sqrt|frac)\b[^,;]*",
        r"\b(?:sin|cos|tan)\s+[A-Za-z0-9()^+\-*/\\ ]+",
        r"\b\d+\s*\^\s*[A-Za-z0-9+\-]+",
    ]

    matches = []
    for pattern in math_patterns:
        for m in re.finditer(pattern, text, flags=re.IGNORECASE):
            s, e = m.span()
            # Trim English continuation phrases from the candidate.
            frag = text[s:e]
            cut = re.search(
                r"(?i)\s+(?:understand|determine|solve|use|know|sketch|find|calculate|"
                r"principle|principal|values|angles|students|and\s+determine)\b",
                frag,
            )
            if cut:
                e = s + cut.start()
            if e > s:
                matches.append((s, e))

    # Merge overlaps.
    matches.sort()
    merged = []
    for s, e in matches:
        if not merged or s > merged[-1][1]:
            merged.append([s, e])
        else:
            merged[-1][1] = max(merged[-1][1], e)

    if not merged:
        st.markdown(text)
        return

    cursor = 0
    for s, e in merged:
        if s > cursor:
            prose = text[cursor:s].strip()
            if prose:
                st.markdown(prose)

        maths = text[s:e].strip(" ,;")
        if maths:
            # Clean Python-like operators and spacing before MathIO.
            maths = maths.replace("**", "^").replace("*", "")
            maths = re.sub(r"\s{2,}", " ", maths)
            render_mathio(_normalise_math_variable_italics(maths))

        cursor = e

    if cursor < len(text):
        prose = text[cursor:].strip()
        if prose:
            st.markdown(prose)



def _offline_statistics_graph_spec(question):
    payload=getattr(question,"statistics_graph",None)
    if not payload:
        return None
    if isinstance(payload,dict):
        return SimpleNamespace(**payload)
    return payload


# ---------- Coverage ----------
if role_mode == "For Teacher":
    with syllabus_tab:
        info = selected_track_info(track_label)
        if info["year"] == 2027:
            st.subheader(f'2027 SEC {info["level"]} {info["subject"]} syllabus coverage')
            st.caption(
                f'Official SEC subject code: {info["subject_code"]} · '
                f'2026-and-earlier reference code: {info["reference_2026"]}'
            )
            st.write(
                "The tutor uses the selected SEC subject level as context for question feasibility, mathematical verification, "
                "student-working analysis, guided solving, and adaptive practice."
            )
            cols = st.columns(len(info["strands"]))
            for col, strand in zip(cols, info["strands"]):
                col.metric("Syllabus strand", strand)

            if info["subject"] == "Additional Mathematics":
                st.info(
                    "2027 SEC Additional Mathematics is organised into Algebra, Geometry and Trigonometry, and Calculus. "
                    "Gemini online mode is enabled for the full selected subject context. The no-credit deterministic question "
                    "generator has not yet been expanded to Additional Mathematics, so it is intentionally disabled for this track."
                )
            else:
                selected = topics_for_track(tcode)
                strong = sum(1 for t in selected if _topic_offline_support(t) == "Strong")
                partial = sum(1 for t in selected if _topic_offline_support(t) == "Partial")
                c1, c2, c3 = st.columns(3)
                c1.metric("Topics mapped", len(selected))
                c2.metric("Outcome-mapped offline topics", strong)
                c3.metric("Legacy/partial topics", partial)
                for strand in ("Number and Algebra", "Geometry and Measurement", "Statistics and Probability"):
                    st.markdown(f"### {strand}")
                    for t in [x for x in selected if x.strand == strand]:
                        badge = "✅ Strong" if _topic_offline_support(t) == "Strong" else "🟡 Partial"
                        with st.expander(f"{official_topic_code(tcode, t.code)} · {t.name} — {badge}"):
                            st.write(_topic_coverage_note(t))
        else:
            st.subheader("2026 Singapore Mathematics syllabus coverage")
            st.write(
                "Offline generated practice spans Number and Algebra, Geometry and Measurement, and Statistics and Probability. "
                "Gemini online mode broadens interpretation to uploaded handwriting, diagrams, PDFs, word problems and alternative methods."
            )
            selected = topics_for_track(tcode)
            strong = sum(1 for t in selected if _topic_offline_support(t) == "Strong")
            partial = sum(1 for t in selected if _topic_offline_support(t) == "Partial")
            c1, c2, c3 = st.columns(3)
            c1.metric("Topics mapped", len(selected))
            c2.metric("Outcome-mapped offline topics", strong)
            c3.metric("Legacy/partial topics", partial)
            for strand in ("Number and Algebra", "Geometry and Measurement", "Statistics and Probability"):
                st.markdown(f"### {strand}")
                for t in [x for x in selected if x.strand == strand]:
                    badge = "✅ Strong" if _topic_offline_support(t) == "Strong" else "🟡 Partial"
                    with st.expander(f"{official_topic_code(tcode, t.code)} · {t.name} — {badge}"):
                        st.write(_topic_coverage_note(t))

        st.warning(
            "Coverage means the tutor has support for these areas; it does not guarantee perfect interpretation of every examination question. "
            "For high-stakes assessment decisions, verify AI feedback against the official syllabus/marking scheme or a teacher."
        )

# ---------- Progress ----------
if role_mode == "For Teacher":
    st.markdown("---")
    


def render_text_with_mathio(value: str) -> None:
    """Render saved student notes using the app's existing mixed text/MathIO renderer."""
    text = str(value or "").strip()
    if not text:
        return
    try:
        render_guidance_mixed_mathio(text)
    except Exception:
        st.write(text)



def _normalise_student_image_bytes(data: bytes) -> bytes | None:
    """Validate an uploaded/camera image and convert it to PNG for reliable export."""
    try:
        if not data or len(data) < 32:
            return None
        with Image.open(BytesIO(data)) as image:
            image.load()
            # Convert formats/modes that python-docx/reportlab may not embed reliably.
            if image.mode not in ("RGB", "RGBA"):
                image = image.convert("RGB")
            output = BytesIO()
            image.save(output, format="PNG", optimize=True)
            png = output.getvalue()
            return png if len(png) >= 32 else None
    except Exception:
        return None


def _valid_student_image_bytes(data: bytes) -> bool:
    return _normalise_student_image_bytes(data) is not None


def _student_picture_input_version() -> int:
    return int(st.session_state.setdefault("student_picture_input_version", 0))


def _reset_student_picture_inputs() -> None:
    # Do not mutate widget-bound session keys after the widgets have been created.
    # Incrementing the version makes Streamlit create fresh camera/uploader widgets on rerun.
    st.session_state.student_picture_input_version = _student_picture_input_version() + 1


def _student_download_basename() -> str:
    return f"{date.today().isoformat()}_Math Lesson"



def _student_notes():
    return st.session_state.setdefault("student_lesson_notes", [])


def _student_notes_docx():
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)
    doc.add_heading("Lesson Notes", 1)

    for item in _student_notes():
        if item.get("kind") == "text":
            paragraph = doc.add_paragraph()
            try:
                append_word_mixed_math(paragraph, str(item.get("content", "")))
            except Exception:
                paragraph.add_run(str(item.get("content", "")))
        elif item.get("kind") == "image":
            image_bytes = _normalise_student_image_bytes(item.get("content", b""))
            if image_bytes is not None:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    p.add_run().add_picture(BytesIO(image_bytes), width=Cm(14))
                except Exception:
                    doc.add_paragraph("[Image could not be embedded]")
                if item.get("caption"):
                    cp = doc.add_paragraph(str(item["caption"]))
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_paragraph("[Saved image could not be read]")
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _student_notes_pdf():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm as rcm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RImage

    bio = BytesIO()
    doc = SimpleDocTemplate(
        bio,
        pagesize=A4,
        rightMargin=1.5*rcm,
        leftMargin=1.5*rcm,
        topMargin=1.5*rcm,
        bottomMargin=1.5*rcm,
    )
    styles = getSampleStyleSheet()
    caption_style = styles["BodyText"].clone("StudentNoteCaption")
    caption_style.fontSize = 9
    caption_style.leading = 11
    caption_style.textColor = colors.HexColor("#6b7280")
    caption_style.spaceBefore = 2
    caption_style.spaceAfter = 8
    story = [Paragraph("Lesson Notes", styles["Title"]), Spacer(1, 8)]

    def _pdf_safe_text(value: str) -> str:
        return (
            str(value or "")
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    for item in _student_notes():
        if item.get("kind") == "text":
            safe = (
                str(item.get("content", ""))
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace("\n", "<br/>")
            )
            story += [Paragraph(safe, styles["BodyText"]), Spacer(1, 8)]
        elif item.get("kind") == "image":
            image_bytes = _normalise_student_image_bytes(item.get("content", b""))
            if image_bytes is not None:
                try:
                    story += [
                        RImage(
                            BytesIO(image_bytes),
                            width=15*rcm,
                            height=10*rcm,
                            kind="proportional",
                        ),
                        Spacer(1, 6),
                    ]
                except Exception:
                    story += [Paragraph("[Image could not be embedded]", styles["BodyText"]), Spacer(1, 6)]
                if item.get("caption"):
                    story += [Paragraph(_pdf_safe_text(item["caption"]), caption_style), Spacer(1, 8)]
    doc.build(story)
    return bio.getvalue()


def _clear_student_notes_after_download():
    st.session_state.student_lesson_notes = []
    st.session_state.pop("student_note_draft", None)
    _reset_student_picture_inputs()




def _student_ask_reset() -> None:
    st.session_state.student_ask_guided_solution = None
    st.session_state.student_ask_error = ""
    st.session_state.student_ask_hint_count = 0
    st.session_state.student_ask_show_full_solution = False
    st.session_state.student_ask_reveal_step = 0
    st.session_state.student_ask_signature = ""


def _student_ask_next_hint(total_hints: int) -> None:
    current = int(st.session_state.get("student_ask_hint_count", 0))
    st.session_state.student_ask_hint_count = min(current + 1, total_hints)


def _student_ask_show_full() -> None:
    st.session_state.student_ask_show_full_solution = True
    if int(st.session_state.get("student_ask_reveal_step", 0)) == 0:
        st.session_state.student_ask_reveal_step = 1


def _student_ask_back_to_hints() -> None:
    st.session_state.student_ask_show_full_solution = False


def _student_ask_next_step(total_steps: int) -> None:
    current = int(st.session_state.get("student_ask_reveal_step", 0))
    st.session_state.student_ask_reveal_step = min(current + 1, total_steps)


def _student_ask_show_all_steps(total_steps: int) -> None:
    st.session_state.student_ask_reveal_step = total_steps


def _student_ask_question_text(prose: str, latex_lines: list[str]) -> str:
    used = [str(line).strip() for line in latex_lines if str(line).strip()]
    maths = "\n".join(rf"\({line}\)" for line in used)
    return (str(prose or "").strip() + ("\n\n" + maths if maths else "")).strip()


def _compact_mathio_guidance_text(value: str) -> None:
    """Compact Ask guidance: prose as text, maths through the real MathIO renderer."""
    text = re.sub(r"\s+", " ", clean_guidance_text(value)).strip()
    if not text:
        return

    # Normalise common Gemini wording.
    text = re.sub(r"(?i)\bdegrees?\b", "°", text)

    # Build proper MathIO delimiters. render_mathio_mixed() then sends only these
    # fragments through the registered MathLive/MathIO component.
    text = re.sub(
        r"(?i)\bangle\s+([A-Za-z]{3})\s*=\s*(-?\d+(?:\.\d+)?)\s*°",
        lambda m: rf"\(\angle {m.group(1)}={m.group(2)}^\circ\)",
        text,
    )
    text = re.sub(
        r"\b([A-Z]{3})\s*=\s*(-?\d+(?:\.\d+)?)\s*°",
        lambda m: rf"\(\angle {m.group(1)}={m.group(2)}^\circ\)",
        text,
    )
    text = re.sub(
        r"\b([A-Za-z]{1,4})\s*=\s*(-?\d+(?:\.\d+)?)\s*cm\b",
        lambda m: rf"\({m.group(1)}={m.group(2)}\text{{ cm}}\)",
        text,
        flags=re.I,
    )

    # The mixed renderer keeps ordinary words outside MathIO and mathematical
    # fragments inside MathIO, so prose spacing remains natural.
    render_mathio_mixed(text)



def _compact_given_items(items) -> list[str]:
    cleaned = []
    seen = set()
    for item in (items or []):
        text = re.sub(r"\s+", " ", clean_guidance_text(item)).strip(" ,;")
        if text and text.lower() not in seen:
            seen.add(text.lower())
            cleaned.append(text)
    return cleaned



def _ask_guidance_mathio_text(value: str) -> None:
    """Render Ask Math Advisor guidance with prose + MathIO consistently."""
    text = re.sub(r"\s+", " ", clean_guidance_text(value)).strip()
    if not text:
        return

    # Normalise common mathematical wording returned by Gemini.
    text = re.sub(r"(?i)\bdegrees?\b", "°", text)

    # angle PQR = 90° / PQR = 90°
    text = re.sub(
        r"(?i)\bangle\s+([A-Za-z]{3})\s*=\s*(-?\d+(?:\.\d+)?)\s*°",
        lambda m: rf"\(\angle {m.group(1)}={m.group(2)}^\circ\)",
        text,
    )
    text = re.sub(
        r"\b([A-Z]{3})\s*=\s*(-?\d+(?:\.\d+)?)\s*°",
        lambda m: rf"\(\angle {m.group(1)}={m.group(2)}^\circ\)",
        text,
    )

    # Lengths / assignments.
    text = re.sub(
        r"\b([A-Za-z]{1,4})\s*=\s*(-?\d+(?:\.\d+)?)\s*cm\b",
        lambda m: rf"\({m.group(1)}={m.group(2)}\text{{ cm}}\)",
        text,
        flags=re.I,
    )
    text = re.sub(
        r"\b([A-Za-z])\s*=\s*(-?\d+(?:\.\d+)?)\b",
        lambda m: rf"\({m.group(1)}={m.group(2)}\)",
        text,
    )

    # Common algebraic expressions not already delimited.
    # Examples: 2x + 5 = 17, x^2 - 9, 3(a+2)
    if r"\(" not in text:
        expr_pat = re.compile(
            r"(?<!\w)"
            r"((?:-?\d+(?:\.\d+)?\s*)?[A-Za-z](?:\^\{?[-+]?\d+\}?)?"
            r"(?:\s*[+\-*/=]\s*(?:-?\d+(?:\.\d+)?|[A-Za-z](?:\^\{?[-+]?\d+\}?)?))+)"
            r"(?!\w)"
        )
        text = expr_pat.sub(lambda m: rf"\({m.group(1)}\)", text)

    # Fractions and simple numeric calculations.
    text = re.sub(
        r"(?<![\w\\])(-?\d+(?:\.\d+)?\s*[+\-*/=]\s*-?\d+(?:\.\d+)?(?:\s*[+\-*/=]\s*-?\d+(?:\.\d+)?)*)",
        lambda m: rf"\({m.group(1)}\)",
        text,
    )

    render_mathio_mixed(text)


def _ask_guidance_item_mathio(value: str) -> None:
    """Bullet-style Ask guidance item using the MathIO-aware renderer."""
    with st.container():
        _ask_guidance_mathio_text(value)


def _ask_guidance_step_mathio(index: int, step) -> None:
    """Worked step with MathIO-aware prompt/reasoning/equations."""
    with st.container(border=True):
        st.markdown(f"**Step {index}**")
        prompt = str(getattr(step, "prompt", "") or getattr(step, "instruction", "") or "").strip()
        if prompt:
            _ask_guidance_mathio_text(prompt)

        explanation = str(
            getattr(step, "explanation", "")
            or getattr(step, "reasoning", "")
            or getattr(step, "why", "")
            or ""
        ).strip()
        if explanation:
            _ask_guidance_mathio_text(explanation)

        equations = getattr(step, "equations", None) or getattr(step, "math", None) or []
        if isinstance(equations, str):
            equations = [equations]
        for equation in equations:
            if str(equation or "").strip():
                render_mathio(str(equation).strip())


def render_student_ask_guided_solution(g: GuidedSolution) -> None:
    """Compact hint-first student rendering with MathIO."""
    with st.container(border=True):
        st.markdown("### 🎯 What the question is asking")
        goal_text = re.sub(r"\s+", " ", str(g.interpreted_goal or "")).strip()
        if len(goal_text) > 220:
            # Keep the first meaningful clause; the exact measurements remain under Given.
            clauses = re.split(r"(?<=[.;])\s+|;\s*", goal_text)
            goal_text = clauses[0].strip() if clauses else goal_text[:220].rstrip() + "…"
        _compact_mathio_guidance_text(goal_text)

        known = _compact_given_items(g.known_information)
        if known:
            st.markdown("**Given:**")
            cols = st.columns(2 if len(known) >= 2 else 1)
            for index, item in enumerate(known[:6]):
                with cols[index % len(cols)]:
                    _ask_guidance_mathio_text(item)

        concepts = _compact_given_items(g.concepts_to_use)
        if concepts:
            with st.expander("Useful ideas / formulae", expanded=False):
                for item in concepts:
                    _ask_guidance_mathio_text(item)

    show_full = bool(st.session_state.get("student_ask_show_full_solution", False))

    if not show_full:
        st.markdown("### 💡 Guided hints")
        st.info(
            "Try each hint before revealing the next one. "
            "The full worked solution is available whenever you need it."
        )

        if str(g.first_question_for_student or "").strip():
            with st.container(border=True):
                st.markdown("**Start here**")
                _ask_guidance_mathio_text(g.first_question_for_student)

        hint_count = int(st.session_state.get("student_ask_hint_count", 0))
        total_hints = len(g.hint_ladder or [])

        for index, hint in enumerate((g.hint_ladder or [])[:hint_count], 1):
            with st.container(border=True):
                st.markdown(f"**Hint {index}**")
                _ask_guidance_mathio_text(hint)

        c1, c2 = st.columns(2)
        with c1:
            if hint_count < total_hints:
                st.button(
                    "Show next hint",
                    key="student_ask_next_hint_button",
                    type="primary" if hint_count == 0 else "secondary",
                    use_container_width=True,
                    on_click=_student_ask_next_hint,
                    args=(total_hints,),
                )
            else:
                st.success("All available hints are shown.")

        with c2:
            st.button(
                "Show full solution",
                key="student_ask_full_solution_button",
                use_container_width=True,
                on_click=_student_ask_show_full,
            )

    else:
        steps = list(g.guided_steps or [])
        total_steps = len(steps)
        reveal = int(st.session_state.get("student_ask_reveal_step", 1))
        reveal = max(1, min(reveal, total_steps)) if total_steps else 0

        st.markdown("### ✅ Full worked solution")

        c1, c2, c3 = st.columns(3)
        with c1:
            st.button(
                "← Back to hints",
                key="student_ask_back_hints_button",
                use_container_width=True,
                on_click=_student_ask_back_to_hints,
            )
        with c2:
            if reveal < total_steps:
                st.button(
                    "Reveal next step",
                    key="student_ask_next_step_button",
                    use_container_width=True,
                    on_click=_student_ask_next_step,
                    args=(total_steps,),
                )
        with c3:
            if reveal < total_steps:
                st.button(
                    "Show all steps",
                    key="student_ask_show_all_button",
                    use_container_width=True,
                    on_click=_student_ask_show_all_steps,
                    args=(total_steps,),
                )

        if not steps:
            st.warning("No worked solution steps were returned. Try asking the question again.")
        else:
            for index, step in enumerate(steps, 1):
                if index <= reveal:
                    _ask_guidance_step_mathio(index, step)

            if reveal >= total_steps:
                with st.container(border=True):
                    st.markdown("### Final answer")
                    if str(g.final_answer_mathio or "").strip():
                        render_mathio(str(g.final_answer_mathio).strip())
                    else:
                        st.caption("No separate final-answer expression was returned.")

                if g.common_pitfalls:
                    with st.expander("Common mistakes to avoid", expanded=False):
                        for pitfall in g.common_pitfalls:
                            _ask_guidance_mathio_text(pitfall)


if role_mode == "For Student":
    with student_practice_tab:
        st.subheader("No-credit syllabus-generated practice")
        st.caption("This tab never calls Gemini. It keeps working even if the API key is missing or a free-tier quota is reached.")
        st.caption("Questions are varied by level, topic, learning outcome and cognitive demand using the compiled G1/G2/G3 learning outcomes.")
        available = topics_for_track(tcode)
        topic_labels = {f"{official_topic_code(tcode, t.code)} · {t.name}": t.code for t in available}
        c1, c2, c3 = st.columns([1.6, 1, 1])
        with c1:
            topic_label = st.selectbox("Topic", list(topic_labels.keys()), key="topic_choice")
        with c2:
            difficulty = st.selectbox(
                "Difficulty",
                ["Foundation", "Similar", "Stretch"],
                index=1,
                help=(
                    "Foundation: direct concept/fluency. Similar: standard syllabus application. "
                    "Stretch: reverse, multi-step, reasoning or less-familiar application."
                ),
            )
        with c3:
            st.write("")
            st.write("")
            if st.button("Generate question", type="primary", use_container_width=True):
                make_new_question(tcode, topic_labels[topic_label], difficulty)
                st.rerun()

        question: Question | None = st.session_state.question
        if question is None or question.track != tcode:
            st.info("Choose a topic and click **Generate question**.")
        else:
            st.markdown(f"### {official_topic_code(question.track, question.topic_code)} · {question.topic_name}")
            st.caption(f"{question.strand} · {question.difficulty}")

            # Keep the complete expression in one MathIO block so operators and powers
            # appear horizontally instead of as separate rows with large blank spaces.
            with st.container(border=True):
                render_offline_practice_prompt(question.prompt)
                show_context_image_for_text(question.prompt)

                offline_graph_spec = _offline_statistics_graph_spec(question)
                if offline_graph_spec is not None:
                    show_statistics_graph(
                        offline_graph_spec,
                        caption="Cumulative frequency curve" if getattr(offline_graph_spec, "graph_type", "") == "cumulative_frequency" else "",
                        completed=True,
                    )

            st.markdown("**Learning outcome focus:**")
            render_learning_outcome_mixed_mathio(question.target_skill)

            if st.button("Show next hint", key="show_hint"):
                st.session_state.hint_level = min(len(question.hints), st.session_state.hint_level + 1)
            for i in range(st.session_state.hint_level):
                st.markdown(f"**Hint {i+1}:**")
                render_mathio_mixed(question.hints[i])
            if st.session_state.hint_level == 0:
                st.caption("Try the question before revealing a hint.")

            student_scientific_calculator(key_base="offline_practice_calculator")
            geogebra_external_tools(
                question_text=str(getattr(question, "question", "") or ""),
                key_base="offline_practice_geogebra",
            )
            working, working_mode, working_offline = working_input(
                "Your working and answer",
                text_key="practice_working",
                format_key="practice_working_format",
                height=190,
                plain_placeholder="Show the important steps, one line at a time where possible.",
            )
            working_to_check = working_offline if working_mode == "Equation editor" else working
            if st.button("Check my reasoning offline", type="primary", use_container_width=True):
                if not working_to_check.strip():
                    st.error("Enter your working and answer first.")
                else:
                    result = evaluate_attempt(question, working_to_check)
                    st.session_state.attempt_result = result
                    record_history(question, result)
                    st.rerun()

            result: AttemptResult | None = st.session_state.attempt_result
            if result is not None:
                render_attempt(result)

            st.markdown("---")
            reveal = st.checkbox("Reveal verified answer and worked solution", key="reveal_solution")
            if reveal:
                st.markdown("**Answer**")
                render_guidance_mixed_mathio(question.answer_display)
                st.markdown("**Worked solution**")
                for i, line in enumerate(question.worked_solution, 1):
                    st.caption(f"Step {i}")
                    render_guidance_mixed_mathio(line)

            cnext1, cnext2 = st.columns(2)
            with cnext1:
                if st.button("Generate a similar question", use_container_width=True):
                    seed = int(datetime.now().timestamp() * 1000) + st.session_state.seed_counter
                    st.session_state.seed_counter += 1
                    st.session_state.question = generate_similar(question, seed=seed, difficulty="Similar")
                    reset_current_question()
                    st.rerun()
            with cnext2:
                if st.button("Generate a stretch question", use_container_width=True):
                    seed = int(datetime.now().timestamp() * 1000) + st.session_state.seed_counter
                    st.session_state.seed_counter += 1
                    st.session_state.question = generate_similar(question, seed=seed, difficulty="Stretch")
                    reset_current_question()
                    st.rerun()

    with student_ask_tab:
        st.subheader("💬 Ask Math Advisor")
        st.caption(
            "Ask a Mathematics question by typing it with the math keyboard, writing it with an iPad/Apple Pencil, or uploading/taking a picture. "
            "Math Advisor gives hints first; the full worked solution is shown only when you choose it."
        )

        with st.container(border=True):
            st.markdown("#### Enter your question")
            student_ask_prose = st.text_area(
                "Question text",
                key="student_ask_prose",
                height=120,
                placeholder="Example: Find the value of x, giving your answer to 3 significant figures.",
            )
            student_ask_math = question_math_editor(
                "Question mathematics",
                key="student_ask_math_keyboard",
            )

            st.markdown("#### Or write the question by hand")
            st.caption(
                "Use an iPad/Apple Pencil or another stylus. "
                "Save the handwriting in the pad before pressing Get hints."
            )
            student_ask_handwriting = handwriting_pad(
                key=f"student_ask_handwriting_{int(st.session_state.get('student_ask_handwriting_version', 0))}"
            )

            st.markdown("#### Or add a picture")
            student_ask_camera = st.camera_input(
                "Take a picture of the question",
                key="student_ask_camera",
            )
            student_ask_uploads = st.file_uploader(
                "Upload question image / screenshot / PDF",
                type=["png", "jpg", "jpeg", "webp", "pdf"],
                accept_multiple_files=True,
                key="student_ask_uploads",
            )

        student_ask_question = _student_ask_question_text(
            student_ask_prose,
            [student_ask_math] if student_ask_math else [],
        )

        browser_question_files = list(student_ask_uploads or [])
        if student_ask_camera is not None:
            browser_question_files.insert(0, student_ask_camera)
        student_ask_handwriting_asset = student_ask_handwriting

        if student_ask_question:
            st.markdown("#### Question preview")
            render_mathio_mixed(student_ask_question)

        # Useful mathematical tools remain available, but do not distract from hint-first tutoring.
        geogebra_external_tools(
            question_text=student_ask_question,
            key_base="student_ask_geogebra",
        )

        st.markdown("#### Ask for help")
        st.caption(
            "Typed, handwritten and uploaded questions are sent to Gemini only when you press **Get hints**. "
            "Avoid including names or other personal information in the picture."
        )

        ask_ready = bool(student_ask_question.strip() or browser_question_files or student_ask_handwriting_asset)

        if st.button(
            "Get hints",
            key="student_ask_get_hints",
            type="primary",
            use_container_width=True,
            disabled=not ask_ready,
        ):
            _student_ask_reset()
            try:
                student_assets = uploaded_assets(browser_question_files)
                if student_ask_handwriting_asset is not None:
                    student_assets.insert(0, student_ask_handwriting_asset)
                explicit_key = get_api_key()
                model = DEFAULT_MODEL

                # Use the selected app/exam track if available; otherwise provide a neutral SEC label.
                student_track_label = str(
                    st.session_state.get("track_label")
                    or st.session_state.get("selected_track_label")
                    or "Singapore SEC Mathematics"
                )

                with st.spinner("Reading the question and preparing hints..."):
                    verification = verify_question_math(
                        track_label=student_track_label,
                        question_text=student_ask_question,
                        question_assets=student_assets,
                        api_key=explicit_key,
                        model=model,
                    )
                    guided = generate_guided_solution(
                        track_label=student_track_label,
                        question_text=student_ask_question,
                        question_assets=student_assets,
                        api_key=explicit_key,
                        model=model,
                        verification=verification,
                    )

                st.session_state.student_ask_guided_solution = guided
                st.session_state.student_ask_hint_count = 0
                st.session_state.student_ask_show_full_solution = False
                st.session_state.student_ask_reveal_step = 0
                st.rerun()

            except GeminiTutorError as exc:
                st.session_state.student_ask_error = str(exc)
                st.rerun()
            except Exception as exc:
                st.session_state.student_ask_error = (
                    "Math Advisor could not process this question. "
                    f"{type(exc).__name__}: {exc}"
                )
                st.rerun()

        if st.session_state.get("student_ask_error"):
            st.error(st.session_state.student_ask_error)

        student_guided = st.session_state.get("student_ask_guided_solution")
        if student_guided is not None:
            render_student_ask_guided_solution(student_guided)

            if _is_construction_question(student_ask_question):
                show_construction_animation(
                    student_ask_question,
                    key_base="student_ask_construction",
                )

            st.markdown("---")
            if st.button(
                "Ask a new question",
                key="student_ask_new_question",
                use_container_width=True,
            ):
                _student_ask_reset()
                st.session_state.student_ask_handwriting_version = int(
                    st.session_state.get("student_ask_handwriting_version", 0)
                ) + 1
                for key in (
                    "student_ask_prose",
                    "student_ask_camera",
                    "student_ask_uploads",
                    "student_ask_math_keyboard__saved_payload",
                ):
                    st.session_state.pop(key, None)
                st.rerun()



    with student_whiteboard_tab:
        st.subheader("📝 Lesson whiteboard")
        st.caption(
            "Save notes and pictures throughout the lesson. "
            "Saved items stay in this browser session until you download the collection."
        )

        note_text = st.text_area(
            "Write lesson notes",
            key="student_note_draft",
            height=170,
            placeholder="Type your lesson notes here…",
        )
        if st.button("💾 Save note", key="student_save_note", use_container_width=True):
            if note_text.strip():
                _student_notes().append({"kind": "text", "content": note_text.strip()})
                st.session_state.pop("student_note_draft", None)
                st.rerun()

        st.markdown("#### Handwrite lesson notes")
        st.caption(
            "Use an iPad/Apple Pencil or another stylus. "
            "Save the handwriting in the pad, then add it to your lesson notes."
        )
        handwriting_version = int(st.session_state.get("student_lesson_handwriting_version", 0))
        student_lesson_handwriting = handwriting_pad(
            key=f"student_lesson_handwriting_{handwriting_version}"
        )
        handwriting_caption = st.text_input(
            "Handwriting caption (optional)",
            key=f"student_lesson_handwriting_caption_{handwriting_version}",
        )
        if st.button(
            "💾 Save handwritten note",
            key=f"student_save_handwriting_{handwriting_version}",
            use_container_width=True,
        ):
            if student_lesson_handwriting is None:
                st.warning("Write and save something in the handwriting pad first.")
            else:
                handwriting_png = _normalise_student_image_bytes(student_lesson_handwriting.data)
                if handwriting_png is None:
                    st.error("The handwritten note could not be saved. Please try again.")
                else:
                    _student_notes().append(
                        {
                            "kind": "image",
                            "content": handwriting_png,
                            "caption": handwriting_caption.strip() or "Handwritten lesson note",
                            "source_name": "ipad-handwritten-note.png",
                        }
                    )
                    st.session_state.student_lesson_handwriting_version = handwriting_version + 1
                    st.rerun()

        st.markdown("#### Add a picture")
        picture_version = _student_picture_input_version()
        photo = st.camera_input(
            "Take a picture",
            key=f"student_camera_{picture_version}",
        )
        uploaded = st.file_uploader(
            "Or upload a picture / screenshot",
            type=["png", "jpg", "jpeg", "webp"],
            key=f"student_picture_{picture_version}",
        )
        picture_caption = st.text_input(
            "Picture caption (optional)",
            key=f"student_picture_caption_{picture_version}",
        )
        if st.button("💾 Save picture", key="student_save_picture", use_container_width=True):
            # Prefer an explicitly uploaded screenshot over an older camera capture.
            chosen = uploaded if uploaded is not None else photo
            if chosen is None:
                st.warning("Take a picture or upload an image before saving.")
            else:
                original_bytes = chosen.getvalue()
                image_bytes = _normalise_student_image_bytes(original_bytes)
                if image_bytes is None:
                    st.error(
                        "The selected file could not be read as an image. "
                        "Please retake the picture or choose another PNG/JPG/WebP image."
                    )
                else:
                    _student_notes().append(
                        {
                            "kind": "image",
                            "content": image_bytes,
                            "caption": picture_caption.strip(),
                            "source_name": str(getattr(chosen, "name", "") or ""),
                        }
                    )
                    _reset_student_picture_inputs()
                    st.success("Picture saved to the lesson notes.")
                    st.rerun()


        st.markdown("#### Working tools")
        geogebra_external_tools(
            question_text="",
            key_base="student_whiteboard_geogebra",
        )
        student_scientific_calculator(
            key_base="student_whiteboard_calculator"
        )
        st.info(
            "GeoGebra opens in a separate tab. Export your GeoGebra drawing as an image, "
            "then upload and save that image above so it becomes part of your lesson notes."
        )

        st.markdown("#### Saved lesson notes")
        if not _student_notes():
            st.caption("No saved notes yet.")
        else:
            for index, item in enumerate(_student_notes(), 1):
                with st.container(border=True):
                    st.markdown(f"**Item {index}**")
                    if item.get("kind") == "text":
                        render_guidance_mixed_mathio(str(item.get("content", "")))
                    elif item.get("kind") == "image":
                        st.image(item.get("content", b""), width=420)
                        if item.get("caption"):
                            st.caption(str(item["caption"]))

        if _student_notes():
            st.markdown("#### Download lesson notes")
            st.caption(
                "Download a copy when you are ready. Saved notes are kept on screen "
                "until you choose **Clear saved notes**."
            )

            export_error = ""
            try:
                docx_data = _student_notes_docx()
            except Exception as exc:
                docx_data = None
                export_error = f"Word export could not be prepared: {type(exc).__name__}: {exc}"

            try:
                pdf_data = _student_notes_pdf()
            except Exception as exc:
                pdf_data = None
                if not export_error:
                    export_error = f"PDF export could not be prepared: {type(exc).__name__}: {exc}"

            if export_error:
                st.warning(export_error)

            d1, d2 = st.columns(2)
            with d1:
                if docx_data is not None:
                    st.download_button(
                        "⬇️ Download notes as Word",
                        data=docx_data,
                        file_name=f"{_student_download_basename()}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key="student_notes_docx_download",
                    )
                else:
                    st.button(
                        "Word download unavailable",
                        disabled=True,
                        use_container_width=True,
                        key="student_notes_docx_unavailable",
                    )

            with d2:
                if pdf_data is not None:
                    st.download_button(
                        "⬇️ Download notes as PDF",
                        data=pdf_data,
                        file_name=f"{_student_download_basename()}.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                        key="student_notes_pdf_download",
                    )
                else:
                    st.button(
                        "PDF download unavailable",
                        disabled=True,
                        use_container_width=True,
                        key="student_notes_pdf_unavailable",
                    )

            if st.button(
                "🗑️ Clear saved notes",
                key="student_notes_clear_saved",
                use_container_width=True,
            ):
                _clear_student_notes_after_download()
                st.rerun()

st.caption(
        f"Educational tool, not an official SEAB/MOE product. Gemini default model: {DEFAULT_MODEL}. "
        "Generated questions are original and are not past-year examination questions."
    )
